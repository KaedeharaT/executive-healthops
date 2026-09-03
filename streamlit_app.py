"""Executive Chronic Care HealthOps workbench for the human-reviewed V0.2 demo."""

from __future__ import annotations

import logging
import html
import os
from contextlib import contextmanager
from datetime import date, datetime, time, timedelta
from pathlib import Path
from time import perf_counter
from uuid import UUID, uuid4

import pandas as pd
import altair as alt
import streamlit as st
from sqlalchemy import func, select
from sqlalchemy.exc import SQLAlchemyError

from executive_health_ai.ai.doctor_brief_agent import build_doctor_brief
from executive_health_ai.blood_pressure import TOKYO_TIMEZONE, build_blood_pressure_records
from executive_health_ai.database import SessionLocal
from executive_health_ai.models import (
    Alert, AnnualHealthAccount, AuditLog, Document, DoctorReview, ExecutionBarrier, FollowUp,
    HealthEvent, HealthJourney, HealthProblem, HealthProgram, ManagementPlan, MedicationEvent,
    ExternalIdentity, IngestionJob, KnowledgeChunk, KnowledgeDocument, KnowledgeReviewAudit, KnowledgeSourceRegistry, KnowledgeUseRecord, RiskEvent, RiskRule, EmergencyContact, MedicationPlan, Observation, OutcomeEvaluation, Patient, RawIngestionRecord, ProgramPhase, ReportExtractionCandidate, ReportExtractionRun, Task, WeeklyReview, SleepSession,
    ExternalReferral, HealthAssessment, ManagementRule, ManagementSignal, MemberDeviceAssignment,
    ServiceCatalogItem, ServicePlan, MemberEntitlement, ServiceRequest,
)
from executive_health_ai.services.timeline import build_patient_timeline
from executive_health_ai.services.risk_triage import RiskEvaluationService
from executive_health_ai.services.risk_operations import RiskOperationsService
from executive_health_ai.integrations.service import ingest, manually_correct_record
from executive_health_ai.ui.localization.zh_cn import (
    BARRIER, OBSERVATION, PRIORITY, PROVIDER, STATUS, TYPE, display_datetime,
    observation as display_observation, priority as display_priority, provider as display_provider,
    status as display_status, type_label, program_type as display_program_type, program_phase as display_program_phase,
    risk_level as display_risk_level, device_class as display_device_class,
    knowledge_category as display_knowledge_category, knowledge_review_status as display_knowledge_review_status,
)
from executive_health_ai.ui.display import (
    get_audit_action_display, get_entity_type_display, get_event_type_display,
    get_provider_display, get_quality_display, get_risk_display, get_role_display,
    get_source_type_display, get_status_display, humanize_source_name,
)
from executive_health_ai.ui.pages.shell import render_more_workspace_shell, render_portfolio_landing
from executive_health_ai.services.knowledge import KnowledgeService
from executive_health_ai.services.knowledge_retrieval import KnowledgeRetrievalService
from executive_health_ai.services.knowledge_sources import KnowledgeProviderError
from executive_health_ai.services.health_data_summary import HealthDataSummaryService
from executive_health_ai.services.operational_worklist import OperationalWorklistService
from executive_health_ai.services.task_transitions import TaskTransitionService
from executive_health_ai.services.longitudinal import (
    HealthAssessmentService, HealthDataCategoryRegistry, HealthTimelineService,
    TimelineV4Service, TimelineViewport, InterventionOutcomeService, ReportComparisonService, ReportRiskSummaryService, OversightRiskSummaryService,
)
from executive_health_ai.services.report_parsing import ReportParseProgress, ReportParsingService
from executive_health_ai.services.member_services import MemberServiceOperations
from executive_health_ai.services.chronic_care import apply_outcome_decision, complete_outcome_doctor_review
from executive_health_ai.services.workflow import (
    close_alert_as_false_positive, complete_follow_up, confirm_alert_as_manager,
    create_operational_task, record_doctor_review,
)


st.set_page_config(page_title="企业高管健康运营中心", page_icon="🩺", layout="wide")
LOGGER = logging.getLogger(__name__)
NAVIGATION_PROFILE_ENABLED = os.getenv("HEALTHOPS_PROFILE_NAV", "").lower() in {"1", "true", "yes"}
PORTFOLIO_DEMO_ENABLED = os.getenv("PORTFOLIO_DEMO", "").lower() in {"1", "true", "yes"}
TECHNICAL_DETAILS_ENABLED = (
    not PORTFOLIO_DEMO_ENABLED
    and os.getenv("HEALTHOPS_TECHNICAL_DETAILS", "").lower() in {"1", "true", "yes"}
)

DATA_STATUS_LABELS = {
    "normal": "数据完整，可进行趋势分析",
    "insufficient_data": "数据不足，暂无法判断趋势",
    "needs_remeasurement": "建议重新测量",
    "needs_clinician_review": "存在数据冲突，建议人工复核",
    "abstain": "证据不足，暂不判断",
}

STATUS_LABELS = {
    "NEW": "新建", "AI_SCREENED": "规则已筛查", "WAITING_MANAGER_REVIEW": "待管理师核实",
    "MANAGER_CONFIRMED": "管理师已确认", "WAITING_DOCTOR_REVIEW": "待医生复核",
    "IN_FOLLOW_UP": "随访中", "CLOSED": "已闭环", "OPEN": "处理中",
    "ACTIVE": "执行中", "PENDING": "待执行", "IN_PROGRESS": "进行中",
    "COMPLETED": "已完成", "CANCELLED": "已取消", "OVERDUE": "已逾期",
    "ACKNOWLEDGED": "已接手", "IN_REVIEW": "处理中", "MONITORING": "处理中", "WAITING_MEMBER": "等待成员",
    "ESCALATED_TO_DOCTOR": "等待医生", "FOLLOW_UP": "待随访", "DISMISSED_DATA_ISSUE": "数据问题已关闭",
    "CONFIRMED": "已确认", "ASSESSMENT": "筛查评估", "90_DAY_PROGRAM": "90天阶段管理",
    "STABILIZATION": "半年稳定管理", "ANNUAL_MANAGEMENT": "年度健康账户",
    "FAMILY_EXTENSION": "家庭健康管理（占位）", "PLANNED": "已规划", "PAUSED": "已暂停",
    "NEEDS_REASSESSMENT": "需要复评", "ESCALATED_TO_MEDICAL_CARE": "医疗优先",
    "ADJUSTED": "已调整", "RESOLVED": "已解决", "IMPROVED": "改善",
    "STABLE": "稳定", "WORSENED": "恶化", "INSUFFICIENT_DATA": "数据不足",
    "NEEDS_MEDICAL_REVIEW": "需要医生复核",
}
STATUS_LABELS.update(STATUS)
SEVERITY_LABELS = PRIORITY
TYPE_LABELS = TYPE
ROUTE_LABELS = {"SELF_MANAGEMENT": "成员持续管理", "HEALTH_MANAGER": "健康管理师处理", "INTERNAL_DOCTOR": "内部医生复核", "EXTERNAL_DOCTOR": "外部医生协同", "EMERGENCY_MANUAL_ACTION": "紧急人工处置"}
PENDING_NAVIGATION_KEY = "_pending_navigation"


def _inject_style() -> None:
    """Stable, low-noise visual tokens shared by both product surfaces."""
    st.markdown(
        """
        <style>
        :root {--ink:#18263a;--muted:#68778a;--faint:#95a2b2;--line:#dfe6ed;--canvas:#f3f5f7;--card:#fff;--blue:#205c9e;--blue-hover:#174b83;--blue-soft:#e8f1fa;--teal:#167b78;--green:#2c7a5f;--amber:#a76513;--red:#b74747;--radius:14px;}
        html, body, [class*="css"] {font-family:"Segoe UI","Microsoft YaHei",system-ui,sans-serif;}
        [data-testid="stAppViewContainer"] {background:var(--canvas); color:var(--ink);}
        [data-testid="stAppViewContainer"] .main .block-container {
            max-width:1260px !important; margin:0 auto; padding:2.4rem 2rem 5.5rem;
        }
        h1 {font-size:1.95rem !important; line-height:1.2; letter-spacing:-.045em; margin:0 0 .45rem !important; font-weight:720 !important; color:var(--ink);}
        h2 {font-size:1.26rem !important; line-height:1.3; letter-spacing:-.022em; margin:2.35rem 0 .55rem !important; font-weight:700 !important; color:var(--ink);}
        h3 {font-size:1rem !important; line-height:1.4; letter-spacing:-.012em; margin:.4rem 0 !important;}
        [data-testid="stCaptionContainer"] {color:var(--muted); font-size:.84rem; line-height:1.55;}
        [data-testid="stMetric"] {background:transparent; border:0; padding:.2rem 0;}
        [data-testid="stMetricLabel"] {font-size:.76rem; color:var(--muted); font-weight:600;}
        [data-testid="stMetricValue"] {font-size:1.55rem; font-weight:720; letter-spacing:-.035em; color:var(--ink);}
        [data-testid="stVerticalBlockBorderWrapper"] {border:1px solid var(--line) !important; border-radius:var(--radius) !important; box-shadow:0 2px 7px rgba(24,38,58,.035); background:var(--card);}
        [data-testid="stVerticalBlockBorderWrapper"] > div {padding:.8rem .85rem;}
        .section-frame-title {font-size:1.08rem; font-weight:720; letter-spacing:-.018em; margin:.15rem 0 .18rem; color:var(--ink);}
        [data-testid="stExpander"] {border:1px solid var(--line) !important; border-radius:12px !important; background:#fff; margin:.5rem 0;}
        [data-testid="stDivider"] {margin:1.75rem 0 !important; border-color:var(--line);}
        .stButton > button {border-radius:9px; min-height:2.35rem; font-size:.88rem; font-weight:650; border-color:#cad5e1; background:#fff; color:var(--ink);}
        .stButton > button:hover {border-color:#9eb6cf; color:var(--blue); background:#f9fbfd;}
        .stButton > button[kind="primary"] {background:var(--blue); border-color:var(--blue); color:#fff;}
        .stButton > button[kind="primary"]:hover {background:var(--blue-hover); border-color:var(--blue-hover); color:#fff;}
        [data-testid="stDataFrame"] {border:1px solid var(--line); border-radius:12px; overflow:hidden; background:#fff;}
        [data-testid="stFileUploader"] {border:1px dashed #b9c8d7; border-radius:12px; padding:.4rem; background:#fbfcfd;}
        [data-testid="stSidebar"] {background:#fff; border-right:1px solid var(--line);}
        [data-testid="stSidebar"] > div:first-child {padding:1.35rem .7rem 2rem;}
        [data-testid="stSidebar"] h2 {font-size:1.08rem !important; margin:.25rem .65rem .15rem !important;}
        [data-testid="stSidebar"] [data-testid="stRadio"] label {padding:.67rem .75rem; margin:.14rem 0; border-radius:9px; transition:background .12s ease, color .12s ease;}
        [data-testid="stSidebar"] [data-testid="stRadio"] label:hover {background:#f2f6f9;}
        [data-testid="stSidebar"] [data-testid="stRadio"] label:has(input:checked) {background:var(--blue-soft); color:var(--blue); font-weight:700;}
        [data-testid="stSidebar"] [data-testid="stRadio"] div[role="radiogroup"][aria-orientation="horizontal"] {background:#edf1f5;padding:3px;border-radius:10px;gap:2px;}
        [data-testid="stSidebar"] [data-testid="stRadio"] div[role="radiogroup"][aria-orientation="horizontal"] label {flex:1;text-align:center;padding:.45rem .35rem;}
        [data-testid="stSidebar"] [data-testid="stRadio"] div[role="radiogroup"][aria-orientation="horizontal"] label:has(input:checked) {background:#fff;box-shadow:0 1px 3px rgba(24,38,58,.10);}
        [data-testid="stSidebar"] [data-testid="stRadio"] label > div:first-child {display:none;}
        [data-testid="stSidebar"] [data-testid="stRadio"] label > div:last-child {padding-left:0;}
        [data-testid="stRadio"] div[role="radiogroup"] {gap:.35rem;}
        [data-testid="stMain"] [data-testid="stRadio"] div[role="radiogroup"] {background:#edf1f5;padding:3px;border-radius:10px;gap:2px;}
        [data-testid="stMain"] [data-testid="stRadio"] label {border-radius:7px;padding:.38rem .7rem;margin:0;min-height:2rem;}
        [data-testid="stMain"] [data-testid="stRadio"] label > div:first-child {display:none;}
        [data-testid="stMain"] [data-testid="stRadio"] label > div:last-child {padding-left:0;}
        [data-testid="stMain"] [data-testid="stRadio"] label:has(input:checked) {background:#fff;box-shadow:0 1px 3px rgba(24,38,58,.10);color:var(--blue);font-weight:700;}
        [data-testid="stSegmentedControl"] {background:#edf1f5; border-radius:10px; padding:3px; width:100%;}
        [data-testid="stSegmentedControl"] button {border-radius:7px !important; font-size:.83rem !important; font-weight:650 !important;}
        .surface-label {font-size:.7rem; font-weight:750; color:var(--faint); letter-spacing:.08em; text-transform:uppercase; margin:.2rem .65rem .45rem;}
        .page-header {margin:0 0 1.35rem; max-width:740px;}
        .page-header .eyebrow {font-size:.71rem; font-weight:750; color:var(--blue); letter-spacing:.08em; margin-bottom:.38rem; text-transform:uppercase;}
        .page-header + h1 {margin-top:0 !important;}.page-header + h1 + [data-testid="stCaptionContainer"] {font-size:.93rem; max-width:650px; margin-bottom:1.25rem;}
        .status-strip {display:flex; gap:0; background:var(--card); border:1px solid var(--line); border-radius:12px; overflow:hidden; margin:.35rem 0 1.8rem; box-shadow:0 2px 7px rgba(24,38,58,.025);}
        .status-strip > div {flex:1; padding:.82rem .95rem; border-right:1px solid var(--line);}
        .status-strip > div:last-child {border-right:0;}
        .status-strip b {font-size:1.32rem; display:block; color:var(--ink); letter-spacing:-.03em;}
        .status-strip span {font-size:.74rem; color:var(--muted);}
        .status-strip .urgent b {color:var(--red);} .status-strip .attention b {color:var(--amber);} .status-strip .action b {color:var(--blue);} .status-strip .neutral b {color:var(--teal);}
        .section-kicker {font-size:.74rem; color:var(--muted); margin-bottom:.25rem;}
        .member-hero,.client-hero {background:var(--card); border:1px solid var(--line); border-radius:16px; padding:1.5rem 1.65rem; margin:.15rem 0 1.25rem; box-shadow:0 2px 8px rgba(24,38,58,.035);}
        .member-hero {border-left:4px solid var(--blue);}.client-hero {border-left:4px solid var(--teal);}
        .member-hero h1,.client-hero h1 {font-size:1.78rem !important; margin:0 !important;}.member-hero p,.client-hero p{margin:.24rem 0 .65rem;color:var(--muted);}
        .hero-facts{display:flex;gap:1.5rem;flex-wrap:wrap;margin-top:1.15rem;padding-top:1rem;border-top:1px solid var(--line);}.hero-fact{min-width:100px;}.hero-fact b{display:block;font-size:1.08rem;color:var(--ink);letter-spacing:-.02em;}.hero-fact span{display:block;font-size:.74rem;color:var(--muted);margin-top:.15rem;}
        .quiet-list {list-style:none; margin:0; padding:0;}.quiet-list li {padding:.68rem 0; border-bottom:1px solid var(--line);}.quiet-list li:last-child{border-bottom:0;}
        .empty-state {text-align:center; max-width:440px; margin:1.1rem auto; padding:1.5rem; color:var(--muted); background:#fafbfd; border:1px dashed #cbd7e2; border-radius:12px;}
        .empty-state strong {display:block; color:var(--ink); font-size:1rem; margin-bottom:.35rem;}
        .summary-note {font-size:.83rem; color:var(--muted); margin:.15rem 0 .75rem;}
        .status-badge {display:inline-flex;align-items:center;border-radius:999px;padding:.24rem .58rem;font-size:.75rem;font-weight:700;line-height:1.2;background:#eef2f5;color:#516174;}.status-badge.urgent{background:#fbecec;color:#aa3838;}.status-badge.attention{background:#fff4df;color:#936019;}.status-badge.action{background:#e8f1fa;color:#205c9e;}.status-badge.stable{background:#e7f4ef;color:#266d55;}.status-badge.neutral{background:#eef2f5;color:#516174;}
        .metric-tile {padding:1rem 1.05rem;border:1px solid var(--line);border-radius:12px;background:#fff;min-height:104px;}.metric-tile .label{font-size:.78rem;font-weight:650;color:var(--muted);}.metric-tile .value{font-size:1.48rem;font-weight:720;letter-spacing:-.03em;color:var(--ink);margin:.28rem 0 .12rem;}.metric-tile .note{font-size:.76rem;color:var(--muted);}
        .entry-card {border:1px solid var(--line);border-radius:14px;background:#fff;padding:1.05rem;min-height:142px;box-shadow:0 2px 7px rgba(24,38,58,.025);}.entry-card .entry-title{font-weight:720;font-size:1rem;color:var(--ink);}.entry-card .entry-value{font-size:.8rem;color:var(--blue);font-weight:650;margin:.55rem 0 .2rem;}.entry-card .entry-copy{font-size:.8rem;line-height:1.55;color:var(--muted);}
        .work-item {border:1px solid var(--line);border-left:4px solid var(--blue);border-radius:12px;background:#fff;padding:1rem 1.05rem;margin:.55rem 0;}.work-item .work-member{font-size:.86rem;font-weight:720;color:var(--ink);}.work-item .work-title{font-size:1.02rem;font-weight:720;color:var(--ink);margin:.6rem 0;}.work-item .work-label{font-size:.73rem;font-weight:720;color:var(--muted);margin-bottom:.1rem;}.work-item .work-copy{font-size:.84rem;color:#46576b;line-height:1.45;}
        .member-card {border:1px solid var(--line);border-radius:14px;background:#fff;padding:1.1rem;min-height:255px;box-shadow:0 2px 8px rgba(24,38,58,.03);}.member-card .member-name{font-size:1.12rem;font-weight:730;color:var(--ink);}.member-card .member-meta{font-size:.8rem;color:var(--muted);margin:.24rem 0 1rem;}.member-card .member-label{font-size:.72rem;font-weight:720;color:var(--muted);margin-top:.68rem;}.member-card .member-value{font-size:.9rem;color:var(--ink);margin-top:.12rem;}
        .detail-panel {background:#fff;border:1px solid var(--line);border-radius:14px;padding:1.1rem;margin:.45rem 0;}
        .portfolio-landing {max-width:760px;margin:8vh auto 0;padding:2.1rem 2.2rem;background:#fff;border:1px solid var(--line);border-radius:18px;box-shadow:0 8px 24px rgba(24,38,58,.06);}
        .portfolio-landing .portfolio-kicker{font-size:.74rem;font-weight:760;letter-spacing:.1em;text-transform:uppercase;color:var(--teal);}.portfolio-landing h1{font-size:2.2rem !important;margin:.45rem 0 .7rem !important;}.portfolio-landing p{max-width:625px;color:var(--muted);line-height:1.75;margin:0 0 1.5rem;}
        .focus-row{display:grid;grid-template-columns:30px 1fr auto;gap:.75rem;align-items:start;padding:.8rem 0;border-bottom:1px solid var(--line);}.focus-row:last-child{border-bottom:0;}.focus-index{font-size:.75rem;font-weight:760;color:var(--blue);padding-top:.14rem;}.focus-title{font-size:.94rem;font-weight:720;color:var(--ink);}.focus-copy{font-size:.79rem;color:var(--muted);margin-top:.15rem;}.next-row{padding:.75rem 0;border-bottom:1px solid var(--line);}.next-row:last-child{border-bottom:0;}.next-date{font-size:.74rem;color:var(--blue);font-weight:720;}.timeline-preview{display:grid;grid-template-columns:92px 1fr;gap:.7rem;padding:.6rem 0;border-bottom:1px solid var(--line);}.timeline-preview:last-child{border-bottom:0;}.timeline-date{font-size:.77rem;color:var(--muted);font-weight:650;}.timeline-title{font-size:.87rem;color:var(--ink);font-weight:680;}.timeline-copy{font-size:.78rem;color:var(--muted);margin-top:.15rem;}
        @media (max-width: 900px) {[data-testid="stAppViewContainer"] .main .block-container{padding:1.35rem 1rem 3rem;} .status-strip{flex-wrap:wrap;}.status-strip > div{min-width:45%;}.member-hero,.client-hero{padding:1.2rem;}}
        </style>
        """,
        unsafe_allow_html=True,
    )


def _fmt_dt(value: datetime | None) -> str:
    if value is None:
        return "未记录"
    return display_datetime(value)


def _age(patient: Patient) -> str:
    if patient.birth_date is None:
        return "年龄未录入"
    today = date.today()
    years = today.year - patient.birth_date.year - ((today.month, today.day) < (patient.birth_date.month, patient.birth_date.day))
    return f"{years} 岁"


def _label(value: str | None, *, context: str | None = None) -> str:
    return get_status_display(value, context=context)


def _role_label(value: str | None, *, name: str | None = None) -> str:
    return get_role_display(value, name=name)


def _severity(value: str | None) -> str:
    return display_priority(value)


def _status_badge(status: str, severity: str | None = None) -> str:
    return f"【{_label(status)}】"


def _severity_badge(severity: str) -> str:
    return f"【{_severity(severity)}优先级】"


def _status_pill(status: str) -> str:
    """Backward-compatible alias for the shared visual status component."""
    return status_badge(status)


def _risk_text(level: str) -> str:
    """Use one business-facing label for risk states; never expose enums in UI."""
    return get_risk_display(level)


def _metric_display_name(metric_code: str | None, raw_name: str | None = None) -> str:
    """Return a member-facing metric label without leaking canonical placeholders.

    The raw parser and device layers may legitimately retain an unmapped code.
    That is useful provenance, but it is not a useful label in a normal health
    screen.  Keep any human source name when it is available; otherwise use the
    central registry's neutral fallback and log the mapping gap without PHI.
    """
    code = (metric_code or "").strip()
    source_name = (raw_name or "").strip()
    placeholder_values = {"", "unknown", "none", "null", "unmapped", "other"}
    if code.lower() not in placeholder_values and code in OBSERVATION:
        return display_observation(code)
    if source_name and source_name.lower() not in placeholder_values:
        return source_name
    if code.lower() not in placeholder_values:
        LOGGER.warning("ui_unmapped_metric_display metric_code=%s", code)
    return "健康数据"


_TECHNICAL_DETAIL_FIELDS = {
    "id", "uuid", "canonical_code", "provider_code", "external_id", "content_hash",
    "document_id", "member_id", "patient_id", "project_id", "risk_event_id", "review_id",
    "service_request_id", "chunk_id", "source_record_id", "raw_record_id", "ingestion_job_id",
    "raw_payload", "payload_json", "metadata_json", "raw_metadata", "trace_id", "run_id",
    "parser_run_id", "extraction_run_id", "evaluation_run_id", "source_id", "entity_id",
}

_BUSINESS_DETAIL_LABELS = {
    "title": "名称", "name": "名称", "label": "名称", "summary": "摘要",
    "description": "说明", "metric": "指标", "value": "数值", "unit": "单位",
    "status": "状态", "quality": "数据状态", "quality_flag": "数据状态",
    "provider": "数据来源", "source": "来源", "source_type": "来源类型", "observed_at": "记录时间",
    "recorded_at": "记录时间", "date": "日期", "reference_range": "参考范围",
    "abnormal_flag": "异常提示", "department": "科室", "organization": "机构",
    "owner": "负责人", "responsible_role": "负责人", "event_type": "记录类型",
}


def _is_technical_detail_field(key: object) -> bool:
    normalized = str(key).strip().lower()
    return (
        normalized in _TECHNICAL_DETAIL_FIELDS
        or normalized.endswith(("_id", "_uuid", "_record_id", "_job_id", "_run_id"))
        or normalized.startswith(("raw_", "canonical_", "ingestion_", "metadata_"))
    )


def _is_uuid_value(value: object) -> bool:
    if isinstance(value, UUID):
        return True
    try:
        UUID(str(value))
    except (TypeError, ValueError, AttributeError):
        return False
    return True


def _business_detail_row(value: dict[object, object]) -> dict[str, object]:
    """Prepare ad-hoc snapshot data for UI without accidentally dumping provenance."""
    result: dict[str, object] = {}
    for key, item in value.items():
        normalized = str(key).strip().lower()
        if _is_technical_detail_field(key) or _is_uuid_value(item):
            continue
        if item is None or (isinstance(item, str) and item.strip().lower() in {"", "none", "null", "unknown", "n/a"}):
            display_value: object = "未记录"
        elif normalized in {"provider", "source_system"}:
            display_value = get_provider_display(str(item))
        elif normalized == "source_type":
            display_value = get_source_type_display(str(item))
        elif normalized in {"status", "quality", "quality_flag"}:
            display_value = get_quality_display(str(item)) if "quality" in normalized else _label(str(item))
        elif normalized in {"owner", "responsible_role"}:
            display_value = _role_label(str(item))
        elif normalized == "event_type":
            display_value = get_event_type_display(str(item))
        elif normalized in {"metric", "metric_code"}:
            display_value = _metric_display_name(str(item))
        elif isinstance(item, list):
            display_value = "、".join(str(part) for part in item if not _is_uuid_value(part)) or "未记录"
        elif isinstance(item, dict):
            display_value = "；".join(f"{label}：{content}" for label, content in _business_detail_row(item).items()) or "未记录"
        else:
            display_value = item
        label = _BUSINESS_DETAIL_LABELS.get(normalized, str(key).replace("_", " "))
        result[label] = display_value
    return result


def _timeline_event_badge(event) -> str:
    """A quiet category label for normal timeline events, never a risk light."""
    label = event.event_type_label or "健康记录"
    return f"<span class='timeline-event-badge'>{html.escape(label)}</span>"


def _timeline_risk_indicator(event) -> str:
    """Accessible traffic-light markup used only by formal risk events."""
    if event.risk_indicator == "TRAFFIC_LIGHT" and event.risk_level:
        level = event.risk_level.lower()
        return (
            f"<span class='timeline-risk-indicator timeline-risk-{html.escape(level)}'>"
            f"<span class='timeline-risk-dot'></span>{html.escape(event.risk_label or _risk_text(event.risk_level))}</span>"
        )
    if event.risk_indicator == "NEUTRAL":
        return "<span class='timeline-risk-indicator timeline-risk-neutral'><span class='timeline-risk-dot'></span>暂无正式风险评估</span>"
    return ""


def _active_program(ctx: dict[str, list[object]]) -> HealthProgram | None:
    return next((item for item in ctx["programs"] if item.status == "ACTIVE"), None)  # type: ignore[return-value]


def _program_day(program: HealthProgram | None) -> int | None:
    if program is None or program.program_type != "NINETY_DAY":
        return None
    return max(1, min((date.today() - program.start_date).days + 1, 90))


def _member_display(member: Patient | None) -> str:
    if member is None:
        return "未匹配成员"
    return member.display_name or "未命名成员"


def _render_timed(page_name: str, renderer) -> None:
    """Development-only timing log; it is never displayed in the product UI."""
    started = perf_counter()
    renderer()
    duration_ms = (perf_counter() - started) * 1000
    LOGGER.debug("ui_render page=%s duration_ms=%.1f", page_name, duration_ms)
    if NAVIGATION_PROFILE_ENABLED:
        LOGGER.warning("[PERF] renderer:%s %.1f ms", page_name, duration_ms)


def _render_sidebar_navigation() -> str:
    st.sidebar.markdown("## 健康管理平台")
    st.sidebar.caption("健康管理师工作区")
    st.sidebar.divider()
    legacy = {"工作台": "今日", "今日工作台": "今日", "数据设备": "更多", "协同": "医疗协同", "风险与医疗协同": "医疗协同"}
    selected = st.session_state.get("ops-navigation")
    if selected in legacy:
        st.session_state["ops-navigation"] = legacy[selected]
    return st.sidebar.radio(
        "工作区", ["今日", "成员", "医疗协同", "服务运营", "更多"],
        key="ops-navigation", label_visibility="collapsed",
    )


def _render_surface_switcher() -> str:
    """Development-only surface switcher; no authentication is implied."""
    st.sidebar.markdown("<div class='surface-label'>切换视图</div>", unsafe_allow_html=True)
    surface = st.sidebar.radio("当前视图", ["运营后台", "成员健康中心"], key="surface-mode", label_visibility="collapsed", horizontal=True)
    st.sidebar.caption("预览视图")
    st.sidebar.divider()
    return surface


def _render_member_center_navigation() -> str:
    st.sidebar.markdown("## 成员健康中心")
    st.sidebar.caption("看状态、健康资料、计划与服务")
    st.sidebar.divider()
    page = st.sidebar.radio(
        "成员健康中心导航", ["首页", "健康", "历程", "计划", "服务"],
        key="member-center-navigation", label_visibility="collapsed",
        on_change=lambda: st.session_state.pop("member-profile-open", None),
    )
    st.sidebar.divider()
    if st.sidebar.button("个人设置", key="member-profile-open-button", width="stretch"):
        st.session_state["member-profile-open"] = True
        st.rerun()
    return "个人设置" if st.session_state.get("member-profile-open") else page


def _render_portfolio_landing() -> None:
    """Compatibility wrapper for the extracted portfolio entry page."""
    render_portfolio_landing(request_navigation)


def _empty_state(title: str, guidance: str) -> None:
    """One calm, consistent empty state for normal product pages."""
    st.markdown(f"<div class='empty-state'><strong>{html.escape(title)}</strong><span>{html.escape(guidance)}</span></div>", unsafe_allow_html=True)


def empty_state(title: str, guidance: str) -> None:
    """Public design-system empty state for product-facing renderers."""
    _empty_state(title, guidance)


def _page_header(title: str, guidance: str, *, eyebrow: str | None = None) -> None:
    """Backward-compatible renderer alias for the shared page header."""
    page_header(title, guidance, eyebrow=eyebrow)


def page_header(title: str, guidance: str, *, eyebrow: str | None = None) -> None:
    """One calm, product-facing first screen for every primary renderer."""
    eyebrow_html = f"<div class='eyebrow'>{html.escape(eyebrow)}</div>" if eyebrow else ""
    st.markdown(f"<div class='page-header'>{eyebrow_html}</div>", unsafe_allow_html=True)
    st.title(title)
    st.caption(guidance)


def _section_header(title: str, guidance: str | None = None) -> None:
    st.subheader(title)
    if guidance:
        st.caption(guidance)


@contextmanager
def section_frame(title: str, guidance: str | None = None):
    """A deliberate white surface for one business decision, not every record."""
    with st.container(border=True):
        st.markdown(f"<div class='section-frame-title'>{html.escape(title)}</div>", unsafe_allow_html=True)
        if guidance:
            st.caption(guidance)
        yield


def summary_metric(label: str, value: str | int, note: str | None = None) -> None:
    note_html = f"<div class='note'>{html.escape(note)}</div>" if note else ""
    st.markdown(
        f"<div class='metric-tile'><div class='label'>{html.escape(label)}</div>"
        f"<div class='value'>{html.escape(str(value))}</div>{note_html}</div>",
        unsafe_allow_html=True,
    )


def status_badge(label: str) -> str:
    """Use traffic-light colour for risk only; workflow status remains neutral."""
    high_risk = {"高风险", "紧急风险"}
    medium_risk = {"中风险"}
    low_risk = {"低风险"}
    action = {"今日跟进", "等待医生", "处理中", "待安排", "进行中", "已批准", "已安排"}
    style = "urgent" if label in high_risk else "attention" if label in medium_risk else "stable" if label in low_risk else "action" if label in action else "neutral"
    return f"<span class='status-badge {style}'>{html.escape(label)}</span>"


def risk_badge(level: str | None) -> str:
    """Render the traffic-light treatment only for a formal risk conclusion."""
    labels = {
        "RED": "高风险", "HIGH": "高风险", "高风险": "高风险",
        "YELLOW": "中风险", "MEDIUM": "中风险", "中风险": "中风险",
        "GREEN": "低风险", "LOW": "低风险", "低风险": "低风险",
    }
    return status_badge(labels.get(str(level or "").upper(), "暂无正式风险评估"))


def health_metric_card(label: str, value: str, note: str = "") -> None:
    summary_metric(label, value, note)


def entry_card(title: str, value: str, description: str) -> None:
    """A lightweight doorway into a deeper health record, never a data wall."""
    st.markdown(
        f"<div class='entry-card'><div class='entry-title'>{html.escape(title)}</div>"
        f"<div class='entry-value'>{html.escape(value)}</div>"
        f"<div class='entry-copy'>{html.escape(description)}</div></div>",
        unsafe_allow_html=True,
    )


def primary_action(label: str, *, key: str, on_click=None, args: tuple = (), disabled: bool = False, width: str = "stretch") -> bool:
    return st.button(label, key=key, type="primary", on_click=on_click, args=args, disabled=disabled, width=width)


def secondary_action(label: str, *, key: str, on_click=None, args: tuple = (), disabled: bool = False, width: str = "stretch") -> bool:
    return st.button(label, key=key, on_click=on_click, args=args, disabled=disabled, width=width)


@contextmanager
def detail_panel(title: str | None = None, note: str | None = None):
    with st.container(border=True):
        if title:
            st.markdown(f"<div class='section-frame-title'>{html.escape(title)}</div>", unsafe_allow_html=True)
        if note:
            st.caption(note)
        yield


def work_item_card(member_name: str, status: str, title: str, reason: str, next_step: str, *, key: str, owner: str | None = None, due_at: datetime | None = None, on_click=None, args: tuple = ()) -> None:
    """A compact operational item with one decision and one action."""
    with st.container():
        content, action = st.columns([5, 1])
        with content:
            st.markdown(
                f"<div class='work-item'><div class='work-member'>{html.escape(member_name)}　{status_badge(status)}</div>"
                f"<div class='work-title'>{html.escape(title)}</div><div class='work-label'>最近变化</div>"
                f"<div class='work-copy'>{html.escape(reason)}</div><div class='work-label' style='margin-top:.6rem'>下一步</div>"
                f"<div class='work-copy'>{html.escape(next_step)}</div>"
                f"<div class='work-label' style='margin-top:.6rem'>负责人</div><div class='work-copy'>{html.escape(owner or '待分配')}</div>"
                f"<div class='work-label' style='margin-top:.6rem'>截止时间</div><div class='work-copy'>{html.escape(_fmt_dt(due_at) if due_at else '暂无截止时间')}</div></div>",
                unsafe_allow_html=True,
            )
        with action:
            if on_click:
                primary_action("处理", key=key, on_click=on_click, args=args)


def member_card(name: str, meta: str, status: str, managed: str, change: str, next_step: str, *, key: str, on_click=None, args: tuple = ()) -> None:
    """A stable two-column member-list card; no database-table presentation."""
    with st.container():
        st.markdown(
            f"<div class='member-card'><div class='member-name'>{html.escape(name)}　{status_badge(status)}</div>"
            f"<div class='member-meta'>{html.escape(meta)}</div><div class='member-label'>管理中</div>"
            f"<div class='member-value'>{html.escape(managed)}</div><div class='member-label'>最近变化</div>"
            f"<div class='member-value'>{html.escape(change)}</div><div class='member-label'>下一步</div>"
            f"<div class='member-value'>{html.escape(next_step)}</div></div>", unsafe_allow_html=True,
        )
        if on_click:
            primary_action("查看成员", key=key, on_click=on_click, args=args)


EVIDENCE_TYPE_LABELS = {
    "TEXT": "原始内容",
    "REPORT_TEXT": "原始内容",
    "TABLE": "原始表格",
    "REPORT_TABLE": "原始表格行",
    "IMAGE_REGION": "原始图片",
    "CELL_RANGE": "单元格",
    "EXCEL": "表格内容",
    "OBSERVATION": "健康数据",
    "DEVICE_DATA": "触发数据",
    "RISK": "触发数据",
    "DOCTOR_REVIEW": "医生意见",
    "MEMBER_REPORTED": "成员自述内容",
    "MANAGER_CONFIRMED": "确认记录",
    "KNOWLEDGE_SOURCE": "医学资料",
}
EVIDENCE_STATUS_LABELS = {
    "COMPLETE": "依据完整",
    "PARTIAL": "依据部分缺失",
    "MISSING": "暂无足够依据",
    "MISMATCH": "依据与结果不一致",
}
REPORT_SECTION_LABELS = {
    "LAB": "生化检查", "VITALS": "一般检查", "IMAGING": "影像检查",
    "ECG": "心电检查", "PULMONARY_FUNCTION": "肺功能", "BODY_COMPOSITION": "人体成分",
    "RECOMMENDATION": "报告建议", "SUMMARY": "报告小结", "OTHER": "体检报告",
}


def _source_display_name(document: Document | None = None, fallback: str | None = None) -> str:
    """Return a human-friendly source name without exposing test file names."""
    return humanize_source_name(document.title if document is not None else fallback)


def _evidence_location_text(
    *, page: int | None = None, section: str | None = None, sheet: str | None = None,
    cell_range: str | None = None, time_window: str | None = None,
) -> str:
    """Format only persisted source positions; never invent a location."""
    parts: list[str] = []
    if page is not None:
        parts.append(f"第 {page} 页")
    if section:
        parts.append(section)
    if sheet:
        parts.append(f"Sheet：{sheet}")
    if cell_range:
        parts.append(f"位置：{cell_range}")
    if time_window:
        parts.append(f"统计时间：{time_window}")
    return " · ".join(parts) or "当前未保存精确位置"


def _evidence_status_label(status: str | None) -> str:
    return EVIDENCE_STATUS_LABELS.get((status or "MISSING").upper(), EVIDENCE_STATUS_LABELS["MISSING"])


def _report_section_display(value: str | None) -> str:
    return REPORT_SECTION_LABELS.get((value or "").upper(), value or "报告内容")


def _candidate_evidence_status(candidate: ReportExtractionCandidate) -> str:
    structured = candidate.structured_data_json or {}
    integrity = str(structured.get("integrity_reason") or "").lower()
    if integrity == "evidence_mismatch" or candidate.status == "EVIDENCE_MISMATCH":
        return "MISMATCH"
    if candidate.candidate_type == "INCOMPLETE" or candidate.status in {"NEEDS_MANUAL_REVIEW", "AMBIGUOUS"}:
        return "PARTIAL"
    return "COMPLETE" if (candidate.evidence_text or "").strip() else "MISSING"


def _candidate_evidence_payload(candidate: ReportExtractionCandidate, document: Document | None) -> dict[str, object]:
    structured = candidate.structured_data_json or {}
    evidence_type = str(structured.get("evidence_type") or ("TABLE" if candidate.extraction_method == "TABLE" else "TEXT"))
    sheet = structured.get("sheet_name") or structured.get("sheet")
    cell_range = structured.get("cell_range") or structured.get("cells")
    return {
        "document": document,
        "source_name": _source_display_name(document),
        "location": _evidence_location_text(
            page=None if sheet else candidate.source_page, section=_report_section_display(candidate.source_section),
            sheet=str(sheet) if sheet else None, cell_range=str(cell_range) if cell_range else None,
        ),
        "evidence_type": evidence_type,
        "raw_evidence": candidate.evidence_text,
        "table_header": structured.get("table_header"),
        "table_row": structured.get("table_row"),
        "structured_interpretation": candidate.summary or _report_candidate_label(candidate),
        "confirmation_status": _label(candidate.status, context="report_candidate"),
        "evidence_status": _candidate_evidence_status(candidate),
        "ocr": candidate.extraction_method == "OCR" or evidence_type == "IMAGE_REGION",
        "bounding_box": structured.get("bounding_box"),
        "image_region": structured.get("image_region"),
        "has_bounding_box": bool(structured.get("bounding_box")),
    }


def _document_for_member_evidence(session, patient_id: UUID, document_id: UUID | str | None) -> Document | None:
    """Resolve a source document only inside the selected member's scope."""
    if not document_id:
        return None
    try:
        document = session.get(Document, UUID(str(document_id)))
    except (TypeError, ValueError):
        return None
    return document if document is not None and document.patient_id == patient_id else None


def _candidate_for_member_evidence(session, patient_id: UUID, candidate_id: UUID | str | None) -> ReportExtractionCandidate | None:
    """Resolve a parsed snippet only for its owning member."""
    if not candidate_id:
        return None
    try:
        candidate = session.get(ReportExtractionCandidate, UUID(str(candidate_id)))
    except (TypeError, ValueError):
        return None
    return candidate if candidate is not None and candidate.patient_id == patient_id else None


def _render_candidate_evidence_by_id(session, patient_id: UUID, candidate_id: UUID | str | None, *, key_scope: str, client_view: bool = False) -> None:
    candidate = _candidate_for_member_evidence(session, patient_id, candidate_id)
    if candidate is None:
        _render_evidence_action({"source_type": "MANAGER_CONFIRMED", "source_note": "该资料需要人工核对。", "location": None, "evidence_type": "MANAGER_CONFIRMED", "raw_evidence": None, "structured_interpretation": "当前未保存可展示的原文片段。", "confirmation_status": "待人工确认", "evidence_status": "MISSING"}, key_scope=key_scope, client_view=client_view)
        return
    document = _document_for_member_evidence(session, patient_id, candidate.document_id)
    _render_evidence_action(_candidate_evidence_payload(candidate, document), key_scope=key_scope, client_view=client_view)


def _risk_evidence_payload(session, patient_id: UUID, risk_id: UUID | str | None) -> dict[str, object]:
    """Build a business-facing deterministic-risk evidence view, never an AI rationale."""
    if not risk_id:
        return {"data_source": "健康数据", "location": None, "evidence_type": "RISK", "raw_evidence": None, "source_note": "当前未保存可定位的触发数据。", "structured_interpretation": "风险处理信息待补充。", "confirmation_status": "待人工确认", "evidence_status": "MISSING"}
    try:
        event = session.get(RiskEvent, UUID(str(risk_id)))
    except (TypeError, ValueError):
        event = None
    if event is None or event.patient_id != patient_id:
        return {"data_source": "健康数据", "location": None, "evidence_type": "RISK", "raw_evidence": None, "source_note": "当前成员暂无可访问的触发数据。", "structured_interpretation": "风险资料不可用。", "confirmation_status": "待人工确认", "evidence_status": "MISSING"}
    evidence = event.evidence_json or {}
    matches = evidence.get("matches") if isinstance(evidence.get("matches"), list) else []
    snippets = [
        " ".join(str(part) for part in (item.get("value"), item.get("unit")) if part not in {None, ""})
        for item in matches[-3:] if isinstance(item, dict)
    ]
    metric = _metric_display_name(event.canonical_code or str(evidence.get("metric") or ""))
    raw = "；".join(snippets)
    rule = session.get(RiskRule, event.risk_rule_id)
    is_demo_rule = rule is not None and rule.scope in {"TEST", "DEMO"}
    rule_scope = "测试规则" if rule is not None and rule.scope in {"TEST", "DEMO"} else "已审核规则"
    rule_source = rule.source_reference if rule is not None and rule.source_reference else "当前未保存规则来源"
    knowledge = None
    if rule is not None and rule.scope == "CLINICAL" and rule.source_reference:
        knowledge = {"名称": rule.source_reference, "状态": "已审核" if rule.review_status == "APPROVED" else "待审核", "版本": rule.version}
    return {
        "data_source": _risk_event_source(event),
        "location": _evidence_location_text(time_window=str((evidence.get("window") or {}).get("label") or "") or None),
        "evidence_type": "RISK",
        "raw_evidence": raw or None,
        "source_note": None if raw else "当前未保存可展示的触发数据样本。",
        "structured_interpretation": f"{'演示风险；' if is_demo_rule else ''}触发指标：{metric}；使用：{rule_scope}“{rule.name if rule else '风险规则'}”。",
        "confirmation_status": _label(event.status, context="risk_event"),
        "evidence_status": "COMPLETE" if matches else "PARTIAL",
        "knowledge_source": knowledge,
        "rule_reference": {"名称": rule.name if rule else "风险规则", "类型": rule_scope, "来源": rule_source, "版本": rule.version if rule else None},
        "technical": {"rule_code": rule.code if rule else None, "risk_event_id": str(event.id)},
    }


def _baseline_evidence_payload(session, patient_id: UUID, assessment: HealthAssessment) -> dict[str, object]:
    """Describe the recorded baseline provenance without recomputing its snapshot."""
    refs = assessment.source_references_json or {}
    report_ids = refs.get("source_report_ids") if isinstance(refs.get("source_report_ids"), list) else []
    document = _document_for_member_evidence(session, patient_id, report_ids[0] if report_ids else None)
    source_types: list[str] = []
    if document is not None:
        source_types.append("体检报告")
    if refs.get("source_observation_ids"):
        source_types.append("健康数据")
    if refs.get("source_problem_ids"):
        source_types.append("健康问题记录")
    if refs.get("source_medication_ids"):
        source_types.append("用药记录")
    if refs.get("source_procedure_ids"):
        source_types.append("医疗记录")
    snapshot = assessment.baseline_json or {}
    if snapshot.get("member_reported"):
        source_types.append("成员自述")
    return {
        "document": document,
        "source_name": _source_display_name(document, "健康基线来源资料") if document else None,
        "source_type": "MANAGER_CONFIRMED",
        "source_note": "健康基线汇总自：" + ("、".join(source_types) if source_types else "当前未保存来源资料"),
        "location": None,
        "evidence_type": "MANAGER_CONFIRMED",
        "raw_evidence": None,
        "structured_interpretation": assessment.summary,
        "confirmation_status": "已由健康管理团队确认" if assessment.status == "CONFIRMED" else "等待健康管理团队确认",
        "evidence_status": "COMPLETE" if source_types else "PARTIAL",
        "show_no_knowledge": True,
    }


def _render_snapshot_item_evidence(patient_id: UUID, value: object, *, key_scope: str, client_view: bool = False) -> None:
    """Keep Baseline's report-derived metrics/findings linked to their Candidate."""
    if not isinstance(value, list):
        return
    rows = [item for item in value if isinstance(item, dict) and item.get("source_candidate_id")]
    if not rows:
        return
    with st.expander("逐项查看依据"):
        with SessionLocal() as session:
            for index, item in enumerate(rows):
                label = str(item.get("metric") or item.get("summary") or "健康资料")
                st.markdown(f"**{_metric_display_name(label) if item.get('metric') else label}**")
                _render_candidate_evidence_by_id(
                    session, patient_id, item.get("source_candidate_id"),
                    key_scope=f"{key_scope}-{index}", client_view=client_view,
                )


def _timeline_evidence_payload(session, patient: Patient, event) -> dict[str, object]:
    """Map a story event to its persisted, member-scoped evidence reference."""
    details = event.expandable_details or {}
    if event.event_type == "risk":
        return _risk_evidence_payload(session, patient.id, event.related_entity)
    if event.event_type == "report":
        document = _document_for_member_evidence(session, patient.id, details.get("document_id") or event.related_entity)
        candidate = None
        if document is not None:
            candidate = session.scalar(
                select(ReportExtractionCandidate)
                .where(
                    ReportExtractionCandidate.patient_id == patient.id,
                    ReportExtractionCandidate.document_id == document.id,
                    ReportExtractionCandidate.status == "CONFIRMED",
                )
                .order_by(ReportExtractionCandidate.source_page, ReportExtractionCandidate.created_at)
            )
        if candidate is not None:
            payload = _candidate_evidence_payload(candidate, document)
            payload["structured_interpretation"] = event.summary
            return payload
        return {
            "document": document,
            "source_name": _source_display_name(document, "体检报告"),
            "location": None,
            "evidence_type": "REPORT_TEXT",
            "raw_evidence": None,
            "source_note": "该报告摘要来自已确认的体检资料；当前未关联可展示的原文片段。",
            "structured_interpretation": "报告摘要仅汇总当前已确认资料。",
            "confirmation_status": _label(str(details.get("review_state") or "PENDING_REVIEW"), context="report_candidate"),
            "evidence_status": "COMPLETE" if details.get("review_state") == "已确认" else "PARTIAL",
        }
    if event.event_type == "assessment":
        try:
            assessment = session.get(HealthAssessment, UUID(event.related_entity)) if event.related_entity else None
        except (TypeError, ValueError):
            assessment = None
        if assessment is not None and assessment.patient_id == patient.id:
            return _baseline_evidence_payload(session, patient.id, assessment)
    if event.event_type == "health_data_summary":
        window = " ～ ".join(str(item) for item in (details.get("window_start"), details.get("window_end")) if item)
        metrics = details.get("metrics") if isinstance(details.get("metrics"), list) else []
        raw = "；".join(f"{item.get('label')}：{item.get('samples')} 条" for item in metrics[:6] if isinstance(item, dict)) or "暂无汇总样本"
        return {"data_source": "健康数据", "location": _evidence_location_text(time_window=window or None), "evidence_type": "DEVICE_DATA", "raw_evidence": raw if metrics else None, "source_note": None if metrics else "当前未保存可展示的汇总样本。", "structured_interpretation": event.summary, "confirmation_status": "系统按有效健康数据汇总", "evidence_status": "COMPLETE" if metrics else "MISSING"}
    if event.event_type == "outcome":
        raw = f"干预前：{details.get('before', '—')} {details.get('unit', '')}；干预后：{details.get('after', '—')} {details.get('unit', '')}"
        return {"data_source": "健康数据", "location": None, "evidence_type": "DEVICE_DATA", "raw_evidence": raw, "source_note": "前后观察窗口未保存为精确时间范围。", "structured_interpretation": "仅记录观察到的变化，不表示任何干预造成该变化。", "confirmation_status": "已记录阶段评估", "evidence_status": "PARTIAL"}
    source_labels = {
        "doctor_review": "内部医生复核", "medication_change": "正式用药记录",
        "procedure": "医疗记录", "surgery": "手术记录", "hospitalization": "住院记录",
        "service": "服务执行记录", "external_referral": "外部医疗记录",
        "program_start": "健康管理计划", "program_adjustment": "健康管理复盘",
    }
    evidence_type = "DOCTOR_REVIEW" if event.event_type == "doctor_review" else "MANAGER_CONFIRMED"
    source_label = source_labels.get(event.event_type)
    source_note = (
        f"该信息来自已确认的{source_label}。"
        if source_label
        else "该信息来自已确认的业务记录；当前未保存更细的来源说明。"
    )
    return {"source_type": evidence_type, "source_note": source_note, "location": None, "evidence_type": evidence_type, "raw_evidence": None, "structured_interpretation": event.summary, "confirmation_status": _label(str(details.get("status") or "CONFIRMED"), context="doctor_review" if evidence_type == "DOCTOR_REVIEW" else None), "evidence_status": "PARTIAL", "show_no_knowledge": True}


def render_evidence_panel(evidence: dict[str, object], *, key_scope: str, client_view: bool = False) -> None:
    """Render one Chinese evidence panel from persisted references and snippets.

    This intentionally does not read document bytes until the user expands the
    panel and requests the complete file.  LLM/parser metadata is neither
    presented as medical evidence nor shown in the client surface.
    """
    with st.container(border=True):
        st.markdown("#### 依据")
        evidence_type = str(evidence.get("evidence_type") or "TEXT").upper()
        document = evidence.get("document")
        file_evidence = evidence_type in {"TEXT", "REPORT_TEXT", "TABLE", "REPORT_TABLE", "IMAGE_REGION", "CELL_RANGE", "EXCEL"} and isinstance(document, Document)
        data_evidence = evidence_type in {"OBSERVATION", "DEVICE_DATA", "RISK"}
        if data_evidence:
            st.markdown("**数据来源**")
            st.write(str(evidence.get("data_source") or evidence.get("source_name") or "健康数据"))
        elif file_evidence:
            st.markdown("**来源文件**")
            st.write(_source_display_name(document))
        else:
            st.markdown(f"**{get_source_type_display(str(evidence.get('source_type') or evidence_type))}**")
            st.write(str(evidence.get("source_note") or "来源信息待补充"))
        if isinstance(document, Document):
            path = Path(document.storage_reference)
            if path.is_file():
                st.download_button("查看完整文件", path.read_bytes(), file_name=_source_display_name(document), key=f"evidence-file-{key_scope}")
            else:
                st.caption("来源文件当前不可直接打开。")
        st.markdown("**来源位置**")
        st.write(str(evidence.get("location") or "当前未保存精确位置"))
        if file_evidence and evidence.get("source_note"):
            st.caption("来源说明：" + str(evidence["source_note"]))
        st.markdown(f"**{EVIDENCE_TYPE_LABELS.get(evidence_type, '原始内容')}**")
        raw_value = evidence.get("raw_evidence")
        raw = str(raw_value).strip() if raw_value is not None else ""
        snippet = raw[:1200] + ("…" if len(raw) > 1200 else "")
        if not raw:
            st.caption("当前未保存可展示的原文片段。")
        elif evidence_type in {"TABLE", "REPORT_TABLE", "CELL_RANGE", "EXCEL"}:
            header = str(evidence.get("table_header") or "").strip()
            row = str(evidence.get("table_row") or snippet).strip()
            if header:
                header_cells = [item.strip() for item in header.split("\t")]
                row_cells = [item.strip() for item in row.split("\t")]
                if len(header_cells) == len(row_cells) and header_cells:
                    st.dataframe(pd.DataFrame([dict(zip(header_cells, row_cells))]), hide_index=True, width="stretch")
                else:
                    st.dataframe(pd.DataFrame({"相关表格行": [row]}), hide_index=True, width="stretch")
            else:
                st.dataframe(pd.DataFrame({"相关表格行": [row]}), hide_index=True, width="stretch")
        elif evidence_type == "IMAGE_REGION":
            image_region = evidence.get("image_region")
            if image_region:
                st.image(image_region, caption="支持当前结论的报告图片区域")
            else:
                st.caption("当前已保存识别文字，但尚未保存精确图片区域定位。")
            st.write(snippet)
        elif raw:
            st.write(snippet)
        if evidence.get("ocr") and not evidence.get("has_bounding_box"):
            st.caption("当前仅保存识别文字，暂无原始区域定位。")
        st.markdown("**系统整理**")
        st.write(str(evidence.get("structured_interpretation") or "待人工整理"))
        st.markdown("**确认状态**")
        st.write(f"{evidence.get('confirmation_status') or '待人工确认'} · {_evidence_status_label(str(evidence.get('evidence_status') or 'MISSING'))}")
        knowledge = evidence.get("knowledge_source")
        if isinstance(knowledge, dict):
            st.markdown("**参考规范**")
            st.write(" · ".join(f"{key}：{value}" for key, value in knowledge.items() if value))
        elif evidence.get("show_no_knowledge"):
            st.caption("本结果基于当前健康资料整理，未使用额外医疗规范。")
        rule_reference = evidence.get("rule_reference")
        if isinstance(rule_reference, dict):
            st.markdown("**规则依据**")
            st.write(" · ".join(f"{key}：{value}" for key, value in rule_reference.items() if value))
        if not client_view:
            if TECHNICAL_DETAILS_ENABLED:
                technical = evidence.get("technical")
                if isinstance(technical, dict) and any(value not in {None, ""} for value in technical.values()):
                    with st.expander("高级信息"):
                        st.dataframe(pd.DataFrame([{
                            {"rule_code": "规则编号", "risk_event_id": "风险记录编号"}.get(key, key): value
                            for key, value in technical.items() if value not in {None, ""}
                        }]), hide_index=True, width="stretch")


def evidence_action(evidence: dict[str, object], *, key_scope: str, client_view: bool = False) -> None:
    """Secondary, on-demand evidence entry; technical context stays out of level 1."""
    open_key = f"evidence-open-{key_scope}"
    if st.button("查看依据", key=f"evidence-action-{key_scope}"):
        st.session_state[open_key] = not bool(st.session_state.get(open_key))
    if st.session_state.get(open_key):
        render_evidence_panel(evidence, key_scope=key_scope, client_view=client_view)


def _render_evidence_action(evidence: dict[str, object], *, key_scope: str, client_view: bool = False) -> None:
    """Backward-compatible alias for the shared evidence action."""
    evidence_action(evidence, key_scope=key_scope, client_view=client_view)


def _render_related_knowledge(query: str, *, key_scope: str) -> None:
    """Offer approved reference material without turning it into a medical decision."""
    normalized = (query or "").strip()
    if not normalized:
        return
    state_key = f"related-knowledge-{key_scope}"
    if secondary_action("查看相关医学参考", key=f"{state_key}-open", width="content"):
        with SessionLocal() as session:
            hits = KnowledgeRetrievalService().search(session, normalized, limit=3)
        st.session_state[state_key] = hits
    hits = st.session_state.get(state_key)
    if hits is None:
        return
    with st.expander("相关医学参考", expanded=True):
        st.caption("仅显示已批准、未归档且未过期的资料，用于解释与人工参考；不会改变风险、诊断、处方或医疗规则。")
        if not hits:
            st.info("当前没有可引用的已批准资料。")
            return
        for hit in hits:
            citation = hit.citation()
            st.markdown(f"**{citation['title']}**")
            st.caption(f"{citation['source'] or '来源待补充'} · {citation['location'] or '位置待补充'} · 获取时间：{citation['retrieved_at'] or '未记录'}")
            st.write(citation["excerpt"] or "当前未保存可展示的资料片段。")
            official = _knowledge_source_link(citation["source_url"])
            if official:
                st.link_button("查看官方来源", official, key=f"{state_key}-source-{hit.chunk.id}")


@contextmanager
def _section_frame(title: str, guidance: str | None = None):
    """Visible business-section boundary; use for major modules only."""
    with st.container(border=True):
        st.markdown(f"<div class='section-frame-title'>{html.escape(title)}</div>", unsafe_allow_html=True)
        if guidance:
            st.caption(guidance)
        yield


def _status_strip(*items: tuple[str, int, str]) -> None:
    """A light operational status bar, deliberately quieter than a KPI dashboard."""
    blocks = "".join(
        f"<div class='{html.escape(style)}'><b>{count}</b><span>{html.escape(label)}</span></div>"
        for label, count, style in items
    )
    st.markdown(f"<div class='status-strip'>{blocks}</div>", unsafe_allow_html=True)


def _navigation_stage(name: str, callback):
    """Emit opt-in, non-PHI timing only when a developer enables profiling."""
    if not NAVIGATION_PROFILE_ENABLED:
        return callback()
    started = perf_counter()
    result = callback()
    LOGGER.warning("[PERF] %s %.1f ms", name, (perf_counter() - started) * 1000)
    return result


def request_navigation(
    *,
    surface: str | None = None,
    ops_page: str | None = None,
    member_page: str | None = None,
    member_id: UUID | str | None = None,
    member_section: str | None = None,
    archive_view: str | None = None,
    report_document_id: UUID | str | None = None,
    health_data_window: dict[str, str | None] | None = None,
    rerun: bool = True,
) -> None:
    """Request a route for the next rerun without mutating live widget keys.

    Streamlit forbids writing a widget's session-state key after that widget
    has been instantiated.  Renderers therefore write only this independent
    router payload; ``apply_pending_navigation`` applies it at the start of
    the following rerun, before any navigation control exists.
    """
    st.session_state[PENDING_NAVIGATION_KEY] = {
        "surface": surface, "ops_page": ops_page, "member_page": member_page,
        "member_id": str(member_id) if member_id is not None else None,
        "member_section": member_section, "archive_view": archive_view,
        "report_document_id": str(report_document_id) if report_document_id is not None else None,
        "health_data_window": health_data_window,
    }
    if rerun:
        st.rerun()


def apply_pending_navigation() -> None:
    """Apply a requested route before Streamlit instantiates navigation widgets."""
    pending = st.session_state.pop(PENDING_NAVIGATION_KEY, None)
    # One-time migration of old page labels also happens before any widget is
    # created.  It avoids the Streamlit state mutation error on existing
    # development sessions without allowing a renderer to mutate live keys.
    legacy_member_page = {"健康档案": "健康", "健康数据": "健康", "我的服务": "服务", "我的": "首页"}
    if st.session_state.get("member-center-navigation") in legacy_member_page:
        st.session_state["member-center-navigation"] = legacy_member_page[st.session_state["member-center-navigation"]]
    if not isinstance(pending, dict):
        return
    if pending.get("surface"):
        st.session_state["surface-mode"] = pending["surface"]
    if pending.get("ops_page"):
        st.session_state["ops-navigation"] = {
            "今日工作台": "今日", "工作台": "今日",
            "协同": "医疗协同", "风险与医疗协同": "医疗协同",
        }.get(pending["ops_page"], pending["ops_page"])
    if pending.get("member_page"):
        st.session_state["member-center-navigation"] = {
            "健康档案": "健康", "健康数据": "健康", "数据": "健康",
            "我的服务": "服务",
        }.get(pending["member_page"], pending["member_page"])
    member_id = pending.get("member_id")
    if member_id:
        st.session_state["focused_member_id"] = member_id
        if pending.get("member_section"):
            st.session_state[f"member-section-{member_id}"] = {"数据": "健康", "档案": "健康"}.get(pending["member_section"], pending["member_section"])
        if pending.get("archive_view") is not None:
            health_view = {
                "健康数据": "健康数据", "体检与检查": "体检", "健康基线": "基线", "基线": "基线",
                "健康史": "医疗档案", "医疗资料": "医疗档案", "用药与医疗": "医疗档案",
                "健康时间轴": "历程", "健康历程": "历程", "报告对比": "体检",
            }.get(str(pending["archive_view"]))
            if health_view:
                if health_view == "历程":
                    st.session_state["member-center-navigation"] = "历程"
                    st.session_state[f"member-section-{member_id}"] = "历程"
                else:
                    st.session_state[f"client-health-view-{member_id}"] = health_view
                    st.session_state[f"member-health-view-{member_id}"] = health_view
        if pending.get("health_data_window"):
            st.session_state[f"health-data-window-{member_id}"] = pending["health_data_window"]
    if pending.get("report_document_id"):
        st.session_state["report-review-document-id"] = pending["report_document_id"]


def _open_risk_events(patient_id: UUID) -> list[RiskEvent]:
    with SessionLocal() as session:
        return list(session.scalars(
            select(RiskEvent).where(
                RiskEvent.patient_id == patient_id,
                RiskEvent.status.in_(["NEW", "ACKNOWLEDGED", "IN_REVIEW", "MONITORING", "ESCALATED_TO_DOCTOR", "WAITING_MEMBER", "FOLLOW_UP", "ESCALATED"]),
            ).order_by(RiskEvent.created_at.desc())
        ))


def _risk_event_source(event: RiskEvent) -> str:
    # The source marker remains in the stored evidence for audit purposes, but
    # the operational surface names the understandable source category only.
    return "健康数据自动监测"


def _risk_event_evidence_caption(event: RiskEvent) -> str:
    """Keep governed rule evidence understandable without exposing JSON internals."""
    evidence = event.evidence_json or {}
    with SessionLocal() as session:
        rule = session.get(RiskRule, event.risk_rule_id)
    is_demo_rule = rule is not None and rule.scope in {"TEST", "DEMO"}
    rule_name = "演示规则" if is_demo_rule else (rule.name if rule else "已审核风险规则")
    metric = _metric_display_name(event.canonical_code or str(evidence.get("metric") or ""))
    window = evidence.get("window") if isinstance(evidence.get("window"), dict) else {}
    lookback = window.get("lookback_minutes") if isinstance(window, dict) else None
    matched = evidence.get("matched_count")
    parts = [
        f"触发来源：{_risk_event_source(event)}",
        f"触发规则：{rule_name}",
        f"建议路由：{ROUTE_LABELS.get(event.recommended_route, '人工处理')}",
    ]
    if is_demo_rule:
        parts.append("演示规则，仅用于演示流程")
    if metric and metric != (event.canonical_code or ""):
        parts.append(f"相关指标：{metric}")
    if matched:
        parts.append(f"匹配记录：{matched} 项")
    if lookback:
        parts.append(f"数据窗口：最近 {lookback} 分钟")
    matches = evidence.get("matches")
    if isinstance(matches, list) and matches and isinstance(matches[-1], dict):
        observed_at = matches[-1].get("observed_at")
        if observed_at:
            parts.append(f"相关数据时间：{observed_at}")
    return " · ".join(parts)


def _member_risk_state(patient_id: UUID, ctx: dict[str, list[object]]) -> tuple[str, str | None, datetime | None]:
    events = _open_risk_events(patient_id)
    if any(event.risk_level == "RED" for event in events):
        event = next(event for event in events if event.risk_level == "RED")
        return "高风险", event.summary, event.created_at
    if any(event.risk_level == "YELLOW" for event in events):
        event = next(event for event in events if event.risk_level == "YELLOW")
        return "中风险", event.summary, event.created_at
    status, _ = _overall_status(ctx)
    if status == "存在待处理健康运营事项":
        alert = next((item for item in ctx["alerts"] if item.status != "CLOSED"), None)
        return "需要关注", alert.finding if alert else "存在待处理健康管理事项。", getattr(alert, "updated_at", None)
    return "暂无正式风险评估", "暂无适用的正式临床风险规则覆盖；这不等于低风险。", None


def _render_current_risk_actions(patient: Patient) -> None:
    """Expose existing audited risk actions without adding new triage logic."""
    events = _open_risk_events(patient.id)
    if not events:
        return
    event = next((item for item in events if item.risk_level == "RED"), events[0])
    if event.risk_level == "RED":
        st.subheader("紧急处置")
        st.error("医疗处置优先。请先确认成员当前安全情况并按实际情况寻求医疗帮助。")
        st.write(f"关键原因：{event.summary}")
        st.caption(f"发现时间：{_fmt_dt(event.created_at)} · {_risk_event_evidence_caption(event)}")
        with SessionLocal() as session:
            _render_evidence_action(_risk_evidence_payload(session, patient.id, event.id), key_scope=f"risk-red-{event.id}")
        contacts: list[EmergencyContact]
        with SessionLocal() as session:
            contacts = list(session.scalars(select(EmergencyContact).where(EmergencyContact.patient_id == patient.id).order_by(EmergencyContact.is_primary.desc())))
        if contacts:
            st.caption(f"紧急联系人：{contacts[0].name}（{contacts[0].relationship}，合成演示联系人）")
        first, second = st.columns(2)
        first.link_button("使用设备拨打120", "tel:120", type="primary", width="stretch")
        if second.button("记录已开始紧急处置", key=f"emergency-action-{event.id}", width="stretch"):
            with SessionLocal() as session:
                stored = session.get(RiskEvent, event.id)
                RiskEvaluationService().emergency_action(session, stored, "健康管理师", "已人工开始紧急处置")
                session.commit()
            st.success("已记录紧急处置开始；系统不会自动拨打120或联系任何真实联系人。")
            st.rerun()
        st.caption("请使用可拨号设备操作。如无法直接拨号，请使用手机拨打120；系统不会自动拨号或联系任何人。")
        if event.status == "ESCALATED":
            with st.form(f"red-risk-close-{event.id}"):
                close_reason = st.text_area("关闭原因", placeholder="说明为何本次人工处置可以结束")
                final_action = st.text_area("最终人工处置", placeholder="记录已完成的人工处理或后续交接")
                if st.form_submit_button("记录结果并关闭风险事项"):
                    try:
                        with SessionLocal() as session:
                            stored = session.get(RiskEvent, event.id)
                            RiskEvaluationService().close_manual_event(session, stored, "健康管理师", close_reason, final_action)
                            session.commit()
                        st.success("已记录人工处置结果并关闭风险事项。")
                        st.rerun()
                    except ValueError as error:
                        st.error(str(error))
    elif event.risk_level == "YELLOW":
        render_yellow_risk_operations(patient, event)


def render_yellow_risk_operations(patient: Patient, event: RiskEvent) -> None:
    """Compact human-only action surface for an observation-driven Yellow event."""
    st.subheader("需要关注")
    st.caption(f"{_risk_event_source(event)} · {_label(event.status)}")
    st.write(event.summary)
    st.caption(f"触发时间：{_fmt_dt(event.created_at)} · {_risk_event_evidence_caption(event)}")
    with SessionLocal() as session:
        _render_evidence_action(_risk_evidence_payload(session, patient.id, event.id), key_scope=f"risk-yellow-{event.id}")
    st.markdown("#### 建议下一步")
    action = st.radio("处理方式", ["继续观察", "记录联系成员", "数据有误", "调整健康管理", "请医生复核"], horizontal=True, key=f"yellow-action-{event.id}")
    with st.form(f"yellow-action-form-{event.id}"):
        actor = st.text_input("健康管理师", value="健康管理师")
        if action == "继续观察":
            reason = st.text_area("观察原因 / 备注", placeholder="说明为何持续观察", key=f"yellow-monitor-note-{event.id}")
            due_date = st.date_input("下次复核日期", value=date.today() + timedelta(days=1), key=f"yellow-monitor-date-{event.id}")
            submit = st.form_submit_button("保存并创建复核任务", type="primary")
        elif action == "记录联系成员":
            method = st.selectbox("联系方式", ["电话", "微信", "当面", "其他"], key=f"yellow-contact-method-{event.id}")
            result = st.selectbox("联系结果", ["已联系", "未接通", "待回访"], key=f"yellow-contact-result-{event.id}")
            reason = st.text_area("联系记录 / 备注", key=f"yellow-contact-note-{event.id}")
            due_date = st.date_input("下次复核日期（可选）", value=date.today() + timedelta(days=2), key=f"yellow-contact-date-{event.id}")
            submit = st.form_submit_button("保存人工联系记录", type="primary")
        elif action == "数据有误":
            reason = st.text_area("数据问题原因", placeholder="设备错误、录入错误或成员确认无效等", key=f"yellow-data-note-{event.id}")
            submit = st.form_submit_button("记录数据问题并关闭事件", type="primary")
        elif action == "调整健康管理":
            reason = st.text_area("调整原因", key=f"yellow-adjust-reason-{event.id}")
            adjustment = st.text_area("调整内容 / 下一步任务", key=f"yellow-adjust-content-{event.id}")
            due_date = st.date_input("复核日期（可选）", value=date.today() + timedelta(days=7), key=f"yellow-adjust-date-{event.id}")
            submit = st.form_submit_button("创建健康管理任务", type="primary")
        else:
            reason = st.text_area("希望医生确认什么？", placeholder="由健康管理师明确提出需要医生回答的问题", key=f"yellow-doctor-question-{event.id}")
            department = st.selectbox("建议科室", ["全科/健康管理", "心内科", "内分泌科", "其他"], key=f"yellow-doctor-department-{event.id}")
            submit = st.form_submit_button("提交医生复核", type="primary")
        if submit:
            try:
                with SessionLocal() as session:
                    operations = RiskOperationsService()
                    if action == "继续观察":
                        operations.continue_monitoring(session, event.id, actor, reason, datetime.combine(due_date, time(9, 0), tzinfo=TOKYO_TIMEZONE))
                    elif action == "记录联系成员":
                        operations.record_contact(session, event.id, actor, method, result, reason, datetime.combine(due_date, time(9, 0), tzinfo=TOKYO_TIMEZONE))
                    elif action == "数据有误":
                        operations.mark_data_issue(session, event.id, actor, reason)
                    elif action == "调整健康管理":
                        operations.adjust_management(session, event.id, actor, adjustment, reason, datetime.combine(due_date, time(9, 0), tzinfo=TOKYO_TIMEZONE))
                    else:
                        operations.escalate_to_doctor(session, event.id, actor, reason, department)
                    session.commit()
                st.success("已保存人工处置记录。系统未发送真实消息，也未自动作出医疗决定。")
                st.rerun()
            except ValueError as error:
                st.error(str(error))
    if event.status in {"FOLLOW_UP", "MONITORING"}:
        with st.form(f"yellow-followup-{event.id}"):
            outcome = st.text_area("记录跟进结果", key=f"yellow-followup-outcome-{event.id}")
            if st.form_submit_button("完成本次跟进"):
                try:
                    with SessionLocal() as session:
                        operations = RiskOperationsService()
                        if event.status == "MONITORING":
                            operations.complete_monitoring_task(session, event.id, "健康管理师", outcome)
                        else:
                            operations.record_follow_up(session, event.id, "健康管理师", outcome)
                        session.commit()
                    st.success("已记录跟进结果。")
                    st.rerun()
                except ValueError as error:
                    st.error(str(error))
    if event.status in {"FOLLOW_UP", "MONITORING", "IN_REVIEW", "ACKNOWLEDGED", "WAITING_MEMBER"}:
        with st.form(f"yellow-close-{event.id}"):
            close_reason = st.text_area("关闭原因", key=f"yellow-close-reason-{event.id}")
            if st.form_submit_button("完成跟进后关闭事件"):
                try:
                    with SessionLocal() as session:
                        RiskOperationsService().close(session, event.id, "健康管理师", close_reason)
                        session.commit()
                    st.success("已关闭本次需要关注事件。")
                    st.rerun()
                except ValueError as error:
                    st.error(str(error))
    with SessionLocal() as session:
        actions = list(session.scalars(select(AuditLog).where(AuditLog.entity_type == "RiskEvent", AuditLog.entity_id == str(event.id)).order_by(AuditLog.created_at.asc()).limit(30)))
    if actions:
        st.markdown("#### 处理记录")
        for item in actions:
            if item.action.startswith("yellow_") or item.action.startswith("risk_event"):
                st.caption(f"{_fmt_dt(item.created_at)} · {_role_label(item.actor_role, name=item.actor)} · {get_audit_action_display(item.action)}")


def _members() -> list[Patient]:
    try:
        with SessionLocal() as session:
            return list(session.scalars(select(Patient).order_by(Patient.display_name, Patient.created_at)))
    except SQLAlchemyError:
        return []


def _context(patient_id: UUID | None = None) -> dict[str, list[object]]:
    filters = [] if patient_id is None else [HealthProblem.patient_id == patient_id]
    with SessionLocal() as session:
        def rows(model: object, order: object, model_filters: list[object] | None = None) -> list[object]:
            statement = select(model).order_by(order)
            if model_filters:
                statement = statement.where(*model_filters)
            return list(session.scalars(statement))

        return {
            "alerts": rows(Alert, Alert.created_at.desc(), [] if patient_id is None else [Alert.patient_id == patient_id]),
            "problems": rows(HealthProblem, HealthProblem.opened_at.desc(), filters),
            "plans": rows(ManagementPlan, ManagementPlan.created_at.desc(), [] if patient_id is None else [ManagementPlan.patient_id == patient_id]),
            "tasks": rows(Task, Task.created_at.desc(), [] if patient_id is None else [Task.patient_id == patient_id]),
            "reviews": rows(DoctorReview, DoctorReview.reviewed_at.desc(), [] if patient_id is None else [DoctorReview.patient_id == patient_id]),
            "followups": rows(FollowUp, FollowUp.created_at.desc(), [] if patient_id is None else [FollowUp.patient_id == patient_id]),
            "audits": rows(AuditLog, AuditLog.created_at.desc(), [] if patient_id is None else [AuditLog.patient_id == patient_id]),
            "med_plans": rows(MedicationPlan, MedicationPlan.created_at.desc(), [] if patient_id is None else [MedicationPlan.patient_id == patient_id]),
            "med_events": rows(MedicationEvent, MedicationEvent.scheduled_at.desc(), [] if patient_id is None else [MedicationEvent.patient_id == patient_id]),
            "observations": rows(Observation, Observation.observed_at.desc(), [] if patient_id is None else [Observation.patient_id == patient_id]),
            "journeys": rows(HealthJourney, HealthJourney.updated_at.desc(), [] if patient_id is None else [HealthJourney.patient_id == patient_id]),
            "programs": rows(HealthProgram, HealthProgram.start_date.desc(), [] if patient_id is None else [HealthProgram.patient_id == patient_id]),
            "phases": rows(ProgramPhase, ProgramPhase.start_date.desc()),
            "barriers": rows(ExecutionBarrier, ExecutionBarrier.detected_at.desc(), [] if patient_id is None else [ExecutionBarrier.patient_id == patient_id]),
            "weekly_reviews": rows(WeeklyReview, WeeklyReview.reviewed_at.desc()),
            "outcomes": rows(OutcomeEvaluation, OutcomeEvaluation.created_at.desc(), [] if patient_id is None else [OutcomeEvaluation.patient_id == patient_id]),
            "annual_accounts": rows(AnnualHealthAccount, AnnualHealthAccount.updated_at.desc(), [] if patient_id is None else [AnnualHealthAccount.patient_id == patient_id]),
            "identities": rows(ExternalIdentity, ExternalIdentity.updated_at.desc(), [] if patient_id is None else [ExternalIdentity.patient_id == patient_id]),
            "ingestion_jobs": rows(IngestionJob, IngestionJob.started_at.desc(), [] if patient_id is None else [IngestionJob.patient_id == patient_id]),
        }


def _dashboard_context() -> dict[str, list[object]]:
    """Load only the operational summaries required on the workbench."""
    with SessionLocal() as session:
        return {
            "alerts": list(session.scalars(select(Alert).where(Alert.status != "CLOSED").order_by(Alert.created_at.desc()).limit(50))),
            "problems": list(session.scalars(select(HealthProblem).where(HealthProblem.status != "CLOSED").order_by(HealthProblem.opened_at.desc()).limit(50))),
            "tasks": list(session.scalars(select(Task).where(Task.status.not_in(["COMPLETED", "CANCELLED"])).order_by(Task.due_at, Task.created_at.desc()).limit(100))),
            "barriers": list(session.scalars(select(ExecutionBarrier).where(ExecutionBarrier.status != "RESOLVED").order_by(ExecutionBarrier.detected_at.desc()).limit(50))),
            "programs": list(session.scalars(select(HealthProgram).where(HealthProgram.status == "ACTIVE").order_by(HealthProgram.start_date.desc()).limit(50))),
            "management_signals": list(session.scalars(select(ManagementSignal).where(ManagementSignal.status.in_(("OPEN", "IN_PROGRESS"))).order_by(ManagementSignal.last_detected_at.desc()).limit(50))),
        }


def _member_summary_context(patient_id: UUID) -> dict[str, list[object]]:
    """Small current-state query; intentionally excludes CGM, raw data and audit history."""
    with SessionLocal() as session:
        return {
            "alerts": list(session.scalars(select(Alert).where(Alert.patient_id == patient_id).order_by(Alert.created_at.desc()).limit(20))),
            "problems": list(session.scalars(select(HealthProblem).where(HealthProblem.patient_id == patient_id).order_by(HealthProblem.opened_at.desc()).limit(20))),
            "tasks": list(session.scalars(select(Task).where(Task.patient_id == patient_id).order_by(Task.created_at.desc()).limit(20))),
            "journeys": list(session.scalars(select(HealthJourney).where(HealthJourney.patient_id == patient_id).order_by(HealthJourney.updated_at.desc()).limit(1))),
            "programs": list(session.scalars(select(HealthProgram).where(HealthProgram.patient_id == patient_id).order_by(HealthProgram.start_date.desc()).limit(10))),
            "observations": list(session.scalars(select(Observation).where(Observation.patient_id == patient_id).order_by(Observation.observed_at.desc()).limit(40))),
            "management_signals": list(session.scalars(select(ManagementSignal).where(ManagementSignal.patient_id == patient_id, ManagementSignal.status.in_(("OPEN", "IN_PROGRESS"))).order_by(ManagementSignal.last_detected_at.desc()).limit(10))),
        }


def _member_management_context(patient_id: UUID) -> dict[str, list[object]]:
    with SessionLocal() as session:
        programs = list(session.scalars(select(HealthProgram).where(HealthProgram.patient_id == patient_id).order_by(HealthProgram.start_date.desc())))
        program_ids = [program.id for program in programs]
        return {
            "programs": programs,
            "phases": list(session.scalars(select(ProgramPhase).where(ProgramPhase.program_id.in_(program_ids)).order_by(ProgramPhase.sequence))) if program_ids else [],
            "problems": list(session.scalars(select(HealthProblem).where(HealthProblem.patient_id == patient_id).order_by(HealthProblem.priority_rank))),
            "tasks": list(session.scalars(select(Task).where(Task.patient_id == patient_id).order_by(Task.created_at.desc()))),
            "alerts": list(session.scalars(select(Alert).where(Alert.patient_id == patient_id).order_by(Alert.created_at.desc()))),
            "plans": list(session.scalars(select(ManagementPlan).where(ManagementPlan.patient_id == patient_id).order_by(ManagementPlan.created_at.desc()))),
            "reviews": list(session.scalars(select(DoctorReview).where(DoctorReview.patient_id == patient_id).order_by(DoctorReview.reviewed_at.desc()))),
            "followups": list(session.scalars(select(FollowUp).where(FollowUp.patient_id == patient_id).order_by(FollowUp.created_at.desc()))),
            "barriers": list(session.scalars(select(ExecutionBarrier).where(ExecutionBarrier.patient_id == patient_id).order_by(ExecutionBarrier.detected_at.desc()))),
            "weekly_reviews": list(session.scalars(select(WeeklyReview).where(WeeklyReview.program_id.in_(program_ids)).order_by(WeeklyReview.reviewed_at.desc()))) if program_ids else [],
            "outcomes": list(session.scalars(select(OutcomeEvaluation).where(OutcomeEvaluation.program_id.in_(program_ids)).order_by(OutcomeEvaluation.created_at.desc()))) if program_ids else [],
            "management_signals": list(session.scalars(select(ManagementSignal).where(ManagementSignal.patient_id == patient_id, ManagementSignal.status.in_(("OPEN", "IN_PROGRESS"))).order_by(ManagementSignal.last_detected_at.desc()).limit(10))),
        }


def _member_medical_context(patient_id: UUID) -> dict[str, list[object]]:
    with SessionLocal() as session:
        return {
            "reviews": list(session.scalars(select(DoctorReview).where(DoctorReview.patient_id == patient_id).order_by(DoctorReview.reviewed_at.desc()).limit(30))),
            "med_plans": list(session.scalars(select(MedicationPlan).where(MedicationPlan.patient_id == patient_id).order_by(MedicationPlan.created_at.desc()).limit(30))),
            "med_events": list(session.scalars(select(MedicationEvent).where(MedicationEvent.patient_id == patient_id).order_by(MedicationEvent.scheduled_at.desc()).limit(30))),
            "followups": list(session.scalars(select(FollowUp).where(FollowUp.patient_id == patient_id).order_by(FollowUp.created_at.desc()).limit(30))),
            "audits": list(session.scalars(select(AuditLog).where(AuditLog.patient_id == patient_id).order_by(AuditLog.created_at.desc()).limit(30))),
            "problems": list(session.scalars(select(HealthProblem).where(HealthProblem.patient_id == patient_id).order_by(HealthProblem.opened_at.desc()).limit(30))),
            "plans": list(session.scalars(select(ManagementPlan).where(ManagementPlan.patient_id == patient_id).order_by(ManagementPlan.created_at.desc()).limit(30))),
            "barriers": [], "journeys": [], "programs": [], "outcomes": [],
        }


def _member_list_summaries(member_ids: list[UUID]) -> dict[UUID, dict[str, object]]:
    """Batch lightweight card summaries; deliberately never reads Observations."""
    if not member_ids:
        return {}
    active_risk = ["NEW", "ACKNOWLEDGED", "IN_REVIEW", "MONITORING", "ESCALATED_TO_DOCTOR", "FOLLOW_UP", "ESCALATED"]
    # The member landing page is a navigation list, not a complete caseload
    # export.  Keep every aggregate bounded without loading observations.
    limit = max(100, len(member_ids) * 5)
    with SessionLocal() as session:
        risks = list(session.scalars(select(RiskEvent).where(RiskEvent.patient_id.in_(member_ids), RiskEvent.status.in_(active_risk)).order_by(RiskEvent.created_at.desc()).limit(limit)))
        programs = list(session.scalars(select(HealthProgram).where(HealthProgram.patient_id.in_(member_ids), HealthProgram.status == "ACTIVE").order_by(HealthProgram.start_date.desc()).limit(limit)))
        problems = list(session.scalars(select(HealthProblem).where(HealthProblem.patient_id.in_(member_ids), HealthProblem.status != "CLOSED").order_by(HealthProblem.opened_at.desc()).limit(limit)))
        tasks = list(session.scalars(select(Task).where(Task.patient_id.in_(member_ids), Task.status.not_in(["COMPLETED", "CANCELLED"])).order_by(Task.due_at, Task.created_at.desc()).limit(limit)))
    result: dict[UUID, dict[str, object]] = {member_id: {"risk": "正常", "problems": [], "program": None, "next_task": None} for member_id in member_ids}
    for item in risks:
        if result[item.patient_id]["risk"] == "正常":
            result[item.patient_id]["risk"] = "紧急风险" if item.risk_level == "RED" else "需要关注"
    for item in programs:
        if result[item.patient_id]["program"] is None:
            result[item.patient_id]["program"] = item
    for item in problems:
        titles = result[item.patient_id]["problems"]
        if len(titles) < 3:
            titles.append(item.title)
    for item in tasks:
        if result[item.patient_id]["next_task"] is None:
            result[item.patient_id]["next_task"] = item
    return result


def _audit_context(patient_id: UUID) -> dict[str, list[object]]:
    """Audit page needs audit rows only; avoid legacy full member context."""
    with SessionLocal() as session:
        return {"audits": list(session.scalars(select(AuditLog).where(AuditLog.patient_id == patient_id).order_by(AuditLog.created_at.desc()).limit(50)))}


def _member_doctor_context(patient_id: UUID) -> dict[str, list[object]]:
    context = _member_medical_context(patient_id)
    with SessionLocal() as session:
        context["alerts"] = list(session.scalars(select(Alert).where(Alert.patient_id == patient_id).order_by(Alert.created_at.desc()).limit(30)))
        context["observations"] = list(session.scalars(select(Observation).where(Observation.patient_id == patient_id).order_by(Observation.observed_at.desc()).limit(50)))
    return context


def _patient_map() -> dict[UUID, Patient]:
    return {member.id: member for member in _members()}


def _overall_status(ctx: dict[str, list[object]]) -> tuple[str, str]:
    alerts = ctx["alerts"]
    problems = ctx["problems"]
    if any(item.status != "CLOSED" and item.severity == "CRITICAL" for item in alerts):
        return "紧急运营处置", "critical"
    if any(item.status != "CLOSED" for item in alerts) or any(item.status != "CLOSED" for item in problems):
        return "存在待处理健康运营事项", "high"
    if any(item.status == "CLOSED" for item in problems):
        return "最近已完成闭环", "closed"
    return "等待数据与人工复核", ""


def _queue_items(ctx: dict[str, list[object]], patients: dict[UUID, Patient]) -> list[dict[str, object]]:
    now = datetime.now(TOKYO_TIMEZONE)
    items: list[dict[str, object]] = []
    for alert in ctx["alerts"]:
        if alert.status == "CLOSED":
            continue
        next_action = "管理师核实数据" if alert.status in {"NEW", "AI_SCREENED", "WAITING_MANAGER_REVIEW"} else "提交医生复核" if alert.status == "WAITING_DOCTOR_REVIEW" else "安排或跟进随访"
        items.append({"kind": "alert", "member": patients.get(alert.patient_id), "title": alert.title, "severity": alert.severity, "status": alert.status, "owner": alert.owner or alert.responsible_role, "created_at": alert.created_at, "due_at": alert.due_at, "next_action": next_action, "object_id": alert.id})
    for task in ctx["tasks"]:
        if task.status in {"COMPLETED", "CANCELLED"}:
            continue
        overdue = task.due_at is not None and task.due_at < now
        items.append({"kind": "task", "member": patients.get(task.patient_id), "title": task.title, "severity": "OVERDUE" if overdue else task.priority, "status": "OVERDUE" if overdue else task.status, "owner": task.assignee or task.responsible_role, "created_at": task.created_at, "due_at": task.due_at, "next_action": "跟进成员执行" if task.responsible_role == "member" else "完成或重新分派任务", "object_id": task.id})
    for problem in ctx["problems"]:
        if problem.status == "CLOSED":
            continue
        if not any(item.get("kind") == "alert" and getattr(item.get("member"), "id", None) == problem.patient_id for item in items):
            items.append({"kind": "problem", "member": patients.get(problem.patient_id), "title": problem.title, "severity": problem.severity, "status": problem.status, "owner": problem.owner or problem.responsible_role, "created_at": problem.opened_at, "due_at": None, "next_action": "检查计划、任务与随访", "object_id": problem.id})
    for barrier in ctx["barriers"]:
        if barrier.status == "RESOLVED":
            continue
        items.append({"kind": "execution_risk", "member": patients.get(barrier.patient_id), "title": f"执行中断：{barrier.reason}", "severity": "HIGH", "status": barrier.status, "owner": barrier.confirmed_by, "created_at": barrier.detected_at, "due_at": None, "next_action": "核实原因并调整计划或任务", "object_id": barrier.id})
    rank = {"CRITICAL": 0, "HIGH": 1, "OVERDUE": 2, "WAITING_DOCTOR_REVIEW": 3}
    return sorted(items, key=lambda item: (rank.get(str(item["severity"]), 4 if item["owner"] == "member" else 5), item["due_at"] or item["created_at"]))


def _render_kpis(ctx: dict[str, list[object]]) -> None:
    now = datetime.now(TOKYO_TIMEZONE)
    alerts = ctx["alerts"]
    tasks = ctx["tasks"]
    problems = ctx["problems"]
    followups = ctx["followups"]
    high = [item for item in alerts if item.status != "CLOSED" and item.severity in {"HIGH", "CRITICAL"}]
    manager_wait = [item for item in alerts if item.status in {"NEW", "AI_SCREENED", "WAITING_MANAGER_REVIEW"}]
    doctor_wait = [item for item in alerts if item.status == "WAITING_DOCTOR_REVIEW"]
    overdue = [item for item in tasks if item.status not in {"COMPLETED", "CANCELLED"} and item.due_at and item.due_at < now]
    next_week = [item for item in followups if item.status != "COMPLETED" and item.due_at and now <= item.due_at <= now + timedelta(days=7)]
    open_problems = [item for item in problems if item.status != "CLOSED"]
    month_closed = [item for item in problems if item.status == "CLOSED" and item.closed_at and item.closed_at.astimezone(TOKYO_TIMEZONE).strftime("%Y-%m") == now.strftime("%Y-%m")]
    response_hours = [(item.reviewed_at - item.created_at).total_seconds() / 3600 for item in alerts if item.reviewed_at]
    first, second = st.columns(4), st.columns(4)
    first[0].metric("高优先级事件", len(high))
    first[1].metric("待健康管理师核实", len(manager_wait))
    first[2].metric("待医生复核", len(doctor_wait))
    first[3].metric("逾期任务", len(overdue))
    second[0].metric("7日内待复查", len(next_week))
    second[1].metric("当前未闭环健康问题", len(open_problems))
    second[2].metric("本月已闭环健康问题", len(month_closed))
    second[3].metric("平均响应时间", "—" if not response_hours else f"{sum(response_hours) / len(response_hours):.1f} h")


def _render_program_funnel(ctx: dict[str, list[object]]) -> None:
    programs: list[HealthProgram] = ctx["programs"]  # type: ignore[assignment]
    journeys: list[HealthJourney] = ctx["journeys"]  # type: ignore[assignment]
    stages = [
        ("ASSESSMENT", "筛查评估"), ("90_DAY_PROGRAM", "90天健康管理"),
        ("STABILIZATION", "稳定管理"), ("ANNUAL_MANAGEMENT", "年度健康管理"),
    ]
    st.subheader("成员健康管理阶段")
    cols = st.columns(4)
    for col, (stage, label) in zip(cols, stages):
        col.metric(label, sum(j.current_stage == stage for j in journeys))
    outcomes = ctx["outcomes"]
    outcome_programs = {item.program_id for item in outcomes}
    metrics = st.columns(5)
    metrics[0].metric("已进入管理计划", len(programs))
    metrics[1].metric("已完成管理计划", sum(p.status == "COMPLETED" for p in programs))
    metrics[2].metric("已完成阶段效果评估", f"{len(outcome_programs)} / {len(programs)}" if programs else "暂无数据")
    metrics[3].metric("已升级医疗处理", sum(p.status == "ESCALATED_TO_MEDICAL_CARE" for p in programs))
    metrics[4].metric("进入稳定管理", sum(p.program_type == "STABILIZATION" for p in programs))


def _open_member(member_id: UUID) -> None:
    """Widget callback: queue navigation for the next Streamlit rerun."""
    request_navigation(ops_page="成员", member_id=member_id, rerun=False)


def _open_member_service(member_id: UUID) -> None:
    """Open the member's service detail from the operations worklist."""
    request_navigation(ops_page="成员", member_id=member_id, member_section="服务", rerun=False)


def _open_member_management(member_id: UUID) -> None:
    request_navigation(ops_page="成员", member_id=member_id, member_section="管理", rerun=False)


def _open_report_review_from_worklist(member_id: UUID, document_id: UUID) -> None:
    """Open the existing manager review workspace from a report work item."""
    request_navigation(
        ops_page="成员", member_id=member_id, member_section="健康",
        archive_view="体检与检查", report_document_id=document_id, rerun=False,
    )


def render_manager_dashboard() -> None:
    patients = _patient_map()
    _page_header("今日", "今天需要优先处理的健康事项。", eyebrow="今日健康运营")
    with SessionLocal() as session:
        work_items = OperationalWorklistService().list_items(session, datetime.now(TOKYO_TIMEZONE))
    _status_strip(
        ("高风险", sum(item.status == "高风险" for item in work_items), "urgent"),
        ("中风险", sum(item.status == "中风险" for item in work_items), "attention"),
        ("今日跟进", sum(item.status in {"今日跟进", "逾期", "建议健康管理", "待随访"} for item in work_items), "action"),
        ("等待医生", sum(item.status == "等待医生" for item in work_items), "action"),
    )

    def render_dashboard_item(item, key: str) -> None:
        member = patients.get(item.member_id)
        if member is None:
            return
        if item.source_type == "report_review" and item.document_id:
            callback, args = _open_report_review_from_worklist, (member.id, item.document_id)
        elif item.route_target == "member_management":
            callback, args = _open_member_management, (member.id,)
        elif item.route_target == "member_service":
            callback, args = _open_member_service, (member.id,)
        elif item.route_target == "doctor_review":
            callback, args = _open_member, (member.id,)
        else:
            callback, args = _open_member, (member.id,)
        work_item_card(_member_display(member), item.status, item.title, item.reason, item.next_action, key=key, owner=item.owner, due_at=item.due_at, on_click=callback, args=args)

    with section_frame("优先处理", "每项只保留发生原因与下一步，进入后再查看完整成员资料。"):
        if not work_items:
            _empty_state("暂无待处理事项", "今天没有需要您处理的健康运营事项。")
            return
        for item in work_items[:5]:
            render_dashboard_item(item, f"today-{item.source_type}-{item.source_id}")
        if len(work_items) > 5:
            with st.expander(f"查看其余 {len(work_items) - 5} 项"):
                for item in work_items[5:]:
                    render_dashboard_item(item, f"today-more-{item.source_type}-{item.source_id}")


def _render_member_header(patient: Patient, ctx: dict[str, list[object]]) -> None:
    risk_label, _, _ = _member_risk_state(patient.id, ctx)
    open_problems = [item for item in ctx["problems"] if item.status != "CLOSED"]
    open_tasks = [item for item in ctx["tasks"] if item.status not in {"COMPLETED", "CANCELLED"}]
    next_task = next((item for item in open_tasks if item.due_at), None)
    name = html.escape(patient.display_name or "未命名成员")
    doctor_reviews = len([item for item in ctx["alerts"] if item.status == "WAITING_DOCTOR_REVIEW"])
    next_review = _fmt_dt(next_task.due_at) if next_task and next_task.due_at else "暂无安排"
    st.markdown(
        f"<div class='member-hero'><h1>{name}</h1><p>{html.escape(_age(patient))} · 成员健康档案</p>"
        f"{risk_badge(risk_label)}<div class='hero-facts'>"
        f"<div class='hero-fact'><b>{len(open_problems)}</b><span>当前管理问题</span></div>"
        f"<div class='hero-fact'><b>{len(open_tasks)}</b><span>待办事项</span></div>"
        f"<div class='hero-fact'><b>{doctor_reviews}</b><span>医生待复核</span></div>"
        f"<div class='hero-fact'><b>{html.escape(next_review)}</b><span>下次复核</span></div>"
        f"</div></div>",
        unsafe_allow_html=True,
    )
    with SessionLocal() as session:
        baseline = HealthAssessmentService().latest_baseline(session, patient.id)
    if baseline is None:
        st.caption("健康基线：尚未建立")
    elif baseline.status == "DRAFT":
        left, right = st.columns([5, 1])
        left.caption("健康基线初稿等待确认")
        right.button("处理", key=f"member-baseline-draft-{patient.id}", on_click=_open_baseline, args=(patient.id,))
    else:
        st.caption(f"健康基线：已建立 · {_fmt_dt(baseline.confirmed_at or baseline.assessed_at)}")


def _open_baseline(member_id: UUID) -> None:
    request_navigation(
        ops_page="成员", member_id=member_id, member_section="健康",
        archive_view="基线", rerun=False,
    )


def _open_member_timeline(member_id: UUID) -> None:
    request_navigation(
        ops_page="成员", member_id=member_id, member_section="历程", rerun=False,
    )


def _problem_related(ctx: dict[str, list[object]], problem: HealthProblem) -> dict[str, list[object]]:
    return {
        "alerts": [item for item in ctx["alerts"] if item.health_problem_id == problem.id],
        "plans": [item for item in ctx["plans"] if item.health_problem_id == problem.id],
        "tasks": [item for item in ctx["tasks"] if item.health_problem_id == problem.id],
        "reviews": [item for item in ctx["reviews"] if item.health_problem_id == problem.id],
        "followups": [item for item in ctx["followups"] if item.health_problem_id == problem.id],
    }


def _render_problem_card(patient: Patient, problem: HealthProblem, ctx: dict[str, list[object]]) -> None:
    related = _problem_related(ctx, problem)
    with st.container(border=True):
        st.markdown(f"### {problem.title}　{_severity_badge(problem.severity)} {_status_badge(problem.status, problem.severity)}")
        st.caption(f"负责人：{_role_label(problem.responsible_role, name=problem.owner)} · 创建：{_fmt_dt(problem.opened_at)} · 截止：{_fmt_dt(next((item.due_at for item in related['tasks'] if item.due_at), None))}")
        st.write(problem.description)
        steps = [
            ("发现健康问题", True), ("建立管理方案", bool(related["plans"])), ("执行任务", bool(related["tasks"])),
            ("医生复核", bool(related["reviews"])), ("随访", bool(related["followups"])), ("完成", problem.status == "CLOSED"),
        ]
        st.markdown("　→　".join(f"{'✅' if done else '○'} {name}" for name, done in steps))
        columns = st.columns(2)
        with columns[0]:
            st.markdown("**关联健康异常**")
            for item in related["alerts"]:
                st.write(f"• {item.title} · {_label(item.status)}")
            st.markdown("**当前管理方案**")
            for item in related["plans"]:
                st.write(f"• {item.title} · {_label(item.status)}")
                st.caption(item.content)
            if not related["plans"]:
                st.caption("尚未形成经医生确认的管理计划。")
            st.markdown("**执行任务**")
            for item in related["tasks"]:
                st.write(f"• {item.title} · {_label(item.status)} · {_role_label(item.responsible_role, name=item.assignee)}")
        with columns[1]:
            st.markdown("**医生复核**")
            for item in related["reviews"]:
                st.write(f"• {item.department} · {item.doctor_name} · {_fmt_dt(item.reviewed_at)}")
                st.caption(item.opinion)
            if not related["reviews"]:
                st.caption("尚待医生复核。")
            st.markdown("**随访与闭环**")
            for item in related["followups"]:
                st.write(f"• {_label(item.status)} · {_fmt_dt(item.completed_at)} · {item.reviewed_by or '—'}")
                if item.outcome:
                    st.caption(item.outcome)
            if problem.status == "CLOSED":
                st.success(f"闭环结果：已由人工随访关闭（{_fmt_dt(problem.closed_at)}）。")
        if problem.status != "CLOSED":
            with st.expander("记录随访并关闭健康问题"):
                task_options = {"不关联具体任务": None, **{f"{item.title} ({_label(item.status)})": item.id for item in related["tasks"]}}
                with st.form(f"followup-{problem.id}"):
                    reviewer = st.text_input("随访记录人", value="健康管理师")
                    outcome = st.text_area("随访结果（人工记录）", placeholder="记录已完成的联系、复查或交接结果；不填写诊断或改药指令。")
                    task_label = st.selectbox("关联执行任务", list(task_options), key=f"followup-task-{problem.id}")
                    submit = st.form_submit_button("保存随访并关闭")
                if submit:
                    if not reviewer.strip() or not outcome.strip():
                        st.error("请填写随访记录人与结果。")
                    else:
                        with SessionLocal() as session:
                            stored = session.get(HealthProblem, problem.id)
                            task = session.get(Task, task_options[task_label]) if task_options[task_label] else None
                            complete_follow_up(session, stored, reviewer.strip(), outcome.strip(), task)
                            session.commit()
                        st.success("已写入随访、关闭关联事项并保留审计记录。")
                        st.rerun()


def render_problems(patient: Patient, ctx: dict[str, list[object]]) -> None:
    st.subheader("健康问题")
    problems = ctx["problems"]
    if not problems:
        st.info("暂无健康问题。仅在人工确认需要持续处理时创建。")
    for problem in problems:
        _render_problem_card(patient, problem, ctx)


def _render_create_task(patient: Patient, alert: Alert | None = None, problem: HealthProblem | None = None) -> None:
    with st.expander("创建执行任务"):
        with st.form(f"task-{alert.id if alert else problem.id if problem else patient.id}"):
            title = st.text_input("任务内容", value="完成健康管理跟进")
            instruction = st.text_area("执行说明", placeholder="填写经人工确认的运营动作；系统不会自动生成治疗或用药决定。")
            priority = st.selectbox("优先级", ["HIGH", "MEDIUM", "LOW"], format_func=_severity)
            assignee = st.text_input("执行人/角色", value="health_manager")
            due_day = st.date_input("截止日期", value=date.today() + timedelta(days=7))
            submit = st.form_submit_button("创建任务")
        if submit:
            if not title.strip() or not instruction.strip() or not assignee.strip():
                st.error("请完整填写任务内容、说明和执行人。")
            else:
                due_at = datetime.combine(due_day, time(17, 0), tzinfo=TOKYO_TIMEZONE)
                with SessionLocal() as session:
                    stored_alert = session.get(Alert, alert.id) if alert else None
                    stored_problem = session.get(HealthProblem, problem.id) if problem else None
                    create_operational_task(session, patient.id, title.strip(), instruction.strip(), priority, assignee.strip(), "健康管理师", due_at, stored_alert, stored_problem)
                    session.commit()
                st.success("已创建可追溯的执行任务。")
                st.rerun()


def _render_alert_card(patient: Patient, alert: Alert, ctx: dict[str, list[object]]) -> None:
    with st.container(border=True):
        st.markdown(f"### {alert.title}　{_severity_badge(alert.severity)} {_status_badge(alert.status, alert.severity)}")
        st.caption(f"负责人：{_role_label(alert.responsible_role, name=alert.owner)} · 创建：{_fmt_dt(alert.created_at)} · 截止：{_fmt_dt(alert.due_at)}")
        st.write(alert.finding)
        dates = alert.evidence_json.get("local_dates", []) if alert.evidence_json else []
        if dates:
            st.caption("筛查证据日期：" + "、".join(dates) + "（用于人工核实，不构成诊断）")
        if alert.review_note:
            st.info(f"核实记录：{alert.review_note}")
        if alert.status in {"NEW", "AI_SCREENED", "WAITING_MANAGER_REVIEW"}:
            open_problems = [item for item in ctx["problems"] if item.status != "CLOSED"]
            choices = {"新建健康问题": None, **{f"关联：{item.title}": item.id for item in open_problems}}
            with st.form(f"alert-review-{alert.id}"):
                reviewer = st.text_input("健康管理师", value="健康管理师", key=f"reviewer-{alert.id}")
                disposition = st.radio("核实结果", ["确认数据有效，升级医生", "确认数据无效，关闭误报"], key=f"disposition-{alert.id}")
                note = st.text_area("核实记录", placeholder="记录测量来源、测量方式或联系结果。", key=f"note-{alert.id}")
                link_label = st.selectbox("健康问题处理", list(choices), key=f"link-{alert.id}")
                submit = st.form_submit_button("保存人工核实")
            if submit:
                if not reviewer.strip() or not note.strip():
                    st.error("请填写健康管理师姓名和核实记录。")
                else:
                    with SessionLocal() as session:
                        stored = session.get(Alert, alert.id)
                        if disposition == "确认数据无效，关闭误报":
                            close_alert_as_false_positive(session, stored, reviewer.strip(), note.strip())
                        else:
                            confirm_alert_as_manager(session, stored, reviewer.strip(), note.strip(), choices[link_label])
                        session.commit()
                    st.success("人工核实已保存，并已写入审计记录。")
                    st.rerun()
        elif alert.status == "WAITING_DOCTOR_REVIEW":
            st.warning("已完成健康管理师确认，下一步为医生复核。请在“医生复核工作台”填写人工意见。")
        elif alert.status == "IN_FOLLOW_UP":
            st.info("医生复核与管理方案已形成，正在等待执行任务和随访完成。")
        _render_create_task(patient, alert=alert)


def render_alerts(patient: Patient, ctx: dict[str, list[object]]) -> None:
    st.subheader("健康异常与事件")
    st.caption("处理路径：新建 → 规则筛查完成 → 管理师核实 → 医生复核 → 随访 → 已关闭")
    if not ctx["alerts"]:
        st.info("暂无健康异常。筛查发现的异常必须先由人工核实。")
    for alert in ctx["alerts"]:
        _render_alert_card(patient, alert, ctx)


def render_tasks(ctx: dict[str, list[object]]) -> None:
    st.subheader("执行任务")
    tasks = ctx["tasks"]
    if not tasks:
        st.info("暂无执行任务。")
        return
    now = datetime.now(TOKYO_TIMEZONE)
    for task in tasks:
        overdue = task.status not in {"COMPLETED", "CANCELLED"} and task.due_at and task.due_at < now
        severity = "OVERDUE" if overdue else task.priority
        columns = st.columns([6, 1])
        columns[0].markdown(f"**{task.title}** {_severity_badge(severity)} {_status_badge('OVERDUE' if overdue else task.status, severity)}  \n\n执行人：{_role_label(task.responsible_role, name=task.assignee)} · 截止：{_fmt_dt(task.due_at)}  \n\n{task.instruction}")
        if task.status not in {"COMPLETED", "CANCELLED"} and columns[1].button("标记完成", key=f"complete-{task.id}"):
            with SessionLocal() as session:
                try:
                    TaskTransitionService().complete(
                        session,
                        task.id,
                        actor=task.assignee or "健康管理师",
                        outcome="已在健康运营工作台记录任务完成。",
                    )
                    session.commit()
                except ValueError as error:
                    session.rollback()
                    st.error(str(error))
                else:
                    st.rerun()


def _recent_observation_table(observations: list[Observation]) -> pd.DataFrame:
    return pd.DataFrame([
        {
            "记录时间": _fmt_dt(item.observed_at),
            "指标": _metric_display_name(item.metric_code),
            "数值": float(item.value_numeric) if item.value_numeric is not None else "暂无数据",
            "单位": item.unit or "未记录",
            "数据来源": get_provider_display(item.source),
            "数据质量": get_quality_display(item.quality_flag),
        }
        for item in observations[:200]
    ])


def _format_observation_value(observation: Observation) -> str:
    if observation.value_numeric is None:
        return "暂无数据"
    try:
        value = float(observation.value_numeric)
    except (TypeError, ValueError):
        return "暂无数据"
    if observation.metric_code == "sleep_duration" and observation.unit in {"minutes", "min"}:
        hours, minutes = divmod(round(value), 60)
        return f"{hours}小时{minutes}分钟"
    if observation.metric_code == "steps":
        return f"{round(value):,} 步"
    return f"{value:g} {observation.unit}" if observation.unit else f"{value:g}"


def render_doctor_reviews(patient: Patient, ctx: dict[str, list[object]]) -> None:
    st.subheader("医生复核")
    yellow_pending = [item for item in ctx["reviews"] if item.status == "PENDING" and item.risk_event_id]
    if yellow_pending:
        st.markdown("#### 来自健康数据自动监测的待复核")
    for review in yellow_pending:
        with SessionLocal() as session:
            event = session.get(RiskEvent, review.risk_event_id)
        with st.container(border=True):
            st.markdown("**需要关注 · 健康数据自动监测**")
            st.write(review.doctor_brief)
            st.markdown("**健康管理师希望医生确认什么**")
            st.write(review.question_for_doctor)
            st.caption(f"提交：{_fmt_dt(review.created_at)} · {_risk_event_evidence_caption(event) if event else '相关数据已归档'}")
            with SessionLocal() as session:
                _render_evidence_action(
                    _risk_evidence_payload(session, patient.id, review.risk_event_id),
                    key_scope=f"doctor-risk-{review.id}",
                )
            _render_related_knowledge(review.question_for_doctor or review.doctor_brief, key_scope=f"doctor-review-{review.id}")
            with st.form(f"yellow-doctor-review-{review.id}"):
                doctor = st.text_input("医生姓名", value="演示医生", key=f"yellow-doctor-name-{review.id}")
                department = st.text_input("科室", value=review.department, key=f"yellow-doctor-dept-{review.id}")
                opinion = st.text_area("医生人工意见", key=f"yellow-doctor-opinion-{review.id}")
                instruction = st.text_area("后续跟进任务", placeholder="记录需由健康管理师完成的下一步", key=f"yellow-doctor-task-{review.id}")
                due = st.date_input("建议跟进日期", value=date.today() + timedelta(days=7), key=f"yellow-doctor-due-{review.id}")
                submit = st.form_submit_button("保存医生复核并创建跟进任务")
            if submit:
                try:
                    with SessionLocal() as session:
                        RiskOperationsService().complete_doctor_review(session, review.id, doctor, department, opinion, instruction, datetime.combine(due, time(9, 0), tzinfo=TOKYO_TIMEZONE))
                        session.commit()
                    st.success("已保存医生人工复核，并创建关联跟进任务。")
                    st.rerun()
                except ValueError as error:
                    st.error(str(error))
    outcome_pending = [item for item in ctx["reviews"] if item.status == "PENDING" and item.risk_event_id is None]
    if outcome_pending:
        st.markdown("#### 来自阶段结果的待复核")
    for review in outcome_pending:
        with detail_panel("阶段结果医生复核", "仅请医生完成人工医学判断；后续执行由健康管理团队承接。"):
            st.markdown("**提交原因**")
            st.write(review.doctor_brief)
            st.markdown("**需要医生回答什么**")
            st.write(review.question_for_doctor)
            _render_related_knowledge(review.question_for_doctor or review.doctor_brief, key_scope=f"doctor-outcome-{review.id}")
            with st.form(f"outcome-doctor-review-{review.id}"):
                doctor = st.text_input("医生姓名", value="演示医生", key=f"outcome-doctor-name-{review.id}")
                department = st.text_input("科室", value=review.department, key=f"outcome-doctor-dept-{review.id}")
                opinion = st.text_area("医生人工意见", key=f"outcome-doctor-opinion-{review.id}")
                instruction = st.text_area("后续跟进任务", placeholder="记录需由健康管理师完成的下一步", key=f"outcome-doctor-task-{review.id}")
                due = st.date_input("建议跟进日期", value=date.today() + timedelta(days=7), key=f"outcome-doctor-due-{review.id}")
                submit = st.form_submit_button("保存医生复核并创建跟进任务")
            if submit:
                try:
                    with SessionLocal() as session:
                        complete_outcome_doctor_review(session, session.get(DoctorReview, review.id), doctor, department, opinion, instruction, datetime.combine(due, time(9, 0), tzinfo=TOKYO_TIMEZONE))
                        session.commit()
                    st.success("已保存医生人工复核，并创建关联跟进任务。")
                    st.rerun()
                except ValueError as error:
                    st.error(str(error))
    pending = [item for item in ctx["alerts"] if item.status == "WAITING_DOCTOR_REVIEW" and item.health_problem_id]
    if pending:
        st.markdown("#### 需要医生回答的问题")
    for alert in pending:
        with SessionLocal() as session:
            problem = session.get(HealthProblem, alert.health_problem_id)
            brief = build_doctor_brief(session, patient.id, problem, alert) if problem else "未找到关联健康问题。"
        with st.container(border=True):
            st.markdown(f"**{alert.title}**　{_status_badge(alert.status, alert.severity)}")
            st.markdown("**提交原因**")
            st.write(alert.finding)
            left, right = st.columns(2)
            with left:
                st.markdown("**相关健康问题与异常**")
                st.write(problem.title if problem else "—")
                st.caption(alert.finding)
                st.markdown("**当前用药信息**")
                meds = [item for item in ctx["med_plans"] if item.status == "active"]
                st.write("；".join(f"{item.drug_name} {item.dose}{item.dose_unit}" for item in meds) or "暂无当前用药记录")
            with right:
                st.markdown("**近期健康数据（最近 12 条）**")
                st.dataframe(_recent_observation_table(ctx["observations"]).head(12), hide_index=True, width="stretch")
            with st.form(f"doctor-review-{alert.id}"):
                doctor = st.text_input("医生姓名", value="演示医生")
                department = st.selectbox("科室", ["心内科", "内分泌科", "睡眠相关咨询", "全科/健康管理"], key=f"department-{alert.id}")
                further = st.text_input("建议进一步检查（可留空）", placeholder="记录医生人工建议")
                follow_cycle = st.text_input("建议随访周期（可留空）", placeholder="例如：7 天后由医疗团队复核")
                opinion = st.text_area("医生意见 / 确认", placeholder="仅记录医生人工确认的意见；不得由系统自动生成诊断或用药调整。")
                submit = st.form_submit_button("确认医生复核并创建管理方案与任务")
            if submit:
                if not doctor.strip() or not opinion.strip():
                    st.error("请填写医生姓名和人工意见。")
                else:
                    additions = [opinion.strip()]
                    if further.strip():
                        additions.append(f"进一步检查：{further.strip()}")
                    if follow_cycle.strip():
                        additions.append(f"随访周期：{follow_cycle.strip()}")
                    with SessionLocal() as session:
                        stored_problem = session.get(HealthProblem, alert.health_problem_id)
                        stored_alert = session.get(Alert, alert.id)
                        record_doctor_review(session, stored_problem, doctor.strip(), department, "\n".join(additions), stored_alert)
                        session.commit()
                    st.success("已记录医生人工复核，并创建关联管理方案与执行任务。")
                    st.rerun()
            with st.expander("查看整理摘要与相关数据"):
                st.write(brief)
            _render_evidence_action(
                {
                    "source_name": "医生复核提交资料",
                    "location": "健康管理师提交记录",
                    "evidence_type": "TEXT",
                    "raw_evidence": alert.finding,
                    "structured_interpretation": "提交医生复核的问题与相关健康资料。",
                    "confirmation_status": _label(alert.status),
                    "evidence_status": "PARTIAL",
                    "show_no_knowledge": True,
                },
                key_scope=f"doctor-alert-{alert.id}",
            )
    if not pending:
        st.info("当前没有待医生复核的健康异常事项。以下为已确认的医生复核记录。")
    if ctx["reviews"]:
        st.markdown("#### 已完成医生复核")
        for review in ctx["reviews"]:
            with st.expander(f"{_fmt_dt(review.reviewed_at)} · {review.department} · {review.doctor_name}"):
                st.markdown("**需要医生回答的问题**")
                st.write(review.question_for_doctor)
                st.markdown("**人工确认意见**")
                st.write(review.opinion)
                with st.expander("查看当时整理摘要"):
                    st.write(review.doctor_brief)


def render_observations(ctx: dict[str, list[object]]) -> None:
    st.subheader("关键健康趋势")
    st.caption("先查看近期趋势；详细记录和技术来源收纳在页面下方。")
    observations = ctx["observations"]
    if not observations:
        st.info("暂无健康数据。")
        return
    latest_by_code: dict[str, Observation] = {}
    for observation in observations:
        latest_by_code.setdefault(observation.metric_code, observation)
    trend_columns = st.columns(min(5, max(1, len(latest_by_code))))
    for column, code in zip(trend_columns, ("systolic_bp", "weight", "sleep_duration", "steps", "heart_rate")):
        item = latest_by_code.get(code)
        if item:
            column.metric(display_observation(code), _format_observation_value(item))

    source_options = sorted({item.source for item in observations})
    quality_options = sorted({item.quality_flag for item in observations})
    filters = st.columns(3)
    selected_sources = filters[0].multiselect("数据来源", source_options, default=source_options, format_func=display_provider)
    selected_quality = filters[1].multiselect("数据质量", quality_options, default=quality_options, format_func=get_quality_display)
    date_from = filters[2].date_input("记录时间从", value=min(item.observed_at.date() for item in observations))
    observations = [item for item in observations if item.source in selected_sources and item.quality_flag in selected_quality and item.observed_at.date() >= date_from]
    frame = _recent_observation_table(observations)
    if frame.empty:
        st.info("当前筛选条件下暂无健康数据。")
        return
    bp = frame[frame["指标"].isin(["收缩压", "舒张压"])]
    if not bp.empty:
        chart = bp.pivot_table(index="记录时间", columns="指标", values="数值", aggfunc="first").sort_index()
        st.line_chart(chart.tail(60))
    with st.expander("查看详细记录"):
        selected = st.multiselect("显示指标", sorted(frame["指标"].unique()), default=sorted(frame["指标"].unique())[:4])
        st.dataframe(frame[frame["指标"].isin(selected)] if selected else frame.iloc[0:0], hide_index=True, width="stretch")


def _downsample_for_chart(frame: pd.DataFrame, max_points: int = 320) -> pd.DataFrame:
    if len(frame) <= max_points:
        return frame
    step = max(1, (len(frame) + max_points - 1) // max_points)
    return frame.iloc[::step].copy()


def _observation_text(observation: Observation | None) -> str:
    if observation is None:
        return "暂无数据"
    value = float(observation.value_numeric)
    if observation.metric_code == "sleep_duration" and observation.unit in {"minutes", "min"}:
        hours, minutes = divmod(round(value), 60)
        return f"{hours}小时{minutes}分钟"
    if observation.metric_code == "steps":
        return f"{round(value):,} 步"
    return f"{value:g} {observation.unit}"


def _observation_when(observation: Observation | None, *, today_label: str = "最近记录") -> str:
    if observation is None:
        return "最近暂无记录"
    local = observation.observed_at.astimezone(TOKYO_TIMEZONE)
    label = "今日" if local.date() == datetime.now(TOKYO_TIMEZONE).date() else today_label
    return f"{label} {local.strftime('%m月%d日 %H:%M')}"


def _observation_freshness(observation: Observation | None) -> str:
    """A product freshness hint, not a clinical threshold or risk judgement."""
    if observation is None:
        return "暂无足够数据"
    age = datetime.now(TOKYO_TIMEZONE) - observation.observed_at.astimezone(TOKYO_TIMEZONE)
    days = max(0, age.days)
    medical = observation.metric_code in {"systolic_bp", "diastolic_bp", "blood_glucose", "cgm_glucose", "spo2"}
    recent_limit = 2 if medical else 3
    stale_limit = 7 if medical else 14
    if days <= recent_limit:
        return "数据较新"
    if days <= stale_limit:
        return f"最后记录：{days} 天前 · 较久未更新"
    return f"最后记录：{days} 天前 · 数据较旧，建议补测或检查设备"


def _frame_from_observations(observations: list[Observation], codes: tuple[str, ...]) -> pd.DataFrame:
    return pd.DataFrame([
        {"记录时间": item.observed_at, "指标": display_observation(item.metric_code), "数值": float(item.value_numeric), "单位": item.unit}
        for item in observations if item.metric_code in codes
    ])


def _daily_frame(observations: list[Observation], code: str) -> pd.DataFrame:
    return pd.DataFrame([
        {"日期": item.observed_at.astimezone(TOKYO_TIMEZONE).date(), "数值": float(item.value_numeric)}
        for item in observations
        if item.metric_code == code
    ])


def _seven_day_change(frame: pd.DataFrame, unit: str) -> str | None:
    if frame.empty or len(frame) < 2:
        return None
    values = frame.sort_values("日期")["数值"].tolist()
    recent = sum(values[-7:]) / min(7, len(values))
    previous_values = values[-14:-7]
    if not previous_values:
        return f"过去7天平均：{recent:,.0f} {unit}"
    previous = sum(previous_values) / len(previous_values)
    if previous == 0:
        return f"过去7天平均：{recent:,.0f} {unit}"
    change = (recent - previous) / previous * 100
    arrow = "↑" if change > 0 else "↓" if change < 0 else "→"
    return f"过去7天平均：{recent:,.0f} {unit} · 较前7天 {arrow} {abs(change):.0f}%"


def _long_metric_summary(observations: list[Observation], code: str) -> tuple[str, str] | None:
    records = [item for item in observations if item.metric_code == code]
    if not records:
        return None
    start, current = records[0], records[-1]
    start_value, current_value = float(start.value_numeric), float(current.value_numeric)
    if code == "sleep_duration" and current.unit in {"minutes", "min"}:
        start_text = _observation_text(start)
        current_text = _observation_text(current)
        change = current_value - start_value
        return f"{start_text} → {current_text}", f"{'↑' if change > 0 else '↓' if change < 0 else '→'} {abs(change):.0f}分钟"
    change = current_value - start_value
    return f"{start_value:g} → {current_value:g} {current.unit}", f"{'↑' if change > 0 else '↓' if change < 0 else '→'} {abs(change):g} {current.unit}"


def _blood_pressure_long_summary(observations: list[Observation]) -> tuple[str, str] | None:
    systolic = [item for item in observations if item.metric_code == "systolic_bp"]
    diastolic = [item for item in observations if item.metric_code == "diastolic_bp"]
    if not systolic or not diastolic:
        return None
    start_sys, current_sys = float(systolic[0].value_numeric), float(systolic[-1].value_numeric)
    start_dia, current_dia = float(diastolic[0].value_numeric), float(diastolic[-1].value_numeric)
    delta = current_sys - start_sys
    return f"{start_sys:g}/{start_dia:g} → {current_sys:g}/{current_dia:g} mmHg", f"{'↑' if delta > 0 else '↓' if delta < 0 else '→'} {abs(delta):g} mmHg"


def _render_realtime_section(patient_id: UUID, service: HealthDataSummaryService) -> list[Observation]:
    st.subheader("健康监测")
    st.caption("只显示最近有效的医疗监测与生命体征；趋势按需展开。")
    with SessionLocal() as session:
        realtime = service.get_realtime_summary(session, patient_id)
        glucose_event = session.scalar(select(RiskEvent).where(
            RiskEvent.patient_id == patient_id,
            RiskEvent.canonical_code == "glucose",
            RiskEvent.status.not_in(("CLOSED", "DISMISSED_DATA_ISSUE")),
        ).order_by(RiskEvent.updated_at.desc()).limit(1))
    if realtime.cgm_current:
        with st.container(border=True):
            st.markdown("**动态血糖（CGM）**")
            st.metric("当前", _observation_text(realtime.cgm_current), _observation_when(realtime.cgm_current, today_label="更新于"))
            st.caption("状态：需要审核" if glucose_event else "状态：正常")
            if st.button("查看血糖详情", key=f"cgm-detail-open-{patient_id}"):
                st.session_state[f"cgm-detail-visible-{patient_id}"] = True
    else:
        _empty_state("暂无动态血糖数据", "连接医疗监测设备后会自动显示最近血糖情况。")
    left, right = st.columns(2)
    with left:
        with st.container(border=True):
            if realtime.latest_systolic and realtime.latest_diastolic:
                st.metric("最近血压", f"{float(realtime.latest_systolic.value_numeric):g} / {float(realtime.latest_diastolic.value_numeric):g} mmHg")
                st.caption(_observation_when(realtime.latest_systolic))
            else:
                st.metric("最近血压", "暂无数据")
    with right:
        with st.container(border=True):
            st.metric("最近心率", _observation_text(realtime.latest_heart_rate))
            st.caption(_observation_when(realtime.latest_heart_rate))
    if not st.session_state.get(f"cgm-detail-visible-{patient_id}"):
        return []
    _section_header("血糖详情")
    window_label = st.radio("动态血糖时间", ["24小时", "7天", "14天"], horizontal=True, label_visibility="collapsed", key=f"cgm-window-{patient_id}")
    hours = {"24小时": 24, "7天": 24 * 7, "14天": 24 * 14}[window_label]
    with SessionLocal() as session:
        cgm_series = service.get_cgm_series(session, patient_id, hours=hours)
    frame = _frame_from_observations(cgm_series, ("glucose",))
    if not frame.empty:
        st.caption(f"最近{window_label}动态曲线")
        st.line_chart(_downsample_for_chart(frame).set_index("记录时间")[["数值"]], height=230)
    else:
        _empty_state("暂无血糖趋势数据", "当前时间范围内尚无连续血糖记录。")
    return cgm_series


def _render_lifestyle_section(patient_id: UUID, service: HealthDataSummaryService) -> list[Observation]:
    st.subheader("今天的生活状态")
    st.caption("睡眠、活动和恢复数据来自已连接或演示的数据来源。")
    with SessionLocal() as session:
        lifestyle = service.get_lifestyle_summary(session, patient_id, days=14)
        sleep_sessions = service.get_sleep_sessions(session, patient_id, days=1)
        open_signals = list(session.scalars(select(ManagementSignal).where(
            ManagementSignal.patient_id == patient_id,
            ManagementSignal.status == "OPEN",
            ManagementSignal.metric_code.in_(("steps", "exercise_minutes", "sleep_duration", "deep_sleep_duration")),
        ).order_by(ManagementSignal.last_detected_at.desc()).limit(20)))
    signal_by_metric = {signal.metric_code: signal for signal in open_signals}

    st.markdown("#### 今日活动")
    steps = lifestyle.latest.get("steps")
    calories = lifestyle.latest.get("active_calories")
    exercise = lifestyle.latest.get("exercise_minutes")
    columns = st.columns(3)
    with columns[0]:
        with st.container(border=True):
            st.metric("步数", _observation_text(steps))
            st.caption(_lifestyle_status_caption(signal_by_metric.get("steps"), steps))
    with columns[1]:
        with st.container(border=True):
            st.metric("活动消耗", _observation_text(calories) if calories else "暂无数据")
            st.caption(_observation_when(calories, today_label="设备提供") if calories else "暂无设备活动消耗数据")
    with columns[2]:
        with st.container(border=True):
            st.metric("运动时间", _observation_text(exercise))
            st.caption(_lifestyle_status_caption(signal_by_metric.get("exercise_minutes"), exercise))
    if steps:
        steps_frame = _daily_frame(lifestyle.daily_values.get("steps", []), "steps")
        if not steps_frame.empty:
            recent = steps_frame.tail(7)
            previous = steps_frame.iloc[-14:-7]
            average = recent["数值"].mean()
            delta = "数据不足"
            if not previous.empty and previous["数值"].mean():
                change = (average - previous["数值"].mean()) / previous["数值"].mean() * 100
                delta = f"相比前7天 {'+' if change >= 0 else ''}{change:.0f}%"
            st.caption(f"7天平均：{average:.0f} 步 · {delta}")
            if st.button("查看步数趋势", key=f"steps-detail-open-{patient_id}"):
                st.session_state[f"steps-detail-visible-{patient_id}"] = True
    if sleep_sessions:
        latest_sleep = sleep_sessions[0]
        sleep_cards = st.columns(2)
        with sleep_cards[0]:
            with st.container(border=True):
                st.metric("睡眠", _sleep_duration_text(latest_sleep.total_sleep_minutes))
        with sleep_cards[1]:
            with st.container(border=True):
                st.metric("深度睡眠", _sleep_duration_text(latest_sleep.deep_sleep_minutes))
        st.caption(_lifestyle_status_caption(signal_by_metric.get("sleep_duration"), None))
        if st.button("查看睡眠详情", key=f"sleep-detail-open-{patient_id}"):
            st.session_state[f"sleep-detail-visible-{patient_id}"] = True
    else:
        _empty_state("暂无睡眠数据", "连接设备后会自动显示总睡眠和睡眠阶段。")
    recovery_columns = st.columns(2)
    for column, (code, label) in zip(recovery_columns, (("resting_heart_rate", "静息心率"), ("spo2", "血氧"))):
        observation = lifestyle.latest.get(code)
        with column:
            with st.container(border=True):
                st.metric(label, _observation_text(observation))
                st.caption(_observation_when(observation, today_label="最近有效值") if observation else "暂无数据")
    if st.session_state.get(f"steps-detail-visible-{patient_id}") and steps:
        _section_header("步数趋势")
        steps_frame = _daily_frame(lifestyle.daily_values.get("steps", []), "steps")
        if not steps_frame.empty:
            st.line_chart(steps_frame.tail(7).set_index("日期")[["数值"]], height=150)
    if st.session_state.get(f"sleep-detail-visible-{patient_id}"):
        with SessionLocal() as session:
            detail_sessions = service.get_sleep_sessions(session, patient_id, days=90)
        _render_sleep_detail(detail_sessions, service, signal_by_metric.get("sleep_duration"))
    return [item for records in lifestyle.daily_values.values() for item in records]


def _sleep_duration_text(minutes: int | None) -> str:
    if minutes is None:
        return "—"
    return f"{minutes // 60}h {minutes % 60:02d}m"


def _lifestyle_status_caption(signal: ManagementSignal | None, observation: Observation | None) -> str:
    if signal:
        label = {"NORMAL": "正常", "WATCH": "持续观察", "ACTION_NEEDED": "建议健康管理"}.get(signal.severity, "建议健康管理")
        return f"状态：{label} · {ROUTE_LABELS.get(signal.recommended_route, signal.recommended_route)}"
    return "状态：正常" if observation else "暂无数据"


def _sleep_stage_timeline(session_item: SleepSession, service: HealthDataSummaryService) -> None:
    intervals = service.sleep_stage_intervals(session_item)
    if not intervals:
        st.info("当前设备仅提供总睡眠时间，暂无睡眠阶段数据。")
        return
    labels = {"DEEP": "深度睡眠", "LIGHT": "浅睡", "REM": "REM", "AWAKE": "清醒"}
    colors = {"DEEP": "#244a8f", "LIGHT": "#75a6e5", "REM": "#8762c6", "AWAKE": "#e89a51"}
    start = min(item.start_at for item in intervals)
    end = max(item.end_at for item in intervals)
    total = max((end - start).total_seconds(), 1)
    blocks: list[str] = []
    for index, item in enumerate(intervals):
        left = (item.start_at - start).total_seconds() / total * 100
        width = max(item.duration_minutes * 60 / total * 100, 0.8)
        lane = {"AWAKE": 0, "REM": 1, "LIGHT": 2, "DEEP": 3}[item.stage]
        blocks.append(f"<span title='{html.escape(labels[item.stage])} {item.start_at.astimezone(TOKYO_TIMEZONE).strftime('%H:%M')}–{item.end_at.astimezone(TOKYO_TIMEZONE).strftime('%H:%M')}' style='position:absolute;left:{left:.3f}%;width:{width:.3f}%;top:{lane * 29 + 3}px;height:22px;background:{colors[item.stage]};border-radius:3px;'></span>")
    lane_labels = "".join(f"<span style='height:29px;display:block'>{label}</span>" for label in ("清醒", "REM", "浅睡", "深睡"))
    st.markdown(
        f"<div style='display:grid;grid-template-columns:62px 1fr;gap:8px;font-size:.82rem'><div>{lane_labels}</div><div><div style='position:relative;height:116px;border:1px solid #dce5ef;border-radius:6px;background:#fff'>{''.join(blocks)}</div><div style='display:flex;justify-content:space-between;margin-top:3px'><span>{start.astimezone(TOKYO_TIMEZONE).strftime('%H:%M')}</span><span>{end.astimezone(TOKYO_TIMEZONE).strftime('%H:%M')}</span></div></div></div>",
        unsafe_allow_html=True,
    )
    st.caption("深度睡眠 · 浅睡 · REM · 清醒（设备实际提供的睡眠阶段）")


def _render_sleep_detail(sessions: list[SleepSession], service: HealthDataSummaryService, signal: ManagementSignal | None) -> None:
    st.markdown("#### 睡眠详情")
    if not sessions:
        st.caption("暂无设备睡眠记录。")
        return
    latest = sessions[0]
    stage_intervals = service.sleep_stage_intervals(latest)
    stage_minutes = {stage: sum(item.duration_minutes for item in stage_intervals if item.stage == stage) for stage in ("DEEP", "LIGHT", "REM", "AWAKE")}
    # A source may provide phases only in the interval payload.  These are
    # direct source durations, not estimates from total sleep.
    deep = latest.deep_sleep_minutes if latest.deep_sleep_minutes is not None else (round(stage_minutes["DEEP"]) if stage_minutes["DEEP"] else None)
    rem = latest.rem_sleep_minutes if latest.rem_sleep_minutes is not None else (round(stage_minutes["REM"]) if stage_minutes["REM"] else None)
    light = sum(item.duration_minutes for item in stage_intervals if item.stage == "LIGHT") or None
    interruption_count = sum(item.stage == "AWAKE" for item in stage_intervals) or None
    deep_share = f"{(deep / latest.total_sleep_minutes * 100):.0f}%" if deep is not None and latest.total_sleep_minutes else "—"
    columns = st.columns(5)
    columns[0].metric("总睡眠", _sleep_duration_text(latest.total_sleep_minutes))
    columns[1].metric("深度睡眠", _sleep_duration_text(deep))
    columns[2].metric("深睡占比", deep_share)
    columns[3].metric("入睡时间", latest.sleep_start.astimezone(TOKYO_TIMEZONE).strftime("%H:%M"))
    columns[4].metric("醒来时间", latest.sleep_end.astimezone(TOKYO_TIMEZONE).strftime("%H:%M"))
    secondary = st.columns(4)
    secondary[0].metric("REM", _sleep_duration_text(rem))
    secondary[1].metric("浅睡", _sleep_duration_text(round(light) if light is not None else None))
    awake = latest.awake_minutes if latest.awake_minutes is not None else (round(stage_minutes["AWAKE"]) if stage_minutes["AWAKE"] else None)
    secondary[2].metric("夜间清醒", _sleep_duration_text(awake))
    secondary[3].metric("夜间中断", f"{interruption_count} 次" if interruption_count is not None else "—")
    st.caption(_lifestyle_status_caption(signal, None))
    _sleep_stage_timeline(latest, service)
    window = st.radio("睡眠趋势时间", ["7天", "30天", "3个月"], horizontal=True, label_visibility="collapsed", key=f"sleep-range-{latest.patient_id}")
    days = {"7天": 7, "30天": 30, "3个月": 90}[window]
    relevant = [item for item in sessions if item.sleep_end >= datetime.now(TOKYO_TIMEZONE) - timedelta(days=days)]
    if relevant:
        summary = service.sleep_trend(relevant)
        average_bedtime = "—" if summary.average_bedtime_minutes is None else f"{int(summary.average_bedtime_minutes) // 60:02d}:{int(summary.average_bedtime_minutes) % 60:02d}"
        summary_columns = st.columns(5)
        summary_columns[0].metric(f"{window}平均总睡眠", _sleep_duration_text(round(summary.average_total_minutes) if summary.average_total_minutes is not None else None))
        summary_columns[1].metric("平均深度睡眠", _sleep_duration_text(round(summary.average_deep_minutes) if summary.average_deep_minutes is not None else None))
        summary_columns[2].metric("平均深睡占比", "—" if summary.average_deep_ratio is None else f"{summary.average_deep_ratio:.1f}%")
        summary_columns[3].metric("平均REM", _sleep_duration_text(round(summary.average_rem_minutes) if summary.average_rem_minutes is not None else None))
        summary_columns[4].metric("平均入睡时间", average_bedtime)
        values = pd.DataFrame([{"日期": item.sleep_end.astimezone(TOKYO_TIMEZONE).date(), "总睡眠(分钟)": item.total_sleep_minutes, "深睡(分钟)": item.deep_sleep_minutes, "深睡占比": (item.deep_sleep_minutes / item.total_sleep_minutes * 100) if item.deep_sleep_minutes is not None and item.total_sleep_minutes else None, "夜间中断": item.awake_minutes} for item in reversed(relevant)])
        st.line_chart(values.set_index("日期")[["总睡眠(分钟)", "深睡(分钟)"]], height=150)


def _render_long_term_section(
    patient_id: UUID,
    service: HealthDataSummaryService,
    *,
    start_at: datetime | None = None,
    end_at: datetime | None = None,
) -> list[Observation]:
    st.subheader("长期健康趋势")
    st.caption("查看可调整时间范围内的重要健康变化。")
    range_label = st.radio("长期趋势时间", ["24小时", "7天", "30天", "3个月", "1年", "全部"], horizontal=True, label_visibility="collapsed", key=f"long-range-{patient_id}")
    days = {"24小时": 1, "7天": 7, "30天": 30, "3个月": 90, "1年": 365}.get(range_label)
    with SessionLocal() as session:
        observations = service.get_long_term_observations(session, patient_id, days=days, start_at=start_at, end_at=end_at)
    summaries = []
    blood_pressure = _blood_pressure_long_summary(observations)
    if blood_pressure:
        summaries.append({"指标": "血压", "变化": blood_pressure[1], "起始与当前": blood_pressure[0], "code": "systolic_bp"})
    for code, title in (("weight", "体重"), ("sleep_duration", "睡眠"), ("steps", "步数"), ("resting_heart_rate", "静息心率"), ("glucose", "血糖")):
        summary = _long_metric_summary(observations, code)
        if summary:
            summaries.append({"指标": title, "变化": summary[1], "起始与当前": summary[0], "code": code})
    if not summaries:
        st.caption("当前时间范围内暂无长期趋势数据。")
        return observations
    columns = st.columns(3)
    for column, summary in zip(columns, summaries[:3]):
        with column:
            health_metric_card(str(summary["指标"]), str(summary["变化"]), str(summary["起始与当前"]))
    for summary in summaries[:3]:
        frame = _frame_from_observations(observations, (str(summary["code"]),))
        if not frame.empty:
            st.caption(f"{summary['指标']}趋势")
            st.line_chart(_downsample_for_chart(frame).set_index("记录时间")[["数值"]], height=140)
    with st.expander("查看详细趋势"):
        for code, title in (("weight", "体重"), ("sleep_duration", "睡眠"), ("steps", "步数"), ("resting_heart_rate", "静息心率"), ("systolic_bp", "收缩压"), ("glucose", "血糖")):
            frame = _frame_from_observations(observations, (code,))
            if frame.empty:
                continue
            st.caption(f"{title}趋势")
            st.line_chart(_downsample_for_chart(frame).set_index("记录时间")[["数值"]], height=180)
    return observations


def render_health_data(patient_id: UUID) -> None:
    """A member-readable health-data surface; provenance stays out of this view."""
    _page_header("健康数据", "查看日常活动、睡眠、医疗监测和长期趋势。", eyebrow="成员健康中心")
    service = HealthDataSummaryService()
    timeline_window = st.session_state.get(f"health-data-window-{patient_id}")
    start_at: datetime | None = None
    end_at: datetime | None = None
    if isinstance(timeline_window, dict) and timeline_window.get("start") and timeline_window.get("end"):
        try:
            start_at = datetime.fromisoformat(str(timeline_window["start"])).replace(tzinfo=TOKYO_TIMEZONE)
            end_at = datetime.fromisoformat(str(timeline_window["end"])).replace(hour=23, minute=59, second=59, tzinfo=TOKYO_TIMEZONE)
            notice, clear = st.columns([6, 1])
            notice.info(f"正在查看时间轴选择的时间段：{start_at.date()} 至 {end_at.date()}")
            if clear.button("返回默认", key=f"health-data-clear-window-{patient_id}"):
                st.session_state.pop(f"health-data-window-{patient_id}", None)
                st.rerun()
        except ValueError:
            st.session_state.pop(f"health-data-window-{patient_id}", None)
    with SessionLocal() as session:
        realtime_summary = service.get_realtime_summary(session, patient_id)
        lifestyle_summary = service.get_lifestyle_summary(session, patient_id, days=14)
        sleep_sessions = service.get_sleep_sessions(session, patient_id, days=1)
        supplementary = service.latest_for_codes(session, patient_id, ("weight",))
        latest_records = list(session.scalars(
            select(Observation).where(
                Observation.patient_id == patient_id,
                Observation.excluded_from_analysis.is_(False),
                Observation.source_deleted.is_(False),
            ).order_by(Observation.observed_at.desc()).limit(160)
        ))

    _section_header("健康数据概览", "缺失或陈旧数据不会被当作正常状态；请按需要补测或连接设备。")
    sleep = sleep_sessions[0] if sleep_sessions else None
    bp_value = "—"
    if realtime_summary.latest_systolic and realtime_summary.latest_diastolic:
        bp_value = f"{float(realtime_summary.latest_systolic.value_numeric):g} / {float(realtime_summary.latest_diastolic.value_numeric):g}"
    glucose = realtime_summary.cgm_current

    activity_cards = [
        ("步数", lifestyle_summary.latest.get("steps")),
        ("活动消耗", lifestyle_summary.latest.get("active_calories")),
        ("运动时间", lifestyle_summary.latest.get("exercise_minutes")),
    ]
    with section_frame("日常活动", "步数、活动消耗和运动时间会在连接支持设备后自动显示。"):
        available = [(label, item) for label, item in activity_cards if item is not None]
        if not available:
            _empty_state("尚未接入可用活动数据", "连接 Apple Health 或其他支持设备后，这里将显示步数、活动消耗与运动时间。")
        else:
            columns = st.columns(len(available))
            for column, (label, item) in zip(columns, available):
                with column:
                    health_metric_card(label, _observation_text(item), _observation_freshness(item))

    with section_frame("睡眠", "重点查看总睡眠与深度睡眠；趋势会在下方按时间范围展开。"):
        if sleep is None:
            _empty_state("暂无睡眠数据", "连接支持睡眠记录的设备后，这里将显示总睡眠时间、睡眠阶段与趋势。")
        else:
            sleep_age = datetime.now(TOKYO_TIMEZONE) - sleep.sleep_end.astimezone(TOKYO_TIMEZONE)
            sleep_note = "昨晚记录" if sleep_age.days <= 2 else f"最后记录：{sleep_age.days} 天前 · 数据较旧，建议检查设备"
            cards = [("睡眠时长", _sleep_duration_text(sleep.total_sleep_minutes), sleep_note)]
            if sleep.deep_sleep_minutes is not None:
                cards.append(("深度睡眠", _sleep_duration_text(sleep.deep_sleep_minutes), "设备记录"))
            columns = st.columns(len(cards))
            for column, (label, value, note) in zip(columns, cards):
                with column:
                    health_metric_card(label, value, note)

    pulse_cards = [
        ("心率", realtime_summary.latest_heart_rate),
        ("静息心率", lifestyle_summary.latest.get("resting_heart_rate")),
        ("血氧", lifestyle_summary.latest.get("spo2")),
    ]
    with section_frame("心率与血氧", "用于查看已同步的生命体征；缺失数据不会被推断。"):
        available = [(label, item) for label, item in pulse_cards if item is not None]
        if not available:
            _empty_state("暂无近期数据", "连接支持心率或血氧记录的设备后，会在这里自动同步。")
        else:
            columns = st.columns(len(available))
            for column, (label, item) in zip(columns, available):
                with column:
                    health_metric_card(label, _observation_text(item), _observation_freshness(item))

    medical_cards = [
        ("血压", bp_value, realtime_summary.latest_systolic) if realtime_summary.latest_systolic and realtime_summary.latest_diastolic else None,
        ("血糖", _observation_text(glucose), glucose) if glucose is not None else None,
        ("体重", _observation_text(supplementary.get("weight")), supplementary.get("weight")) if supplementary.get("weight") is not None else None,
    ]
    with section_frame("医疗监测", "血压、血糖与体重等记录需结合专业人员的人工判断。"):
        available = [item for item in medical_cards if item is not None]
        if not available:
            _empty_state("暂无医疗监测数据", "连接血压、血糖或体重设备后，这里将显示最近一次可用记录。")
        else:
            columns = st.columns(len(available))
            for column, (label, value, item) in zip(columns, available):
                with column:
                    health_metric_card(label, value, _observation_freshness(item))

    _section_header("周 / 月 / 年汇总", "选择时间跨度后，只展示最关键的长期变化；其他指标按需查看。")
    with section_frame("最近趋势", "睡眠趋势会同时保留总睡眠和深度睡眠，帮助理解睡眠结构。"):
        long_term = _render_long_term_section(patient_id, service, start_at=start_at, end_at=end_at)
    realtime = [item for item in (realtime_summary.cgm_current, realtime_summary.latest_systolic, realtime_summary.latest_diastolic, realtime_summary.latest_heart_rate) if item]
    lifestyle = [item for records in lifestyle_summary.daily_values.values() for item in records]
    detail_records = sorted({item.id: item for item in [*realtime, *lifestyle, *long_term]}.values(), key=lambda item: item.observed_at, reverse=True)
    with st.expander("查看全部健康数据"):
        st.dataframe(_recent_observation_table(detail_records), hide_index=True, width="stretch")
    with st.expander("查看监测详情"):
        _render_realtime_section(patient_id, service)
        _render_lifestyle_section(patient_id, service)
    source_latest: dict[str, Observation] = {}
    for item in latest_records:
        source_latest.setdefault(item.source, item)
    with section_frame("数据来源", "展示最近同步状态；原始追溯记录由授权人员在设备管理中核对。"):
        if not source_latest:
            _empty_state("尚未接入数据来源", "连接设备或录入健康资料后，这里会显示数据来源与最近更新时间。")
        else:
            rows = [
                {"数据来源": get_provider_display(source), "最近更新": _observation_freshness(item), "数据状态": get_quality_display(item.quality_flag)}
                for source, item in list(source_latest.items())[:6]
            ]
            st.dataframe(pd.DataFrame(rows), hide_index=True, width="stretch")


def render_medications(ctx: dict[str, list[object]]) -> None:
    st.subheader("用药信息")
    st.caption("此处展示既有医生管理计划与用户记录；系统不提供自动处方、停药、换药或剂量调整。")
    plans = ctx["med_plans"]
    if plans:
        st.dataframe(pd.DataFrame([{"药物": item.drug_name or "待补充", "剂量": " ".join(part for part in (item.dose, item.dose_unit) if part) or "待补充", "频率": item.frequency or "待补充", "途径": item.route or "待补充", "开立医生": item.prescriber_name or "待补充", "科室": item.department or "待补充", "状态": _label(item.status)} for item in plans]), hide_index=True, width="stretch")
    else:
        st.info("暂无已记录的用药信息。")
    events = ctx["med_events"]
    if events:
        st.markdown("**最近执行记录**")
        st.dataframe(pd.DataFrame([{"计划时间": _fmt_dt(item.scheduled_at), "实际记录": _fmt_dt(item.taken_at), "状态": _label(item.status)} for item in events[:50]]), hide_index=True, width="stretch")


def render_timeline(patient: Patient, ctx: dict[str, list[object]]) -> None:
    st.subheader("健康时间线")
    rows: list[dict[str, object]] = []
    with SessionLocal() as session:
        for item in build_patient_timeline(session, patient.id, days=60):
            rows.append({"at": item.occurred_at, "type": item.category, "title": item.title, "detail": item.detail, "source": item.source})
    for problem in ctx["problems"]:
        rows.append({"at": problem.opened_at, "type": "problem", "title": "创建健康问题", "detail": problem.title, "source": problem.owner or problem.responsible_role})
        if problem.closed_at:
            rows.append({"at": problem.closed_at, "type": "problem", "title": "健康问题已关闭", "detail": problem.title, "source": "健康管理师"})
    for plan in ctx["plans"]:
        rows.append({"at": plan.created_at, "type": "management_plan", "title": "建立管理方案", "detail": plan.title, "source": plan.owner or plan.source})
    for journey in ctx["journeys"]:
        rows.append({"at": journey.created_at, "type": "program_review", "title": "健康评估", "detail": f"{_label(journey.current_stage)} · {journey.main_focus}", "source": journey.owner or "健康管理师"})
    for program in ctx["programs"]:
        rows.append({"at": program.created_at, "type": "program_review", "title": "创建健康管理计划", "detail": f"{program.title} · {program.main_goal}", "source": program.owner or "健康管理师"})
    for barrier in ctx["barriers"]:
        rows.append({"at": barrier.detected_at, "type": "execution_risk", "title": f"执行障碍：{BARRIER.get(barrier.reason, barrier.reason)}", "detail": f"{barrier.description} · 调整：{barrier.resolution or '待处理'}", "source": barrier.confirmed_by})
    for outcome in ctx["outcomes"]:
        rows.append({"at": datetime.combine(outcome.evaluation_date, time.min, tzinfo=TOKYO_TIMEZONE), "type": "outcome", "title": f"阶段效果评估：{display_observation(outcome.metric)}", "detail": f"{outcome.baseline_value}{outcome.unit} → {outcome.current_value}{outcome.unit} · {_label(outcome.result)}", "source": outcome.evaluator})
    for audit in ctx["audits"]:
        if audit.action in {"confirmed_alert", "recorded_doctor_review", "closed_follow_up"}:
            rows.append({"at": audit.created_at, "type": "audit", "title": "人工处理记录", "detail": get_audit_action_display(audit.action), "source": _role_label(audit.actor_role, name=audit.actor)})
    for item in sorted(rows, key=lambda row: row["at"], reverse=True):
        label = TYPE_LABELS.get(str(item["type"]), str(item["type"]).replace("_", " ").title())
        st.markdown(f"**{label} · {_fmt_dt(item['at'])} · {item['title']}**  \\n+{item['detail']} · {item['source']}")


def render_audit(ctx: dict[str, list[object]]) -> None:
    st.subheader("操作记录")
    audits = ctx["audits"]
    if not audits:
        st.info("暂无审计记录。")
        return
    st.dataframe(pd.DataFrame([
        {
            "时间": _fmt_dt(item.created_at),
            "操作者": item.actor or "未记录",
            "角色": _role_label(item.actor_role),
            "处理动作": get_audit_action_display(item.action),
            "处理事项": get_entity_type_display(item.entity_type),
            "处理结果": "已记录",
        }
        for item in audits
    ]), hide_index=True, width="stretch")


def render_overview(patient: Patient, ctx: dict[str, list[object]]) -> None:
    st.subheader("健康概览")
    status, _ = _overall_status(ctx)
    st.info(f"运营摘要：{status}。本系统展示 HealthOps 工作状态，健康数值仅作为人工复核的事实依据。")
    closed = [item for item in ctx["problems"] if item.status == "CLOSED"]
    if closed:
        problem = closed[0]
        related = _problem_related(ctx, problem)
        st.markdown("#### Demo Executive A 的演示医疗子闭环")
        st.markdown("连续 5 天血压数据需人工核实　→　健康管理师核实　→　医生复核　→　健康管理方案　→　跟进任务　→　随访　→　已关闭")
        st.success(f"已完成：{problem.title}，{len(related['reviews'])} 次医生复核、{len(related['tasks'])} 个跟进任务、{len(related['followups'])} 条随访。请在成员历程中查看每一步。")
    else:
        st.caption("闭环将在管理师确认、医生复核与人工随访后逐步出现。")


def render_programs(patient: Patient, ctx: dict[str, list[object]]) -> None:
    st.subheader("健康管理")
    st.caption("把阶段目标、执行任务、调整和复盘放在同一处查看。")
    programs = ctx["programs"]
    if not programs:
        st.info("尚未建立健康评估或健康管理计划。")
        return
    for program in programs:
        phases = sorted([p for p in ctx["phases"] if p.program_id == program.id], key=lambda p: p.sequence)
        problems = sorted([p for p in ctx["problems"] if p.program_id == program.id], key=lambda p: p.priority_rank or 99)
        tasks = [t for t in ctx["tasks"] if t.program_id == program.id]
        barriers = [b for b in ctx["barriers"] if b.program_id == program.id]
        outcomes = [o for o in ctx["outcomes"] if o.program_id == program.id]
        reviews = [r for r in ctx["weekly_reviews"] if r.program_id == program.id]
        with st.container(border=True):
            st.markdown(f"### {display_program_type(program.program_type)}　{_status_badge(program.status)}")
            day = _program_day(program)
            if day:
                st.progress(day / 90, text=f"第{day}天 / 共90天")
            st.markdown(f"**当前阶段：** {display_program_phase(program.current_phase)}")
            st.markdown(f"**主目标：** {program.main_goal}")
            if program.supporting_goals_json:
                st.caption("支持目标：" + " · ".join(program.supporting_goals_json))
            st.markdown("**本周任务**")
            if tasks:
                for task in tasks[:5]:
                    st.write(f"{'✓' if task.status == 'COMPLETED' else '○'} {task.title}")
            else:
                st.caption("暂无执行任务。")
            cols = st.columns(2)
            cols[0].markdown("**重点关注**  \n" + ("  \n".join(problem.title for problem in problems[:3]) or "—"))
            next_phase = next((phase for phase in phases if phase.status == "PLANNED"), None)
            cols[1].markdown(f"**下一次复盘**  \n{_fmt_dt(datetime.combine(next_phase.start_date, time.min, tzinfo=TOKYO_TIMEZONE)) if next_phase else '按计划持续跟进'}")
            if barriers:
                latest = barriers[0]
                st.markdown("**最近调整**")
                st.write(f"因{BARRIER.get(latest.reason, latest.reason)}，{latest.description}")
                st.caption(f"已调整为：{latest.resolution or '等待健康管理师确认'}")
            if reviews:
                st.caption(f"最近复盘：{reviews[0].key_changes} · 下周重点：{reviews[0].next_week_focus}")
            if outcomes:
                st.success("阶段结果：" + "；".join(f"{display_observation(o.metric)} {o.baseline_value}{o.unit} → {o.current_value}{o.unit}（{_label(o.result)}）" for o in outcomes))
                latest_outcome = outcomes[0]
                with st.expander("确认阶段结果后的下一步"):
                    decision = st.radio(
                        "后续管理决定", ["继续当前方案", "调整方案", "进入稳定期", "提交医生复核"],
                        horizontal=True, key=f"outcome-decision-{latest_outcome.id}",
                    )
                    note = st.text_area("说明（可选）", key=f"outcome-decision-note-{latest_outcome.id}", placeholder="记录人工确认的后续管理安排。")
                    if primary_action("确认后续管理决定", key=f"outcome-decision-save-{latest_outcome.id}", width="content"):
                        try:
                            choice = {"继续当前方案": "CONTINUE", "调整方案": "ADJUST", "进入稳定期": "STABILIZE", "提交医生复核": "DOCTOR_REVIEW"}[decision]
                            with SessionLocal() as session:
                                outcome = session.get(OutcomeEvaluation, latest_outcome.id)
                                apply_outcome_decision(session, outcome, choice, "健康管理师", note)
                                session.commit()
                            st.success("已记录阶段结果后的人工管理决定，并创建对应下一步。")
                            st.rerun()
                        except ValueError as error:
                            st.error(str(error))
            if program.next_decision:
                st.info(f"下一步安排：{_label(program.next_decision)}")


def render_member_management_signals(patient: Patient, ctx: dict[str, list[object]]) -> None:
    """Keep lifestyle follow-up actionable without presenting it as medical risk."""
    signals: list[ManagementSignal] = ctx.get("management_signals", [])  # type: ignore[assignment]
    if not signals:
        return
    _section_header("需要持续关注")
    for signal in signals[:3]:
        with st.container(border=True):
            left, action = st.columns([6, 1.2])
            with left:
                state = {"ACTION_NEEDED": "建议健康管理", "WATCH": "持续观察", "NORMAL": "正常"}.get(signal.severity, "持续观察")
                st.markdown(f"**{display_observation(signal.metric_code)} · {state}**")
                st.write(signal.summary)
                st.caption(f"最近发现：{_fmt_dt(signal.last_detected_at)} · 建议：{ROUTE_LABELS.get(signal.recommended_route, signal.recommended_route)}")
            with action:
                if st.button("创建任务", key=f"member-signal-task-{signal.id}", type="primary", width="stretch"):
                    with SessionLocal() as session:
                        current = session.get(ManagementSignal, signal.id)
                        existing = session.scalar(select(Task).where(
                            Task.patient_id == patient.id,
                            Task.source == "management_signal",
                            Task.title == f"跟进：{current.summary[:80]}",
                            Task.status.not_in(("COMPLETED", "CANCELLED")),
                        ))
                        if existing is None:
                            task = create_operational_task(
                                session, patient.id, f"跟进：{current.summary[:80]}",
                                "查看近期健康趋势，记录人工跟进并决定继续观察或调整健康管理。",
                                "MEDIUM", "health_manager", "health_manager",
                            )
                            task.source = "management_signal"
                            current.status = "IN_PROGRESS"
                            session.add(AuditLog(patient_id=current.patient_id, actor="health_manager", actor_role="health_manager", action="management_signal_task_created", entity_type="ManagementSignal", entity_id=str(current.id), detail_json={"task_id": str(task.id)}))
                            session.commit()
                            st.success("已创建健康管理跟进任务。")
                        else:
                            st.info("该信号已有进行中的跟进任务。")


def render_outcomes(ctx: dict[str, list[object]]) -> None:
    st.subheader("阶段效果评估")
    st.caption("仅描述基线与当前可观测变化；不构成疾病治愈或自动医学结论。")
    outcomes = ctx["outcomes"]
    if not outcomes:
        st.info("尚未完成阶段效果评估。")
        return
    st.dataframe(pd.DataFrame([{
        "指标": item.metric, "基线": f"{item.baseline_value} {item.unit}", "当前": f"{item.current_value} {item.unit}",
        "目标": item.target_value or "—", "方向": item.direction, "结果": _label(item.result),
        "评价日期": str(item.evaluation_date), "评价人": item.evaluator,
    } for item in outcomes]), hide_index=True, width="stretch")


def render_data_sources(ctx: dict[str, list[object]]) -> None:
    st.subheader("数据来源")
    st.caption("展示成员外部身份与最近同步状态；原始数据仅在数据接入中心的高级信息中查看。")
    identities = ctx["identities"]
    jobs = ctx["ingestion_jobs"]
    if not identities:
        st.info("尚未配置外部身份映射。系统不会自动猜测成员。")
        return
    st.dataframe(pd.DataFrame([{
        "数据来源": get_provider_display(item.provider),
        "连接状态": _label(item.status),
        "最近同步": _fmt_dt(next((job.completed_at for job in jobs if job.source_system == item.provider), None)),
        "同步状态": _label(next((job.status for job in jobs if job.source_system == item.provider), None)),
    } for item in identities]), hide_index=True, width="stretch")


def render_simple_member_overview(patient: Patient, ctx: dict[str, list[object]]) -> None:
    """The member's five-second summary, independent of backend object names."""
    program = _active_program(ctx)
    tasks = ctx["tasks"]
    problems = sorted([item for item in ctx["problems"] if item.status != "CLOSED"], key=lambda item: item.priority_rank or 99)
    open_tasks = [task for task in tasks if task.status not in {"COMPLETED", "CANCELLED"}]
    primary, next_steps = st.columns([1.8, 1])
    with primary:
        with section_frame("当前重点", "优先展示正在影响本周管理的三项信息。"):
            if problems:
                rows = "".join(
                    f"<div class='focus-row'><div class='focus-index'>0{index}</div><div><div class='focus-title'>{html.escape(problem.title)}</div>"
                    f"<div class='focus-copy'>持续管理中 · {_severity(problem.severity)}</div></div>{status_badge(_severity(problem.severity))}</div>"
                    for index, problem in enumerate(problems[:3], start=1)
                )
                st.markdown(rows, unsafe_allow_html=True)
            else:
                _empty_state("暂无重点问题", "健康管理师确认后会在这里显示。")
    with next_steps:
        with section_frame("下一步", "仅保留近期需要人工处理的事项。"):
            if open_tasks:
                rows = "".join(
                    f"<div class='next-row'><div class='next-date'>{html.escape(_fmt_dt(task.due_at) if task.due_at else '待安排')}</div>"
                    f"<div class='focus-title'>{html.escape(task.title)}</div><div class='focus-copy'>{html.escape(task.instruction)}</div></div>"
                    for task in open_tasks[:3]
                )
                st.markdown(rows, unsafe_allow_html=True)
                primary_action("进入管理", key=f"member-next-{patient.id}", on_click=_open_member_management, args=(patient.id,))
            else:
                _empty_state("暂无下一步", "当前没有待办任务。")

    with section_frame("最近变化", "只保留近期最值得继续关注的 3 项健康变化。"):
        observations = ctx.get("observations", [])
        if observations:
            rows = "".join(
                f"<div class='timeline-preview'><div class='timeline-date'>{html.escape(_fmt_dt(item.observed_at))}</div>"
                f"<div><div class='timeline-title'>{html.escape(_metric_display_name(item.metric_code))}</div>"
                f"<div class='timeline-copy'>最近记录：{html.escape(_format_observation_value(item))}</div></div></div>"
                for item in observations[:3]
            )
            st.markdown(rows, unsafe_allow_html=True)
        else:
            _empty_state("暂无近期变化", "连接健康数据或确认体检资料后，这里会显示重要变化。")

    with section_frame("最近健康历程", "只预览最近的重要节点；完整信息在“历程”中按日期查看。"):
        with SessionLocal() as session:
            events = HealthTimelineService().get_timeline(session, patient.id, limit=4)
        if events:
            rows = "".join(
                f"<div class='timeline-preview'><div class='timeline-date'>{html.escape(_fmt_dt(event.occurred_at))}</div>"
                f"<div><div class='timeline-title'>{html.escape(event.title)}</div><div class='timeline-copy'>{html.escape(event.summary or '已记录重要健康事件')}</div></div></div>"
                for event in events
            )
            st.markdown(rows, unsafe_allow_html=True)
        elif program:
            st.caption(f"当前正在执行：{program.main_goal or display_program_type(program.program_type)}")
        else:
            _empty_state("暂无重要健康事件", "确认体检、健康数据或管理记录后会在这里显示。")
        secondary_action("查看完整健康历程", key=f"member-overview-timeline-{patient.id}", on_click=_open_member_timeline, args=(patient.id,), width="content")


def render_simple_health_problems(patient: Patient, ctx: dict[str, list[object]]) -> None:
    st.subheader("健康问题")
    problems = [problem for problem in ctx["problems"] if problem.status != "CLOSED"]
    if not problems:
        st.success("当前没有需要持续跟进的健康问题。")
        return
    for problem in problems:
        related = _problem_related(ctx, problem)
        latest_alert = related["alerts"][0] if related["alerts"] else None
        plan = related["plans"][0] if related["plans"] else None
        followup = related["followups"][0] if related["followups"] else None
        with st.container(border=True):
            st.markdown(f"### {problem.title}　{_severity_badge(problem.severity)}")
            st.write(f"发现：{latest_alert.finding if latest_alert else problem.description}")
            done = ["健康管理师已核实" if latest_alert and latest_alert.status not in {"NEW", "AI_SCREENED", "WAITING_MANAGER_REVIEW"} else None,
                    "医生已复核" if related["reviews"] else None]
            st.caption("已完成：" + " · ".join(item for item in done if item) if any(done) else "当前等待健康管理师处理")
            st.write(f"当前方案：{plan.content if plan else '待建立管理方案'}")
            next_step = followup.due_at if followup and followup.due_at else next((task.due_at for task in related["tasks"] if task.status != "COMPLETED"), None)
            st.write(f"下一步：{_fmt_dt(next_step) if next_step else '继续跟进'}")
            if latest_alert and latest_alert.status in {"NEW", "AI_SCREENED", "WAITING_MANAGER_REVIEW"}:
                st.button("处理", key=f"simple-problem-{problem.id}", on_click=_open_member, args=(patient.id,))


def render_simple_medical_records(patient: Patient, ctx: dict[str, list[object]]) -> None:
    _section_header("体检与检查")
    st.caption("上传、查看和人工确认体检报告；技术解析信息默认隐藏。")
    if st.button("上传体检报告", key=f"medical-report-open-{patient.id}", type="primary"):
        st.session_state[f"medical-report-visible-{patient.id}"] = True
    if st.session_state.get(f"medical-report-visible-{patient.id}"):
        render_report_upload(patient, key_prefix=f"medical-{patient.id}")

    _section_header("健康问题 / 病史")
    problems = [item for item in ctx["problems"] if item.status != "CLOSED"]
    if problems:
        st.write(" · ".join(item.title for item in problems[:3]))
    else:
        st.caption("暂无需要持续跟进的健康问题。")

    _section_header("手术与住院")
    with SessionLocal() as session:
        events = list(session.scalars(select(HealthEvent).where(HealthEvent.patient_id == patient.id, HealthEvent.event_type.in_(("surgery", "hospitalization"))).order_by(HealthEvent.start_at.desc()).limit(5)))
    if events:
        for event in events:
            st.caption(f"{_fmt_dt(event.start_at)} · {event.description}")
    else:
        st.caption("暂无手术或住院记录。")

    _section_header("用药")
    plans = ctx["med_plans"]
    if plans:
        st.dataframe(pd.DataFrame([{"药品名称": plan.drug_name, "剂量": f"{plan.dose} {plan.dose_unit}", "频次": plan.frequency, "使用方式": plan.route} for plan in plans[:10]]), hide_index=True, width="stretch")
    else:
        st.caption("暂无已记录的用药信息。")

    _section_header("医生意见")
    reviews = ctx["reviews"]
    if reviews:
        for review in reviews[:5]:
            with st.container(border=True):
                st.markdown(f"**{_fmt_dt(review.reviewed_at)} · {review.doctor_name}医生完成复核**")
                st.write(review.opinion)
                if review.question_for_doctor:
                    st.caption(f"复核问题：{review.question_for_doctor}")
    else:
        st.caption("当前没有需要查看的医生意见。")

    _section_header("转诊 / 外部医疗")
    with SessionLocal() as session:
        referrals = list(session.scalars(select(ExternalReferral).where(ExternalReferral.patient_id == patient.id).order_by(ExternalReferral.created_at.desc()).limit(5)))
    if referrals:
        for referral in referrals:
            st.caption(f"{referral.specialty or '专科待确认'} · {_label(referral.status)} · {referral.organization or '机构待确认'}")
    else:
        st.caption("暂无外部医疗协同记录。")


def render_members_workspace(members: list[Patient]) -> None:
    _page_header("成员", "查看成员当前状态、最近变化和下一步。", eyebrow="成员管理")
    query = st.text_input("搜索成员", placeholder="输入姓名或职位")
    visible = [member for member in members if not query or query.lower() in _member_display(member).lower()]
    if not visible:
        _empty_state("未找到匹配成员", "请调整搜索条件，或稍后再试。")
        return
    summaries = _member_list_summaries([member.id for member in visible])
    for start in range(0, len(visible), 2):
        columns = st.columns(2)
        for column, member in zip(columns, visible[start:start + 2]):
            summary = summaries.get(member.id, {})
            program = summary.get("program")
            display_risk = summary.get("risk", "正常")
            problems = summary.get("problems", [])
            next_task = summary.get("next_task")
            with column:
                stage = display_program_type(program.program_type) if program else "健康评估"
                member_card(
                    _member_display(member), f"{_age(member)} · 健康管理成员", "稳定" if display_risk == "正常" else display_risk,
                    f"{len(problems)} 项", str(problems[0]) if problems else stage,
                    next_task.title if next_task else "等待阶段复盘", key=f"member-card-{member.id}", on_click=_open_member, args=(member.id,),
                )


KNOWLEDGE_CATEGORIES = ["PATIENT_EDUCATION", "MEDICATION", "LAB_TEST", "CLINICAL_GUIDELINE", "TEXTBOOK_REFERENCE", "INTERNAL_SOP", "COMMUNICATION", "SERVICE_SOP", "AI_SAFETY", "PRIVACY"]
KNOWLEDGE_SOURCE_TYPES = ["GUIDELINE", "TEXTBOOK", "INTERNAL_SOP", "AUTHORIZED_UPLOAD", "MANUAL_ENTRY"]


@st.cache_data(ttl=30, show_spinner=False)
def _knowledge_stats() -> dict[str, int]:
    """Short cache for low-frequency library counts; live risk data is never cached here."""
    with SessionLocal() as session:
        rows = session.execute(
            select(KnowledgeDocument.category, func.count(KnowledgeDocument.id)).where(KnowledgeDocument.is_active.is_(True)).group_by(KnowledgeDocument.category)
        ).all()
    return {str(category): int(total) for category, total in rows}


def _save_knowledge_upload(uploaded_file) -> str:
    safe_name = Path(uploaded_file.name).name.replace(" ", "_")
    target_dir = Path(__file__).resolve().parent / "knowledge_uploads"
    target_dir.mkdir(parents=True, exist_ok=True)
    target = target_dir / f"{uuid4().hex}_{safe_name}"
    target.write_bytes(uploaded_file.getvalue())
    return str(target.relative_to(Path(__file__).resolve().parent))


def _knowledge_uploaded_text(uploaded_file) -> str | None:
    suffix = Path(uploaded_file.name).suffix.lower()
    payload = uploaded_file.getvalue()
    if suffix in {".txt", ".md"}:
        return payload.decode("utf-8", errors="replace")
    try:
        if suffix == ".docx":
            from io import BytesIO
            from docx import Document as DocxDocument
            return "\n".join(paragraph.text for paragraph in DocxDocument(BytesIO(payload)).paragraphs if paragraph.text.strip()) or None
        if suffix == ".pdf":
            from io import BytesIO
            from pypdf import PdfReader
            reader = PdfReader(BytesIO(payload))
            pages = []
            for page_number, page in enumerate(reader.pages[:500], start=1):
                page_text = (page.extract_text() or "").strip()
                if page_text:
                    # The chunker treats headings as source locations, so PDF
                    # excerpts remain traceable to a page without inventing a
                    # finer location than the parser actually captured.
                    pages.append(f"# 第 {page_number} 页\n{page_text}")
            return "\n\n".join(pages).strip() or None
    except Exception:
        LOGGER.exception("knowledge upload text extraction failed")
    return None


KNOWLEDGE_SOURCE_TYPE_LABELS = {
    "PATIENT_EDUCATION": "疾病与健康教育",
    "TERMINOLOGY": "术语与标准化",
    "REGULATORY_DRUG_INFORMATION": "监管资料",
    "CLASSIFICATION": "疾病分类",
    "GUIDELINE": "医疗指南",
    "TEXTBOOK": "授权教材",
    "INTERNAL_SOP": "内部SOP",
}
KNOWLEDGE_SOURCE_PROVIDER_LABELS = {
    "MEDLINEPLUS": "美国国家医学图书馆 / NIH",
    "RXNORM": "美国国家医学图书馆 / NIH",
    "OPENFDA": "美国食品药品监督管理局 / FDA",
    "WHO_ICD11": "世界卫生组织 / WHO",
}
KNOWLEDGE_SOURCE_CARD_CODES = ("MEDLINEPLUS", "RXNORM", "OPENFDA", "WHO_ICD11")
KNOWLEDGE_DOCUMENT_SOURCE_TYPES = {
    "GUIDELINE": "医疗指南",
    "TEXTBOOK": "教材与参考书",
    "INTERNAL_SOP": "内部 SOP",
    "AUTHORIZED_UPLOAD": "授权上传资料",
    "MANUAL_ENTRY": "人工整理资料",
}


def _knowledge_source_status(source: KnowledgeSourceRegistry) -> str:
    if source.enabled and source.review_status == "APPROVED_SOURCE":
        return "可查询"
    if source.source_code == "WHO_ICD11":
        return "需要 API 凭证"
    return "暂未启用"


def _knowledge_source_display(source_code: str | None, source_names: dict[str, str]) -> str:
    return source_names.get(source_code or "", "")


def _knowledge_source_link(url: str | None) -> str | None:
    """Only surface a human-readable official page, never an API/REST endpoint."""
    if not url or "/REST/" in url or "api.fda.gov" in url or url.endswith(".json"):
        return None
    return url


def _select_knowledge_source(source_code: str) -> None:
    """Prepare a source change before the next rerun creates its widgets."""
    st.session_state["knowledge_search_provider_pending"] = source_code
    if str(st.session_state.get("knowledge-search-query-ui") or "").strip():
        st.session_state["knowledge_search_autorun"] = True


def _show_who_icd11_setup() -> None:
    st.session_state["knowledge_search_who_setup"] = True


def _knowledge_result_key(result) -> str:
    return f"{result.provider_code}:{result.external_id}"


def _render_knowledge_search_detail(result) -> None:
    """Human-facing preview; provider payloads and REST paths never render here."""
    with detail_panel("资料预览", "预览在当前页面打开；保存后仍需人工审核。"):
        st.markdown(f"### {result.title}")
        st.caption(result.subtitle)
        st.markdown("**来源**")
        st.write(f"{result.source_name} · {result.source_organization}")
        st.caption(f"获取时间：{_fmt_dt(result.retrieved_at)}")
        if result.provider_code == "RXNORM":
            st.markdown("**标准化信息**")
            st.write("该资料用于统一药物名称与同义名称，不构成用药建议。")
            if result.structured_metadata.get("synonym"):
                st.write("同义名称/品牌映射：" + str(result.structured_metadata["synonym"]))
        elif result.provider_code == "OPENFDA":
            st.markdown("**标签信息**")
            for label, key in (("通用名", "generic_name"), ("品牌名", "brand_name"), ("生产企业", "manufacturer")):
                if result.structured_metadata.get(key):
                    st.write(f"{label}：{result.structured_metadata[key]}")
            st.caption("FDA/openFDA 监管资料，不构成个体化诊疗或用药建议。")
        if result.summary:
            st.markdown("**摘要**")
            st.write(result.summary)
        # Only a human-readable official resource may be opened outside HealthOps.
        if "/REST/" not in result.official_url and "api.fda.gov" not in result.official_url:
            st.link_button("查看官方来源", result.official_url, key=f"knowledge-result-source-{_knowledge_result_key(result)}")


def _render_knowledge_detail(document: KnowledgeDocument, source_names: dict[str, str], *, key_scope: str) -> None:
    source_name = _knowledge_source_display(document.source_provider, source_names) or document.source_name or "来源信息待补充"
    service = KnowledgeService()
    with SessionLocal() as session:
        chunks = list(session.scalars(select(KnowledgeChunk).where(KnowledgeChunk.knowledge_document_id == document.id).order_by(KnowledgeChunk.chunk_index)))
        audits = service.review_audits(session, document.id)
        usages = list(session.scalars(select(KnowledgeUseRecord).where(KnowledgeUseRecord.knowledge_document_id == document.id).order_by(KnowledgeUseRecord.created_at.desc()).limit(5)))
    st.markdown(f"### {document.title}")
    st.caption(f"{display_knowledge_category(document.category)} · 版本：{document.source_version or document.version} · 获取时间：{_fmt_dt(document.retrieved_at or document.updated_at)}")
    if document.summary:
        st.markdown("**摘要**")
        st.write(document.summary)
    st.markdown("**来源**")
    st.write(f"来源机构：{source_name}")
    official_link = _knowledge_source_link(document.source_url)
    if official_link:
        st.link_button("查看官方来源", official_link, key=f"knowledge-source-link-{key_scope}-{document.id}")
    elif document.source_url:
        st.caption("原始来源链接仅供高级核对，不在普通页面直接打开接口地址。")
    if document.attribution:
        st.caption("署名要求：" + document.attribution)
    if document.license_note:
        st.caption("许可/使用说明：" + document.license_note)
    publication = document.metadata_json.get("publication") if document.metadata_json else None
    if isinstance(publication, dict):
        publication_parts = [
            str(publication.get(key)).strip()
            for key in ("author", "publisher", "publication_year")
            if publication.get(key)
        ]
        if publication_parts:
            st.caption("出版信息：" + " · ".join(publication_parts))
    ai_eligible = service._eligible_for_formal_ai(document)
    st.markdown("**审核状态**")
    st.markdown(status_badge(display_knowledge_review_status(document.review_status)), unsafe_allow_html=True)
    st.write("AI 可引用：" + ("是" if ai_eligible else "否"))
    if document.review_due_at:
        st.caption(f"下次复核：{document.review_due_at.isoformat()}" + ("（已到期，AI 暂不使用）" if not ai_eligible else ""))
    if document.reviewed_by:
        st.caption(f"审核人：{document.reviewed_by} · 审核时间：{_fmt_dt(document.reviewed_at)}")
    if document.review_comment:
        st.caption("审核说明：" + document.review_comment)
    if chunks:
        st.caption(f"已建立 {len(chunks)} 个可追溯知识片段；仅批准后可被检索。")
    if usages:
        st.markdown("**引用记录**")
        for usage in usages:
            st.caption(f"{usage.feature or usage.output_type} · {_fmt_dt(usage.created_at)}")
    if document.review_status in {"DRAFT", "PENDING_REVIEW"}:
        comment = st.text_area("审核说明（可选）", key=f"knowledge-review-comment-{key_scope}-{document.id}", placeholder="例如：已核对来源、版本与许可说明")
        approve, reject = st.columns(2)
        if approve.button("批准并允许 AI 引用", type="primary", key=f"knowledge-approve-{key_scope}-{document.id}"):
            with SessionLocal() as session:
                current = session.get(KnowledgeDocument, document.id)
                if current:
                    KnowledgeService().approve_document(session, current, "授权审核人", comment)
                    session.commit()
            _knowledge_stats.clear()
            st.session_state["knowledge-notice"] = "资料已批准，可作为正式 AI 引用来源。"
            st.rerun()
        if reject.button("退回资料", key=f"knowledge-reject-{key_scope}-{document.id}"):
            with SessionLocal() as session:
                current = session.get(KnowledgeDocument, document.id)
                if current:
                    KnowledgeService().reject_document(session, current, "授权审核人", comment)
                    session.commit()
            _knowledge_stats.clear()
            st.session_state["knowledge-notice"] = "资料已退回，不会作为 AI 正式引用来源。"
            st.rerun()
    elif document.review_status in {"APPROVED", "REJECTED"}:
        with st.expander("版本与归档"):
            st.caption("归档后仍保留审计记录，但不会继续供 AI 正式引用。新版本须再次经过人工审核。")
            if secondary_action("归档此资料", key=f"knowledge-archive-{key_scope}-{document.id}"):
                with SessionLocal() as session:
                    current = session.get(KnowledgeDocument, document.id)
                    if current:
                        KnowledgeService().archive_document(session, current, "授权审核人", "人工归档")
                        session.commit()
                _knowledge_stats.clear()
                st.session_state["knowledge-notice"] = "资料已归档，AI 将不再引用此版本。"
                st.rerun()
            st.markdown("**创建替代版本**")
            with st.form(f"knowledge-replacement-{key_scope}-{document.id}"):
                replacement_version = st.text_input("新版本号", value=f"{document.version}.1", key=f"knowledge-replacement-version-{key_scope}-{document.id}")
                replacement_summary = st.text_area("新版本摘要", value=document.summary or "", key=f"knowledge-replacement-summary-{key_scope}-{document.id}")
                if st.form_submit_button("保存为待审核新版本"):
                    with SessionLocal() as session:
                        current = session.get(KnowledgeDocument, document.id)
                        if current:
                            KnowledgeService().create_document(
                                session, title=current.title, category=current.category,
                                summary=replacement_summary, content_text=current.content_text,
                                source_type=current.source_type, source_name=current.source_name,
                                source_reference=current.source_reference, source_provider=current.source_provider,
                                source_external_id=current.source_external_id, source_url=current.source_url,
                                source_version=current.source_version, retrieved_at=current.retrieved_at,
                                license_note=current.license_note, attribution=current.attribution,
                                version=replacement_version, tags=current.tags,
                                review_status="PENDING_REVIEW", supersedes_id=current.id,
                                metadata_json={"replacement_of": str(current.id)},
                            )
                            session.commit()
                    _knowledge_stats.clear()
                    st.session_state["knowledge-notice"] = "新版本已保存为待审核资料；批准后将归档旧版本。"
                    st.rerun()
    if document.file_reference:
        st.markdown("**原始文件**")
        file_path = Path(__file__).resolve().parent / document.file_reference
        if file_path.is_file():
            st.download_button("下载原始文件", file_path.read_bytes(), file_name=file_path.name, key=f"knowledge-download-{key_scope}-{document.id}")
        else:
            st.caption("原始文件目前不可用。")
    if TECHNICAL_DETAILS_ENABLED:
        with st.expander("高级信息"):
            st.caption("内部标识和原始元数据仅供授权人员核对。")
            if audits:
                st.dataframe(pd.DataFrame([
                    {
                        "原状态": _label(item.previous_status),
                        "新状态": _label(item.new_status),
                        "审核人": item.reviewer or "未记录",
                        "时间": _fmt_dt(item.created_at),
                    }
                    for item in audits
                ]), hide_index=True, width="stretch")
            with st.expander("查看原始技术信息"):
                st.json(document.metadata_json)


def _render_knowledge_search(sources: list[KnowledgeSourceRegistry]) -> None:
    """Run an explicit, same-page search with stable session business state."""
    service = KnowledgeService()
    source_names = {source.source_code: source.display_name for source in sources}
    source_options = ["ALL", *[source.source_code for source in sources if source.source_code in KNOWLEDGE_SOURCE_CARD_CODES]]
    pending_provider = st.session_state.pop("knowledge_search_provider_pending", None)
    if pending_provider in source_options:
        st.session_state["knowledge-search-provider-ui"] = pending_provider
    with section_frame("搜索知识", "按关键词查询官方来源。在线结果不是平台知识资产；只有保存并经人工批准的资料才可供 AI 正式引用。"):
        left, right = st.columns([2, 1])
        query = left.text_input("关键词", placeholder="例如：hypertension / metformin", key="knowledge-search-query-ui")
        selected = right.selectbox(
            "来源", source_options, key="knowledge-search-provider-ui",
            format_func=lambda value: "全部官方来源" if value == "ALL" else source_names.get(value, value),
        )
        search_requested = (
            primary_action("搜索", key="knowledge-source-search", disabled=not query.strip())
            or st.session_state.pop("knowledge_search_retry", False)
            or st.session_state.pop("knowledge_search_autorun", False)
        )
        if search_requested:
            requested = [source for source in sources if source.enabled and source.review_status == "APPROVED_SOURCE"] if selected == "ALL" else [source for source in sources if source.source_code == selected]
            results: list = []
            errors: list[str] = []
            if selected == "WHO_ICD11" and requested and not requested[0].enabled:
                errors.append("当前尚未配置 WHO ICD-11 API 凭证。")
            for source in requested:
                if not source.enabled:
                    if source.source_code != "WHO_ICD11":
                        errors.append(f"{source.display_name}：{_knowledge_source_status(source)}")
                    continue
                try:
                    with SessionLocal() as session:
                        results.extend(service.query_source(session, source.source_code, query.strip(), limit=5))
                except KnowledgeProviderError as error:
                    errors.append(f"{source.display_name}：{error}")
                except Exception:
                    LOGGER.exception("official knowledge source query failed", extra={"source": source.source_code})
                    errors.append(f"{source.display_name}：暂时无法连接，请稍后重试。")
            st.session_state["knowledge_search_results"] = results
            st.session_state["knowledge_search_errors"] = errors
            st.session_state["knowledge_search_last_query"] = query.strip()
            st.session_state.pop("knowledge_search_selected_result", None)
        for error in st.session_state.get("knowledge_search_errors", []):
            st.warning(error)
        if st.session_state.get("knowledge_search_errors") and secondary_action("重试", key="knowledge-search-retry"):
            st.session_state["knowledge_search_retry"] = True
            st.rerun()
        results = st.session_state.get("knowledge_search_results", [])
        last_query = st.session_state.get("knowledge_search_last_query")
        if last_query and not results and not st.session_state.get("knowledge_search_errors"):
            _empty_state(f"没有找到与“{last_query}”相关的资料", "可以尝试更换关键词，或切换知识来源后重新搜索。")
        if results:
            _section_header("搜索结果", f"找到 {len(results)} 项。选择一项查看平台内预览；保存后需要人工审核。")
            selected_result_key = st.session_state.get("knowledge_search_selected_result")
            if selected_result_key not in {_knowledge_result_key(result) for result in results}:
                selected_result_key = _knowledge_result_key(results[0])
            list_column, detail_column = st.columns([1.15, 1], gap="large")
            with list_column:
                for index, result in enumerate(results):
                    result_key = _knowledge_result_key(result)
                    with st.container(border=True):
                        st.markdown(f"**{result.title}**")
                        st.caption(f"{result.source_name} · {result.source_organization}")
                        st.caption(result.subtitle)
                        if result.provider_code == "RXNORM":
                            st.caption("药物标准资料 · 用于统一药物名称与同义名称")
                        elif result.summary:
                            st.caption(result.summary[:220])
                        view, save = st.columns(2)
                        if view.button("查看", key=f"knowledge-search-view-{index}"):
                            st.session_state["knowledge_search_selected_result"] = result_key
                            st.rerun()
                        with SessionLocal() as session:
                            existing = service.cached_source_result(session, result)
                        if existing:
                            save.button("已保存", key=f"knowledge-source-saved-{index}", disabled=True)
                            if st.button("查看已保存资料", key=f"knowledge-source-existing-{index}"):
                                st.session_state["knowledge-selected-id"] = str(existing.id)
                                st.rerun()
                        elif save.button("保存到知识库", key=f"knowledge-source-cache-{index}"):
                            with SessionLocal() as session:
                                cached = service.cache_source_result(session, result)
                                session.commit()
                            _knowledge_stats.clear()
                            st.session_state["knowledge-selected-id"] = str(cached.id)
                            st.session_state["knowledge-notice"] = f"已保存“{cached.title}”为待审核资料。"
                            st.rerun()
            with detail_column:
                selected = next(result for result in results if _knowledge_result_key(result) == selected_result_key)
                _render_knowledge_search_detail(selected)


def _render_knowledge_source_cards(sources: list[KnowledgeSourceRegistry]) -> None:
    source_by_code = {source.source_code: source for source in sources}
    with section_frame("知识来源", "在线来源回答“知识从哪里来”；保存并审核后的资料才进入正式知识库。它们不会自动生成临床风险规则。"):
        cards = [source_by_code[code] for code in KNOWLEDGE_SOURCE_CARD_CODES if code in source_by_code]
        for source in cards:
            name, organization, kind, action = st.columns([1.35, 1.35, 1.3, 0.7])
            name.markdown(f"**{source.display_name}**")
            organization.caption(source.organization or KNOWLEDGE_SOURCE_PROVIDER_LABELS.get(source.source_code, source.provider))
            kind.caption(f"{KNOWLEDGE_SOURCE_TYPE_LABELS.get(source.source_type, '医学资料')} · {_knowledge_source_status(source)}")
            if source.enabled:
                action.button("搜索", key=f"knowledge-source-card-{source.source_code}", on_click=_select_knowledge_source, args=(source.source_code,))
            elif source.source_code == "WHO_ICD11":
                action.button("配置说明", key="knowledge-source-who-setup", on_click=_show_who_icd11_setup)
            else:
                action.caption("暂未启用")
    if st.session_state.pop("knowledge_search_who_setup", False):
        st.info("WHO ICD-11 需要在安全配置中提供官方 API 凭证；未配置前不会发起查询，也不会把分类结果当作诊断。")


def _render_knowledge_add_form() -> None:
    if secondary_action("添加内部资料或授权文件", key="knowledge-add"):
        st.session_state["knowledge-add-open"] = True
    if st.session_state.get("knowledge-add-open"):
        with st.expander("添加资料", expanded=True):
            entry_mode = st.radio("添加方式", ["手工录入", "上传文件"], horizontal=True, key="knowledge-entry-mode")
            if entry_mode == "手工录入":
                with st.form("knowledge-manual-form"):
                    title = st.text_input("标题 *")
                    category = st.selectbox("分类 *", KNOWLEDGE_CATEGORIES, format_func=display_knowledge_category)
                    summary = st.text_area("摘要")
                    content = st.text_area("正文")
                    source_type = st.selectbox("资料来源类型", KNOWLEDGE_SOURCE_TYPES, format_func=lambda value: KNOWLEDGE_DOCUMENT_SOURCE_TYPES.get(value, "其他资料"))
                    source_name = st.text_input("来源名称", value="内部资料")
                    version = st.text_input("版本", value="v1.0")
                    review_due_at = st.date_input("下次复核日期", value=date.today() + timedelta(days=365), key="knowledge-manual-review-due")
                    tags = st.text_input("标签", placeholder="用逗号分隔")
                    st.caption("审核状态：草稿。只有完成审核的资料才能在未来作为 AI 的正式知识来源。")
                    if st.form_submit_button("保存草稿", type="primary"):
                        if not title.strip() or not source_name.strip():
                            st.error("请填写标题和来源名称。")
                        else:
                            with SessionLocal() as session:
                                KnowledgeService().create_document(
                                    session, title=title, category=category, summary=summary, content_text=content,
                                    source_type=source_type, source_name=source_name, version=version,
                                    tags=tags.split(","), review_status="DRAFT", review_due_at=review_due_at, metadata_json={},
                                )
                                session.commit()
                            _knowledge_stats.clear()
                            st.session_state["knowledge-add-open"] = False
                            st.session_state["knowledge-notice"] = "资料已保存为草稿。"
                            st.rerun()
            else:
                with st.form("knowledge-upload-form"):
                    uploaded = st.file_uploader("选择资料文件", type=["pdf", "docx", "txt", "md"])
                    category = st.selectbox("分类 *", KNOWLEDGE_CATEGORIES, format_func=display_knowledge_category, key="knowledge-file-category")
                    title = st.text_input("资料标题（可留空，默认使用文件名）")
                    source_type = st.selectbox("资料来源类型", KNOWLEDGE_SOURCE_TYPES, key="knowledge-file-source-type", format_func=lambda value: KNOWLEDGE_DOCUMENT_SOURCE_TYPES.get(value, "其他资料"))
                    source_name = st.text_input("来源名称", value="上传资料", key="knowledge-file-source-name")
                    version = st.text_input("版本", value="v1.0", key="knowledge-file-version")
                    author = st.text_input("作者（如适用）", key="knowledge-file-author")
                    publisher = st.text_input("出版社（如适用）", key="knowledge-file-publisher")
                    publication_year = st.text_input("出版年份（如适用）", key="knowledge-file-publication-year")
                    upload_origin = st.text_input("上传来源", value="授权资料上传", key="knowledge-file-upload-origin")
                    review_due_at = st.date_input("下次复核日期", value=date.today() + timedelta(days=365), key="knowledge-file-review-due")
                    tags = st.text_input("标签", key="knowledge-file-tags", placeholder="用逗号分隔")
                    legal_note = st.text_area("版权/授权说明 *", key="knowledge-file-license", placeholder="确认您拥有上传、保存及平台内部使用该资料的合法授权。")
                    st.caption("教材与参考书只能上传您拥有合法使用权的文件。文件默认待审核，不会自动作为 AI 来源或医疗规则。")
                    if st.form_submit_button("登记上传资料", type="primary"):
                        if uploaded is None or not source_name.strip() or not legal_note.strip():
                            st.error("请选择文件，填写来源名称和版权/授权说明。")
                        else:
                            reference = _save_knowledge_upload(uploaded)
                            content = _knowledge_uploaded_text(uploaded)
                            with SessionLocal() as session:
                                KnowledgeService().create_document(
                                    session, title=title.strip() or Path(uploaded.name).stem, category=category,
                                    summary=f"上传资料：{uploaded.name}", content_text=content, source_type=source_type,
                                    source_name=source_name, file_reference=reference, version=version,
                                    tags=tags.split(","), review_status="PENDING_REVIEW",
                                    processing_status="TEXT_EXTRACTED" if content else "WAITING_REVIEW",
                                    license_note=legal_note, review_due_at=review_due_at,
                                    metadata_json={
                                        "original_filename": uploaded.name,
                                        "file_kind": Path(uploaded.name).suffix.lower(),
                                        "authorized_upload": True,
                                        "upload_origin": upload_origin,
                                        "publication": {"author": author, "publisher": publisher, "publication_year": publication_year},
                                    },
                                )
                                session.commit()
                            _knowledge_stats.clear()
                            st.session_state["knowledge-add-open"] = False
                            st.session_state["knowledge-notice"] = "资料已登记，正在等待人工审核。"
                            st.rerun()


def _render_saved_knowledge(sources: list[KnowledgeSourceRegistry]) -> None:
    source_names = {source.source_code: source.display_name for source in sources}
    with section_frame("已保存知识", "这是平台本地知识资产区。草稿、待审核和已批准资料均可在这里查看；只有已批准资料可供 AI 正式引用。"):
        filters = st.columns(3)
        search = filters[0].text_input("搜索已保存资料", placeholder="标题、摘要或标签", key="knowledge-saved-search")
        category_filter = filters[1].selectbox("分类", ["全部分类", *KNOWLEDGE_CATEGORIES], format_func=lambda value: "全部分类" if value == "全部分类" else display_knowledge_category(value), key="knowledge-category-filter")
        review_filter = filters[2].selectbox("审核状态", ["全部状态", "DRAFT", "PENDING_REVIEW", "APPROVED", "REJECTED", "ARCHIVED"], format_func=lambda value: "全部状态" if value == "全部状态" else display_knowledge_review_status(value), key="knowledge-review-filter")
        _render_knowledge_add_form()
        with SessionLocal() as session:
            documents = KnowledgeService().search_documents(
                session, search, category=None if category_filter == "全部分类" else category_filter,
                review_status=None if review_filter == "全部状态" else review_filter,
            )
        if not documents:
            _empty_state("还没有保存的知识资料", "请先搜索官方来源并保存需要审核的资料；在线来源仍可直接查询。")
            return
        selected_id = st.session_state.get("knowledge-selected-id")
        if selected_id and not any(str(document.id) == selected_id for document in documents):
            selected_id = None
        left, right = st.columns([1, 2])
        with left:
            st.markdown("**资料列表**")
            for document in documents[:30]:
                source_label = _knowledge_source_display(document.source_provider, source_names) or document.source_name or "来源待补充"
                if st.button(document.title, key=f"knowledge-document-{document.id}", width="stretch"):
                    st.session_state["knowledge-selected-id"] = str(document.id)
                    st.rerun()
                ai_status = "AI 可引用" if KnowledgeService._eligible_for_formal_ai(document) else "AI 暂不使用"
                st.caption(f"{display_knowledge_category(document.category)} · {source_label}")
                st.caption(f"{display_knowledge_review_status(document.review_status)} · {ai_status}")
        with right:
            selected = next((document for document in documents if str(document.id) == selected_id), documents[0])
            with detail_panel("资料详情", "详情在当前页面打开，不新增导航层级。"):
                _render_knowledge_detail(selected, source_names, key_scope="saved")


def _render_pending_knowledge(sources: list[KnowledgeSourceRegistry]) -> None:
    """A visible governance queue; it is not a separate review-center route."""
    source_names = {source.source_code: source.display_name for source in sources}
    with SessionLocal() as session:
        pending = list(session.scalars(select(KnowledgeDocument).where(
            KnowledgeDocument.is_active.is_(True),
            KnowledgeDocument.review_status.in_(["DRAFT", "PENDING_REVIEW"]),
        ).order_by(KnowledgeDocument.updated_at.desc())))
    with section_frame("待审核", "审核来源、版本与许可说明后，资料才可供 AI 正式引用。当前为 UAT 治理流程，不代表生产环境权限体系。"):
        if not pending:
            _empty_state("当前没有待审核资料", "在线搜索结果保存后，或上传内部 SOP、授权教材后，会在这里等待人工审核。")
            return
        selected_id = st.session_state.get("knowledge-review-selected-id")
        if selected_id not in {str(document.id) for document in pending}:
            selected_id = str(pending[0].id)
        list_column, detail_column = st.columns([1, 1.6], gap="large")
        with list_column:
            for document in pending[:20]:
                source_label = _knowledge_source_display(document.source_provider, source_names) or document.source_name or "来源待补充"
                if st.button(document.title, key=f"knowledge-review-document-{document.id}", width="stretch"):
                    st.session_state["knowledge-review-selected-id"] = str(document.id)
                    st.rerun()
                st.caption(f"{source_label} · {_fmt_dt(document.updated_at)}")
        with detail_column:
            selected = next(document for document in pending if str(document.id) == selected_id)
            with detail_panel("审核资料", "核对内容、来源、版本和许可说明；批准后才允许 AI 引用。"):
                _render_knowledge_detail(selected, source_names, key_scope="pending")


def render_knowledge_library_entry() -> None:
    _page_header("医学知识中心", "统一管理医学资料来源、审核状态和 AI 引用依据。公开知识用于解释与引用；可执行医疗规则由独立治理流程发布。", eyebrow="平台工具")
    service = KnowledgeService()
    with SessionLocal() as session:
        sources = service.list_sources(session)
        service.ensure_approved_chunks(session)
        session.commit()

    notice = st.session_state.pop("knowledge-notice", None)
    if notice:
        st.success(notice)
    with SessionLocal() as session:
        review_counts = dict(session.execute(
            select(KnowledgeDocument.review_status, func.count(KnowledgeDocument.id)).group_by(KnowledgeDocument.review_status)
        ).all())
    stats = st.columns(4)
    with stats[0]:
        summary_metric("在线来源", len([source for source in sources if source.source_code in KNOWLEDGE_SOURCE_CARD_CODES]))
    with stats[1]:
        summary_metric("已保存资料", sum(int(value) for value in review_counts.values()))
    with stats[2]:
        summary_metric("待审核", int(review_counts.get("PENDING_REVIEW", 0) + review_counts.get("DRAFT", 0)))
    with stats[3]:
        summary_metric("已批准", int(review_counts.get("APPROVED", 0)))
    _render_knowledge_search(sources)
    _render_knowledge_source_cards(sources)
    _render_saved_knowledge(sources)
    _render_pending_knowledge(sources)


def render_more_workspace() -> None:
    """Compatibility wrapper for the extracted tools-shell page."""
    render_more_workspace_shell(
        page_header=_page_header,
        load_members=_members,
        render_data_gateway=render_data_gateway,
        render_knowledge_library=render_knowledge_library_entry,
        render_risk_rules=render_risk_rules,
        render_audit=render_audit,
        audit_context=_audit_context,
        member_display=_member_display,
    )


def render_oversight_summary() -> None:
    st.title("风险监管摘要")
    st.caption("演示环境 · 仅显示服务汇总，不展示个人疾病、用药、报告或具体指标。")
    with SessionLocal() as session:
        summary = OversightRiskSummaryService().summarize(session)
    _status_strip(("成员覆盖", summary["member_coverage"], "action"), ("紧急", summary["red"], "urgent"), ("需要关注", summary["yellow"], "attention"), ("未处理", summary["unhandled"], "attention"))
    cards = st.columns(3)
    cards[0].metric("医生待复核", summary["doctor_pending"])
    cards[1].metric("风险关闭率", "暂无足够数据")
    cards[2].metric("个人临床信息", "未展示")


def render_collaboration_workspace() -> None:
    """Keep internal medical review distinct from external-care coordination."""
    _page_header("医疗协同", "只处理需要医学判断或外部医疗安排的事项。", eyebrow="运营后台")
    collaboration = st.radio("医疗协同内容", ["内部医生", "外部医疗"], horizontal=True, label_visibility="collapsed", key="collaboration-view")
    members = _members()
    if collaboration == "内部医生":
        render_global_doctor_workspace(members)
    else:
        render_external_doctor_workspace(members)


def render_service_operations_workspace() -> None:
    """A dedicated queue with its service detail in the same page inspector."""
    _page_header("服务运营", "审核服务申请、安排执行并跟进服务结果。", eyebrow="服务工作台")
    members = _patient_map()
    with SessionLocal() as session:
        requests = list(session.scalars(
            select(ServiceRequest).order_by(ServiceRequest.requested_at.desc()).limit(100)
        ))
        service_names = {item.id: item.name for item in session.scalars(select(ServiceCatalogItem))}
    _status_strip(
        ("待审核", sum(item.status in {"REQUESTED", "REVIEWING"} for item in requests), "attention"),
        ("待安排", sum(item.status == "APPROVED" for item in requests), "action"),
        ("进行中", sum(item.status in {"SCHEDULED", "IN_PROGRESS"} for item in requests), "action"),
        ("等待反馈", sum(item.status == "COMPLETED" and not item.result_summary for item in requests), "neutral"),
    )
    filters = {"全部": set(), "待审核": {"REQUESTED", "REVIEWING"}, "待安排": {"APPROVED"}, "进行中": {"SCHEDULED", "IN_PROGRESS"}, "等待反馈": {"COMPLETED"}, "已完成": {"COMPLETED"}}
    selected_filter = st.radio("服务状态筛选", list(filters), horizontal=True, label_visibility="collapsed", key="service-operations-filter")
    visible = requests if not filters[selected_filter] else [item for item in requests if item.status in filters[selected_filter]]
    if not visible:
        _empty_state("暂无待处理服务", "新的服务申请会按申请时间显示在这里。")
        return
    selected_key = "service-operations-selected"
    if st.session_state.get(selected_key) not in {str(item.id) for item in visible}:
        st.session_state[selected_key] = str(visible[0].id)
    left, right = st.columns([1.15, 1], gap="large")
    with left:
        _section_header("服务工作列表", "选择一项后，在右侧直接审核、安排或记录结果。")
        for request in visible[:20]:
            member = members.get(request.patient_id)
            if member and secondary_action(f"{_member_display(member)} · {service_names.get(request.service_item_id, '会员服务')} · {_label(request.status)}", key=f"service-operations-select-{request.id}"):
                st.session_state[selected_key] = str(request.id); st.rerun()
    selected = next(item for item in visible if str(item.id) == st.session_state[selected_key])
    with right:
        member = members.get(selected.patient_id)
        with detail_panel("服务详情", "服务申请不等于自动医疗预约；所有安排均由人工确认。"):
            st.markdown(f"**{service_names.get(selected.service_item_id, '会员服务')} · {_member_display(member)}**")
            st.write(selected.reason or "成员提交服务申请。")
            st.caption(f"当前状态：{_label(selected.status)} · 申请时间：{_fmt_dt(selected.requested_at)}")
            if selected.status in {"REQUESTED", "REVIEWING"}:
                if primary_action("审核并安排", key=f"service-operations-approve-{selected.id}", width="content"):
                    with SessionLocal() as session:
                        MemberServiceOperations().approve(session, selected.id, "健康管理师"); session.commit()
                    st.rerun()
            elif selected.status in {"APPROVED", "SCHEDULED", "IN_PROGRESS"}:
                if primary_action("记录服务完成", key=f"service-operations-complete-{selected.id}", width="content"):
                    with SessionLocal() as session:
                        MemberServiceOperations().complete(session, selected.id, "服务已完成，后续由健康管理团队跟进。", "健康管理师"); session.commit()
                    st.rerun()
            else:
                st.write(selected.result_summary or "服务已完成，等待补充结果。")


def _report_candidate_label(candidate: ReportExtractionCandidate) -> str:
    if candidate.candidate_type == "OBSERVATION":
        return _metric_display_name(candidate.canonical_code, candidate.raw_name)
    label = (candidate.summary or candidate.raw_name or "报告资料").strip()
    if label.lower().startswith(("synthetic_", "demo_", "test_")):
        return "演示资料 · 体检结果"
    return label


def _report_candidate_method(candidate: ReportExtractionCandidate) -> str:
    return {"LLM": "本地AI辅助解析", "RULE": "规则解析", "TABLE": "规则解析", "ADAPTER": "规则解析", "OCR": "文字识别", "MANUAL": "人工整理"}.get(candidate.extraction_method, "规则解析")


def _report_model_name(model: str | None) -> str:
    return "local LLM" if model == "local LLM" else (model or "本地开源大模型")


def _report_parse_mode(run: ReportExtractionRun, candidates: list[ReportExtractionCandidate]) -> str:
    has_local_ai = any(item.extraction_method == "LLM" for item in candidates)
    has_rule = any(item.extraction_method != "LLM" for item in candidates)
    if has_local_ai and has_rule:
        return "混合解析"
    if has_local_ai:
        return "本地AI辅助解析"
    return "规则解析"


def _report_risk_next_step(level: str) -> str:
    """Explain governed risk routing without introducing clinical thresholds."""
    normalized = level.upper()
    if normalized in {"HIGH", "RED"}:
        return "优先由医生与健康管理团队人工处理。"
    if normalized in {"MEDIUM", "YELLOW"}:
        return "由健康管理师先核实并安排后续跟进，必要时提交医生复核。"
    if normalized in {"LOW", "GREEN"}:
        return "以健康提醒和日常自我管理为主；出现新变化时再由团队复核。"
    return "当前没有正式风险结论；健康管理团队会结合已确认资料继续判断。"


def _render_report_parse_method(run: ReportExtractionRun, candidates: list[ReportExtractionCandidate]) -> None:
    """Show persisted parse provenance only; never invoke a model during reruns."""
    st.markdown("#### 解析信息")
    st.markdown(f"**解析方式：{_report_parse_mode(run, candidates)}**")
    st.write("规则解析：已完成")
    if run.llm_status == "USED":
        st.success("本地AI辅助：已使用")
        details = ["AI 辅助整理已完成"]
        if run.llm_processed_sections:
            details.append(f"已整理：{' / '.join(run.llm_processed_sections)}")
        st.caption(" · ".join(details) + "。所有AI辅助结果仍需人工确认后才会入档。")
    elif run.llm_status == "NOT_NEEDED":
        st.info("本地AI辅助：本次未调用")
        st.caption("原因：本报告相关内容已由规则可靠解析。")
    elif not run.llm_enabled:
        st.info("本地AI辅助：未启用")
        st.caption("原因：本地语义模型未启用；规则解析与人工确认仍可正常使用。")
    else:
        st.warning("本地AI辅助：当前不可用")
        st.caption(f"原因：{run.llm_failure_reason or '本地开源大模型 当前不可用'}。规则解析与人工确认仍可正常使用。")
    st.caption(f"解析时间：{_fmt_dt(run.completed_at or run.created_at)}")
    if TECHNICAL_DETAILS_ENABLED:
        with st.expander("高级信息"):
            st.write(f"解析器版本：{run.parser_version}")
            st.write(f"标准指标库版本：{run.canonical_registry_version}")
            st.write(f"规则识别资料：{run.rule_candidate_count} 项")
            if run.llm_used:
                st.write(f"本地AI辅助资料：{run.llm_candidate_count} 项")
            if run.llm_model:
                st.write(f"本地模型：{run.llm_model}")
            st.caption("重新解析会创建新的解析记录，不会覆盖已确认入档的健康数据。")


def _run_report_parse_with_progress(parse_action):
    """Render transient progress while the service performs one explicit parse."""
    started = perf_counter()
    completed_lines: list[str] = []
    with st.status("正在解析体检报告…", expanded=True) as status:
        current = status.empty()
        detail = status.empty()
        history = status.empty()
        progress = status.progress(0, text="正在读取报告")
        elapsed = status.empty()

        def update(event: ReportParseProgress) -> None:
            elapsed_seconds = (event.elapsed_ms or round((perf_counter() - started) * 1000)) / 1000
            elapsed.caption(f"已用时：{elapsed_seconds:.1f} 秒")
            if event.stage == "RULE_PARSE_STARTED":
                current.markdown("**正在进行规则解析…**")
                detail.caption("正在识别结构化健康指标。")
            elif event.stage == "RULE_PARSE_COMPLETED":
                current.markdown("**✓ 规则解析完成**")
                detail.caption(f"识别结构化指标：{event.rule_candidate_count or 0} 项")
                completed_lines.append(f"✓ 规则解析完成 · {event.rule_candidate_count or 0} 项结构化指标")
            elif event.stage == "LLM_STARTED":
                current.markdown("**本地AI辅助解析中**")
                detail.caption(f"模型：local LLM · {event.message}")
            elif event.stage == "LLM_SECTION_STARTED":
                current.markdown("**本地AI辅助解析中**")
                completed = max((event.current or 1) - 1, 0)
                remaining = max((event.total or 0) - completed - 1, 0)
                detail.caption(
                    f"当前：{event.section_name or '复杂检查内容'} · 第 {event.current or 0} / {event.total or 0} 次"
                    f" · 已完成：{completed} · 待处理：{remaining}"
                )
                if event.total:
                    progress.progress((event.current or 0) / event.total, text=f"本地AI调用进度：{event.current or 0} / {event.total}")
            elif event.stage == "LLM_SECTION_COMPLETED":
                duration = f" · {event.call_duration_ms / 1000:.1f} 秒" if event.call_duration_ms is not None else ""
                completed_lines.append(f"✓ {event.section_name or '检查内容'} AI辅助解析完成{duration}")
            elif event.stage == "LLM_SECTION_FAILED":
                completed_lines.append(f"⚠ {event.section_name or '检查内容'} 本地AI辅助解析失败，已保留规则结果")
            elif event.stage == "LLM_UNAVAILABLE":
                current.markdown("**⚠ 本地AI当前不可用**")
                detail.caption("复杂检查内容将保留供人工确认，规则解析继续完成。")
            elif event.stage == "LLM_NOT_NEEDED":
                current.markdown("**规则解析已完成**")
                detail.caption("本报告没有需要本地AI辅助整理的复杂检查内容。")
            elif event.stage == "MERGING":
                current.markdown("**正在合并解析结果…**")
            elif event.stage == "EVIDENCE_VALIDATION":
                current.markdown("**正在校验本地AI结果证据…**")
            elif event.stage == "DEDUPLICATION":
                current.markdown("**正在去重解析结果…**")
            elif event.stage == "SAVING":
                current.markdown("**正在保存解析结果…**")
            elif event.stage == "COMPLETED":
                current.markdown(f"**✓ {event.message}**")
                completion = [
                    "解析方式：混合解析" if (event.llm_call_count or 0) else "解析方式：规则解析",
                    f"规则解析：{event.rule_candidate_count or 0} 项",
                ]
                if event.llm_call_count:
                    completion.append(f"本地AI：local LLM · {event.llm_success_count or 0} / {event.llm_call_count} 次成功")
                completion.extend((
                    f"检查结论：{event.finding_count or 0} 项",
                    f"随访建议：{event.followup_count or 0} 项",
                ))
                detail.caption(" · ".join(completion))
                progress.progress(1.0, text="解析完成")
            else:
                current.markdown(f"**{event.message}**")
            if completed_lines:
                history.markdown("  \n".join(completed_lines[-5:]))

        try:
            result = parse_action(update)
        except Exception:
            status.update(label="体检报告解析未完成", state="error", expanded=True)
            raise
        run = result[1] if isinstance(result, tuple) else result
        status.update(label="体检报告部分完成" if getattr(run, "status", None) == "PARTIAL_SUCCESS" else "体检报告解析完成", state="complete", expanded=True)
        return result


def render_report_review(document_id: UUID) -> None:
    """A concise, human-review workspace for evidence-backed report candidates."""
    with SessionLocal() as session:
        document = session.get(Document, document_id)
        service = ReportParsingService()
        runs = service.runs(session, document_id)
        run = runs[0] if runs else None
        candidates = service.candidates(session, document_id, run.id) if run else []
    if document is None or run is None:
        st.error("未找到体检报告解析记录。")
        return
    _page_header(_source_display_name(document, "体检报告"), "先看本次结论、与上次相比和需要处理事项；解析技术信息默认折叠。", eyebrow="体检与检查")
    if run.status == "NEEDS_OCR":
        st.warning("该报告为扫描件，需要文字识别后继续解析。原文件已保留，系统不会伪造解析结果。")
        return
    st.caption(f"{run.detected_hospital or '体检机构待补充'} · {run.detected_report_date or '体检日期未识别'} · 报告整理完成")
    with SessionLocal() as session:
        report_risk = ReportRiskSummaryService().summarize(session, document.patient_id, document.id)
    with section_frame("本次核心结论", "这是对当前已确认报告资料的整理，不替代医生判断。"):
        st.markdown(risk_badge(report_risk["level"]), unsafe_allow_html=True)
        st.markdown(f"**{report_risk['reason']}**")
        st.caption("处理方式：" + _report_risk_next_step(str(report_risk["level"])))
    with SessionLocal() as session:
        prior_document_ids = list(session.scalars(
            select(ReportExtractionRun.document_id)
            .where(
                ReportExtractionRun.patient_id == document.patient_id,
                ReportExtractionRun.status.in_(("COMPLETED", "PARTIAL_SUCCESS")),
                ReportExtractionRun.document_id != document.id,
            )
            .distinct()
        ))
        previous = session.scalar(
            select(Document)
            .where(Document.patient_id == document.patient_id, Document.id.in_(prior_document_ids))
            .order_by(Document.created_at.desc())
        ) if prior_document_ids else None
        comparison = None
        if previous is not None:
            try:
                comparison = ReportComparisonService().compare(session, document.patient_id, previous.id, document.id)
            except ValueError:
                comparison = None
    with section_frame("与上次相比", "只比较两份已人工确认的健康资料。"):
        if comparison is None:
            _empty_state("这是当前成员第一份体检报告", "已作为后续长期健康比较的基线。")
        else:
            _status_strip(
                ("新增", len(comparison["new_findings"]), "attention"),
                ("持续", len(comparison["persistent_findings"]), "attention"),
                ("未再出现", len(comparison["resolved_findings"]), "action"),
                ("需要处理", len(comparison["new_findings"]) + len(comparison["persistent_findings"]), "attention"),
            )
            if comparison["new_findings"] or comparison["persistent_findings"]:
                st.caption("新增或持续项目不会自动形成医学结论；请在下方按“仅记录、健康管理、医生复核或建议复查”完成人工分流。")
    reparse_message_key = f"report-reparse-message-{document_id}"
    if st.session_state.pop(reparse_message_key, None):
        st.success(f"本次重新解析完成，识别 {run.candidate_count} 项候选资料；所有结果仍需人工确认后才会入档。")
    is_processing = run.status in {"PENDING", "PROCESSING"} or bool(st.session_state.get(f"report-reparse-in-progress-{document_id}"))
    manual = [item for item in candidates if item.status == "NEEDS_MANUAL_REVIEW" or item.candidate_type == "INCOMPLETE"]
    operational = [item for item in candidates if item not in manual]
    findings = [item for item in operational if item.candidate_type == "FINDING"]
    followups = [item for item in operational if item.candidate_type == "FOLLOWUP"]
    observations = [item for item in operational if item.candidate_type == "OBSERVATION"]
    management = [item for item in observations if item.abnormal_flag]
    general = [item for item in operational if item not in findings and item not in followups and item not in management]
    with section_frame("需要处理", "报告整理摘要：按角色明确下一步；进入具体事项后可查看原始依据。"):
        _status_strip(
            ("医生复核", len(findings), "attention"),
            ("健康管理", len(management), "action"),
            ("建议复查", len(followups), "action"),
            ("人工核对", len(manual), "urgent"),
        )
    with section_frame("主要结果", "优先查看需要处理的项目；完整指标和技术资料按需展开。"):
        st.caption("异常指标、影像检查和随访建议均保留来源依据；人工确认前不会作为正式医疗结论。")
    with st.expander(f"需要医生复核 · 影像与检查 · {len(findings)} 项", expanded=bool(findings)):
        _render_report_candidate_group("影像与检查 · 需要医生复核", findings[:3], "FINDING", document, key_scope="report-doctor")
        if len(findings) > 3:
            st.caption(f"其余 {len(findings) - 3} 项请在完整报告中继续处理。")
    with st.expander(f"关键健康指标 · {len(observations)} 项"):
        _render_report_candidate_group("关键健康指标", observations, "OBSERVATION", document, key_scope="report-metrics")
    with st.expander(f"主要异常与健康问题 · {len(management)} 项"):
        _render_report_candidate_group("健康管理关注", management, "OBSERVATION", document, key_scope="report-management")
    with st.expander(f"建议复查 · {len(followups)} 项"):
        _render_report_candidate_group("建议复查", followups, "FOLLOWUP", document, key_scope="report-followup")
    with st.expander(f"需要人工核对内容 · {len(manual)} 项"):
        _render_report_manual_review_group(manual, document)
    with st.expander(f"一般记录 · {len(general)} 项"):
        _render_report_candidate_group("一般记录", general, "GENERAL", document, key_scope="report-general")

    with st.expander("查看全部指标"):
        rows = [{"指标": _report_candidate_label(item), "本次": " ".join(part for part in (item.normalized_value or item.raw_value or "—", item.unit or "") if part), "参考范围": item.reference_range or "—", "异常": item.abnormal_flag or "—", "上次": "—", "变化": "—", "状态": {"PENDING_REVIEW":"待确认", "CONFIRMED":"已确认", "REJECTED":"已忽略", "CORRECTED":"已修正"}.get(item.status, "需人工核对")} for item in observations]
        st.dataframe(pd.DataFrame(rows), hide_index=True, width="stretch")
    _render_baseline_draft_action(document, candidates)
    with st.expander("查看解析详情（高级信息）"):
        _render_report_parse_method(run, candidates)
        st.caption("如原始文件已更新，可在此重新整理；已确认资料和长期健康档案不会被覆盖。")
        if secondary_action("重新整理报告" if not is_processing else "报告整理中…", key=f"report-reparse-{document_id}", disabled=is_processing, width="content"):
            st.session_state[f"report-reparse-in-progress-{document_id}"] = True
            try:
                def reparse(progress_callback):
                    with SessionLocal() as session:
                        new_run = ReportParsingService().reparse_document(session, document_id, "health_manager", progress_callback=progress_callback)
                        session.commit()
                        return new_run
                new_run = _run_report_parse_with_progress(reparse)
                st.session_state[reparse_message_key] = str(new_run.id)
                st.rerun()
            except ValueError as error:
                st.error(str(error))
            except Exception:
                LOGGER.exception("report reparse failed", extra={"document_id": str(document_id)})
                st.error("报告重新解析失败。原有解析记录和已确认健康数据均未改变。")
            finally:
                st.session_state[f"report-reparse-in-progress-{document_id}"] = False
        if len(runs) > 1:
            st.caption("查看历史解析")
            st.dataframe(pd.DataFrame([{"解析时间": _fmt_dt(old_run.completed_at or old_run.created_at), "解析方式": "混合解析" if old_run.llm_used else "规则解析", "解析器版本": old_run.parser_version, "本地AI辅助": f"LLM {old_run.llm_call_count} 次" if old_run.llm_used else "未调用", "候选资料": old_run.candidate_count, "状态": "当前" if index == 0 else "历史记录"} for index, old_run in enumerate(runs)]), hide_index=True, width="stretch")
    with st.expander("查看完整文件"):
        path = Path(document.storage_reference)
        if path.is_file():
            st.download_button("查看完整文件", path.read_bytes(), file_name=_source_display_name(document), key=f"report-download-{document.id}")
        else:
            st.warning("原始报告文件不可用。")


def _render_baseline_draft_action(document: Document, candidates: list[ReportExtractionCandidate]) -> None:
    """Manager-only bridge: confirmed report facts may create a draft, never a baseline."""
    with SessionLocal() as session:
        baseline = HealthAssessmentService().latest_baseline(session, document.patient_id)
    if baseline is not None and baseline.status == "CONFIRMED":
        st.caption("该成员已有正式健康基线。本报告会用于长期体检比较，不会覆盖初始基线。")
        return
    if baseline is not None and baseline.status == "DRAFT":
        st.info("健康基线初稿已生成，等待补充资料与健康管理师确认。")
        return
    confirmed = [item for item in candidates if item.status == "CONFIRMED"]
    if not confirmed:
        st.caption("完成至少一项报告资料的人工确认后，可生成健康基线初稿。")
        return
    _section_header("建立健康基线")
    st.caption("初稿只汇总已确认的报告资料和既有健康档案；缺失项目会明确标为待补充。")
    if st.button("生成健康基线初稿", key=f"baseline-draft-{document.id}", type="primary"):
        try:
            with SessionLocal() as session:
                HealthAssessmentService().create_draft_from_report(session, document.patient_id, document.id, created_by="health_manager")
                session.commit()
            st.success("已生成健康基线初稿，请补充资料并确认。")
            st.rerun()
        except ValueError as error:
            st.error(str(error))


def _render_report_candidate_group(title: str, candidates: list[ReportExtractionCandidate], kind: str, document: Document, *, key_scope: str) -> None:
    st.subheader(f"{title} {len(candidates)}")
    if not candidates:
        st.caption("暂无项目。")
        return
    if kind in {"OBSERVATION", "GENERAL"}:
        rows = [{"指标": _report_candidate_label(item), "本次": " ".join(part for part in (item.normalized_value or item.raw_value or "—", item.unit or "") if part), "参考范围": item.reference_range or "—", "异常": item.abnormal_flag or "—", "上次": "—", "变化": "—", "状态": {"PENDING_REVIEW":"待确认", "CONFIRMED":"已确认", "REJECTED":"已忽略", "CORRECTED":"已修正"}.get(item.status, "需人工核对")} for item in candidates if item.candidate_type == "OBSERVATION"]
        if rows:
            st.dataframe(pd.DataFrame(rows), hide_index=True, width="stretch")
        with st.expander("逐项查看依据"):
            for item in [item for item in candidates if item.candidate_type == "OBSERVATION"]:
                st.markdown(f"**{_report_candidate_label(item)}**")
                _render_evidence_action(
                    _candidate_evidence_payload(item, document),
                    key_scope=f"{key_scope}-candidate-{item.id}",
                )
        for item in [item for item in candidates if item.candidate_type == "OBSERVATION" and item.status == "PENDING_REVIEW"]:
            with st.expander(f"处理：{_report_candidate_label(item)}"):
                _render_report_observation_actions(item, key_scope=key_scope)
        return
    for item in candidates:
        if kind == "FINDING":
            with detail_panel(item.summary or "检查结论", "当前处理：等待人工确认或医生复核"):
                st.write("本次：" + (item.summary or "—"))
                st.caption(f"检查：{_report_section_display(item.source_section)} · 第 {item.source_page or '—'} 页")
                _render_report_finding_actions(item, key_scope=key_scope)
                _render_evidence_action(_candidate_evidence_payload(item, document), key_scope=f"{key_scope}-candidate-{item.id}")
                _render_related_knowledge(item.raw_name or item.summary or "", key_scope=f"report-finding-{item.id}")
        elif kind == "FOLLOWUP":
            with detail_panel(item.summary or "建议复查"):
                st.caption(f"来源：第 {item.source_page or '—'} 页 · {_report_section_display(item.source_section)}")
                _render_report_followup_actions(item, key_scope=key_scope)
                _render_evidence_action(_candidate_evidence_payload(item, document), key_scope=f"{key_scope}-candidate-{item.id}")
        else:
            st.markdown(f"**{_report_candidate_label(item)}**")
        if kind not in {"FINDING", "FOLLOWUP"}:
            _render_evidence_action(_candidate_evidence_payload(item, document), key_scope=f"{key_scope}-candidate-{item.id}")


def _render_report_manual_review_group(candidates: list[ReportExtractionCandidate], document: Document) -> None:
    st.subheader(f"需要人工核对内容 {len(candidates)}")
    if not candidates:
        st.caption("没有检测到无法可靠恢复的文本。")
        return
    for item in candidates:
        st.markdown("**原报告文本疑似存在跨行断裂**")
        st.caption("系统未自动生成正式医疗建议；请核对原报告后再人工录入，不会自动创建任务或进入风险评估。")
        st.caption(f"来源：第 {item.source_page or '—'} 页 · {_report_section_display(item.source_section)}")
        _render_evidence_action(_candidate_evidence_payload(item, document), key_scope=f"manual-candidate-{item.id}")
        st.divider()


def _render_report_observation_actions(candidate: ReportExtractionCandidate, *, key_scope: str) -> None:
    if candidate.status != "PENDING_REVIEW":
        st.caption(f"当前状态：{_label(candidate.status, context='report_candidate')}")
        return
    if candidate.candidate_type == "OBSERVATION":
        with SessionLocal() as session:
            duplicate = ReportParsingService().possible_duplicate_observation(session, candidate)
        if duplicate is not None:
            st.warning("该健康指标可能已经入档。为避免重复写入，本次候选仍保留待人工确认。")
            with st.expander("其他处理"):
                st.caption(f"已有记录：{display_observation(duplicate.metric_code)} {duplicate.value_numeric} {duplicate.unit} · {_fmt_dt(duplicate.observed_at)}")
                if st.button("忽略重复项", key=f"{key_scope}-duplicate-reject-{candidate.id}"):
                    with SessionLocal() as session:
                        ReportParsingService().reject_candidate(session, session.get(ReportExtractionCandidate, candidate.id), "health_manager", "与既有健康档案重复")
                        session.commit()
                    st.rerun()
            return
        if st.button("确认入档", key=f"{key_scope}-confirm-{candidate.id}", type="primary"):
            with SessionLocal() as session:
                ReportParsingService().confirm_candidate(session, session.get(ReportExtractionCandidate, candidate.id), "health_manager"); session.commit()
            st.rerun()
        with st.expander("其他处理"):
            if st.button("忽略", key=f"{key_scope}-reject-{candidate.id}"):
                with SessionLocal() as session:
                    ReportParsingService().reject_candidate(session, session.get(ReportExtractionCandidate, candidate.id), "health_manager", "人工忽略"); session.commit()
                st.rerun()
            if st.button("人工修正", key=f"{key_scope}-correct-open-{candidate.id}"):
                st.session_state[f"{key_scope}-correct-visible-{candidate.id}"] = True
            if st.session_state.get(f"{key_scope}-correct-visible-{candidate.id}"):
                with st.form(f"{key_scope}-correct-form-{candidate.id}"):
                    metric_options = list(OBSERVATION)
                    canonical_default = candidate.canonical_code if candidate.canonical_code in metric_options else metric_options[0]
                    canonical = st.selectbox(
                        "对应健康指标", metric_options,
                        index=metric_options.index(canonical_default),
                        format_func=display_observation,
                    )
                    value = st.text_input("数值", value=candidate.normalized_value or candidate.raw_value or "")
                    unit = st.text_input("单位", value=candidate.unit or "")
                    reason = st.text_input("修正原因")
                    if st.form_submit_button("保存修正") and reason:
                        with SessionLocal() as session:
                            ReportParsingService().correct_candidate(session, session.get(ReportExtractionCandidate, candidate.id), "health_manager", canonical=canonical or None, value=value or None, unit=unit or None, reason=reason); session.commit()
                        st.rerun()


def _render_report_finding_actions(candidate: ReportExtractionCandidate, *, key_scope: str) -> None:
    if candidate.status != "PENDING_REVIEW":
        st.caption(f"当前状态：{_label(candidate.status, context='report_candidate')}")
        return
    if st.button("请医生复核", key=f"{key_scope}-doctor-{candidate.id}", type="primary"):
        with SessionLocal() as session:
            ReportParsingService().action_finding(session, session.get(ReportExtractionCandidate, candidate.id), "health_manager", "DOCTOR_REVIEW"); session.commit()
        st.rerun()
    with st.expander("其他处理"):
        left, right = st.columns(2)
        if left.button("纳入健康管理", key=f"{key_scope}-manage-{candidate.id}"):
            with SessionLocal() as session:
                ReportParsingService().action_finding(session, session.get(ReportExtractionCandidate, candidate.id), "health_manager", "MANAGE"); session.commit()
            st.rerun()
        if right.button("仅保留记录", key=f"{key_scope}-record-{candidate.id}"):
            with SessionLocal() as session:
                ReportParsingService().action_finding(session, session.get(ReportExtractionCandidate, candidate.id), "health_manager", "RECORD"); session.commit()
            st.rerun()


def _render_report_followup_actions(candidate: ReportExtractionCandidate, *, key_scope: str) -> None:
    if candidate.status != "PENDING_REVIEW":
        st.caption(f"当前状态：{_label(candidate.status, context='report_candidate')}")
        return
    if st.button("创建随访任务", key=f"{key_scope}-followup-{candidate.id}", type="primary"):
        with SessionLocal() as session:
            ReportParsingService().create_followup_task(session, session.get(ReportExtractionCandidate, candidate.id), "health_manager"); session.commit()
        st.rerun()


def _reset_report_selection_for_new_file(uploaded_file, *, scope: str, selected_key: str) -> None:
    """Keep an old report result from appearing under a newly selected file.

    Streamlit keeps session state across widget reruns.  The selected file is a
    new intake context, so the prior report detail must disappear until this
    file has been explicitly parsed.  Historical runs remain in the database.
    """
    if uploaded_file is None:
        return
    fingerprint = f"{Path(uploaded_file.name).name}:{uploaded_file.size}"
    fingerprint_key = f"report-upload-fingerprint-{scope}"
    if st.session_state.get(fingerprint_key) != fingerprint:
        st.session_state[fingerprint_key] = fingerprint
        st.session_state.pop(selected_key, None)


def render_report_upload(patient: Patient, *, key_prefix: str) -> None:
    st.subheader("上传体检报告")
    st.caption("选择已核对的成员后上传。报告中的姓名仅供人工核对，系统不会据此自动匹配成员。")
    st.caption("开始解析后，系统会用规则识别结构化指标，并使用本地AI辅助整理复杂检查结论。")
    uploaded = st.file_uploader("选择报告文件", type=["pdf", "jpg", "jpeg", "png", "xlsx", "csv", "docx", "txt"], key=f"report-file-{key_prefix}")
    _reset_report_selection_for_new_file(
        uploaded, scope=f"ops-{key_prefix}", selected_key="report-review-document-id",
    )
    parse_key = f"report-parse-in-progress-{key_prefix}"
    in_progress = bool(st.session_state.get(parse_key))
    if uploaded and st.button("解析进行中…" if in_progress else "开始解析报告", type="primary", key=f"report-parse-{key_prefix}", disabled=in_progress):
        st.session_state[parse_key] = True
        try:
            def parse(progress_callback):
                with SessionLocal() as session:
                    document, run, duplicate = ReportParsingService().upload_and_parse(session, patient.id, uploaded.name, uploaded.getvalue(), "health_manager", progress_callback=progress_callback)
                    session.commit()
                    return document, run, duplicate
            document, run, duplicate = _run_report_parse_with_progress(parse)
            st.session_state["report-review-document-id"] = str(document.id)
            st.info("该文件此前已上传过；已使用当前规则和本地AI创建新的解析结果。" if duplicate else "报告已完成本地解析，正在进入人工确认。")
        except Exception:
            LOGGER.exception("operations report intake failed", extra={"member_id": str(patient.id)})
            st.error("报告无法解析。请确认文件完整、格式受支持，或使用人工处理流程。")
        finally:
            st.session_state[parse_key] = False
    selected_id = st.session_state.get("report-review-document-id")
    if selected_id:
        render_report_review(UUID(selected_id))


def render_member_medical_workspace(patient: Patient, ctx: dict[str, list[object]]) -> None:
    """One medical tab with four inline views and no record-detail route."""
    view = st.radio("医疗内容", ["医生复核", "用药", "检查", "手术住院"], horizontal=True, label_visibility="collapsed", key=f"member-medical-view-{patient.id}")
    if view == "医生复核":
        rows: list[DoctorReview] = ctx["reviews"]  # type: ignore[assignment]
        with section_frame("医生复核", "选择记录后在当前区域查看复核问题、意见和依据。"):
            if not rows:
                _empty_state("暂无医生复核", "需要医学判断的事项会在这里显示。")
                return
            selected_key = f"member-medical-review-selected-{patient.id}"
            if st.session_state.get(selected_key) not in {str(item.id) for item in rows}:
                st.session_state[selected_key] = str(rows[0].id)
            left, right = st.columns([1, 1.3], gap="large")
            with left:
                for item in rows[:12]:
                    if secondary_action(f"{_fmt_dt(item.reviewed_at)} · {_label(item.status)}", key=f"member-medical-review-{item.id}"):
                        st.session_state[selected_key] = str(item.id); st.rerun()
            selected = next(item for item in rows if str(item.id) == st.session_state[selected_key])
            with right:
                with detail_panel("医生复核详情", "只显示需要医学判断的记录。"):
                    st.write(selected.opinion or "医生意见待补充")
                    st.caption("复核问题：" + (selected.question_for_doctor or "待补充"))
        return
    if view == "用药":
        render_medications(ctx)
        return
    if view == "检查":
        render_report_upload(patient, key_prefix=f"member-medical-checkup-{patient.id}")
        return
    with SessionLocal() as session:
        events = list(session.scalars(select(HealthEvent).where(HealthEvent.patient_id == patient.id, HealthEvent.event_type.in_(("surgery", "hospitalization"))).order_by(HealthEvent.start_at.desc()).limit(20)))
    with section_frame("手术与住院", "点击一条记录在当前页面展开；未记录的信息不会被推断。"):
        if not events:
            _empty_state("暂无手术或住院记录", "已确认的重要医疗经历会在这里显示。")
            return
        selected_key = f"member-medical-event-selected-{patient.id}"
        if st.session_state.get(selected_key) not in {str(item.id) for item in events}:
            st.session_state[selected_key] = str(events[0].id)
        left, right = st.columns([1, 1.3], gap="large")
        with left:
            for item in events:
                if secondary_action(f"{_fmt_dt(item.start_at)} · {item.description or '医疗记录'}", key=f"member-medical-event-{item.id}"):
                    st.session_state[selected_key] = str(item.id)
                    st.rerun()
        selected = next(item for item in events if str(item.id) == st.session_state[selected_key])
        with right:
            with detail_panel("手术 / 住院详情", "来源、术前术后、用药与随访以已有正式记录为准。"):
                st.markdown(f"**{selected.description or '重要医疗记录'}**")
                st.caption(f"时间：{_fmt_dt(selected.start_at)}")
                st.write("来源：" + (selected.source or "已确认医疗记录"))


def render_member_detail(patient: Patient) -> None:
    if st.button("← 返回运营中心", key="back-to-dashboard"):
        st.session_state.pop("focused_member_id", None)
        st.rerun()
    summary_ctx = _member_summary_context(patient.id)
    _render_member_header(patient, summary_ctx)
    section = st.radio(
        "成员页面", ["概览", "管理", "健康", "医疗", "历程"],
        horizontal=True, label_visibility="collapsed", key=f"member-section-{patient.id}",
    )
    if section == "概览":
        render_simple_member_overview(patient, summary_ctx)
    elif section == "管理":
        management_ctx = _member_management_context(patient.id)
        render_programs(patient, management_ctx)
        render_member_management_signals(patient, management_ctx)
        _section_header("本周任务")
        render_tasks(management_ctx)
        _section_header("最近复盘与结果")
        render_intervention_comparison(patient)
        _section_header("服务摘要")
        render_member_service_management(patient)
    elif section == "健康":
        render_member_archive(patient)
    elif section == "医疗":
        render_member_medical_workspace(patient, _member_medical_context(patient.id))
    else:
        render_longitudinal_timeline(patient, key_scope="member-journey")


def render_member_archive(patient: Patient) -> None:
    """The member-health tab keeps four inline views; timeline is first-level."""
    _section_header("健康")
    key = f"member-health-view-{patient.id}"
    views = ["数据", "体检", "基线", "健康史"]
    if st.session_state.get(key) not in views:
        st.session_state[key] = "数据"
    view = st.radio("成员健康内容", views, horizontal=True, label_visibility="collapsed", key=key)
    if view == "数据":
        render_health_data(patient.id)
        return
    if view == "体检":
        left, right = st.columns([1, 1.7], gap="large")
        with SessionLocal() as session:
            documents = list(session.scalars(select(Document).where(Document.patient_id == patient.id).order_by(Document.created_at.desc()).limit(20)))
        with left:
            with section_frame("报告列表", "选择报告后，在右侧审核与处理。"):
                for document in documents:
                    if secondary_action(f"{_source_display_name(document)} · {_fmt_dt(document.created_at)}", key=f"ops-report-select-{patient.id}-{document.id}"):
                        st.session_state["report-review-document-id"] = str(document.id); st.rerun()
                if not documents:
                    _empty_state("暂无报告", "上传报告后将自动出现在此列表。")
        with right:
            render_report_upload(patient, key_prefix=f"member-health-{patient.id}")
        return
    if view == "基线":
        render_health_assessments(patient)
        return
    if view == "健康史":
        with SessionLocal() as session:
            history = list(session.scalars(select(HealthProblem).where(HealthProblem.patient_id == patient.id).order_by(HealthProblem.opened_at.desc()).limit(20)))
        with section_frame("健康史", "显示已确认的健康问题与既往资料。"):
            if history:
                for item in history:
                    st.markdown(f"<div class='next-row'><div class='focus-title'>{html.escape(item.title)}</div><div class='focus-copy'>{html.escape(item.description or '已确认健康记录')}</div></div>", unsafe_allow_html=True)
            else:
                _empty_state("暂无健康史", "人工确认的既往史和健康问题会在这里长期保留。")
        return
    _empty_state("请从“历程”查看长期健康变化", "健康数据、体检与医疗事件会统一汇入成员的健康历程。")


def _select_archive_timeline(patient_id: UUID) -> None:
    """Compatibility callback for old controls: promote the selection to 历程."""
    request_navigation(member_id=patient_id, member_section="历程", rerun=False)


def render_health_assessments(patient: Patient) -> None:
    st.subheader("健康基线")
    with SessionLocal() as session:
        assessments = HealthAssessmentService().history(session, patient.id)
        baseline = HealthAssessmentService().latest_baseline(session, patient.id)
    if baseline is not None:
        label = "健康基线 · 待确认" if baseline.status == "DRAFT" else "健康基线 · 已建立"
        st.markdown(f"### {label}")
        st.caption(baseline.summary)
        snapshot = baseline.baseline_json or {}
        completeness = snapshot.get("completeness") or {}
        organized = "、".join(completeness.get("organized", [])) or "已确认健康资料"
        pending = "、".join(completeness.get("pending", [])) or "暂无"
        columns = st.columns(2)
        columns[0].write("已整理：" + organized)
        columns[1].write("待补充：" + pending)
        for heading, key in (("基本情况", "basic_information"), ("主要健康问题", "health_problems"), ("关键健康指标", "key_metrics"), ("重要检查结果", "important_findings"), ("当前用药", "current_medications"), ("手术 / 住院史", "procedures_or_hospitalizations"), ("近期生活健康数据", "recent_health_data"), ("当前管理重点", "management_focus")):
            value = snapshot.get(key)
            with st.expander(heading):
                if isinstance(value, list) and value:
                    st.dataframe(pd.DataFrame([_business_detail_row(item) for item in value if isinstance(item, dict)]), hide_index=True, width="stretch") if isinstance(value[0], dict) else st.write("；".join(str(item) for item in value))
                elif isinstance(value, dict):
                    if value.get("label"):
                        st.caption(str(value["label"]))
                    else:
                        st.dataframe(pd.DataFrame([_business_detail_row(value)]), hide_index=True, width="stretch")
                else:
                    st.caption("待补充")
                _render_snapshot_item_evidence(patient.id, value, key_scope=f"baseline-snapshot-{baseline.id}-{key}")
        risk = snapshot.get("risk_summary") or {}
        st.caption("当前风险摘要：" + _risk_text(str(risk.get("level") or "UNKNOWN")))
        with SessionLocal() as session:
            _render_evidence_action(
                _baseline_evidence_payload(session, patient.id, baseline),
                key_scope=f"baseline-{patient.id}-{baseline.id}",
            )
        if baseline.status == "DRAFT":
            st.warning("该初稿仅由已确认报告资料和现有健康档案整理，仍需健康管理师审核后才能成为正式健康基线。")
            if st.button("确认健康基线", key=f"confirm-baseline-{baseline.id}", type="primary"):
                try:
                    with SessionLocal() as session:
                        item = HealthAssessmentService().confirm(session, baseline.id, "health_manager")
                        report_ids = (item.source_references_json or {}).get("source_report_ids", [])
                        for report_id in report_ids:
                            HealthAssessmentService().complete_report_review_task(session, patient.id, UUID(report_id))
                        session.commit()
                    st.success("健康基线已确认，并已进入重大健康时间轴。")
                    st.rerun()
                except ValueError as error:
                    st.error(str(error))
        else:
            st.success(f"已建立 · {_fmt_dt(baseline.confirmed_at or baseline.assessed_at)}")
        if assessments:
            with st.expander("查看健康评估历史"):
                st.dataframe(pd.DataFrame([{"版本": item.version, "类型": {"BASELINE": "健康基线", "REASSESSMENT": "阶段复评", "ANNUAL": "年度评估"}.get(item.assessment_type, item.assessment_type), "状态": "已确认" if item.status == "CONFIRMED" else "待确认", "时间": _fmt_dt(item.confirmed_at or item.assessed_at), "摘要": item.summary, "创建人": item.created_by} for item in assessments]), hide_index=True, width="stretch")
    else:
        _empty_state("尚未建立健康基线", "成员上传最近一次体检报告并完成人工确认后，可生成健康基线初稿。")
    with st.expander("建立新的健康评估"):
        with st.form(f"assessment-{patient.id}"):
            kind = st.selectbox("评估类型", ["INITIAL", "REASSESSMENT"], format_func=lambda item: "初始健康评估" if item == "INITIAL" else "阶段健康复评")
            summary = st.text_area("健康管理摘要", placeholder="基于已确认资料的人工摘要")
            if st.form_submit_button("保存健康评估"):
                if not summary.strip(): st.error("请填写人工健康管理摘要。")
                else:
                    with SessionLocal() as session:
                        HealthAssessmentService().create_assessment(session, patient.id, title="初始健康评估" if kind == "INITIAL" else "阶段健康复评", summary=summary, baseline={}, created_by="健康管理师", assessment_type="BASELINE" if kind == "INITIAL" else "REASSESSMENT", confirmed=True)
                        session.commit()
                    st.success("已保存新的版本化健康评估。")
                    st.rerun()


def render_report_comparison(patient: Patient) -> None:
    st.subheader("体检变化对比")
    st.caption("仅比较两份报告中已人工确认的资料；不会由本地AI决定风险或诊断。")
    with SessionLocal() as session:
        report_document_ids = list(session.scalars(select(ReportExtractionRun.document_id).where(ReportExtractionRun.patient_id == patient.id, ReportExtractionRun.status.in_(("COMPLETED", "PARTIAL_SUCCESS"))).distinct()))
        documents = list(session.scalars(select(Document).where(Document.patient_id == patient.id, Document.id.in_(report_document_ids)).order_by(Document.created_at.desc()))) if report_document_ids else []
    if len(documents) < 2:
        st.caption("至少需要两份已解析并人工确认的体检报告，才能进行变化对比。")
        return
    old, new = st.columns(2)
    old_document = old.selectbox("较早报告", documents, index=1, format_func=lambda item: f"{_source_display_name(item)} · {_fmt_dt(item.created_at)}", key=f"compare-old-{patient.id}")
    new_document = new.selectbox("较新报告", documents, index=0, format_func=lambda item: f"{_source_display_name(item)} · {_fmt_dt(item.created_at)}", key=f"compare-new-{patient.id}")
    if old_document.id == new_document.id:
        st.caption("请选择两份不同报告。")
        return
    try:
        with SessionLocal() as session:
            comparison = ReportComparisonService().compare(session, patient.id, old_document.id, new_document.id)
    except ValueError as error:
        st.caption(str(error)); return
    if comparison["metric_changes"]:
        st.dataframe(pd.DataFrame([{"指标": display_observation(str(item["metric"])), "此前": item["previous"] or "—", "当前": item["current"] or "—", "单位": item["unit"] or "—", "变化": item["delta"] if item["delta"] is not None else "—", "趋势": item["trend"]} for item in comparison["metric_changes"]]), hide_index=True, width="stretch")
        with st.expander("逐项查看依据"):
            with SessionLocal() as session:
                for index, item in enumerate(comparison["metric_changes"]):
                    st.markdown(f"**{_metric_display_name(str(item['metric']))}**")
                    columns = st.columns(2)
                    with columns[0]:
                        st.caption("此前报告")
                        _render_candidate_evidence_by_id(session, patient.id, item.get("previous_candidate_id"), key_scope=f"comparison-old-{patient.id}-{index}")
                    with columns[1]:
                        st.caption("本次报告")
                        _render_candidate_evidence_by_id(session, patient.id, item.get("current_candidate_id"), key_scope=f"comparison-new-{patient.id}-{index}")
    columns = st.columns(3)
    columns[0].markdown("**新增检查结论**\n\n" + ("\n".join(f"• {item}" for item in comparison["new_findings"]) or "—"))
    columns[1].markdown("**持续结论**\n\n" + ("\n".join(f"• {item}" for item in comparison["persistent_findings"]) or "—"))
    columns[2].markdown("**此前未再出现**\n\n" + ("\n".join(f"• {item}" for item in comparison["resolved_findings"]) or "—"))
    finding_evidence = comparison.get("finding_evidence") if isinstance(comparison.get("finding_evidence"), dict) else {}
    if finding_evidence:
        with st.expander("检查结论的查看依据"):
            with SessionLocal() as session:
                for period, document_label in (("old", "此前报告"), ("new", "本次报告")):
                    entries = finding_evidence.get(period) if isinstance(finding_evidence.get(period), dict) else {}
                    for title, candidate_id in entries.items():
                        st.markdown(f"**{document_label} · {title}**")
                        _render_candidate_evidence_by_id(session, patient.id, candidate_id, key_scope=f"comparison-finding-{period}-{patient.id}-{candidate_id}")
    st.caption(comparison["risk_summary"])


def render_intervention_comparison(patient: Patient) -> None:
    st.subheader("干预前后数据变化")
    st.caption("展示观察到的前后变化，不表示本服务或某项干预造成该变化。")
    with SessionLocal() as session:
        latest = session.scalar(select(func.max(Observation.observed_at)).where(Observation.patient_id == patient.id))
        codes = list(session.scalars(select(Observation.metric_code).where(Observation.patient_id == patient.id).distinct().limit(30)))
    if latest is None or not codes:
        st.caption("需要同一指标在干预前后均有可用数据，才能进行比较。")
        return
    default_date = (latest.astimezone(TOKYO_TIMEZONE) - timedelta(days=15)).date()
    cols = st.columns(3)
    metric = cols[0].selectbox("指标", codes, format_func=display_observation, key=f"intervention-metric-{patient.id}")
    started = cols[1].date_input("干预开始日", value=default_date, key=f"intervention-start-{patient.id}")
    days = cols[2].selectbox("比较窗口", [7, 14, 30], format_func=lambda value: f"前后 {value} 天", key=f"intervention-window-{patient.id}")
    pivot = datetime.combine(started, time(12), tzinfo=TOKYO_TIMEZONE)
    with SessionLocal() as session:
        result = InterventionOutcomeService().compare(session, patient.id, metric, pivot, days=days)
    if result and result.get("status") == "READY":
        st.success(f"{display_observation(metric)}：干预前平均 {result['before_summary']:.1f} {result['unit']} → 干预后平均 {result['after_summary']:.1f} {result['unit']}（变化 {result['difference']:+.1f}）")
        st.caption(result["label"])
        _render_evidence_action(
            {
                "source_name": "健康数据观察记录",
                "location": _evidence_location_text(time_window=f"干预日前后各 {days} 天"),
                "evidence_type": "OBSERVATION",
                "raw_evidence": f"干预前样本：{result.get('before_samples', '—')} 条；干预后样本：{result.get('after_samples', '—')} 条。",
                "structured_interpretation": "仅展示干预前后观察到的变化，不表示任何干预造成该变化。",
                "confirmation_status": "健康数据有效记录",
                "evidence_status": "COMPLETE",
            },
            key_scope=f"outcome-comparison-{patient.id}-{metric}-{started.isoformat()}-{days}",
        )
    else:
        st.caption("该窗口前后数据不足，暂不显示前后变化。")


TIMELINE_ZOOM_LEVELS = ("YEAR", "QUARTER", "MONTH", "WEEK")


def _timeline_default_viewport() -> TimelineViewport:
    end = datetime.now(TOKYO_TIMEZONE)
    return TimelineViewport(end - timedelta(days=365), end, "YEAR")


def _timeline_viewport_key(key_scope: str, patient_id: UUID) -> str:
    return f"timeline-viewport-{key_scope}-{patient_id}"


def _timeline_current_viewport(key_scope: str, patient_id: UUID) -> TimelineViewport:
    raw = st.session_state.get(_timeline_viewport_key(key_scope, patient_id))
    if not isinstance(raw, dict):
        return _timeline_default_viewport()
    try:
        start = datetime.fromisoformat(str(raw["start"]))
        end = datetime.fromisoformat(str(raw["end"]))
        if start.tzinfo is None: start = start.replace(tzinfo=TOKYO_TIMEZONE)
        if end.tzinfo is None: end = end.replace(tzinfo=TOKYO_TIMEZONE)
        zoom = str(raw.get("zoom_level", "YEAR")).upper()
        if zoom not in TIMELINE_ZOOM_LEVELS or end <= start:
            raise ValueError("invalid viewport")
        return TimelineViewport(start, end, zoom)
    except (KeyError, TypeError, ValueError):
        return _timeline_default_viewport()


def _set_timeline_viewport(key_scope: str, patient_id: UUID, viewport: TimelineViewport) -> None:
    """Write only router/view state, never a live Streamlit widget key."""
    st.session_state[_timeline_viewport_key(key_scope, patient_id)] = {
        "start": viewport.start.isoformat(), "end": viewport.end.isoformat(), "zoom_level": viewport.zoom_level,
    }


def _timeline_range_slider_key(key_scope: str, patient_id: UUID) -> str:
    return f"timeline-range-slider-ui-{key_scope}-{patient_id}"


def _timeline_pending_range_key(key_scope: str, patient_id: UUID) -> str:
    return f"timeline-range-pending-{key_scope}-{patient_id}"


def _timeline_zoom_level_for_range(start: datetime, end: datetime) -> str:
    days = max(0, (end - start).days)
    if days <= 14:
        return "WEEK"
    if days <= 60:
        return "MONTH"
    if days <= 180:
        return "QUARTER"
    return "YEAR"


def _timeline_viewport_from_dates(start_date: date, end_date: date) -> TimelineViewport:
    start = datetime.combine(start_date, time.min, tzinfo=TOKYO_TIMEZONE)
    end = datetime.combine(end_date, time.max, tzinfo=TOKYO_TIMEZONE)
    return TimelineViewport(start, end, _timeline_zoom_level_for_range(start, end))


def _request_timeline_range(key_scope: str, patient_id: UUID, start: datetime, end: datetime) -> None:
    """Request a range change without mutating an instantiated slider key."""
    st.session_state[_timeline_pending_range_key(key_scope, patient_id)] = {
        "start": start.astimezone(TOKYO_TIMEZONE).date().isoformat(),
        "end": end.astimezone(TOKYO_TIMEZONE).date().isoformat(),
    }


def _apply_pending_timeline_range(key_scope: str, patient_id: UUID) -> TimelineViewport | None:
    """Apply shortcut navigation before the range slider is created."""
    pending = st.session_state.pop(_timeline_pending_range_key(key_scope, patient_id), None)
    if not isinstance(pending, dict):
        return None
    try:
        viewport = _timeline_viewport_from_dates(date.fromisoformat(str(pending["start"])), date.fromisoformat(str(pending["end"])))
    except (KeyError, TypeError, ValueError):
        return None
    _set_timeline_viewport(key_scope, patient_id, viewport)
    # This is safe because this helper is called before st.slider is created.
    st.session_state[_timeline_range_slider_key(key_scope, patient_id)] = (viewport.start.date(), viewport.end.date())
    return viewport


def _trend_value_summary(series, *, hours: bool = False) -> str:
    if not series or not series.points:
        return "暂无数据"
    latest = float(series.points[-1]["value"])
    if hours:
        value = latest / 60
        return f"最新 {value:.1f} 小时"
    return f"最新 {latest:,.0f} {series.unit}"


def _render_timeline_v4_trends(view) -> None:
    """Render health data only; lifecycle events belong exclusively below.

    The range slider synchronizes the data window with the lifecycle axis, but
    the trend frame deliberately contains neither event markers nor overlays.
    This keeps the answer to “身体数据怎么变” visually independent from the
    answer to “这段时间发生了什么”.
    """
    chart_at = lambda value: value.astimezone(TOKYO_TIMEZONE).replace(tzinfo=None)
    domain = [chart_at(view.start), chart_at(view.end)] if view.start is not None else None
    series_by_code = {series.metric_code: series for series in view.metric_series}
    charts = []

    def line_chart(series, color: str):
        frame = pd.DataFrame(series.points)
        frame["at"] = frame["at"].map(chart_at)
        return alt.Chart(frame).mark_line(color=color, strokeWidth=2.3).encode(
            x=alt.X("at:T", title=None, scale=alt.Scale(domain=domain) if domain else alt.Undefined),
            y=alt.Y("value:Q", title=f"{series.display_name}（{series.unit}）"),
            tooltip=[alt.Tooltip("at:T", title="时间"), alt.Tooltip("value:Q", title=series.display_name), alt.Tooltip("samples:Q", title="样本")],
        )

    sleep_total, sleep_deep = series_by_code.get("sleep_duration"), series_by_code.get("deep_sleep_duration")
    if sleep_total is not None:
        total_frame = pd.DataFrame(sleep_total.points)
        total_frame["at"] = total_frame["at"].map(chart_at)
        sleep_layers = [alt.Chart(total_frame).mark_area(color="#c9b7e7", opacity=.42).encode(
            x=alt.X("at:T", title=None, scale=alt.Scale(domain=domain) if domain else alt.Undefined),
            y=alt.Y("value:Q", title="睡眠（分钟）"),
            tooltip=[alt.Tooltip("at:T", title="时间"), alt.Tooltip("value:Q", title="总睡眠"), alt.Tooltip("samples:Q", title="样本")],
        ), alt.Chart(total_frame).mark_line(color="#8264b9", strokeWidth=2.2).encode(x="at:T", y="value:Q")]
        if sleep_deep is not None:
            deep_frame = pd.DataFrame(sleep_deep.points)
            deep_frame["at"] = deep_frame["at"].map(chart_at)
            sleep_layers.append(alt.Chart(deep_frame).mark_line(color="#4b2e83", strokeWidth=3).encode(
                x=alt.X("at:T", scale=alt.Scale(domain=domain) if domain else alt.Undefined), y="value:Q",
                tooltip=[alt.Tooltip("at:T", title="时间"), alt.Tooltip("value:Q", title="深度睡眠"), alt.Tooltip("samples:Q", title="样本")],
            ))
        charts.append(alt.layer(*sleep_layers).properties(title="睡眠趋势 · 总睡眠 / 深度睡眠", height=128))
        st.caption(f"睡眠摘要：{_trend_value_summary(sleep_total, hours=True)} · 深度睡眠 {_trend_value_summary(sleep_deep, hours=True) if sleep_deep else '暂无数据'}")

    steps, calories = series_by_code.get("steps"), series_by_code.get("active_calories")
    if steps is not None:
        steps_frame = pd.DataFrame(steps.points)
        steps_frame["at"] = steps_frame["at"].map(chart_at)
        activity_layers = [alt.Chart(steps_frame).mark_bar(color="#4f9b85", opacity=.72).encode(
            x=alt.X("at:T", title=None, scale=alt.Scale(domain=domain) if domain else alt.Undefined),
            y=alt.Y("value:Q", title="步数"),
            tooltip=[alt.Tooltip("at:T", title="时间"), alt.Tooltip("value:Q", title="步数"), alt.Tooltip("samples:Q", title="样本")],
        )]
        if calories is not None:
            calories_frame = pd.DataFrame(calories.points)
            calories_frame["at"] = calories_frame["at"].map(chart_at)
            activity_layers.append(alt.Chart(calories_frame).mark_line(color="#2d6f9c", strokeWidth=2.4).encode(
                x=alt.X("at:T", scale=alt.Scale(domain=domain) if domain else alt.Undefined),
                y=alt.Y("value:Q", title="活动消耗（kcal）"),
                tooltip=[alt.Tooltip("at:T", title="时间"), alt.Tooltip("value:Q", title="活动消耗"), alt.Tooltip("samples:Q", title="样本")],
            ))
        charts.append(alt.layer(*activity_layers).resolve_scale(y="independent").properties(title="活动趋势 · 步数 / 活动消耗", height=128))
        st.caption(f"活动摘要：{_trend_value_summary(steps)} · 活动消耗 {_trend_value_summary(calories) if calories else '暂无数据'}")

    colors = {"glucose": "#2f6f9f", "systolic_bp": "#267d93", "weight": "#6686a3", "resting_heart_rate": "#ba6c62", "exercise_minutes": "#4f9b85"}
    for code, series in series_by_code.items():
        if code in {"sleep_duration", "deep_sleep_duration", "steps", "active_calories"}:
            continue
        if series.points:
            charts.append(line_chart(series, colors.get(code, "#6686a3")).properties(height=112))
    if charts:
        st.altair_chart(alt.vconcat(*charts).resolve_scale(x="shared"), width="stretch")
    elif not view.metric_series:
        st.caption("当前时间范围内暂无可用于趋势展示的健康数据。")


def _render_timeline_v4_summary(summary) -> None:
    """Window-scoped, factual summary; deltas are observations, not causes."""
    cards = st.columns(4)
    changes = " · ".join(
        f"{item['label']} {item['delta']:+g}{item['unit']}" for item in summary.health_changes[:2]
    ) or "暂无足够连续数据"
    cards[0].markdown("**健康数据**\n\n" + changes + "\n\n观察变化")
    cards[1].markdown(
        "**风险**\n\n"
        f"中风险 {summary.risk_counts.get('medium', 0)} 次 · 高风险 {summary.risk_counts.get('high', 0)} 次"
    )
    medical = " · ".join(f"{name} {count} 次" for name, count in summary.medical_counts.items() if count) or "暂无医疗事件"
    cards[2].markdown("**医疗**\n\n" + medical)
    managed = " · ".join(f"{name} {count} 次" for name, count in {**summary.management_counts, **summary.service_counts}.items() if count) or "暂无管理事件"
    cards[3].markdown("**健康管理**\n\n" + managed)


def _timeline_card_text(event) -> tuple[str, str, str]:
    """Return a compact type badge, title and one-line business summary."""
    title = (event.title or "健康记录").strip()
    normalized = title.lower()
    synthetic_title = normalized.startswith(("demo ", "synthetic_", "synthetic-", "test_")) or "synthetic" in normalized or "合成" in title
    if event.event_type == "medication_change":
        title = title.replace("开始用药记录：", "开始用药").replace("停止用药记录：", "停止用药")
        if title.lower().startswith(("demo ", "synthetic_", "synthetic-", "test_", "开始用药demo")) or "demo medication" in title.lower():
            title = "用药记录"
    elif event.event_type == "report" and ("synthetic" in normalized or "合成" in title or "（较早）" in title):
        title = "年度体检"
    elif normalized.startswith(("demo ", "synthetic_", "synthetic-", "test_")) or "90-day metabolic" in normalized:
        synthetic_titles = {
            "program_start": "90天代谢健康计划", "program_adjustment": "健康管理计划",
            "assessment": "阶段健康评估", "health_data_summary": "健康数据总结",
            "external_referral": "外部医疗记录", "doctor_review": "医生复核",
        }
        title = synthetic_titles.get(event.event_type, "健康记录")
    if title.lower() in {"synthetic_prog", "unknown", "none"}:
        title = "用药记录" if event.event_type == "medication_change" else ("90天代谢健康计划" if event.event_type == "program_start" else "健康管理计划")
    badge = event.risk_label if event.event_type == "risk" else (event.event_type_label or get_event_type_display(event.event_type))
    if synthetic_title and event.event_type != "risk":
        badge = "演示 · " + badge
    summary = (event.summary or "已记录").strip()
    if summary.lower() in {"unknown", "none"} or "synthetic_" in summary.lower() or "demo medication" in summary.lower():
        summary = "已记录健康事件"
    return badge, title, summary


_LEFT_TIMELINE_EVENT_TYPES = frozenset({"program_start", "program_adjustment", "intervention", "medication_change", "service"})
_RIGHT_TIMELINE_EVENT_TYPES = frozenset({
    "risk", "doctor_review", "external_referral", "report", "assessment", "major_problem",
    "procedure", "surgery", "hospitalization", "health_data_summary", "outcome",
})


def _lifecycle_lane(event) -> str:
    """Apply the product's fixed lane contract; the renderer never guesses."""
    if event.event_type in _LEFT_TIMELINE_EVENT_TYPES:
        return "LEFT"
    if event.event_type in _RIGHT_TIMELINE_EVENT_TYPES:
        return "RIGHT"
    return "RIGHT"


def _lifecycle_rows(events) -> list[dict[str, object]]:
    """Group visible major events by local day, with at most one card per lane."""
    by_day: dict[date, list[object]] = {}
    for event in sorted(events, key=lambda item: (item.occurred_at, item.group_key)):
        by_day.setdefault(event.occurred_at.astimezone(TOKYO_TIMEZONE).date(), []).append(event)
    rows: list[dict[str, object]] = []
    for day, day_events in by_day.items():
        left = [event for event in day_events if _lifecycle_lane(event) == "LEFT"]
        right = [event for event in day_events if _lifecycle_lane(event) == "RIGHT"]
        rows.append({"date": day, "left": left, "right": right, "events": day_events})
    return rows


def _timeline_lane_card(events, *, side: str, selected_event_id: str | None, key_scope: str, patient_id: UUID) -> str | None:
    """Render one compact selectable card; no action can exist outside its card."""
    if not events:
        return None
    event = events[-1]
    badge, title, summary = _timeline_card_text(event)
    if len(events) > 1:
        summary = f"{len(events)}项记录"
    selected = event.group_key == selected_event_id
    risk_key = event.risk_level.lower() if event.event_type == "risk" and event.risk_level in {"GREEN", "YELLOW", "RED"} else "neutral"
    card_key = f"timeline-card-{side}-{risk_key}-{'selected' if selected else 'idle'}-{key_scope}-{patient_id}-{event.group_key}"
    # The Streamlit button is intentionally the card itself.  Rendering a
    # markdown card followed by a sibling control previously created a detached
    # action when an empty lane changed the row height.
    label = f"{badge}\n\n**{title}**\n\n{summary}"
    with st.container(key=card_key):
        if st.button(label, key=f"timeline-card-select-{side}-{key_scope}-{patient_id}-{event.group_key}", type="secondary", width="stretch"):
            return event.group_key
    return None


def _timeline_spine_markup(day: date, *, has_left: bool, has_right: bool, risk_level: str | None, selected: bool, has_next: bool) -> str:
    """A small, row-owned date spine. It deliberately has no free coordinates."""
    node_class = (
        f"timeline-spine-risk timeline-spine-risk-{risk_level.lower()}"
        if risk_level in {"GREEN", "YELLOW", "RED"}
        else ("timeline-spine-selected" if selected else "timeline-event-marker")
    )
    left_connector = "<span class='timeline-spine-connector'></span>" if has_left else "<span class='timeline-spine-empty'></span>"
    right_connector = "<span class='timeline-spine-connector'></span>" if has_right else "<span class='timeline-spine-empty'></span>"
    continuation = "timeline-spine-continuation" if has_next else "timeline-spine-continuation timeline-spine-continuation-end"
    return (
        f"<div class='timeline-row-spine'><strong>{day.strftime('%m-%d')}</strong>"
        f"<div class='timeline-spine-node'>{left_connector}<span class='{node_class}'></span>{right_connector}</div>"
        f"<span class='{continuation}'></span></div>"
    )


def _render_lifecycle_grid(rows: list[dict[str, object]], selected_event_id: str | None, *, key_scope: str, patient_id: UUID) -> str | None:
    """Stable three-column lifecycle renderer using ordinary Streamlit columns."""
    if not rows:
        st.caption("该时间范围内暂无重大健康事件。")
        return None
    clicked_event_id: str | None = None
    previous_year: int | None = None
    with st.container(key=f"lifecycle-grid-{key_scope}-{patient_id}"):
        for row_index, row in enumerate(rows):
            day = row["date"]
            assert isinstance(day, date)
            if previous_year != day.year:
                st.markdown(f"<div class='timeline-year-break'><span>{day.year}</span></div>", unsafe_allow_html=True)
            previous_year = day.year
            left_events = row["left"]
            right_events = row["right"]
            assert isinstance(left_events, list) and isinstance(right_events, list)
            row_events = row["events"]
            assert isinstance(row_events, list)
            primary_event = (right_events or left_events)[-1]
            risk_levels = [event.risk_level for event in row_events if event.event_type == "risk"]
            risk_level = next((level for level in ("RED", "YELLOW", "GREEN") if level in risk_levels), None)
            with st.container(key=f"timeline-row-{key_scope}-{patient_id}-{day.isoformat()}"):
                columns = st.columns([2.8, 1, 2.8], gap="small")
                with columns[0]:
                    clicked_event_id = _timeline_lane_card(left_events, side="left", selected_event_id=selected_event_id, key_scope=key_scope, patient_id=patient_id) or clicked_event_id
                with columns[1]:
                    st.markdown(
                        _timeline_spine_markup(
                            day, has_left=bool(left_events), has_right=bool(right_events), risk_level=risk_level,
                            selected=primary_event.group_key == selected_event_id, has_next=row_index < len(rows) - 1,
                        ),
                        unsafe_allow_html=True,
                    )
                with columns[2]:
                    clicked_event_id = _timeline_lane_card(right_events, side="right", selected_event_id=selected_event_id, key_scope=key_scope, patient_id=patient_id) or clicked_event_id
    return clicked_event_id


def render_longitudinal_timeline(patient: Patient, *, key_scope: str = "archive", client_view: bool = False) -> None:
    st.subheader("健康历程")
    st.caption("按日期查看健康管理、用药、风险和医疗大事件。")
    # The lifecycle uses normal document flow: one grid row per calendar day.
    # Risk colour is reserved for formal risk records; all other nodes remain neutral.
    st.markdown(
        """<style>
        .timeline-risk-legend {display:flex; align-items:center; gap:.7rem; flex-wrap:wrap; margin:.2rem 0 1rem; color:#68778a; font-size:.78rem;}
        .timeline-risk-indicator {display:inline-flex; align-items:center; gap:.28rem; font-weight:650; white-space:nowrap;}
        .timeline-risk-dot {display:inline-block; width:10px; height:10px; border-radius:50%; background:#8795a1;}
        .timeline-risk-green .timeline-risk-dot {background:#3c9b62;}.timeline-risk-yellow .timeline-risk-dot {background:#d9951c;}.timeline-risk-red .timeline-risk-dot {background:#d84b4b;}
        .timeline-risk-neutral .timeline-risk-dot {background:transparent; border:1.5px solid #8795a1; box-sizing:border-box;}
        .timeline-event-marker {display:inline-block; width:10px; height:10px; border-radius:50%; background:#fff; border:2px solid #8795a1; box-sizing:border-box;}
        .timeline-event-badge {display:inline-block; margin-bottom:.2rem; padding:1px 6px; border-radius:999px; background:#edf2f7; color:#526173; font-size:.68rem; font-weight:650; line-height:1.4;}
        /* A selectable card is one native control. This prevents a detached action
           button and keeps target/action proximity intact in every lane state. */
        [class*="st-key-timeline-card-"] {max-width:292px; margin-top:.05rem; margin-bottom:.05rem;}
        [class*="st-key-timeline-card-left-"] {margin-left:auto;}
        [class*="st-key-timeline-card-right-"] {margin-right:auto;}
        [class*="st-key-timeline-card-"] .stButton {width:100%;}
        [class*="st-key-timeline-card-"] .stButton > button {min-height:92px; width:100%; padding:.62rem .72rem; border:1px solid #dce5ec; border-radius:9px; background:#fff; color:#26384a; box-shadow:none; text-align:left; align-items:flex-start; overflow-wrap:anywhere; white-space:pre-line;}
        [class*="st-key-timeline-card-"] .stButton > button:hover {border-color:#97abc0; background:#f8fafc; color:#1f3448;}
        [class*="st-key-timeline-card-"] .stButton > button:focus-visible {outline:2px solid #2f6f9f; outline-offset:2px;}
        [class*="st-key-timeline-card-"] .stButton > button p {margin:0; color:#5f7183; font-size:.79rem; font-weight:500; line-height:1.42; text-align:left; white-space:pre-line;}
        [class*="st-key-timeline-card-"] .stButton > button strong {color:#26384a; font-size:.92rem; font-weight:700;}
        [class*="st-key-timeline-card-"][class*="-green-"] .stButton > button {border-left:3px solid #4a9968;}
        [class*="st-key-timeline-card-"][class*="-yellow-"] .stButton > button {border-left:3px solid #d79a28;}
        [class*="st-key-timeline-card-"][class*="-red-"] .stButton > button {border-left:3px solid #d45555;}
        [class*="st-key-timeline-card-"][class*="-selected-"] .stButton > button {border:2px solid #2f6f9f; padding:calc(.62rem - 1px) calc(.72rem - 1px); background:#f8fbff;}
        [class*="st-key-timeline-card-"] .stButton > button[kind="secondary"]:active {background:#f1f5f9;}
        [class*="st-key-timeline-card-"]:has(.stButton > button:focus-visible) .stButton > button {border-color:#2f6f9f;}
        .timeline-row-spine {display:flex; min-height:104px; flex-direction:column; align-items:center; color:#516274; font-size:.74rem; line-height:1.15; text-align:center;}
        .timeline-row-spine strong {font-size:.84rem; font-weight:700; letter-spacing:.01em; white-space:nowrap;}
        .timeline-spine-node {display:grid; grid-template-columns:26px 12px 26px; align-items:center; justify-content:center; min-height:22px; margin-top:.22rem;}
        .timeline-spine-connector {height:1px; background:#c4d0da;}.timeline-spine-empty {display:block; height:1px;}
        .timeline-spine-risk,.timeline-spine-selected,.timeline-event-marker {display:block; width:11px; height:11px; border-radius:50%; box-sizing:border-box;}
        .timeline-spine-risk,.timeline-event-marker {background:#fff; border:2px solid #8795a1;}.timeline-spine-risk-green {border-color:#3c9b62;}.timeline-spine-risk-yellow {border-color:#d9951c;}.timeline-spine-risk-red {border-color:#d84b4b;}.timeline-spine-selected {background:#fff; border:2px solid #2f6f9f;}
        .timeline-spine-continuation {display:block; flex:1 1 auto; min-height:56px; margin-top:3px; border-left:1px solid #c4d0da;}
        .timeline-spine-continuation-end {border-left-color:transparent;}
        .timeline-year-break {display:flex; align-items:center; gap:.5rem; margin:.2rem 0 .48rem; color:#708091; font-size:.76rem; font-weight:700; letter-spacing:.04em;}
        .timeline-year-break::before,.timeline-year-break::after {content:""; flex:1; border-top:1px solid #dce5ec;}
        @media (max-width: 760px) {
            [class*="st-key-timeline-card-"] {max-width:none; margin-left:0; margin-right:0;}
            .timeline-row-spine {min-height:72px; align-items:flex-start; text-align:left;}.timeline-spine-node {justify-content:flex-start;width:auto;}
            .timeline-spine-continuation {min-height:20px; margin-left:5px;}
            [class*="st-key-lifecycle-grid"] [data-testid="stHorizontalBlock"] {flex-direction:column !important; gap:.3rem !important;}
            [class*="st-key-lifecycle-grid"] [data-testid="stHorizontalBlock"] > [data-testid="column"]:nth-child(1) {order:2;}
            [class*="st-key-lifecycle-grid"] [data-testid="stHorizontalBlock"] > [data-testid="column"]:nth-child(2) {order:1;}
            [class*="st-key-lifecycle-grid"] [data-testid="stHorizontalBlock"] > [data-testid="column"]:nth-child(3) {order:3;}
        }
        </style>""",
        unsafe_allow_html=True,
    )
    selected_key = f"timeline-selected-{key_scope}-{patient.id}"
    selected_event_key = f"timeline-selected-event-{key_scope}-{patient.id}"
    with SessionLocal() as session:
        timeline_service = TimelineV4Service()
        history_bounds = timeline_service.available_history_bounds(session, patient.id)
    if history_bounds is None:
        st.info("暂无健康历程数据。上传体检报告或连接健康数据后，这里会展示可调整的长期健康历程。")
        return
    history_start, history_end = history_bounds
    earliest_date, latest_date = history_start.date(), history_end.date()
    initialized_key = f"timeline-range-initialized-{key_scope}-{patient.id}"
    pending_viewport = _apply_pending_timeline_range(key_scope, patient.id)
    if pending_viewport is not None:
        viewport = pending_viewport
    elif not st.session_state.get(initialized_key):
        default_start = max(earliest_date, latest_date - timedelta(days=365))
        viewport = _timeline_viewport_from_dates(default_start, latest_date)
        _set_timeline_viewport(key_scope, patient.id, viewport)
        st.session_state[initialized_key] = True
    else:
        viewport = _timeline_current_viewport(key_scope, patient.id)
    # Never allow stale routing state from another member or an older UI to
    # query outside this member's available history.
    range_start = max(earliest_date, min(viewport.start.date(), latest_date))
    range_end = max(range_start, min(viewport.end.date(), latest_date))
    viewport = _timeline_viewport_from_dates(range_start, range_end)
    _set_timeline_viewport(key_scope, patient.id, viewport)
    shortcut_controls = st.columns([.72, .78, .9, .8, .72, 3.1])
    shortcuts = (("近7天", 7), ("近30天", 30), ("近3个月", 90), ("近1年", 365), ("全部", None))
    for column, (label, days) in zip(shortcut_controls[:5], shortcuts):
        if column.button(label, key=f"timeline-range-shortcut-{label}-{key_scope}-{patient.id}"):
            shortcut_start = earliest_date if days is None else max(earliest_date, latest_date - timedelta(days=days))
            _request_timeline_range(
                key_scope, patient.id,
                datetime.combine(shortcut_start, time.min, tzinfo=TOKYO_TIMEZONE),
                datetime.combine(latest_date, time.max, tzinfo=TOKYO_TIMEZONE),
            )
            st.session_state.pop(selected_key, None)
            st.session_state.pop(selected_event_key, None)
            st.rerun()
    shortcut_controls[5].caption("拖动下方横条两端，直接调整整个健康时间窗口。")
    start, end = viewport.start, viewport.end
    controls = st.columns([1.3, 2.2])
    with SessionLocal() as session:
        v4_service = TimelineV4Service()
        available_metrics = v4_service.available_metric_codes(session, patient.id, start=start, end=end)
    # Companion values are added automatically to the paired sleep/activity
    # charts, rather than requiring members to understand internal metrics.
    selectable_metrics = [code for code in available_metrics if code not in {"deep_sleep_duration", "active_calories"}]
    metric_options = {code: TimelineV4Service._METRICS[code][0] for code in selectable_metrics}
    selected_metric_codes = controls[1].multiselect(
        "健康趋势（最多选择 4 项）", options=list(metric_options), default=list(metric_options)[:3],
        format_func=lambda code: metric_options[code], key=f"timeline-metrics-{key_scope}-{patient.id}",
        placeholder="选择要关联查看的健康指标",
    )
    # Sleep and activity are shown as purposeful paired charts. Selecting the
    # primary metric therefore includes its companion when data exists.
    if "sleep_duration" in selected_metric_codes and "deep_sleep_duration" in metric_options and "deep_sleep_duration" not in selected_metric_codes:
        selected_metric_codes.append("deep_sleep_duration")
    if "steps" in selected_metric_codes and "active_calories" in metric_options and "active_calories" not in selected_metric_codes:
        selected_metric_codes.append("active_calories")
    if len(selected_metric_codes) > 6:
        st.warning("趋势图最多同时展示 4 项指标，当前仅显示前 4 项。")
        selected_metric_codes = selected_metric_codes[:6]
    # One viewport projection feeds trend, summary, semantic clusters and the
    # inspector; changing it never invokes parsing, LLM or risk evaluation.
    with SessionLocal() as session:
        # HealthTimelineService().get_timeline remains the major-event source;
        # TimelineV4Service adds semantic zoom without copying business data.
        view = timeline_service.get_timeline_view(
            session, patient.id, viewport=viewport, metric_codes=tuple(selected_metric_codes),
            event_limit=36 if (end - start).days > 180 else 100,
        )
    category = st.radio("事件筛选", ["全部", "体检", "健康数据", "风险", "用药", "医疗", "健康管理", "服务"], horizontal=True, label_visibility="collapsed", key=f"timeline-category-{key_scope}-{patient.id}")
    category_types = {
        "体检": {"report"}, "健康数据": {"health_data_summary"}, "风险": {"risk"}, "用药": {"medication_change"},
        "医疗": {"doctor_review", "external_referral", "procedure", "surgery", "hospitalization"},
        "健康管理": {"assessment", "program_start", "program_adjustment", "major_problem", "outcome"},
        "服务": {"service"},
    }
    lane_events = [event for event in view.events if category == "全部" or event.event_type in category_types[category]]
    rows = _lifecycle_rows(lane_events)
    selected_event_id = st.session_state.get(selected_event_key) or st.session_state.get(selected_key)
    selected_event = next((item for item in lane_events if item.group_key == selected_event_id), None)
    if selected_event is None and rows:
        latest_row_events = rows[-1]["events"]
        assert isinstance(latest_row_events, list)
        selected_event = latest_row_events[-1]
    slider_key = _timeline_range_slider_key(key_scope, patient.id)
    with _section_frame("健康趋势", "上面看身体数据怎么变；下方生命轴单独说明这段时间发生了什么。"):
        _render_timeline_v4_trends(view)
        st.markdown("**时间范围**")
        if earliest_date < latest_date:
            selected_dates = st.slider(
                "健康历程时间范围", min_value=earliest_date, max_value=latest_date,
                value=None if slider_key in st.session_state else (viewport.start.date(), viewport.end.date()), step=timedelta(days=1),
                format="YYYY-MM-DD", key=slider_key,
                help="拖动左右两个手柄，调整健康趋势、期间总结和健康生命轴的时间范围。",
            )
            st.caption(
                f"起始：{viewport.start.date().isoformat()} · 结束：{viewport.end.date().isoformat()} · "
                f"当前范围：{(viewport.end.date() - viewport.start.date()).days + 1} 天"
            )
            if isinstance(selected_dates, tuple) and len(selected_dates) == 2 and selected_dates != (viewport.start.date(), viewport.end.date()):
                _set_timeline_viewport(key_scope, patient.id, _timeline_viewport_from_dates(selected_dates[0], selected_dates[1]))
                st.session_state.pop(selected_key, None)
                st.session_state.pop(selected_event_key, None)
                st.rerun()
        else:
            st.caption(f"目前只有 {earliest_date.isoformat()} 一项健康记录。等有更多历史后即可调整时间范围。")
    window_title = f"{start.astimezone(TOKYO_TIMEZONE).date().isoformat()} ～ {end.astimezone(TOKYO_TIMEZONE).date().isoformat()}"
    with _section_frame("当前期间总结", f"{window_title}；以下仅为当前窗口内的观察变化与事件统计。"):
        _render_timeline_v4_summary(view.summary)
    # The axis owns the available reading width.  The one inspector follows it
    # in normal document flow, so detailed evidence never squeezes timeline
    # rows or competes with the left/right event lanes.
    lifecycle = st.container()
    inspector = st.container()
    with lifecycle:
        with section_frame("健康生命轴", "按日期查看健康管理、用药、风险和医疗大事件；选择事件后，详情只在右侧更新。"):
            headers = st.columns([2.8, 1, 2.8], gap="small")
            headers[0].caption("健康管理 · 用药 · 服务")
            headers[1].caption("日期")
            headers[2].caption("风险 · 医疗 · 医生")
            clicked_event_id = _render_lifecycle_grid(
                rows, selected_event.group_key if selected_event else None,
                key_scope=key_scope, patient_id=patient.id,
            )
            if clicked_event_id:
                st.session_state[selected_key] = clicked_event_id
                st.session_state[selected_event_key] = clicked_event_id
                st.rerun()
            if any(item.event_type == "health_data_summary" for item in lane_events):
                if st.button("查看当前时间范围健康数据", key=f"timeline-data-{key_scope}-window"):
                    health_window = {
                        "start": start.astimezone(TOKYO_TIMEZONE).date().isoformat() if start else None,
                        "end": end.astimezone(TOKYO_TIMEZONE).date().isoformat(),
                    }
                    if client_view:
                        request_navigation(
                            surface="成员健康中心", member_page="健康数据", member_id=patient.id,
                            health_data_window=health_window,
                        )
                    else:
                        request_navigation(
                            ops_page="成员", member_id=patient.id, member_section="健康",
                            archive_view="健康数据", health_data_window=health_window,
                        )
    if selected_event is None:
        with inspector.container(border=True):
            st.markdown("### 事件详情")
            st.caption("拖动上方时间横条到包含重大健康事件的范围，再点击生命轴事件查看详情。")
        return
    event = selected_event
    with inspector.container(border=True):
        st.markdown(_timeline_event_badge(event), unsafe_allow_html=True)
        _, display_title, _ = _timeline_card_text(event)
        st.markdown(f"### {event.occurred_at.astimezone(TOKYO_TIMEZONE).strftime('%Y-%m-%d')} · {display_title}")
        risk_indicator = _timeline_risk_indicator(event)
        if risk_indicator:
            st.markdown(f"当前风险：{risk_indicator}", unsafe_allow_html=True)
        elif (related_risk := (event.expandable_details or {}).get("related_risk_level")):
            risk_label = _risk_text(str(related_risk))
            st.markdown(f"关联风险：{_status_pill(risk_label)}", unsafe_allow_html=True)
        st.write(event.summary)
        details = event.expandable_details or {}
        with SessionLocal() as session:
            _render_evidence_action(
                _timeline_evidence_payload(session, patient, event),
                key_scope=f"timeline-{key_scope}-{patient.id}-{event.group_key or event.related_entity}",
                client_view=client_view,
            )
        status_context = "doctor_review" if event.event_type == "doctor_review" else "risk_event" if event.event_type == "risk" else "health_problem"
        business_details = {
            "当前状态": _label(str(details.get("status")) if details.get("status") else None, context=status_context),
            "负责人": _role_label(str(details.get("owner")) if details.get("owner") else None),
            "科室": details.get("department"),
            "医生问题": details.get("question"), "机构": details.get("organization"), "外部反馈": details.get("feedback"),
            "开始日期": details.get("start_date"), "结束日期": details.get("end_date"),
        }
        client_visible_keys = {"当前状态", "科室", "机构"}
        visible = {
            key: value
            for key, value in business_details.items()
            if value not in {None, "", "未记录", "负责人待分配"} and (not client_view or key in client_visible_keys)
        }
        if visible:
            st.dataframe(pd.DataFrame([visible]), hide_index=True, width="stretch")
        if event.event_type == "risk" and not client_view:
            matches = details.get("matches") if isinstance(details.get("matches"), list) else []
            st.markdown("**风险处理链路**")
            st.write(f"触发指标：{_metric_display_name(str(details.get('metric') or ''))} · 匹配记录：{details.get('matched_count', len(matches))}")
            if matches:
                st.caption("当时数据：" + "；".join(f"{item.get('value')} {item.get('unit', '')}" for item in matches[-3:] if isinstance(item, dict)))
            with SessionLocal() as session:
                risk_id = UUID(event.related_entity) if event.related_entity else None
                tasks = list(session.scalars(select(Task).where(Task.risk_event_id == risk_id).order_by(Task.created_at.desc()).limit(5))) if risk_id else []
                reviews = list(session.scalars(select(DoctorReview).where(DoctorReview.risk_event_id == risk_id).order_by(DoctorReview.reviewed_at.desc()).limit(3))) if risk_id else []
            if tasks:
                st.write("后续任务：" + "；".join(item.title for item in tasks))
            if reviews:
                st.write("医生意见：" + "；".join(item.opinion for item in reviews))
        if event.event_type == "assessment":
            st.markdown("**当时健康快照**")
            snapshot = {key: value for key, value in details.items() if key not in {"status", "reviewed_by", "source_references"} and value not in {None, "", [], {}}}
            if snapshot:
                basic = snapshot.get("basic_information") if isinstance(snapshot.get("basic_information"), dict) else {}
                columns = st.columns(2)
                columns[0].markdown("**主要健康问题**\n\n" + ("\n".join(f"• {item.get('title', '已确认健康问题')}" for item in snapshot.get("health_problems", [])[:5] if isinstance(item, dict)) or "待补充"))
                columns[1].markdown("**当前用药**\n\n" + ("\n".join(f"• {item.get('name', '已记录用药')}" for item in snapshot.get("current_medications", [])[:5] if isinstance(item, dict)) or "待补充"))
                metric_rows = snapshot.get("key_metrics") if isinstance(snapshot.get("key_metrics"), list) else []
                if metric_rows:
                    st.markdown("**关键指标**")
                    st.dataframe(pd.DataFrame([{
                        "指标": display_observation(str(item.get("metric", "—"))),
                        "数值": f"{item.get('value', '—')} {item.get('unit', '')}",
                    } for item in metric_rows[:8] if isinstance(item, dict)]), hide_index=True, width="stretch")
                focus = snapshot.get("management_focus") if isinstance(snapshot.get("management_focus"), list) else []
                if focus:
                    st.markdown("**初始管理重点**\n\n" + "\n".join(f"• {item}" for item in focus[:3]))
                if basic:
                    safe_basic = _business_detail_row(basic)
                    st.caption("基本情况：" + " · ".join(f"{key}：{value}" for key, value in safe_basic.items() if value not in {None, "", "待补充"}))
                if client_view:
                    st.caption("这次评估汇总了当时已确认的健康问题、健康数据和当前管理重点。")
            else:
                st.caption("该次评估未记录额外结构化快照。")
        if event.event_type == "report":
            st.markdown("**报告摘要**")
            report_details = {
                "报告日期": details.get("report_date"), "体检机构": details.get("hospital"),
                "主要发现": f"{details.get('findings', 0)} 项", "关键指标": f"{details.get('metrics', 0)} 项",
                "建议复查": f"{details.get('followups', 0)} 项", "审核状态": _label(str(details.get("review_state")) if details.get("review_state") else None, context="report_candidate"),
            }
            st.dataframe(pd.DataFrame([{key: value for key, value in report_details.items() if value not in {None, ""}}]), hide_index=True, width="stretch")
            comparison = details.get("comparison") if isinstance(details.get("comparison"), dict) else None
            if comparison:
                st.markdown("**与上次相比**\n\n" + " · ".join((
                    f"新增 {comparison.get('new', 0)}",
                    f"持续 {comparison.get('persistent', 0)}",
                    f"发生变化 {comparison.get('changed', 0)}",
                    f"未复查 {comparison.get('not_rechecked', 0)}",
                )))
            document_id = details.get("document_id")
            if document_id:
                with SessionLocal() as session:
                    report_findings = list(session.scalars(
                        select(ReportExtractionCandidate).where(
                            ReportExtractionCandidate.patient_id == patient.id,
                            ReportExtractionCandidate.document_id == UUID(str(document_id)),
                            ReportExtractionCandidate.candidate_type == "FINDING",
                            ReportExtractionCandidate.status == "CONFIRMED",
                        ).order_by(ReportExtractionCandidate.source_page).limit(6)
                    ))
                    if report_findings:
                        st.markdown("**主要发现**")
                        for finding in report_findings:
                            st.write(f"• {finding.summary or '已确认检查结果'}")
                            _render_evidence_action(
                                _candidate_evidence_payload(finding, _document_for_member_evidence(session, patient.id, finding.document_id)),
                                key_scope=f"timeline-report-finding-{key_scope}-{finding.id}", client_view=client_view,
                            )
            if client_view:
                if st.button("查看体检报告", key=f"timeline-member-report-{event.related_entity}"):
                    st.session_state[f"member-report-document-{patient.id}"] = str(document_id)
                    request_navigation(
                        surface="成员健康中心", member_page="健康档案", member_id=patient.id,
                        archive_view="体检与检查",
                    )
            else:
                actions = st.columns(2)
                if actions[0].button("查看完整体检", key=f"timeline-report-{event.related_entity}"):
                    request_navigation(
                        ops_page="成员", member_id=patient.id, member_section="健康",
                        archive_view="体检与检查", report_document_id=details["document_id"],
                    )
                if actions[1].button("查看新旧报告对比", key=f"timeline-report-comparison-{event.related_entity}"):
                    request_navigation(ops_page="成员", member_id=patient.id, member_section="健康", archive_view="报告对比")
        if event.event_type == "health_data_summary":
            st.markdown("**这段时间的健康数据**")
            metrics = details.get("metrics") if isinstance(details.get("metrics"), list) else []
            if metrics:
                st.dataframe(pd.DataFrame([{
                    "指标": item.get("label"), "平均": item.get("average"), "单位": item.get("unit"),
                    "数据量": item.get("samples"), "本月变化": item.get("direction"),
                } for item in metrics[:6]]), hide_index=True, width="stretch")
            if st.button("查看这段时间的完整健康数据", key=f"timeline-data-{key_scope}-{event.group_key}"):
                health_window = {"start": details.get("window_start"), "end": details.get("window_end")}
                if client_view:
                    request_navigation(
                        surface="成员健康中心", member_page="健康数据", member_id=patient.id,
                        health_data_window=health_window,
                    )
                else:
                    request_navigation(
                        ops_page="成员", member_id=patient.id, member_section="健康",
                        archive_view="健康数据", health_data_window=health_window,
                    )
        if event.event_type in {"procedure", "surgery", "hospitalization"}:
            st.markdown("**医疗履历详情**")
            st.write(f"记录：{event.summary}")
            st.caption("仅显示已有正式记录；没有记录的信息不会被推断。")
            if st.button("查看手术履历详情", key=f"timeline-procedure-{event.group_key}"):
                request_navigation(ops_page="成员", member_id=patient.id, member_section="医疗")
        if event.event_type == "medication_change":
            st.markdown("**用药变化**")
            st.write(f"变化：{details.get('change_type', '用药记录更新')} · 记录来源：{details.get('record_source', '正式记录')}")
            if client_view and st.button("查看用药与医疗", key=f"timeline-member-medication-{event.group_key}"):
                request_navigation(
                    surface="成员健康中心", member_page="健康档案", member_id=patient.id,
                    archive_view="用药与医疗",
                )
        if event.event_type in {"program_start", "program_adjustment", "external_referral", "doctor_review"}:
            st.markdown("**下一步与关联处理**")
            if details.get("goal"):
                st.write(f"管理目标：{details['goal']}")
            if details.get("question"):
                st.write(f"希望医生确认：{details['question']}")
            if details.get("next_focus"):
                st.write(f"下一阶段重点：{details['next_focus']}")
        if event.event_type == "outcome":
            st.markdown("**干预前后观察到的变化**")
            st.write(f"{details.get('metric', '指标')}：{details.get('before', '—')} {details.get('unit', '')} → {details.get('after', '—')} {details.get('unit', '')}")
            st.caption("该比较不表示任何干预造成了该变化。")
            if not client_view and st.button("查看前后对比", key=f"timeline-outcome-{event.group_key}"):
                request_navigation(ops_page="成员", member_id=patient.id, member_section="管理")


def _client_device_status(assignment: MemberDeviceAssignment) -> str:
    """Translate implementation states into client-facing connection language."""
    if assignment.assignment_status == "DISABLED" or assignment.connection_status == "DISABLED":
        return "停用"
    return {
        "CONNECTED": "已连接", "SYNCING": "正在同步", "MOCK": "演示数据",
        "PENDING": "待连接",
    }.get(assignment.connection_status, "待连接")


def render_member_report_upload(patient: Patient) -> None:
    """Member-surface intake reusing the existing parser, with no review controls."""
    _section_header("上传体检报告", "上传后将由健康管理团队审核整理结果；成员本人不能确认医疗资料。")
    uploaded = st.file_uploader(
        "选择最近一次体检报告", type=["pdf", "jpg", "jpeg", "png", "xlsx", "csv", "docx", "txt"],
        key=f"member-report-file-{patient.id}",
    )
    _reset_report_selection_for_new_file(
        uploaded, scope=f"member-{patient.id}", selected_key=f"member-report-document-{patient.id}",
    )
    parsing_key = f"member-report-parsing-{patient.id}"
    if uploaded and st.button("上传并整理报告", key=f"member-report-parse-{patient.id}", type="primary", disabled=bool(st.session_state.get(parsing_key))):
        st.session_state[parsing_key] = True
        try:
            def parse(progress_callback):
                with SessionLocal() as session:
                    document, run, duplicate = ReportParsingService().upload_and_parse(
                        session, patient.id, uploaded.name, uploaded.getvalue(), "member_surface", progress_callback=progress_callback,
                    )
                    HealthAssessmentService().ensure_report_review_task(session, patient.id, document)
                    session.commit()
                    return document, run, duplicate
            document, _, duplicate = _run_report_parse_with_progress(parse)
            st.session_state[f"member-report-document-{patient.id}"] = str(document.id)
            st.success("报告已上传，正在整理体检报告……" if not duplicate else "该报告已上传过，已创建新的整理结果供健康管理团队审核。")
        except Exception:
            LOGGER.exception("member report intake failed")
            st.error("报告整理遇到问题，请稍后重试或联系健康管理团队。")
        finally:
            st.session_state[parsing_key] = False
    document_id = st.session_state.get(f"member-report-document-{patient.id}")
    if not document_id:
        return
    with SessionLocal() as session:
        document = session.get(Document, UUID(document_id))
        runs = ReportParsingService().runs(session, UUID(document_id)) if document else []
        run = runs[0] if runs else None
        candidates = ReportParsingService().candidates(session, document.id, run.id) if document and run else []
    if document is None or run is None:
        return
    if run.status == "NEEDS_OCR":
        st.warning("当前报告暂时无法完整读取，健康管理团队将进一步处理。")
        return
    if run.llm_enabled and not run.llm_available:
        st.caption("本地AI暂不可用，系统已保留规则整理结果，健康管理团队会继续人工确认。")
    metrics = [item for item in candidates if item.candidate_type == "OBSERVATION"]
    findings = [item for item in candidates if item.candidate_type == "FINDING"]
    followups = [item for item in candidates if item.candidate_type == "FOLLOWUP"]
    pending = [item for item in candidates if item.status in {"PENDING_REVIEW", "NEEDS_MANUAL_REVIEW"}]
    with SessionLocal() as session:
        baseline = HealthAssessmentService().latest_baseline(session, patient.id)
    if run.status in {"PROCESSING", "PENDING"}:
        intake_status, intake_next = "正在整理", "系统正在整理报告；完成后将交由健康管理团队核对。"
    elif pending:
        intake_status, intake_next = "等待审核", "报告已收到并完成初步整理，等待健康管理团队核对后入档。"
    elif baseline and baseline.status == "DRAFT":
        intake_status, intake_next = "健康档案正在建立", "报告已审核，健康基线初稿正在等待健康管理团队确认。"
    elif baseline and baseline.status == "CONFIRMED":
        intake_status, intake_next = "审核完成", "报告已纳入健康档案；后续变化将由健康管理团队持续跟进。"
    else:
        intake_status, intake_next = "已收到", "报告已收到，健康管理团队将继续审核整理结果。"
    with section_frame("报告处理进度", "成员只需关注当前进度和下一步，无需理解内部解析状态。"):
        st.markdown(status_badge(intake_status), unsafe_allow_html=True)
        st.write(intake_next)
    with SessionLocal() as session:
        report_risk = ReportRiskSummaryService().summarize(session, patient.id, document.id)
    with section_frame("本次核心结论", "以下为报告整理摘要；健康管理团队会进一步确认后纳入长期健康档案。"):
        st.markdown(risk_badge(report_risk["level"]), unsafe_allow_html=True)
        st.markdown(f"**{report_risk['reason']}**")
        st.caption("处理方式：" + _report_risk_next_step(str(report_risk["level"])))
    with section_frame("报告整理结果", "先看异常、随访和需要人工确认的内容。"):
        _status_strip(
            ("异常指标", len([item for item in metrics if item.abnormal_flag]), "attention"),
            ("检查结果", len(findings), "attention"),
            ("建议随访", len(followups), "action"),
            ("等待确认", len(pending), "neutral"),
        )
        st.caption("报告已保留；如需继续上传，请选择另一份体检报告。")
    with st.expander("查看分类结果"):
        if metrics:
            st.markdown("**关键健康指标**")
            st.dataframe(pd.DataFrame([{
                "指标": _report_candidate_label(item), "本次": " ".join(part for part in (item.normalized_value or item.raw_value or "—", item.unit or "") if part),
                "状态": "等待确认" if item.status != "CONFIRMED" else "已确认",
            } for item in metrics]), hide_index=True, width="stretch")
        if findings:
            st.markdown("**影像与检查**")
        for item in findings[:8]:
            st.write(f"**{item.summary or '检查结果'}**")
        if followups:
            st.markdown("**建议复查**")
        for item in followups[:8]:
            st.write(f"**{item.summary or '建议复查'}**")
        if pending:
            st.markdown(f"**需要人工核对 · {len(pending)} 项**")
        with st.expander("查看依据"):
            for item in candidates[:16]:
                st.markdown(f"**{_report_candidate_label(item)}**")
                _render_evidence_action(
                    _candidate_evidence_payload(item, document),
                    key_scope=f"member-upload-evidence-{patient.id}-{item.id}", client_view=True,
                )


def _render_member_baseline_center(patient: Patient) -> None:
    with SessionLocal() as session:
        baseline = HealthAssessmentService().latest_baseline(session, patient.id)
    if baseline is None:
        _empty_state("尚未建立健康基线", "上传最近一次体检报告后，系统可帮助健康管理团队整理初稿。")
        render_member_report_upload(patient)
        return
    snapshot = baseline.baseline_json or {}
    if baseline.status == "DRAFT":
        st.markdown("### 健康基线 · 待确认")
        st.caption("已由健康管理团队开始整理；成员补充资料后仍需健康管理师确认。")
        st.caption("您补充的内容会标记为成员自述资料，不会自动作为医学确认结论。")
        completeness = snapshot.get("completeness") or {}
        st.write("待补充：" + ("、".join(completeness.get("pending", [])) or "暂无"))
        with st.expander("补充我的健康资料"):
            with st.form(f"member-reported-baseline-{baseline.id}"):
                history = st.text_area("既往史")
                medication = st.text_area("当前用药信息")
                procedure = st.text_area("手术 / 住院史")
                family = st.text_area("家族史")
                lifestyle = st.text_area("生活方式资料")
                if st.form_submit_button("提交给健康管理团队"):
                    with SessionLocal() as session:
                        HealthAssessmentService().update_member_reported(session, baseline.id, {
                            "既往史": history, "当前用药": medication, "手术 / 住院史": procedure,
                            "家族史": family, "生活方式资料": lifestyle,
                        })
                        session.commit()
                    st.success("已提交为成员自述资料，等待健康管理团队审核。")
                    st.rerun()
    else:
        st.markdown("### 健康基线 · 已建立")
        st.caption(f"建立日期：{_fmt_dt(baseline.confirmed_at or baseline.assessed_at)}")
    st.write(baseline.summary)
    with SessionLocal() as session:
        _render_evidence_action(
            _baseline_evidence_payload(session, patient.id, baseline),
            key_scope=f"client-baseline-{patient.id}-{baseline.id}", client_view=True,
        )
    for heading, key in (("基本情况", "basic_information"), ("主要健康问题", "health_problems"), ("关键健康指标", "key_metrics"), ("重要检查结果", "important_findings"), ("当前用药", "current_medications"), ("手术 / 住院史", "procedures_or_hospitalizations"), ("近期生活健康数据", "recent_health_data"), ("当前管理重点", "management_focus")):
        with st.expander(heading):
            value = snapshot.get(key)
            if isinstance(value, list) and value:
                st.dataframe(pd.DataFrame([_business_detail_row(item) for item in value if isinstance(item, dict)]), hide_index=True, width="stretch") if isinstance(value[0], dict) else st.write("；".join(str(item) for item in value))
            elif isinstance(value, dict):
                if value.get("label"):
                    st.caption(str(value["label"]))
                else:
                    st.dataframe(pd.DataFrame([_business_detail_row(value)]), hide_index=True, width="stretch")
            else:
                st.caption("待补充")
            _render_snapshot_item_evidence(patient.id, value, key_scope=f"client-baseline-snapshot-{baseline.id}-{key}", client_view=True)


def _render_member_center_baseline_entry(patient: Patient) -> None:
    """Prominent archive-root entry; upload parsing happens only after click."""
    with SessionLocal() as session:
        assessments = HealthAssessmentService()
        baseline = assessments.latest_baseline(session, patient.id, include_draft=False)
        draft = assessments.latest_baseline(session, patient.id, include_draft=True)
    _section_header("健康基线")
    if baseline is None:
        if draft is not None and draft.status == "DRAFT":
            st.markdown("**健康基线 · 待确认**")
            st.caption("健康管理团队正在补充与审核初稿；确认后会成为正式健康基线。")
            if st.button("查看健康基线初稿", key=f"member-baseline-open-{patient.id}"):
                st.session_state[f"client-archive-{patient.id}"] = "健康基线"
                st.rerun()
        else:
            st.markdown("**尚未建立健康基线**")
            st.caption("上传最近一次体检报告，系统可以帮助整理健康基线初稿。")
            if st.button("上传最近体检报告", key=f"member-baseline-upload-{patient.id}", type="primary"):
                _open_member_report_upload(patient.id)
    else:
        st.markdown("**健康基线 · 已建立**")
        st.caption(f"建立日期：{_fmt_dt(baseline.confirmed_at or baseline.assessed_at)} · 新体检报告可在下方“体检与检查”上传并进行长期比较。")
        if st.button("查看完整健康基线", key=f"member-baseline-open-{patient.id}"):
            st.session_state[f"client-archive-{patient.id}"] = "健康基线"
            st.rerun()
def _open_member_report_upload(patient_id: UUID, *, rerun: bool = True) -> None:
    """Open the member health checkup view without creating a deeper route."""
    request_navigation(
        surface="成员健康中心", member_page="健康", member_id=patient_id,
        archive_view="体检与检查", rerun=rerun,
    )


def _open_member_baseline_upload(patient_id: UUID) -> None:
    _open_member_report_upload(patient_id, rerun=False)


def _open_ops_member_report_upload(patient_id: UUID) -> None:
    """Open the shared report intake in the operations health view."""
    st.session_state[f"medical-report-visible-{patient_id}"] = True
    request_navigation(
        surface="运营后台", ops_page="成员", member_id=patient_id,
        member_section="健康", archive_view="体检与检查",
    )


def _report_upload_label(report_count: int, has_confirmed_baseline: bool) -> str:
    """Keep the permanent CTA human-readable while preserving all states."""
    return "上传新体检报告" if report_count and has_confirmed_baseline else "上传体检报告"


def _report_upload_state(
    report_count: int, latest_title: str | None, has_confirmed_baseline: bool,
) -> tuple[str, str, str]:
    """Return member-facing report copy without ever removing the upload CTA."""
    label = _report_upload_label(report_count, has_confirmed_baseline)
    if report_count == 0:
        return "尚未上传体检报告", "上传最近一次体检报告，可帮助建立初始健康档案。", label
    if has_confirmed_baseline:
        title = latest_title or "体检报告"
        return f"最近体检：{title}", f"历史体检：{report_count} 份。上传新报告后，系统会与历史结果比较并更新长期健康记录。", label
    if latest_title:
        return f"已上传 {report_count} 份报告", f"最近：{latest_title} · 等待健康管理团队确认", label
    return f"已上传 {report_count} 份报告", "当前状态：等待健康管理团队确认", label


def _render_client_report_intake_entry(
    patient: Patient,
    *,
    report_count: int,
    latest_document: Document | None,
    has_confirmed_baseline: bool,
    compact: bool = False,
) -> None:
    """Permanent, state-aware member-facing report entry; one shared intake."""
    title = "最近体检" if compact else "体检与检查"
    _section_header(title)
    with st.container(border=True):
        headline, guidance, label = _report_upload_state(
            report_count,
            _source_display_name(latest_document) if latest_document else None,
            has_confirmed_baseline,
        )
        st.markdown(f"**{headline}**")
        st.caption(guidance)
        left, right = st.columns([1, 1])
        with left:
            if st.button(
                label,
                key=f"client-report-upload-permanent-{'home' if compact else 'archive'}-{patient.id}",
                type="primary",
            ):
                _open_member_report_upload(patient.id)
        if report_count:
            with right:
                if st.button("查看体检与检查", key=f"client-report-open-{'home' if compact else 'archive'}-{patient.id}"):
                    request_navigation(
                        surface="成员健康中心", member_page="健康档案", member_id=patient.id,
                        archive_view="体检与检查",
                    )


def _render_client_home(patient: Patient, ctx: dict[str, list[object]]) -> None:
    risk, reason, _ = _member_risk_state(patient.id, ctx)
    name = html.escape(patient.display_name or "成员")
    st.markdown(
        f"<div class='client-hero'><h1>我的健康</h1><p>欢迎回来，{name}</p>"
        f"{risk_badge(risk)}<div class='focus-copy' style='margin-top:.7rem'>"
        f"{html.escape(reason or '当前没有正式风险评估；其他主要指标会随确认资料持续更新。')}</div></div>",
        unsafe_allow_html=True,
    )
    _section_header("我现在怎么样")
    st.caption("状态会随着已确认健康资料和人工复核结果更新。")
    next_task = next((item for item in ctx["tasks"] if item.status not in {"COMPLETED", "CANCELLED"}), None)
    with SessionLocal() as session:
        lifestyle = HealthDataSummaryService().get_lifestyle_summary(session, patient.id, days=7)
        realtime = HealthDataSummaryService().get_realtime_summary(session, patient.id)
    with section_frame("今天", "只保留今天最重要的健康数据；详细趋势在健康数据中查看。"):
        cards = st.columns(3)
        today_cards = (
            ("睡眠", lifestyle.latest.get("sleep_duration"), "今日"),
            ("深度睡眠", lifestyle.latest.get("deep_sleep_duration"), "今日"),
            ("步数", lifestyle.latest.get("steps"), "今日"),
            ("活动消耗", lifestyle.latest.get("active_calories"), "今日"),
            ("血压", realtime.latest_systolic, "最近一次"),
            ("血糖", realtime.cgm_current, "持续关注"),
        )
        for index, (label, observation, note) in enumerate(today_cards):
            with cards[index % 3]:
                if observation is None:
                    health_metric_card(label, "暂无数据", "连接后自动显示")
                elif label == "血压" and realtime.latest_diastolic:
                    health_metric_card(label, f"{int(observation.value_numeric)} / {int(realtime.latest_diastolic.value_numeric)}", note)
                else:
                    health_metric_card(label, _format_observation_value(observation), note)
        if not any(item for _, item, _ in today_cards):
            _empty_state("暂无今日健康数据", "连接健康数据来源后，这里会自动显示步数、睡眠和健康监测数据。")

    with section_frame("最近变化", "用少量变化帮助你判断今天需要关注什么。"):
        observations = ctx.get("observations", [])
        if observations:
            rows = "".join(
                f"<div class='timeline-preview'><div class='timeline-date'>{html.escape(_metric_display_name(item.metric_code))}</div>"
                f"<div><div class='timeline-title'>{html.escape(_format_observation_value(item))}</div>"
                f"<div class='timeline-copy'>最近一次健康记录</div></div></div>"
                for item in observations[:3]
            )
            st.markdown(rows, unsafe_allow_html=True)
        else:
            _empty_state("暂无近期变化", "后续健康数据会在这里形成趋势。")
    with section_frame("我的下一步", "只显示近期最重要的安排。"):
        if next_task:
            st.markdown(f"<div class='next-row'><div class='next-date'>{html.escape(_fmt_dt(next_task.due_at) if next_task.due_at else '待安排')}</div><div class='focus-title'>{html.escape(next_task.title)}</div><div class='focus-copy'>{html.escape(next_task.instruction)}</div></div>", unsafe_allow_html=True)
        else:
            _empty_state("暂无待办事项", "健康管理师安排的任务、复核和下次服务会在这里显示。")
    with section_frame("快速入口", "需要时进入相应页面继续处理，不在首页展开完整资料。"):
        report, data, service, plan = st.columns(4)
        with report:
            if primary_action("上传体检报告", key=f"client-home-report-{patient.id}"):
                _open_member_report_upload(patient.id)
        with data:
            if secondary_action("查看健康数据", key=f"client-home-data-{patient.id}"):
                request_navigation(surface="成员健康中心", member_page="健康", member_id=patient.id, archive_view="健康数据")
        with service:
            if secondary_action("申请服务", key=f"client-home-service-{patient.id}"):
                request_navigation(surface="成员健康中心", member_page="服务", member_id=patient.id)
        with plan:
            if secondary_action("查看健康计划", key=f"client-home-plan-{patient.id}"):
                request_navigation(surface="成员健康中心", member_page="计划", member_id=patient.id)


def _render_client_plan(patient: Patient, ctx: dict[str, list[object]]) -> None:
    """Three plan views; task and result details stay on the same page."""
    _page_header("计划", "查看当前方案、我的任务和阶段结果。", eyebrow="成员健康中心")
    program = _active_program(ctx)
    active_tasks = [item for item in ctx["tasks"] if item.status not in {"COMPLETED", "CANCELLED"}]
    completed_tasks = [item for item in ctx["tasks"] if item.status == "COMPLETED"]
    view = st.radio("计划内容", ["当前方案", "我的任务", "阶段结果"], horizontal=True, label_visibility="collapsed", key=f"client-plan-view-{patient.id}")
    if view == "当前方案":
        with section_frame("当前方案", "健康管理团队确认后执行；不包含自动医疗处方或诊断。"):
            if program:
                st.markdown(f"### {program.title}")
                st.caption(program.main_goal or "当前目标待健康管理团队补充。")
                day = _program_day(program)
                if day:
                    st.progress(min(day / 90, 1.0), text=f"第 {day} 天 / 90 天")
                choice = st.radio("我的选择", ["接受方案", "希望调整", "暂缓", "和健康管理师讨论"], horizontal=True, key=f"member-plan-choice-{patient.id}")
                if primary_action("记录选择", key=f"member-plan-choice-save-{patient.id}", width="content"):
                    with SessionLocal() as session:
                        MemberServiceOperations().record_choice(session, patient.id, choice)
                        session.commit()
                    st.success("已记录您的选择，健康管理师会后续跟进。")
            else:
                _empty_state("暂无当前健康方案", "健康管理师确认方案后，会在这里说明目标和执行安排。")
        return
    if view == "我的任务":
        total = len(active_tasks) + len(completed_tasks)
        with section_frame("我的任务", "按今天、本周与逾期安排查看；完成情况不代表医疗结果。"):
            summary_metric("本周完成", f"{len(completed_tasks)} / {total}" if total else "暂无任务", "持续完成比一次性冲刺更重要")
            rows = [*completed_tasks, *active_tasks]
            if rows:
                st.markdown("".join(
                    f"<div class='next-row'><div class='focus-title'>{'✓' if task.status == 'COMPLETED' else '○'}　{html.escape(task.title)}</div>"
                    f"<div class='focus-copy'>{html.escape(_fmt_dt(task.due_at) if task.due_at else '本周安排')} · {html.escape(task.instruction)}</div></div>"
                    for task in rows[:10]
                ), unsafe_allow_html=True)
            else:
                _empty_state("暂无待完成任务", "本周需要完成的健康管理任务会在这里显示。")
        return
    with section_frame("阶段结果", "仅呈现已经确认的前后变化，不推断干预原因。"):
        with SessionLocal() as session:
            outcomes = list(session.scalars(select(OutcomeEvaluation).where(OutcomeEvaluation.patient_id == patient.id).order_by(OutcomeEvaluation.evaluation_date.desc()).limit(5)))
        if outcomes:
            for outcome in outcomes:
                st.markdown(f"<div class='next-row'><div class='focus-title'>{html.escape(_metric_display_name(outcome.metric))}：{html.escape(str(outcome.baseline_value))} {html.escape(outcome.unit)} → {html.escape(str(outcome.current_value))} {html.escape(outcome.unit)}</div><div class='focus-copy'>{html.escape(outcome.notes or '已确认阶段结果')}</div></div>", unsafe_allow_html=True)
        else:
            _empty_state("尚无阶段结果", "完成阶段复盘后，确认的结果会显示在这里。")


def render_member_service_management(patient: Patient) -> None:
    with SessionLocal() as session:
        operations = MemberServiceOperations(); plan = operations.ensure_demo_plan(session, patient.id); session.commit()
        rows = operations.member_services(session, patient.id)
        requests = list(session.scalars(select(ServiceRequest).where(ServiceRequest.patient_id == patient.id).order_by(ServiceRequest.requested_at.desc()).limit(20)))
        names = {item.id: item.name for item, _ in rows}
    with _section_frame("服务管理", "审核成员申请、安排服务并记录结果；服务申请不等于自动医疗预约。"):
        st.markdown(f"**当前服务计划：{plan.name.replace('（演示）', '').replace('(演示)', '')}**")
        st.caption(f"可用服务：{len(rows)} 项 · 待处理申请：{sum(item.status != 'COMPLETED' for item in requests)} 项")
        st.dataframe(pd.DataFrame([{"服务": item.name, "使用情况": "不限次" if entitlement.total_quota is None else f"已使用 {entitlement.used_quota} / {entitlement.total_quota} 次"} for item, entitlement in rows]), hide_index=True, width="stretch")
        if not requests:
            _empty_state("暂无服务申请", "成员申请服务后，健康管理师可在这里审核、安排和记录结果。")
            return
        st.markdown("**服务申请与执行**")
        for request in requests:
            with st.container(border=True):
                st.markdown(f"**{names.get(request.service_item_id, '会员服务')}** · {_label(request.status)}")
                st.caption(f"申请原因：{request.reason} · 负责人：{request.assigned_manager or '待分配'} · 安排时间：{_fmt_dt(request.scheduled_at) if request.scheduled_at else '待安排'}")
                if request.status == "REQUESTED" and st.button("审核服务申请", key=f"service-approve-{request.id}", type="primary"):
                    with SessionLocal() as session:
                        MemberServiceOperations().approve(session, request.id, "健康管理师"); session.commit()
                    st.rerun()
                elif request.status == "APPROVED":
                    with st.form(f"service-schedule-{request.id}"):
                        scheduled_day = st.date_input("安排日期", value=date.today() + timedelta(days=3), key=f"service-schedule-date-{request.id}")
                        scheduled_time = st.time_input("安排时间", value=time(10, 0), key=f"service-schedule-time-{request.id}")
                        manager = st.text_input("负责人", value=request.assigned_manager or "健康管理师", key=f"service-schedule-manager-{request.id}")
                        if st.form_submit_button("确认服务安排"):
                            with SessionLocal() as session:
                                MemberServiceOperations().schedule(session, request.id, datetime.combine(scheduled_day, scheduled_time, tzinfo=TOKYO_TIMEZONE), manager)
                                session.commit()
                            st.rerun()
                elif request.status == "SCHEDULED":
                    if st.button("记录开始服务", key=f"service-start-{request.id}"):
                        with SessionLocal() as session:
                            MemberServiceOperations().start(session, request.id, request.assigned_manager or "健康管理师")
                            session.commit()
                        st.rerun()
                elif request.status == "IN_PROGRESS":
                    with st.form(f"service-complete-{request.id}"):
                        result = st.text_area("服务结果摘要", placeholder="仅记录已完成的服务结果与后续安排", key=f"service-result-{request.id}")
                        if st.form_submit_button("记录服务完成"):
                            with SessionLocal() as session:
                                MemberServiceOperations().complete(session, request.id, result, request.assigned_manager or "健康管理师")
                                session.commit()
                            st.rerun()
                elif request.status == "COMPLETED":
                    st.caption("完成结果：" + (request.result_summary or "已完成，结果待补充。"))
                    _render_evidence_action(
                        {"source_name": "服务执行记录", "location": "服务结果记录", "evidence_type": "TEXT", "raw_evidence": request.result_summary or "当前未保存结果摘要。", "structured_interpretation": "服务完成结果由健康管理团队记录。", "confirmation_status": _label(request.status), "evidence_status": "PARTIAL", "show_no_knowledge": True},
                        key_scope=f"ops-service-evidence-{request.id}",
                    )


def _render_client_service(patient: Patient, ctx: dict[str, list[object]]) -> None:
    _page_header("服务", "查看可用服务、我的申请和已完成服务。", eyebrow="成员健康中心")
    with SessionLocal() as session:
        operations = MemberServiceOperations(); plan = operations.ensure_demo_plan(session, patient.id); session.commit()
        services = operations.member_services(session, patient.id)
        requests = list(session.scalars(select(ServiceRequest).where(ServiceRequest.patient_id == patient.id).order_by(ServiceRequest.requested_at.desc()).limit(30)))
    view = st.radio("服务内容", ["可用服务", "我的申请", "服务记录"], horizontal=True, label_visibility="collapsed", key=f"client-service-view-{patient.id}")
    names = {item.id: item.name for item, _ in services}
    if view != "可用服务":
        rows = requests if view == "我的申请" else [item for item in requests if item.status == "COMPLETED"]
        with section_frame(view, "选择一条记录后，同页查看状态、安排和结果。"):
            if not rows:
                _empty_state("暂无服务记录", "提交服务申请后，审核、安排和结果会显示在这里。")
                return
            selected_key = f"client-service-request-selected-{patient.id}"
            if st.session_state.get(selected_key) not in {str(item.id) for item in rows}:
                st.session_state[selected_key] = str(rows[0].id)
            left, right = st.columns([1, 1.3], gap="large")
            with left:
                for request in rows[:12]:
                    if secondary_action(f"{names.get(request.service_item_id, '会员服务')} · {_label(request.status)}", key=f"client-service-select-{request.id}"):
                        st.session_state[selected_key] = str(request.id); st.rerun()
            selected = next(item for item in rows if str(item.id) == st.session_state[selected_key])
            with right:
                with detail_panel(names.get(selected.service_item_id, "会员服务"), "服务状态与结果由健康管理团队确认。"):
                    st.write(selected.reason or "成员服务申请")
                    st.caption(f"当前状态：{_label(selected.status)} · 申请时间：{_fmt_dt(selected.requested_at)}")
                    if selected.scheduled_at:
                        st.write(f"已安排：{_fmt_dt(selected.scheduled_at)}")
                    next_member_action = {
                        "REQUESTED": "健康管理团队将审核您的申请。",
                        "REVIEWING": "健康管理团队正在审核您的申请。",
                        "APPROVED": "申请已通过，正在安排服务时间。",
                        "SCHEDULED": "服务已安排，请按约定时间准备。",
                        "IN_PROGRESS": "服务正在进行，完成后将更新结果。",
                        "COMPLETED": "服务已完成，结果已记录。",
                        "CANCELLED": "本次服务已取消；如仍有需要，可重新申请。",
                    }.get(selected.status, "健康管理团队将更新下一步安排。")
                    st.write(selected.result_summary or next_member_action)
                    if selected.status in {"REQUESTED", "REVIEWING", "APPROVED", "SCHEDULED"}:
                        if secondary_action("取消本次申请", key=f"client-service-cancel-{selected.id}", width="content"):
                            with SessionLocal() as session:
                                MemberServiceOperations().cancel(session, selected.id, "成员取消本次服务申请。")
                                session.commit()
                            st.success("已取消本次服务申请；如仍有需要，可重新申请。")
                            st.rerun()
        return
    with section_frame("当前会员", "当前服务计划与权益以健康管理团队确认记录为准。"):
        st.markdown(f"### {html.escape(plan.name.replace('（演示）', '').replace('(演示)', ''))}")
    categories = ["健管服务", "精准诊疗", "就医协助", "远程问诊", "会员权益"]
    category = st.selectbox("服务分类", categories, key=f"client-service-category-filter-{patient.id}")
    available = [(item, entitlement) for item, entitlement in services if item.category == category]
    selected_key = f"client-service-selected-{patient.id}"
    if st.session_state.get(selected_key) not in {str(item.id) for item, _ in available}:
        st.session_state[selected_key] = str(available[0][0].id) if available else None
    left, right = st.columns([1, 1.3], gap="large")
    with left:
        if not available:
            _empty_state("暂无可使用服务", "当前服务计划下暂无此类别服务。")
        for item, entitlement in available:
            quota = "不限次" if entitlement.total_quota is None else f"剩余 {max(entitlement.total_quota - entitlement.used_quota, 0)} 次"
            if secondary_action(f"{item.name} · {quota}", key=f"client-service-select-item-{item.id}"):
                st.session_state[selected_key] = str(item.id); st.rerun()
    with right:
        selected = next(((item, entitlement) for item, entitlement in available if str(item.id) == st.session_state.get(selected_key)), None)
        if selected:
            item, entitlement = selected
            quota = "不限次" if entitlement.total_quota is None else f"剩余 {max(entitlement.total_quota - entitlement.used_quota, 0)} 次"
            with detail_panel(item.name, quota):
                st.write(item.description)
                if primary_action("申请服务", key=f"member-service-request-{item.id}", width="content"):
                    with SessionLocal() as session:
                        MemberServiceOperations().request(session, patient.id, item.id, f"成员申请：{item.name}")
                        session.commit()
                    st.success("服务申请已提交，健康管理师会审核权益与安排。")
                    st.rerun()


def _render_client_profile(patient: Patient) -> None:
    _page_header("个人设置", "个人资料、数据来源和隐私授权。", eyebrow="成员健康中心")
    view = st.radio("个人设置内容", ["资料", "设备与数据", "隐私授权"], horizontal=True, label_visibility="collapsed", key=f"client-profile-view-{patient.id}")
    if view == "资料":
        with section_frame("个人资料", "需要更正姓名、联系方式或紧急联系人时，请联系健康管理团队。"):
            st.write(f"**{_member_display(patient)}**")
            st.caption("医生确认后的反馈、服务安排和任务提醒会在相应业务页面显示。")
        return
    if view == "设备与数据":
        with SessionLocal() as session:
            assignments = list(session.scalars(select(MemberDeviceAssignment).where(MemberDeviceAssignment.patient_id == patient.id).order_by(MemberDeviceAssignment.assigned_at.desc()).limit(12)))
            apple_job = session.scalar(select(IngestionJob).where(IngestionJob.patient_id == patient.id, IngestionJob.source_system == "apple_health").order_by(IngestionJob.completed_at.desc()))
        with section_frame("我的健康数据来源", "只展示已分配的数据来源和连接状态。"):
            if not assignments:
                _empty_state("暂无健康数据来源", "请联系健康管理师协助连接日常健康设备或医疗监测设备。")
            else:
                for category, items in (("日常健康", [item for item in assignments if item.device_category == "WELLNESS"]), ("医疗监测", [item for item in assignments if item.device_category != "WELLNESS"])):
                    if items:
                        st.markdown(f"**{category}**")
                        st.caption(" · ".join(f"{display_provider(item.provider)}：{_client_device_status(item)}" for item in items))
        with section_frame("Apple 健康", "Apple 健康授权由 iPhone 上的 Executive Health Bridge 发起；数据缺失时不会补造。"):
            assigned = next((item for item in assignments if item.provider == "apple_health" and item.assignment_status != "DISABLED"), None)
            st.write("分配状态：" + ("已分配" if assigned else "尚未分配"))
            if apple_job and apple_job.status == "SUCCESS" and apple_job.created_by == "apple_health_bridge":
                st.write("最近同步：" + _fmt_dt(apple_job.completed_at))
                st.caption("Apple 健康授权已完成；此处仅代表已收到桥接同步，不代表全部数据类型均已授权。")
            else:
                st.write("最近同步：暂无已验证的 iPhone 同步")
                st.caption("请先在 iPhone 安装 Bridge 后连接 Apple 健康。后台同步由 iOS 系统调度，不承诺实时。")
            st.link_button("连接 Apple 健康", "executivehealthbridge://connect", width="content")
            st.caption("首次连接与后端配置请参阅项目内《Apple 健康接入说明》。")
        return
    with section_frame("隐私与授权", "健康资料仅用于已授权的健康管理与医疗协同。"):
        st.write("连接、共享或授权变更需由您确认，并保留记录。")
        st.caption("如需调整授权范围，请联系健康管理团队协助处理。")


def _render_client_health_overview(patient: Patient, ctx: dict[str, list[object]]) -> None:
    """The first of five health views; details expand here instead of routing away."""
    risk, reason, _ = _member_risk_state(patient.id, ctx)
    with SessionLocal() as session:
        baseline = HealthAssessmentService().latest_baseline(session, patient.id, include_draft=False)
    with section_frame("当前健康状态", "先看已经确认的结论；需要医学判断时由医生处理。"):
        st.markdown(risk_badge(risk), unsafe_allow_html=True)
        st.markdown(f"**{reason or '当前没有正式风险评估。'}**")
    with section_frame("健康基线", "基线汇总当前主要问题、管理重点和来源；详情在本页展开。"):
        if baseline is None:
            _empty_state("尚未建立健康基线", "上传最近体检报告并经健康管理团队确认后，可在这里形成健康基线。")
        else:
            st.markdown(f"**{baseline.summary}**")
            with st.expander("查看健康基线详情"):
                snapshot = baseline.baseline_json or {}
                for label, key in (("主要健康问题", "health_problems"), ("关键健康指标", "key_metrics"), ("当前管理重点", "management_focus")):
                    value = snapshot.get(key)
                    st.markdown(f"**{label}**")
                    if isinstance(value, list) and value:
                        st.write("；".join(str(item.get("title") or item.get("metric") or item) if isinstance(item, dict) else str(item) for item in value[:6]))
                    else:
                        st.caption("待补充")
                with SessionLocal() as evidence_session:
                    _render_evidence_action(_baseline_evidence_payload(evidence_session, patient.id, baseline), key_scope=f"client-health-baseline-{baseline.id}", client_view=True)
    with section_frame("当前主要问题", "只显示最需要持续关注的三项。"):
        problems = [item for item in ctx["problems"] if item.status != "CLOSED"]
        if problems:
            for item in problems[:3]:
                st.markdown(f"<div class='next-row'><div class='focus-title'>{html.escape(item.title)}</div><div class='focus-copy'>{html.escape(item.description or '正在由健康管理团队持续跟进。')}</div></div>", unsafe_allow_html=True)
        else:
            _empty_state("暂无持续管理问题", "后续已确认的健康问题会在这里显示。")


def _render_client_checkup_page(patient: Patient) -> None:
    """Report list and selected report live on the same second-level health page."""
    left, right = st.columns([1, 1.7], gap="large")
    with SessionLocal() as session:
        documents = list(session.scalars(select(Document).where(Document.patient_id == patient.id).order_by(Document.created_at.desc()).limit(20)))
    with left:
        with section_frame("体检报告", "上传、选择历史报告或比较变化。"):
            if not documents:
                _empty_state("暂无体检报告", "上传最近一次体检报告后，整理结果会保留在此列表。")
            for document in documents:
                label = f"{_source_display_name(document, '体检报告')} · {_fmt_dt(document.created_at)}"
                if secondary_action(label, key=f"client-checkup-select-{patient.id}-{document.id}"):
                    st.session_state[f"member-report-document-{patient.id}"] = str(document.id)
                    st.rerun()
    with right:
        render_member_report_upload(patient)


def _render_client_medical_archive(patient: Patient) -> None:
    """One medical record page with inline categories, never separate navigation."""
    section = st.radio("医疗档案内容", ["用药", "手术住院", "医生意见", "病史"], horizontal=True, label_visibility="collapsed", key=f"client-medical-section-{patient.id}")
    with SessionLocal() as session:
        if section == "用药":
            rows = list(session.scalars(select(MedicationPlan).where(MedicationPlan.patient_id == patient.id).order_by(MedicationPlan.created_at.desc()).limit(20)))
            if rows:
                st.dataframe(pd.DataFrame([{"药物": item.drug_name or "待补充", "剂量": " ".join(part for part in (item.dose, item.dose_unit) if part) or "待补充", "频率": item.frequency or "待补充", "状态": _label(item.status)} for item in rows]), hide_index=True, width="stretch")
            else:
                _empty_state("暂无已确认用药", "医生或健康管理师确认的用药记录会显示在这里。")
        elif section == "手术住院":
            rows = list(session.scalars(select(HealthEvent).where(HealthEvent.patient_id == patient.id, HealthEvent.event_type.in_(("surgery", "hospitalization"))).order_by(HealthEvent.start_at.desc()).limit(20)))
            if rows:
                selected_key = f"client-medical-event-selected-{patient.id}"
                if st.session_state.get(selected_key) not in {str(item.id) for item in rows}:
                    st.session_state[selected_key] = str(rows[0].id)
                left, right = st.columns([1, 1.25], gap="large")
                with left:
                    for item in rows:
                        if secondary_action(f"{_fmt_dt(item.start_at)} · {item.description or '重要医疗记录'}", key=f"client-medical-event-{item.id}"):
                            st.session_state[selected_key] = str(item.id)
                            st.rerun()
                selected = next(item for item in rows if str(item.id) == st.session_state[selected_key])
                with right:
                    with detail_panel("手术 / 住院详情", "仅展示已有正式记录；未记录的信息不会被推断。"):
                        st.markdown(f"**{selected.description or '重要医疗记录'}**")
                        st.caption(f"时间：{_fmt_dt(selected.start_at)} · 经确认医疗资料")
            else:
                _empty_state("暂无手术或住院记录", "已确认的重要医疗经历会在这里长期保留。")
        elif section == "医生意见":
            rows = list(session.scalars(select(DoctorReview).where(DoctorReview.patient_id == patient.id, DoctorReview.status == "CONFIRMED").order_by(DoctorReview.reviewed_at.desc()).limit(20)))
            if rows:
                for item in rows:
                    with st.expander(f"{_fmt_dt(item.reviewed_at)} · 医生反馈"):
                        st.write(item.opinion or "医生意见待补充")
                        if item.question_for_doctor:
                            st.caption("复核问题：" + item.question_for_doctor)
            else:
                _empty_state("暂无医生反馈", "完成医生复核后，确认可共享的反馈会在这里显示。")
        else:
            rows = list(session.scalars(select(HealthProblem).where(HealthProblem.patient_id == patient.id).order_by(HealthProblem.opened_at.desc()).limit(20)))
            if rows:
                for item in rows:
                    st.markdown(f"<div class='next-row'><div class='focus-title'>{html.escape(item.title)}</div><div class='focus-copy'>{html.escape(item.description or '已确认健康记录')}</div></div>", unsafe_allow_html=True)
            else:
                _empty_state("暂无健康史记录", "既往史、个人史和家族史需要人工确认后才会显示。")


def render_client_health_hub(patient: Patient, ctx: dict[str, list[object]]) -> None:
    """Four second-level health views; the longitudinal story is first-level."""
    _page_header("健康", "查看健康状态、数据、体检和医疗档案。", eyebrow="成员健康中心")
    key = f"client-health-view-{patient.id}"
    allowed = ["健康概览", "健康数据", "体检", "医疗档案"]
    if st.session_state.get(key) not in allowed:
        st.session_state[key] = "健康概览"
    view = st.radio("健康内容", allowed, horizontal=True, label_visibility="collapsed", key=key)
    if view == "健康概览":
        _render_client_health_overview(patient, ctx)
    elif view == "健康数据":
        render_health_data(patient.id)
    elif view == "体检":
        _render_client_checkup_page(patient)
    elif view == "医疗档案":
        _render_client_medical_archive(patient)


def render_member_client_view(patient: Patient, page: str = "首页") -> None:
    """Development preview only; authentication/authorization is intentionally not implied."""
    st.caption("成员健康中心")
    ctx = _member_summary_context(patient.id)
    if page == "首页":
        _render_client_home(patient, ctx)
    elif page in {"健康", "健康档案", "健康数据", "数据"}:
        render_client_health_hub(patient, ctx)
    elif page in {"历程", "健康历程"}:
        render_longitudinal_timeline(patient, key_scope="member-center-journey", client_view=True)
    elif page == "计划":
        _render_client_plan(patient, ctx)
    elif page in {"服务", "我的服务"}:
        _render_client_service(patient, ctx)
    else:
        _render_client_profile(patient)


def render_global_alert_workspace(members: list[Patient]) -> None:
    st.caption("健康管理师 · 异常处理")
    st.title("健康异常处理工作台")
    st.caption("处理顺序：核实数据 → 关联或创建健康问题 → 医生复核 → 管理方案与执行任务 → 随访 → 已关闭。")
    member = st.selectbox("选择成员", members, format_func=_member_display, key="alert-member")
    ctx = _context(member.id)
    render_alerts(member, ctx)


def render_global_doctor_workspace(members: list[Patient]) -> None:
    _page_header("内部医生", "只处理需要医学判断的复核事项。", eyebrow="医疗协同")
    member = st.selectbox("选择成员", members, format_func=_member_display, key="doctor-member")
    render_doctor_reviews(member, _member_doctor_context(member.id))


def render_external_doctor_workspace(members: list[Patient]) -> None:
    _page_header("外部医疗", "管理转诊、预约和外部反馈。", eyebrow="医疗协同")
    with SessionLocal() as session:
        referrals = list(session.scalars(select(ExternalReferral).order_by(ExternalReferral.created_at.desc()).limit(100)))
        reviews = list(session.scalars(select(DoctorReview).where(DoctorReview.status == "CONFIRMED").order_by(DoctorReview.reviewed_at.desc()).limit(100)))
    if referrals:
        counts = {status: sum(item.status == status for item in referrals) for status in ("PENDING", "SCHEDULED", "WAITING_FEEDBACK", "COMPLETED")}
        _status_strip(
            ("待安排", counts["PENDING"], "attention"),
            ("已预约", counts["SCHEDULED"], "action"),
            ("等待反馈", counts["WAITING_FEEDBACK"], "neutral"),
            ("已完成", counts["COMPLETED"], "neutral"),
        )
        filters = {"全部": set(), "待安排": {"PENDING"}, "已预约": {"SCHEDULED"}, "等待反馈": {"WAITING_FEEDBACK"}, "已完成": {"COMPLETED"}}
        selected_filter = st.radio("外部医疗状态", list(filters), horizontal=True, label_visibility="collapsed", key="external-medical-filter")
        visible = referrals if not filters[selected_filter] else [item for item in referrals if item.status in filters[selected_filter]]
        if visible:
            selected_key = "external-medical-selected"
            if st.session_state.get(selected_key) not in {str(item.id) for item in visible}:
                st.session_state[selected_key] = str(visible[0].id)
            member_map = {item.id: item for item in members}
            left, right = st.columns([1, 1.25], gap="large")
            with left:
                with section_frame("外部医疗列表", "选择一项后，在右侧查看安排与反馈。"):
                    for referral in visible[:20]:
                        member = member_map.get(referral.patient_id)
                        if secondary_action(f"{_member_display(member)} · {referral.specialty} · {_label(referral.status)}", key=f"external-medical-select-{referral.id}"):
                            st.session_state[selected_key] = str(referral.id)
                            st.rerun()
            selected = next(item for item in visible if str(item.id) == st.session_state[selected_key])
            with right:
                member = member_map.get(selected.patient_id)
                with detail_panel("外部医疗详情", "外部预约、就医和反馈均由人工协调确认。"):
                    st.markdown(f"**{_member_display(member)} · {selected.specialty}**")
                    st.write(selected.reason or "转诊原因待补充")
                    st.caption(f"当前状态：{_label(selected.status)} · 预约：{_fmt_dt(selected.appointment_at)}")
                    st.write("外部机构：" + (selected.organization or "待确认"))
                    st.write("外部医生：" + (selected.doctor_name or "待确认"))
                    st.write(selected.feedback or "等待外部反馈。")
        else:
            _empty_state("此状态下暂无外部医疗事项", "调整状态筛选，或登记新的外部医疗协同。")
    else:
        _empty_state("暂无外部医生协同", "内部医生建议转诊后，会在这里记录预约和外部反馈。")
    with st.expander("登记外部医生协同"):
        with st.form("external-referral-form"):
            member = st.selectbox("成员", members, format_func=_member_display)
            available_reviews = [item for item in reviews if item.patient_id == member.id]
            review = st.selectbox("关联内部医生复核（可选）", [None, *available_reviews], format_func=lambda item: "不关联" if item is None else f"{item.department} · {_fmt_dt(item.reviewed_at)}")
            specialty = st.text_input("外部专科 *")
            reason = st.text_area("转诊原因 *")
            question = st.text_area("希望外部医生确认什么 *")
            organization = st.text_input("外部机构（可选）")
            doctor_name = st.text_input("外部医生（可选）")
            if st.form_submit_button("登记外部协同"):
                if not specialty.strip() or not reason.strip() or not question.strip(): st.error("请填写专科、转诊原因和协同问题。")
                else:
                    with SessionLocal() as session:
                        session.add(ExternalReferral(patient_id=member.id, doctor_review_id=review.id if review else None, specialty=specialty, reason=reason, question=question, organization=organization or None, doctor_name=doctor_name or None, status="PENDING"))
                        session.commit()
                    st.success("已登记外部医生协同；请由人工完成预约与反馈录入。")
                    st.rerun()


def render_demo_story(members: list[Patient]) -> None:
    st.caption("演示数据 · 管理故事")
    st.title("Demo Executive A：持续健康与代谢管理")
    st.markdown("### 健康评估 → 重点健康问题 → 90天健康管理 → 执行任务 → 阶段效果评估 → 稳定管理")
    st.markdown("同时保留医疗子闭环：连续血压异常 → 健康异常 → 管理师核实 → 医生复核 → 管理方案 → 随访 → 已关闭")
    st.caption("所有数据仅用于演示；系统不做自动诊断、处方、停药或剂量调整。")
    demo = next((member for member in members if member.external_id == "demo-executive-001"), members[0])
    st.button("进入成员的完整处理记录", type="primary", on_click=_open_member, args=(demo.id,))
    render_timeline(demo, _context(demo.id))


def render_data_gateway(members: list[Patient]) -> None:
    _page_header("数据接入与设备", "查看日常健康设备、医疗监测设备和成员设备分配。", eyebrow="平台工具")
    jobs, review_count = _device_overview_snapshot()
    st.caption(f"需要人工复核的数据：{review_count} 条")
    latest_by_source: dict[str, IngestionJob] = {}
    for job in jobs:
        latest_by_source.setdefault(job.source_system, job)
    source_rows: list[dict[str, str]] = []
    for source in ("apple_health", "mock_oura", "mock_yuwell", "mock_cgm"):
        job = latest_by_source.get(source)
        if source == "apple_health":
            simple_status = "等待真机验证" if not job else ("已收到桥接同步" if job.status == "SUCCESS" and job.created_by == "apple_health_bridge" else "演示同步" if job.created_by == "mock_apple_health" else "需要处理")
        elif source == "mock_oura":
            simple_status = "演示 / 未连接" if not job else "演示数据"
        else:
            simple_status = "尚未连接" if not job else ("演示数据" if source.startswith("mock_") else "已连接")
        source_rows.append({"name": display_provider(source), "status": simple_status, "last": _fmt_dt(job.completed_at) if job else "暂无同步", "category": "日常健康" if source in {"apple_health", "mock_oura"} else "医疗监测"})
    for title in ("日常健康设备", "医疗监测设备"):
        _section_header(title)
        rows = [row for row in source_rows if row["category"] == ("日常健康" if title == "日常健康设备" else "医疗监测")]
        columns = st.columns(max(1, len(rows)))
        for column, row in zip(columns, rows):
            with column:
                with st.container(border=True):
                    st.markdown(f"**{row['name']}**")
                    st.markdown(_status_pill(row["status"]), unsafe_allow_html=True)
                    st.caption(f"最近同步：{row['last']}")
                    if row["name"] == "Apple Health":
                        verified = "已收到桥接同步" if row["status"] == "已收到桥接同步" else "未完成"
                        st.caption("后端接收：已就绪 · iOS Bridge：源码已就绪 · 真机验证：" + verified)
    render_member_device_assignments(members)
    device_view = st.radio("数据设备功能", ["设备概览", "数据复核", "上传健康资料"], horizontal=True, label_visibility="collapsed", key="device-workspace")
    if device_view == "设备概览":
        st.caption("选择“数据复核”或“上传健康资料”后才会加载对应队列与表单。")
        return
    with SessionLocal() as session:
        records = list(session.scalars(select(RawIngestionRecord).where(RawIngestionRecord.status.in_(["SUSPECT", "INVALID", "UNMATCHED", "WAITING_REVIEW"])).order_by(RawIngestionRecord.created_at.desc()).limit(50)))
    if device_view == "数据复核":
        _render_data_review_queue(records)
        return
    _render_data_upload_workspace(members, jobs, records)


def _device_overview_snapshot() -> tuple[list[IngestionJob], int]:
    """Bounded overview data only; no sync, provider ping, or health-data scan."""
    with SessionLocal() as session:
        jobs = list(session.scalars(select(IngestionJob).order_by(IngestionJob.started_at.desc()).limit(20)))
        review_count = int(session.scalar(select(func.count(RawIngestionRecord.id)).where(RawIngestionRecord.status.in_(["SUSPECT", "INVALID", "UNMATCHED", "WAITING_REVIEW"]))) or 0)
    return jobs, review_count


def _render_data_upload_workspace(members: list[Patient], jobs: list[IngestionJob], records: list[RawIngestionRecord]) -> None:
    st.subheader("需要处理的数据问题")
    if not records:
        st.success("数据同步正常，当前没有需要人工处理的数据。")
    else:
        st.info(f"有 {len(records)} 条数据需要人工确认或绑定成员。请在下方处理。")
    member = st.selectbox("接入成员", members, format_func=_member_display, key="gateway-member")
    render_report_upload(member, key_prefix="gateway")
    st.subheader("上传健康资料")
    st.caption("支持 CSV、Excel 和 PDF 健康资料。导入前可先预览，系统不会自动作出医疗结论。")
    action, upload = st.columns(2)
    with upload:
        uploaded = st.file_uploader("选择文件", type=["csv", "xlsx", "pdf"], key="gateway-file")
        dry_run = st.checkbox("仅预览，不写入数据库", value=True)
        if uploaded and st.button("上传文件", type="primary"):
            provider = "excel" if uploaded.name.lower().endswith(".xlsx") else "pdf" if uploaded.name.lower().endswith(".pdf") else "csv"
            mapping = {"timestamp": "observed_at", "高压": "systolic_bp", "低压": "diastolic_bp", "heart_rate": "heart_rate"}
            with SessionLocal() as session:
                if provider == "pdf":
                    job = IngestionJob(source_system="pdf", source_type="file", patient_id=member.id, status="PARTIAL_SUCCESS", records_received=1, records_valid=0, records_invalid=0, records_duplicate=0, records_created=0, records_updated=0, error_count=0, created_by="health_manager_upload", completed_at=datetime.now(TOKYO_TIMEZONE))
                    session.add(job); session.flush()
                    session.add(Document(patient_id=member.id, document_type="health_check_pdf", title=uploaded.name, storage_reference=f"gateway-upload://{job.id}/{uploaded.name}", source="pdf_gateway", status="WAITING_REVIEW"))
                    session.add(RawIngestionRecord(job_id=job.id, patient_id=member.id, source_system="pdf", source_type="file", source_record_id=uploaded.name, payload_json={"filename": uploaded.name, "bytes": uploaded.size}, adapter_name="PDFParserInterface", adapter_version="v1", status="WAITING_REVIEW", normalization_json={"message": "等待人工提取"}))
                    session.commit()
                    st.success("PDF 健康报告已登记，正在等待人工复核。")
                    return
                result = ingest(session, provider, uploaded.getvalue(), member_id=member.id, mapping=mapping, dry_run=dry_run, created_by="health_manager_upload")
                session.commit()
            st.success(f"文件导入{_label(result.status)}：新增 {result.created} 条健康数据。")
    if TECHNICAL_DETAILS_ENABLED:
        with action:
            with st.expander("高级信息"):
                st.caption("仅用于合成演示环境。")
                if st.button("同步演示血压设备数据"):
                    with SessionLocal() as session:
                        result = ingest(session, "mock_yuwell", {"user_id": "YUWELL-DEMO-001", "device_id": "BP-DEMO-01", "measure_time": "2026-08-15T07:30:00+09:00", "sys": 148, "dia": 94, "pulse": 76}, member_id=member.id, created_by="health_manager_demo")
                        session.commit()
                    st.success(f"同步{_label(result.status)}：新增 {result.created} 条。")
            if st.button("同步演示 Oura 数据"):
                with SessionLocal() as session:
                    result = ingest(session, "mock_oura", {"user_id": "OURA-DEMO-001", "day": "2026-08-14", "total_sleep_duration": 21120, "score": 72, "resting_heart_rate": 61}, member_id=member.id, created_by="health_manager_demo")
                    session.commit()
                st.success(f"同步{_label(result.status)}：新增 {result.created} 条健康数据。")
            if st.button("同步演示 Apple Health 数据"):
                payload = {"samples": [{"sample_id": "apple-steps-ui-001", "type": "stepCount", "value": 12032, "unit": "count", "start_date": "2026-08-14T00:00:00+09:00", "end_date": "2026-08-14T23:59:00+09:00"}, {"sample_id": "apple-weight-ui-001", "type": "bodyMass", "value": 81, "unit": "kg", "start_date": "2026-08-14T07:30:00+09:00", "end_date": "2026-08-14T07:30:00+09:00"}]}
                with SessionLocal() as session:
                    result = ingest(session, "apple_health", payload, member_id=member.id, created_by="mock_apple_health")
                    session.commit()
                st.success(f"Apple Health 同步{_label(result.status)}：新增 {result.created} 条，已拦截重复 {result.duplicates} 条。")
    if TECHNICAL_DETAILS_ENABLED:
        with st.expander("高级信息：备用导入流程"):
            st.markdown("#### 文件数据导入")
            uploaded = st.file_uploader("选择演示用 CSV、Excel 或 PDF 文件", type=["csv", "xlsx", "pdf"], key="gateway-file-advanced")
            dry_run = st.checkbox("仅预览，不写入数据库", value=True, key="gateway-dry-run-advanced")
            if uploaded and st.button("验证并导入", key="gateway-import-advanced"):
                provider = "excel" if uploaded.name.lower().endswith(".xlsx") else "pdf" if uploaded.name.lower().endswith(".pdf") else "csv"
                mapping = {"timestamp": "observed_at", "高压": "systolic_bp", "低压": "diastolic_bp", "heart_rate": "heart_rate"}
                with SessionLocal() as session:
                    if provider == "pdf":
                        job = IngestionJob(source_system="pdf", source_type="file", patient_id=member.id, status="PARTIAL_SUCCESS", records_received=1, records_valid=0, records_invalid=0, records_duplicate=0, records_created=0, records_updated=0, error_count=0, created_by="health_manager_upload", completed_at=datetime.now(TOKYO_TIMEZONE))
                        session.add(job); session.flush()
                        session.add(Document(patient_id=member.id, document_type="health_check_pdf", title=uploaded.name, storage_reference=f"gateway-upload://{job.id}/{uploaded.name}", source="pdf_gateway", status="WAITING_REVIEW"))
                        session.add(RawIngestionRecord(job_id=job.id, patient_id=member.id, source_system="pdf", source_type="file", source_record_id=uploaded.name, payload_json={"filename": uploaded.name, "bytes": uploaded.size}, adapter_name="PDFParserInterface", adapter_version="v1", status="WAITING_REVIEW", normalization_json={"message": "DEMO / RULE-BASED EXTRACTION NOT RUN; human review required."}))
                        session.commit()
                        st.success("PDF 健康报告已登记，正在等待人工复核；系统不会自动做医学结论。")
                        return
                    result = ingest(session, provider, uploaded.getvalue(), member_id=member.id, mapping=mapping, dry_run=dry_run, created_by="health_manager_upload")
                    session.commit()
                st.success(f"文件导入{_label(result.status)}：收到 {result.received} 条，有效 {result.valid} 条，新增 {result.created} 条，重复 {result.duplicates} 条，无效 {result.invalid} 条。")
    if TECHNICAL_DETAILS_ENABLED:
        with st.expander("高级信息：数据同步记录"):
            if jobs:
                member_names = {member.id: (member.display_name or member.external_id or "未命名成员") for member in members}
                st.dataframe(pd.DataFrame([{
                    "时间": _fmt_dt(job.started_at), "数据来源": display_provider(job.source_system),
                    "成员": member_names.get(job.patient_id, "未匹配成员"), "状态": _label(job.status),
                    "接收记录": job.records_received, "成功创建": job.records_created,
                    "重复数据": job.records_duplicate, "错误记录": job.error_count,
                } for job in jobs[:30]]), hide_index=True, width="stretch")
    _render_data_review_queue(records)


def _render_data_review_queue(records: list[RawIngestionRecord]) -> None:
    st.subheader("数据复核队列")
    if not records:
        st.success("当前没有待核实、无效、未匹配或等待人工确认的数据。")
    for record in records[:20]:
        with st.expander(f"{_label(record.status)} · {display_provider(record.source_system)} · 需要人工确认"):
            st.write(record.error_message or "需要人工确认")
            if TECHNICAL_DETAILS_ENABLED:
                with st.expander("高级信息"):
                    st.caption("技术详情、原始数据和解析结果仅供授权人员核对。")
                    with st.expander("查看原始技术信息"):
                        st.json({"处理信息": record.normalization_json, "原始记录": record.payload_json})
            if record.patient_id and record.status in {"SUSPECT", "INVALID"}:
                new_value = st.text_input("更正数值", key=f"correct-value-{record.id}")
                reason = st.text_input("更正原因", key=f"correct-reason-{record.id}")
                if st.button("人工更正", key=f"correct-{record.id}") and new_value and reason:
                    with SessionLocal() as session:
                        stored = session.get(RawIngestionRecord, record.id)
                        manually_correct_record(session, stored, new_value, reason, "health_manager")
                        session.commit()
                    st.success("已保留原始记录，并创建已人工修正的健康数据。")
                    st.rerun()


def render_member_device_assignments(members: list[Patient]) -> None:
    st.subheader("成员设备分配")
    st.caption("“已分配”与“已连接”是不同状态；演示设备不会被显示为真实连接。")
    member = st.selectbox("选择成员以管理设备", members, format_func=_member_display, key="device-assignment-member")
    with SessionLocal() as session:
        assignments = list(session.scalars(select(MemberDeviceAssignment).where(MemberDeviceAssignment.patient_id == member.id).order_by(MemberDeviceAssignment.assigned_at.desc())))
    grouped = {"日常健康": [item for item in assignments if item.device_category == "WELLNESS"], "医疗监测": [item for item in assignments if item.device_category != "WELLNESS"]}
    available = {
        "日常健康": ("apple_health", "mock_oura"),
        "医疗监测": ("mock_yuwell", "glucose_meter_interface", "mock_cgm"),
    }
    if not assignments:
        _empty_state("尚未分配设备", "为成员分配设备后，可在此查看连接状态。")
    for title, items in grouped.items():
        st.markdown(f"**{title}**")
        assigned_providers = {item.provider for item in items if item.assignment_status != "DISABLED"}
        for item in items:
            status = {"CONNECTED": "已连接", "SYNCING": "同步中", "MOCK": "演示", "PENDING": "待连接", "DISABLED": "停用"}.get(item.connection_status, "待连接")
            st.caption(f"{display_provider(item.provider)} · 已分配 · {status}")
        for provider in available[title]:
            if provider not in assigned_providers:
                st.caption(f"{display_provider(provider)} · 未分配")
    with st.expander("添加或更新设备分配"):
        with st.form(f"device-assignment-{member.id}"):
            provider = st.selectbox("设备", ["apple_health", "mock_oura", "mock_yuwell", "mock_cgm", "glucose_meter_interface"], format_func=display_provider)
            category = st.radio("设备类别", ["WELLNESS", "MEDICAL_MONITOR"], format_func=lambda item: "日常健康" if item == "WELLNESS" else "医疗监测", horizontal=True)
            connection = st.selectbox("连接状态", ["PENDING", "CONNECTED", "SYNCING", "MOCK", "DISABLED"], format_func=lambda item: {"PENDING": "已分配，尚未连接", "CONNECTED": "已连接", "SYNCING": "同步中", "MOCK": "演示数据", "DISABLED": "已停用"}[item])
            notes = st.text_input("说明（可选）")
            if st.form_submit_button("保存设备分配"):
                with SessionLocal() as session:
                    existing = session.scalar(select(MemberDeviceAssignment).where(MemberDeviceAssignment.patient_id == member.id, MemberDeviceAssignment.provider == provider, MemberDeviceAssignment.assignment_status != "DISABLED").order_by(MemberDeviceAssignment.assigned_at.desc()))
                    if existing:
                        existing.connection_status, existing.device_category, existing.notes = connection, category, notes or None
                    else:
                        session.add(MemberDeviceAssignment(patient_id=member.id, provider=provider, device_category=category, assignment_status="ASSIGNED" if connection != "DISABLED" else "DISABLED", connection_status=connection, assigned_by="健康管理师", notes=notes or None))
                    session.commit()
                st.success("成员设备分配已保存。")
                st.rerun()


def render_risk_rules() -> None:
    st.title("风险规则")
    st.caption("规则用于健康监护与风险分流，不替代医生诊断或急救机构判断。仅已审核且启用的规则可进入风险引擎。")
    with SessionLocal() as session:
        rules=list(session.scalars(select(RiskRule).order_by(RiskRule.updated_at.desc())))
        management_rules = list(session.scalars(select(ManagementRule).order_by(ManagementRule.updated_at.desc())))
    st.dataframe(pd.DataFrame([{"规则名称":rule.name,"适用数据":"日常健康设备" if rule.applicable_device_class=="WELLNESS" else "医疗监测设备","风险等级":_risk_text(rule.risk_level),"状态":_label(rule.review_status),"版本":rule.version or "未记录","来源":rule.source_reference or "待补充"} for rule in rules]),hide_index=True,width="stretch")
    st.info("当前仅包含明确标记为演示工作流规则的内容，不作为真实医疗临床阈值。")
    st.subheader("生活方式管理规则")
    st.caption("步数、运动与睡眠使用独立的健康管理信号，不生成医疗风险事件，也不会自动升级医生。只有已审核且启用的规则才会在新数据写入时运行。")
    if management_rules:
        st.dataframe(pd.DataFrame([{
            "规则名称": rule.name, "指标": _metric_display_name(rule.canonical_code),
            "状态": _label(rule.review_status),
            "启用": "是" if rule.is_active else "否", "路由": ROUTE_LABELS.get(rule.recommended_route, "待确认处理路径"),
            "版本": rule.version or "未记录", "来源": rule.source_reference or "待补充",
        } for rule in management_rules]), hide_index=True, width="stretch")
    else:
        st.caption("尚无已配置的生活方式管理规则。系统不会因为步数或睡眠数据自动创建医疗风险。")


def main() -> None:
    started = perf_counter()
    # Must be first: it writes widget values only before their controls exist.
    apply_pending_navigation()
    _navigation_stage("bootstrap:style", _inject_style)
    if PORTFOLIO_DEMO_ENABLED and not st.session_state.get("portfolio-demo-landing-dismissed"):
        _render_portfolio_landing()
        return
    surface = _navigation_stage("surface", _render_surface_switcher)
    if surface == "成员健康中心":
        page = _navigation_stage("member-center-sidebar", _render_member_center_navigation)
        members = _navigation_stage("member-center-member", _members)
        if not members:
            _empty_state("暂无演示成员", "完成演示数据初始化后即可查看成员健康中心。")
            return
        patient = st.sidebar.selectbox("查看成员", members, format_func=_member_display, key="member-center-member-select")
        _render_timed("成员健康中心", lambda: render_member_client_view(patient, page))
        return
    page = _navigation_stage("sidebar", _render_sidebar_navigation)
    focused = st.session_state.get("focused_member_id")
    if page == "今日":
        _render_timed("今日", render_manager_dashboard)
        if NAVIGATION_PROFILE_ENABLED: LOGGER.warning("[PERF] total %.1f ms", (perf_counter() - started) * 1000)
        return
    if page == "医疗协同":
        _render_timed("医疗协同", render_collaboration_workspace)
        if NAVIGATION_PROFILE_ENABLED: LOGGER.warning("[PERF] total %.1f ms", (perf_counter() - started) * 1000)
        return
    if page == "服务运营":
        _render_timed("服务运营", render_service_operations_workspace)
        if NAVIGATION_PROFILE_ENABLED: LOGGER.warning("[PERF] total %.1f ms", (perf_counter() - started) * 1000)
        return
    if page == "更多":
        _render_timed("更多", render_more_workspace)
        if NAVIGATION_PROFILE_ENABLED: LOGGER.warning("[PERF] total %.1f ms", (perf_counter() - started) * 1000)
        return
    # Member records are unnecessary for the default workbench and are loaded only when needed.
    members = _navigation_stage("member list", _members)
    if not members:
        st.title("企业高管健康运营中心")
        st.warning("尚未初始化演示数据。请先完成数据库迁移和演示数据初始化。")
        if NAVIGATION_PROFILE_ENABLED: LOGGER.warning("[PERF] total %.1f ms", (perf_counter() - started) * 1000)
        return
    if page == "成员" and focused:
        patient = next((item for item in members if str(item.id) == focused), None)
        if patient:
            _render_timed("成员详情", lambda: render_member_detail(patient))
            if NAVIGATION_PROFILE_ENABLED: LOGGER.warning("[PERF] total %.1f ms", (perf_counter() - started) * 1000)
            return
        st.session_state.pop("focused_member_id", None)
    if page == "成员":
        _render_timed("成员", lambda: render_members_workspace(members))
    else:
        _render_timed("更多", render_more_workspace)
    if NAVIGATION_PROFILE_ENABLED:
        LOGGER.warning("[PERF] total %.1f ms", (perf_counter() - started) * 1000)


if __name__ == "__main__":
    main()
