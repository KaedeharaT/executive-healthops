"""Universal, local-first health-check report parsing with mandatory human review.

The parser deliberately produces evidence-backed *candidates*, never diagnoses
or risk events.  A hospital-specific adapter may improve extraction later, but
all adapters must return the same candidate shape.
"""

from __future__ import annotations

import csv
import hashlib
import io
import os
import re
import logging
import time as clock
from dataclasses import dataclass, field
from datetime import date, datetime, time, timezone
from decimal import Decimal, InvalidOperation
from pathlib import Path
from typing import Callable, Protocol
from uuid import UUID, uuid4

import pandas as pd
from openpyxl.utils import get_column_letter
from docx import Document as DocxDocument
from pypdf import PdfReader
from sqlalchemy import select
from sqlalchemy.orm import Session

from executive_health_ai.integrations.codes import ObservationCode, canonical_code
from executive_health_ai.integrations.normalization import normalize_unit, quality_for
from executive_health_ai.models import AuditLog, Document, HealthProblem, Observation, ReportExtractionCandidate, ReportExtractionRun, Task
from executive_health_ai.models.base import utc_now
from executive_health_ai.llm.prompts.health_report import FINDING_EXTRACTION_SYSTEM_PROMPT, finding_extraction_prompt
from executive_health_ai.llm.local_llm_client import LocalLLMClient, LocalLLMHealth, LocalLLMUnavailable, sanitize_for_llm
from executive_health_ai.services.risk_triage import RiskEvaluationService
from executive_health_ai.services.longitudinal import ManagementRoutingService

PARSER_VERSION = "universal-report-parser-v2"
CANONICAL_REGISTRY_VERSION = "registry-v1"
ALLOW_EXTERNAL_PHI_LLM = os.getenv("ALLOW_EXTERNAL_PHI_LLM", "false").lower() == "true"
MAX_REPORT_UPLOAD_BYTES = max(1, int(os.getenv("MAX_REPORT_UPLOAD_BYTES", str(25 * 1024 * 1024))))
MISSING_VALUES = {"", "-", "—", "/", "未测", "无法检测", "n/a", "na"}
logger = logging.getLogger(__name__)


@dataclass(frozen=True)
class DocumentPreflight:
    file_type: str
    mime_type: str
    file_hash: str
    page_count: int | None
    has_text_layer: bool
    is_probably_scanned: bool
    detected_hospital: str | None
    detected_report_type: str | None
    detected_report_date: date | None
    detected_language: str
    possible_template: str | None


@dataclass(frozen=True)
class ExtractedPage:
    page_number: int
    text: str
    # A reconstructed paragraph can begin on one page and finish on the next.
    # Keep this transient provenance for candidate evidence without changing the
    # raw document or inventing text.
    page_span: tuple[int, ...] = ()
    # File-specific source metadata is carried into Candidate.  This is
    # provenance only: parsing and medical confirmation semantics are unchanged.
    source_metadata: dict[str, object] = field(default_factory=dict)
    line_locations: dict[str, dict[str, object]] = field(default_factory=dict)


@dataclass(frozen=True)
class CandidateDraft:
    candidate_type: str
    raw_name: str | None
    raw_value: str | None
    canonical_code: str | None
    normalized_value: str | None
    unit: str | None
    reference_range: str | None
    abnormal_flag: str | None
    summary: str | None
    structured_data: dict[str, object]
    confidence: str
    extraction_method: str
    source_page: int
    source_section: str
    evidence_text: str


@dataclass(frozen=True)
class ReportParseProgress:
    """A transient, frontend-agnostic update for one explicit parse command."""

    stage: str
    message: str
    current: int | None = None
    total: int | None = None
    section_name: str | None = None
    rule_candidate_count: int | None = None
    call_duration_ms: int | None = None
    elapsed_ms: int | None = None
    candidate_count: int | None = None
    finding_count: int | None = None
    followup_count: int | None = None
    llm_call_count: int | None = None
    llm_success_count: int | None = None
    llm_failure_count: int | None = None


ProgressCallback = Callable[[ReportParseProgress], None]


class OCRProvider(Protocol):
    """Optional provider contract. No OCR provider is bundled in this release."""

    def extract(self, content: bytes, file_type: str) -> list[ExtractedPage]: ...


class BaseReportAdapter(Protocol):
    adapter_name: str
    adapter_version: str

    def can_handle(self, context: DocumentPreflight) -> bool: ...
    def normalize_sections(self, pages: list[ExtractedPage]) -> list[ExtractedPage]: ...
    def extract_candidates(self, pages: list[ExtractedPage]) -> list[CandidateDraft]: ...


def _file_type(filename: str) -> str:
    suffix = Path(filename).suffix.lower().lstrip(".")
    return {"jpg": "image", "jpeg": "image", "png": "image", "xlsx": "xlsx", "csv": "csv", "docx": "docx", "txt": "txt", "pdf": "pdf"}.get(suffix, "unknown")


def _mime_type(file_type: str) -> str:
    return {"pdf": "application/pdf", "image": "image/*", "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", "csv": "text/csv", "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document", "txt": "text/plain"}.get(file_type, "application/octet-stream")


def _safe_upload_filename(filename: str) -> str:
    """Keep the displayed name, but never use an uploaded path as a path."""
    clean = Path(filename.replace("\\", "/")).name.strip()
    if not clean or clean in {".", ".."}:
        raise ValueError("报告文件名无效。")
    return clean


def _validate_report_upload(filename: str, content: bytes) -> tuple[str, str]:
    """Reject malformed or unsupported intake before any file is persisted."""
    safe_name = _safe_upload_filename(filename)
    file_type = _file_type(safe_name)
    if file_type == "unknown":
        raise ValueError("当前仅支持 PDF、图片、DOCX、XLSX、CSV 或 TXT 格式的报告。")
    if not content:
        raise ValueError("报告文件为空，请重新选择完整文件。")
    if len(content) > MAX_REPORT_UPLOAD_BYTES:
        raise ValueError(f"报告文件超过当前 {MAX_REPORT_UPLOAD_BYTES // (1024 * 1024)} MB 上传限制。")
    if file_type == "pdf" and not content.startswith(b"%PDF-"):
        raise ValueError("文件扩展名为 PDF，但文件内容不是有效 PDF。")
    if file_type == "image" and not (content.startswith(b"\x89PNG\r\n\x1a\n") or content.startswith(b"\xff\xd8\xff")):
        raise ValueError("图片报告格式无法识别，请上传 PNG 或 JPEG 文件。")
    if file_type in {"docx", "xlsx"} and not content.startswith(b"PK"):
        raise ValueError("Office 报告文件格式无法识别，请重新导出后上传。")
    return safe_name, file_type


def _date_from_text(text: str) -> date | None:
    matched = re.search(r"(20\d{2})[年./-](\d{1,2})[月./-](\d{1,2})", text)
    if not matched:
        return None
    try:
        return date(*map(int, matched.groups()))
    except ValueError:
        return None


def _hospital_from_text(text: str) -> str | None:
    known = ("四川大学华西医院", "北京协和医院", "瑞金医院", "爱康", "美年")
    for name in known:
        if name in text:
            return name
    matched = re.search(r"([^\n]{2,40}(?:医院|体检中心|健康管理中心))", text)
    return matched.group(1).strip() if matched else None


class DocumentPreflightService:
    def inspect(self, filename: str, content: bytes) -> tuple[DocumentPreflight, list[ExtractedPage]]:
        file_type = _file_type(filename)
        pages = self._extract_pages(file_type, content)
        joined = "\n".join(page.text for page in pages)
        has_text = bool(joined.strip())
        is_scanned = file_type in {"pdf", "image"} and not has_text
        return DocumentPreflight(
            file_type=file_type, mime_type=_mime_type(file_type), file_hash=hashlib.sha256(content).hexdigest(),
            page_count=len(pages) if file_type == "pdf" else 1, has_text_layer=has_text,
            is_probably_scanned=is_scanned, detected_hospital=_hospital_from_text(joined),
            detected_report_type="健康检查报告" if any(token in joined for token in ("体检", "健康检查", "检查报告")) else None,
            detected_report_date=_date_from_text(joined), detected_language="zh-CN" if re.search(r"[\u4e00-\u9fff]", joined) else "unknown",
            possible_template=_hospital_from_text(joined),
        ), pages

    def _extract_pages(self, file_type: str, content: bytes) -> list[ExtractedPage]:
        if file_type == "pdf":
            reader = PdfReader(io.BytesIO(content))
            return [ExtractedPage(index + 1, page.extract_text() or "") for index, page in enumerate(reader.pages)]
        if file_type == "docx":
            document = DocxDocument(io.BytesIO(content))
            rows = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
            for table in document.tables:
                rows.extend("\t".join(cell.text.strip() for cell in row.cells) for row in table.rows)
            return [ExtractedPage(1, "\n".join(rows))]
        if file_type == "csv":
            frame = pd.read_csv(io.BytesIO(content), dtype=str)
            return [self._tabular_page(frame, page_number=1, source_metadata={"evidence_type": "TABLE", "source_format": "CSV"})]
        if file_type == "xlsx":
            workbook = pd.read_excel(io.BytesIO(content), sheet_name=None, dtype=str)
            return [
                self._tabular_page(
                    frame, page_number=index + 1,
                    source_metadata={"evidence_type": "CELL_RANGE", "source_format": "XLSX", "sheet_name": str(sheet_name)},
                )
                for index, (sheet_name, frame) in enumerate(workbook.items())
            ] or [ExtractedPage(1, "")]
        if file_type == "txt":
            return [ExtractedPage(1, content.decode("utf-8-sig", errors="replace"))]
        return [ExtractedPage(1, "")]

    @staticmethod
    def _tabular_page(frame: pd.DataFrame, *, page_number: int, source_metadata: dict[str, object]) -> ExtractedPage:
        """Retain only row-level coordinates needed to explain parsed values."""
        columns = ["" if pd.isna(value) else str(value) for value in frame.columns]
        header = "\t".join(columns)
        lines = [header]
        locations: dict[str, dict[str, object]] = {}
        last_column = get_column_letter(max(1, len(columns)))
        if header.strip():
            locations[header] = {**source_metadata, "row_index": 1, "cell_range": f"A1:{last_column}1", "table_header": header}
        for row_index, row in enumerate(frame.itertuples(index=False, name=None), start=2):
            line = "\t".join("" if pd.isna(value) else str(value) for value in row)
            lines.append(line)
            locations[line] = {
                **source_metadata, "row_index": row_index, "cell_range": f"A{row_index}:{last_column}{row_index}",
                "table_header": header, "table_row": line,
            }
        return ExtractedPage(page_number, "\n".join(lines), source_metadata=source_metadata, line_locations=locations)


class ReportSectionClassifier:
    _sections = (("LAB", ("检验", "生化", "血脂", "血常规")), ("VITALS", ("一般检查", "血压", "身高", "体重")), ("IMAGING", ("彩超", "ct", "影像", "结节", "超声")), ("ECG", ("心电图",)), ("PULMONARY_FUNCTION", ("肺功能",)), ("BODY_COMPOSITION", ("人体成分", "体脂", "骨骼肌")), ("RECOMMENDATION", ("建议", "复查", "门诊")), ("SUMMARY", ("小结", "总结")))

    def classify(self, text: str) -> str:
        lowered = text.lower()
        for section, terms in self._sections:
            if any(term in lowered for term in terms):
                return section
        return "OTHER"


class ReportTextReconstructor:
    """Restore conservative paragraph continuity from PDF extraction lines.

    PDF text extraction frequently turns visual wrapping into hard newlines.
    We only join lines when the preceding line has an explicit unfinished
    construction.  Tables, headings, and complete sentences remain separate.
    This layer uses only adjacent source text; it never fills missing content.
    """

    _unfinished_end = re.compile(r"(?:[，、：；,;:]|(?:若|如|建议|需|请|考虑|复查|就诊|诊)\s*)$")
    _fragment_start = re.compile(r"^(?:定有临床意义|进一步|并随访|治[。；;]?|时[，,]?|到.+?(?:内科|外科|专科)诊$)")

    @classmethod
    def reconstruct(cls, pages: list[ExtractedPage]) -> list[ExtractedPage]:
        rebuilt: dict[int, list[str]] = {page.page_number: [] for page in pages}
        spans: dict[int, set[int]] = {page.page_number: {page.page_number} for page in pages}
        pending: tuple[int, str] | None = None

        def flush() -> None:
            nonlocal pending
            if pending is None:
                return
            page_number, text = pending
            rebuilt.setdefault(page_number, []).append(text)
            spans.setdefault(page_number, {page_number})
            pending = None

        for page in pages:
            for raw_line in page.text.splitlines():
                line = raw_line.strip()
                if not line:
                    flush()
                    continue
                if pending is not None and cls._should_continue(pending[1], line):
                    start_page, previous = pending
                    pending = (start_page, cls._join(previous, line))
                    spans.setdefault(start_page, {start_page}).add(page.page_number)
                    continue
                flush()
                pending = (page.page_number, line)
            # Do not flush an explicitly unfinished final line: its context may
            # continue at the beginning of the next PDF page.
            if pending is not None and not cls._needs_continuation(pending[1]) and not re.search(r"(?:复查|检查|胸部|腹部|肺部)$", pending[1]):
                flush()
        flush()
        return [
            ExtractedPage(
                page.page_number, "\n".join(rebuilt.get(page.page_number, [])),
                tuple(sorted(spans.get(page.page_number, {page.page_number}))),
                page.source_metadata, page.line_locations,
            )
            for page in pages
        ]

    @classmethod
    def _should_continue(cls, previous: str, following: str) -> bool:
        if cls._looks_like_table(previous) or cls._looks_like_table(following):
            return False
        if cls._looks_like_heading(following):
            return False
        continuation_token = bool(re.match(r"^(?:CT|MRI|X线|超声|彩超|检查|复查|诊治)[。；;，,]?$", following, re.IGNORECASE))
        return cls._needs_continuation(previous) or bool(cls._fragment_start.match(following)) or (continuation_token and bool(re.search(r"(?:复查|检查|胸部|腹部|肺部)$", previous)))

    @classmethod
    def _needs_continuation(cls, text: str) -> bool:
        return bool(cls._unfinished_end.search(text.strip()))

    @staticmethod
    def _looks_like_table(text: str) -> bool:
        return "\t" in text or bool(re.search(r"\S+\s+[-+]?\d+(?:\.\d+)?\s+(?:mmol|mg|ng|U/|IU/|%|cm|kg)", text, re.IGNORECASE))

    @staticmethod
    def _looks_like_heading(text: str) -> bool:
        return bool(re.match(r"^(?:第?\d+[页、.]?|[一二三四五六七八九十]+、|(?:检查)?(?:结论|建议|小结|报告))", text))

    @staticmethod
    def _join(previous: str, following: str) -> str:
        return f"{previous.rstrip()} {following.lstrip()}".replace("  ", " ")


class SentenceCompletenessValidator:
    """Reject clinical candidates whose evidence is visibly incomplete."""

    _incomplete_end = re.compile(r"[，、：；,;:]$|(?:若|如).{0,36}时[，,]?$|(?:内科|外科|专科)诊$")
    _incomplete_start = re.compile(r"^(?:定有临床意义|进一步|并随访|治[。；;]?|时[，,]?|到.+?(?:内科|外科|专科)诊$)")

    @classmethod
    def reason(cls, text: str) -> str | None:
        value = re.sub(r"\s+", " ", text or "").strip()
        if len(value) < 6:
            return "evidence_too_short"
        if cls._incomplete_end.search(value):
            return "unfinished_sentence"
        if cls._incomplete_start.search(value):
            return "missing_sentence_context"
        return None

    @classmethod
    def is_complete(cls, text: str) -> bool:
        return cls.reason(text) is None


def _manual_review_draft(*, page: ExtractedPage, section: str, evidence: str, reason: str, method: str = "RULE") -> CandidateDraft:
    return CandidateDraft(
        "INCOMPLETE", None, None, None, None, None, None, None,
        "原报告文本疑似存在跨行断裂，系统未自动生成正式医疗建议。",
        {"candidate_quality": "INCOMPLETE", "integrity_reason": reason, "page_span": list(page.page_span or (page.page_number,)), **_evidence_metadata(page, evidence)},
        "LOW", method, page.page_number, section, evidence,
    )


def _evidence_metadata(page: ExtractedPage, evidence: str, *, table_row: str | None = None, table_header: str | None = None) -> dict[str, object]:
    """Carry exact persisted file coordinates into a Candidate's JSON payload."""
    location = page.line_locations.get(table_row or evidence) or page.line_locations.get(evidence) or {}
    metadata = {**page.source_metadata, **location}
    if table_header:
        metadata.setdefault("table_header", table_header)
    if table_row:
        metadata.setdefault("table_row", table_row)
    if metadata.get("source_format") in {"CSV", "XLSX"}:
        metadata.setdefault("evidence_type", "CELL_RANGE" if metadata.get("sheet_name") else "TABLE")
    return metadata


class GenericReportParser:
    """Rule/table parser. It never calls an external LLM."""

    _line = re.compile(r"^\s*([^\t:：]{2,64}?)\s*(?:\t+|[:：]\s*|\s{2,})([^\t\s]+)\s*([^\t\s]*)\s*(.*?)\s*$")

    def __init__(self) -> None:
        self.sections = ReportSectionClassifier()

    def extract(self, pages: list[ExtractedPage]) -> list[CandidateDraft]:
        drafts: list[CandidateDraft] = []
        for page in pages:
            section = self.sections.classify(page.text)
            drafts.extend(self._observations(page, section))
            drafts.extend(self._narrative(page, section))
        return self._deduplicate(drafts)

    def _observations(self, page: ExtractedPage, section: str) -> list[CandidateDraft]:
        drafts: list[CandidateDraft] = []
        table_header: str | None = None
        for line in page.text.splitlines():
            if "\t" in line and (
                line.strip().startswith(("项目\t", "指标\t", "检查项目\t"))
                or ("结果" in line and "参考范围" in line)
            ):
                table_header = line.strip()
                continue
            multi_metric = self._multi_metric_drafts(line, page, section, table_header=table_header)
            if multi_metric:
                drafts.extend(multi_metric)
                continue
            match = self._line.match(line)
            if not match:
                continue
            raw_name, raw_value, raw_unit, remainder = (part.strip() for part in match.groups())
            if raw_name in {"项目", "检查项目", "结果", "单位", "参考范围"} or raw_value.lower() in MISSING_VALUES:
                continue
            if not re.fullmatch(r"[+-]?\d+(?:\.\d+)?", raw_value):
                continue
            drafts.append(self._observation_draft(raw_name, raw_value, raw_unit, remainder, page, section, line.strip(), table_row=line.strip(), table_header=table_header if "\t" in line else None))
        return drafts

    @staticmethod
    def _metric_names() -> tuple[str, ...]:
        # Keep aliases centralised in the canonical registry.  Sorting longest
        # first avoids matching "胆固醇" inside a longer reported label.
        from executive_health_ai.integrations.codes import REGISTRY
        names = {name for code in REGISTRY.values() for name in (code.canonical_code, *code.aliases)}
        names.add("胆固醇")  # common report shorthand; retained as unmapped if needed
        return tuple(sorted(names, key=len, reverse=True))

    def _multi_metric_drafts(self, line: str, page: ExtractedPage, section: str, *, table_header: str | None = None) -> list[CandidateDraft]:
        """Split a collapsed table row into one candidate per reported metric."""
        names = "|".join(re.escape(name) for name in self._metric_names())
        matcher = re.compile(
            rf"(?P<name>{names})\s*(?:[:：]|\t+|\s+)\s*(?P<value>[+-]?\d+(?:\.\d+)?)\s*(?P<unit>(?:[a-zA-Zμ/%²0-9]+(?:/[a-zA-Z0-9.²]+)?|mmol/L|ng/mL|U/L|IU/L)?)",
            re.IGNORECASE,
        )
        matches = list(matcher.finditer(line))
        recognised = [match for match in matches if canonical_code(match.group("name"))]
        if len(recognised) < 2:
            return []
        drafts: list[CandidateDraft] = []
        for match in recognised:
            raw_name, raw_value, raw_unit = match.group("name"), match.group("value"), match.group("unit")
            drafts.append(self._observation_draft(raw_name, raw_value, raw_unit, "", page, section, match.group(0).strip(), table_row=line.strip(), table_header=table_header))
        return drafts

    @staticmethod
    def _observation_draft(raw_name: str, raw_value: str, raw_unit: str, remainder: str, page: ExtractedPage, section: str, evidence: str, *, table_row: str | None = None, table_header: str | None = None) -> CandidateDraft:
        code = canonical_code(raw_name)
        normalized, unit, confidence = raw_value, raw_unit or None, "MEDIUM"
        structured: dict[str, object] = {"raw_name": raw_name, "raw_value": raw_value, "raw_unit": raw_unit or None, "candidate_quality": "COMPLETE", "page_span": list(page.page_span or (page.page_number,)), **_evidence_metadata(page, evidence, table_row=table_row, table_header=table_header)}
        if code:
            try:
                amount, normalized_unit = normalize_unit(code, raw_value, raw_unit or code.default_unit)
                normalized, unit, confidence = str(amount), normalized_unit, "HIGH"
            except ValueError:
                confidence = "MEDIUM"
        abnormal = next((flag for flag in ("↑", "↓", "H", "L", "阳性", "可疑") if flag in remainder), None)
        reference = remainder.strip() or None
        extraction_method = "TABLE" if "\t" in (table_row or evidence) or structured.get("evidence_type") in {"TABLE", "CELL_RANGE"} else "RULE"
        return CandidateDraft("OBSERVATION", raw_name, raw_value, code.canonical_code if code else None, normalized, unit, reference, abnormal, None, structured, confidence, extraction_method, page.page_number, section, evidence)

    def _narrative(self, page: ExtractedPage, section: str) -> list[CandidateDraft]:
        drafts: list[CandidateDraft] = []
        for line in page.text.splitlines():
            text = line.strip()
            if len(text) < 5:
                continue
            if re.search(r"(?:建议|宜).{0,40}(?:复查|门诊|就诊|检查)", text):
                reason = SentenceCompletenessValidator.reason(text)
                if reason:
                    drafts.append(_manual_review_draft(page=page, section="RECOMMENDATION", evidence=text, reason=reason))
                else:
                    drafts.append(CandidateDraft("FOLLOWUP", None, None, None, None, None, None, None, self._followup_title(text), {"candidate_quality": "COMPLETE", "page_span": list(page.page_span or (page.page_number,)), **_evidence_metadata(page, text)}, "MEDIUM", "RULE", page.page_number, "RECOMMENDATION", text))
            elif re.match(r"(?:检查结论|结论|提示|所见)[:：]", text) or (section in {"IMAGING", "ECG", "PULMONARY_FUNCTION"} and any(token in text for token in ("结节", "脂肪肝", "异常", "改变", "障碍"))):
                summary = re.split(r"[:：]", text, maxsplit=1)[-1].strip()
                if SentenceCompletenessValidator.reason(summary):
                    drafts.append(_manual_review_draft(page=page, section=section, evidence=text, reason="unfinished_finding"))
                    continue
                for finding in self._atomic_findings(summary):
                    drafts.append(CandidateDraft("FINDING", None, None, None, None, None, None, None, finding, {"exam_name": section, "candidate_quality": "COMPLETE", "page_span": list(page.page_span or (page.page_number,)), **_evidence_metadata(page, finding)}, "MEDIUM", "RULE", page.page_number, section, finding))
        return drafts

    @staticmethod
    def _atomic_findings(summary: str) -> list[str]:
        """Preserve each report-supported finding instead of collapsing a CT paragraph."""
        parts = [part.strip(" ，；;。") for part in re.split(r"(?<=[。；;])", summary) if part.strip(" ，；;。")]
        clinical = [part for part in parts if not re.match(r"^(?:建议|宜|请|需)", part)]
        return clinical or [summary]

    @staticmethod
    def _followup_title(text: str) -> str:
        """A compact, source-backed title; the complete text remains evidence."""
        match = re.search(r"(?:建议|宜)\s*[^，。；;]{0,48}(?:复查|门诊|就诊|检查)[^，。；;]*", text)
        return match.group(0).strip() if match else "报告建议复查"

    def _deduplicate(self, drafts: list[CandidateDraft]) -> list[CandidateDraft]:
        unique: dict[tuple[object, ...], CandidateDraft] = {}
        for draft in drafts:
            key = (draft.candidate_type, draft.canonical_code or draft.raw_name, draft.normalized_value or draft.summary)
            if key not in unique or unique[key].confidence != "HIGH":
                unique[key] = draft
        return list(unique.values())


@dataclass(frozen=True)
class SemanticFallbackResult:
    drafts: list[CandidateDraft]
    used: bool
    enabled: bool
    available: bool
    provider: str
    model: str
    status: str
    call_count: int
    success_count: int
    failure_count: int
    total_duration_ms: int
    processed_sections: list[str]
    failure_reason: str | None = None


def _normalised_evidence(value: str) -> str:
    return re.sub(r"[\s，。；;：:、,.()（）\[\]{}]", "", value or "").lower()


def _evidence_in_source(evidence: str, source: str) -> bool:
    normalised = _normalised_evidence(evidence)
    return bool(normalised) and normalised in _normalised_evidence(source)


def _field_is_directly_supported(value: str, evidence: str) -> bool:
    """Clinical targets must appear in the cited evidence, never be inferred."""
    return not value.strip() or _normalised_evidence(value) in _normalised_evidence(evidence)


def _action_is_directly_supported(action: str, evidence: str) -> bool:
    """Allow a concise action label only when its stated operation is explicit."""
    if _field_is_directly_supported(action, evidence):
        return True
    source = _normalised_evidence(evidence)
    action = _normalised_evidence(action)
    verbs = ("建议", "复查", "就诊", "门诊", "检查", "随访", "转诊")
    mentioned = [token for token in verbs if token in action]
    return bool(mentioned) and all(token in source for token in mentioned)


class ReportSemanticFallback:
    """Optional, evidence-bound local LLM semantic enrichment for narratives.

    This runs alongside deterministic extraction in every normal parse run.  It
    is not a retry path for rule-parser failure: structured measurements remain
    rule-owned while eligible narrative sections are sent to local LLM once.
    """

    supported_sections = {"IMAGING", "ULTRASOUND", "ECG", "PULMONARY_FUNCTION", "RECOMMENDATION", "OTHER"}
    section_labels = {
        "IMAGING": "影像检查",
        "ULTRASOUND": "超声检查",
        "ECG": "心电图",
        "PULMONARY_FUNCTION": "肺功能",
        "RECOMMENDATION": "随访建议",
        "OTHER": "复杂检查结论",
    }
    _heading_patterns = (
        r"胸(?:部)?\s*CT",
        r"腹部(?:彩超|超声)",
        r"甲状腺(?:彩超|超声)",
        r"肺功能",
        r"心电图",
        r"(?:健康|随访|复查)建议",
    )

    def __init__(self, client: LocalLLMClient | None = None) -> None:
        self.client = client or LocalLLMClient()

    def extract(self, *, pages: list[ExtractedPage], existing: list[CandidateDraft], document_id: UUID, progress_callback: ProgressCallback | None = None) -> SemanticFallbackResult:
        # Select narrative work first.  A lab-only report must remain a clean
        # rule-only run even when the optional local model is switched off.
        classifier = ReportSectionClassifier()
        fragments = [fragment for page in pages for fragment in self._semantic_fragments(page)]
        eligible = [(page, self._classify_fragment(page, classifier)) for page in fragments]
        eligible = [(page, section) for page, section in eligible if self._requires_semantic_assistance(page, section)]
        incomplete_drafts: list[CandidateDraft] = []
        complete_eligible: list[tuple[ExtractedPage, str]] = []
        for page, section in eligible:
            reason = SentenceCompletenessValidator.reason(page.text)
            if reason:
                incomplete_drafts.append(_manual_review_draft(page=page, section=section, evidence=page.text, reason=reason, method="LLM"))
            else:
                complete_eligible.append((page, section))
        eligible = complete_eligible
        selected: list[tuple[ExtractedPage, str, str]] = []
        seen_sections: set[str] = set()
        for page, section in eligible:
            section_key = self._section_key(page, section)
            if section_key in seen_sections:
                continue
            seen_sections.add(section_key)
            selected.append((page, section, self._section_label(page, section)))
        health = self.client.health_check()
        if not selected:
            self._emit_progress(progress_callback, "LLM_NOT_NEEDED", "本报告没有需要本地AI辅助整理的复杂检查内容。")
            return self._result(health, "NOT_NEEDED", drafts=incomplete_drafts)
        self._emit_progress(progress_callback, "LLM_STARTED", f"预计由本地AI辅助整理 {len(selected)} 个检查段。", total=len(selected))
        if not health.enabled:
            self._emit_progress(progress_callback, "LLM_UNAVAILABLE", "本地AI未启用，复杂检查内容将保留供人工确认。", total=len(selected))
            return self._result(health, "UNAVAILABLE", drafts=incomplete_drafts, failure_reason=health.reason)
        if not health.available:
            self._emit_progress(progress_callback, "LLM_UNAVAILABLE", "本地AI当前不可用，复杂检查内容将保留供人工确认。", total=len(selected))
            return self._result(health, "UNAVAILABLE", drafts=incomplete_drafts, failure_reason=health.reason)
        drafts: list[CandidateDraft] = list(incomplete_drafts)
        call_count = success_count = failure_count = total_duration_ms = 0
        processed_sections: list[str] = []
        for current, (page, section, label) in enumerate(selected, start=1):
            call_count += 1
            if label not in processed_sections:
                processed_sections.append(label)
            self._emit_progress(progress_callback, "LLM_SECTION_STARTED", f"本地AI正在解析：{label}", current=current, total=len(selected), section_name=label)
            sanitized = sanitize_for_llm(page.text)
            started = clock.perf_counter()
            try:
                payload = self.client.generate_structured(task="report_semantic_fallback", system_prompt=FINDING_EXTRACTION_SYSTEM_PROMPT, user_prompt=finding_extraction_prompt(sanitized), document_id=str(document_id), page=page.page_number)
            except LocalLLMUnavailable as error:
                elapsed = round((clock.perf_counter() - started) * 1000)
                total_duration_ms += elapsed
                failure_count += 1
                logger.warning("report_llm_call provider=%s model=%s document_id=%s page=%s section=%s latency_ms=%s status=failed error_type=%s", health.provider, health.model, document_id, page.page_number, section, elapsed, type(error).__name__)
                self._emit_progress(progress_callback, "LLM_SECTION_FAILED", f"{label} 的本地AI辅助解析失败，将保留规则结果供人工确认。", current=current, total=len(selected), section_name=label, call_duration_ms=elapsed)
                continue
            elapsed = round((clock.perf_counter() - started) * 1000)
            total_duration_ms += elapsed
            success_count += 1
            logger.info("report_llm_call provider=%s model=%s document_id=%s page=%s section=%s latency_ms=%s status=success", health.provider, health.model, document_id, page.page_number, section, elapsed)
            drafts.extend(self._validated_drafts(payload, sanitized, page, section))
            self._emit_progress(progress_callback, "LLM_SECTION_COMPLETED", f"{label} 的本地AI辅助解析完成。", current=current, total=len(selected), section_name=label, call_duration_ms=elapsed)
        if success_count:
            return self._result(health, "USED", drafts=drafts, call_count=call_count, success_count=success_count, failure_count=failure_count, total_duration_ms=total_duration_ms, processed_sections=processed_sections)
        return self._result(health, "UNAVAILABLE", call_count=call_count, failure_count=failure_count, total_duration_ms=total_duration_ms, processed_sections=processed_sections, failure_reason="本地开源大模型 调用失败")

    def _requires_semantic_assistance(self, page: ExtractedPage, section: str) -> bool:
        """Use LLM for complex clinical narratives, not structured measurements."""
        text = page.text.strip()
        sentence_count = sum(text.count(mark) for mark in ("。", "；", ";", "\n"))
        if len(text) < 24 or sentence_count < 1:
            return False
        if section in self.supported_sections - {"OTHER"}:
            return True
        return bool(re.search(r"(?:检查结论|影像|所见|提示|异常|改变|结节|彩超|超声|肺功能|复查|门诊|随访)", text, re.IGNORECASE))

    @classmethod
    def _semantic_fragments(cls, page: ExtractedPage) -> list[ExtractedPage]:
        """Split a multi-section extracted page into minimal named narratives."""
        lines = [line.strip() for line in page.text.splitlines() if line.strip()]
        fragments: list[str] = []
        current: list[str] = []
        for line in lines:
            is_heading = any(re.match(pattern, line, re.IGNORECASE) for pattern in cls._heading_patterns)
            if is_heading and current:
                fragments.append("\n".join(current))
                current = []
            current.append(line)
        if current:
            fragments.append("\n".join(current))
        return [ExtractedPage(page.page_number, text, page.page_span) for text in fragments] or [page]

    @staticmethod
    def _classify_fragment(page: ExtractedPage, classifier: ReportSectionClassifier) -> str:
        if re.match(r"^(?:健康|随访|复查)建议", page.text.strip(), re.IGNORECASE):
            return "RECOMMENDATION"
        return classifier.classify(page.text)

    @classmethod
    def _section_label(cls, page: ExtractedPage, section: str) -> str:
        text = page.text
        if section == "RECOMMENDATION":
            return cls.section_labels[section]
        for label, pattern in (("胸部CT", r"胸(?:部)?\s*CT"), ("腹部彩超", r"腹部(?:彩超|超声)"), ("甲状腺彩超", r"甲状腺(?:彩超|超声)"), ("肺功能", r"肺功能"), ("心电图", r"心电图")):
            if re.search(pattern, text, re.IGNORECASE):
                return label
        return cls.section_labels.get(section, "复杂检查结论")

    @classmethod
    def _section_key(cls, page: ExtractedPage, section: str) -> str:
        label = cls._section_label(page, section)
        # A named examination may span several extracted pages; send that section
        # once. Generic unnamed narratives remain page-scoped to avoid merging
        # unrelated exams.
        if section in {"PULMONARY_FUNCTION", "ECG", "RECOMMENDATION"}:
            return f"{section}:{label}"
        return f"{section}:{label}" if label != cls.section_labels.get(section) else f"{section}:page-{page.page_number}"

    @staticmethod
    def _emit_progress(callback: ProgressCallback | None, stage: str, message: str, *, current: int | None = None, total: int | None = None, section_name: str | None = None, call_duration_ms: int | None = None) -> None:
        if callback is not None:
            callback(ReportParseProgress(stage, message, current, total, section_name, call_duration_ms=call_duration_ms))

    @staticmethod
    def _result(health: LocalLLMHealth, status: str, *, drafts: list[CandidateDraft] | None = None, call_count: int = 0, success_count: int = 0, failure_count: int = 0, total_duration_ms: int = 0, processed_sections: list[str] | None = None, failure_reason: str | None = None) -> SemanticFallbackResult:
        return SemanticFallbackResult(drafts or [], status == "USED", health.enabled, health.available, health.provider, health.model, status, call_count, success_count, failure_count, total_duration_ms, processed_sections or [], failure_reason)

    def _validated_drafts(self, payload: dict[str, object], source: str, page: ExtractedPage, section: str) -> list[CandidateDraft]:
        result: list[CandidateDraft] = []
        exam_name = str(payload.get("exam_name") or "")
        for item in payload.get("findings", []) if isinstance(payload.get("findings"), list) else []:
            if not isinstance(item, dict): continue
            evidence = str(item.get("evidence") or "")
            summary = str(item.get("summary") or "")
            reason = SentenceCompletenessValidator.reason(evidence)
            if reason:
                result.append(_manual_review_draft(page=page, section=section, evidence=evidence, reason=reason, method="LLM"))
                continue
            if not summary or not _evidence_in_source(evidence, source):
                continue
            result.append(CandidateDraft("FINDING", None, None, None, None, None, None, None, summary, {"exam_name": exam_name, "body_system": str(item.get("body_system") or ""), "reported_change": str(item.get("reported_change") or ""), "reported_severity": str(item.get("reported_severity") or ""), "candidate_quality": "COMPLETE", "page_span": list(page.page_span or (page.page_number,)), **_evidence_metadata(page, evidence)}, "MEDIUM", "LLM", page.page_number, section, evidence))
        for item in payload.get("recommendations", []) if isinstance(payload.get("recommendations"), list) else []:
            if not isinstance(item, dict): continue
            evidence = str(item.get("evidence") or "")
            action = str(item.get("action") or "")
            department = str(item.get("department") or "")
            interval = str(item.get("interval_text") or "")
            reason = SentenceCompletenessValidator.reason(evidence)
            if reason:
                result.append(_manual_review_draft(page=page, section="RECOMMENDATION", evidence=evidence, reason=reason, method="LLM"))
                continue
            if not action or not _evidence_in_source(evidence, source):
                continue
            # Department, revisit interval, procedure and medication wording are
            # safety-critical: an LLM may only expose text that is explicitly
            # present in the cited complete evidence span.
            if not all((_action_is_directly_supported(action, evidence), _field_is_directly_supported(department, evidence), _field_is_directly_supported(interval, evidence))):
                result.append(_manual_review_draft(page=page, section="RECOMMENDATION", evidence=evidence, reason="evidence_mismatch", method="LLM"))
                continue
            result.append(CandidateDraft("FOLLOWUP", None, None, None, None, None, None, None, action, {"department": department, "interval_text": interval, "candidate_quality": "COMPLETE", "page_span": list(page.page_span or (page.page_number,)), **_evidence_metadata(page, evidence)}, "MEDIUM", "LLM", page.page_number, "RECOMMENDATION", evidence))
        return result


class ReportParsingService:
    storage_root = Path("report_uploads")

    def __init__(self, preflight: DocumentPreflightService | None = None, parser: GenericReportParser | None = None, semantic_fallback: ReportSemanticFallback | None = None) -> None:
        self.preflight = preflight or DocumentPreflightService()
        self.parser = parser or GenericReportParser()
        self.semantic_fallback = semantic_fallback or ReportSemanticFallback()

    def upload_and_parse(self, session: Session, patient_id: UUID, filename: str, content: bytes, actor: str, progress_callback: ProgressCallback | None = None) -> tuple[Document, ReportExtractionRun, bool]:
        """Intake a document and execute one explicit, fresh parse command.

        File hashes deduplicate immutable source documents only.  They never
        deduplicate parsing commands: every user click on "start parsing" gets
        a new extraction run and fresh rule/LLM candidates.
        """
        started = clock.perf_counter()
        self._emit_progress(progress_callback, started, "READING_REPORT", "正在读取报告…")
        safe_filename, _ = _validate_report_upload(filename, content)
        context, pages = self.preflight.inspect(safe_filename, content)
        self._emit_progress(progress_callback, started, "PREFLIGHT_COMPLETED", "文档预处理完成。")
        existing = session.scalar(select(ReportExtractionRun).where(ReportExtractionRun.patient_id == patient_id, ReportExtractionRun.file_hash == context.file_hash).order_by(ReportExtractionRun.created_at.desc()))
        if existing:
            document = session.get(Document, existing.document_id)
            assert document is not None
            self._ensure_not_processing(session, document.id)
            return document, self._parse_document(session, document, context, pages, actor, is_reparse=False, progress_callback=progress_callback, operation_started=started), True
        self.storage_root.mkdir(parents=True, exist_ok=True)
        stored_name = f"{uuid4()}-{safe_filename}"
        stored_path = self.storage_root / stored_name
        stored_path.write_bytes(content)
        document = Document(patient_id=patient_id, document_type="health_check_report", title=safe_filename, storage_reference=str(stored_path), source="report_upload", status="PARSING")
        session.add(document); session.flush()
        run = self._parse_document(session, document, context, pages, actor, is_reparse=False, progress_callback=progress_callback, operation_started=started)
        return document, run, False

    def reparse_document(self, session: Session, document_id: UUID, actor: str, progress_callback: ProgressCallback | None = None) -> ReportExtractionRun:
        """Create a new immutable extraction run from the original document.

        Re-parsing never changes candidates or observations created by an earlier
        run.  The new run starts with fresh, human-review-only candidates.
        """
        started = clock.perf_counter()
        self._emit_progress(progress_callback, started, "READING_REPORT", "正在读取原始报告…")
        document = session.get(Document, document_id)
        if document is None or document.document_type != "health_check_report":
            raise ValueError("未找到可重新解析的体检报告。")
        source_path = Path(document.storage_reference)
        if not source_path.is_file():
            raise ValueError("原始报告文件不可用，无法重新解析。")
        self._ensure_not_processing(session, document.id)
        context, pages = self.preflight.inspect(document.title, source_path.read_bytes())
        self._emit_progress(progress_callback, started, "PREFLIGHT_COMPLETED", "文档预处理完成。")
        return self._parse_document(session, document, context, pages, actor, is_reparse=True, progress_callback=progress_callback, operation_started=started)

    @staticmethod
    def _ensure_not_processing(session: Session, document_id: UUID) -> None:
        processing = session.scalar(select(ReportExtractionRun).where(
            ReportExtractionRun.document_id == document_id,
            ReportExtractionRun.status.in_(("PENDING", "PROCESSING")),
        ).order_by(ReportExtractionRun.created_at.desc()))
        if processing is not None:
            raise ValueError("该报告正在解析，请等待当前解析完成后再试。")

    def _parse_document(self, session: Session, document: Document, context: DocumentPreflight, pages: list[ExtractedPage], actor: str, *, is_reparse: bool, progress_callback: ProgressCallback | None = None, operation_started: float | None = None) -> ReportExtractionRun:
        started = operation_started or clock.perf_counter()
        run = ReportExtractionRun(document_id=document.id, patient_id=document.patient_id, status="PROCESSING", parser_version=PARSER_VERSION, canonical_registry_version=CANONICAL_REGISTRY_VERSION, file_hash=context.file_hash, file_type=context.file_type, detected_hospital=context.detected_hospital, detected_report_type=context.detected_report_type, detected_report_date=context.detected_report_date, page_count=context.page_count, has_text_layer=context.has_text_layer, is_scanned=context.is_probably_scanned, template_fingerprint=context.possible_template, llm_used=False, metadata_json={"detected_language": context.detected_language, "external_phi_llm_enabled": ALLOW_EXTERNAL_PHI_LLM})
        session.add(run); session.flush()
        if context.is_probably_scanned:
            run.status, document.status = "NEEDS_OCR", "OCR_REQUIRED"
            self._emit_progress(progress_callback, started, "COMPLETED", "报告需要文字识别后才能继续解析。")
        else:
            # One shared parse pipeline for first upload and reparse: rules own
            # structured facts, while local LLM enriches only complex narrative
            # sections in this same extraction run.
            self._emit_progress(progress_callback, started, "TEXT_RECONSTRUCTION", "正在重建报告段落与跨页句子…")
            pages = ReportTextReconstructor.reconstruct(pages)
            self._emit_progress(progress_callback, started, "RULE_PARSE_STARTED", "正在进行规则解析…")
            generic_drafts = self.parser.extract(pages)
            self._emit_progress(progress_callback, started, "RULE_PARSE_COMPLETED", "规则解析完成。", rule_candidate_count=len(generic_drafts))
            semantic_result = self.semantic_fallback.extract(
                pages=pages,
                existing=generic_drafts,
                document_id=document.id,
                progress_callback=lambda event: self._forward_progress(progress_callback, started, event),
            )
            if semantic_result.success_count:
                self._emit_progress(progress_callback, started, "EVIDENCE_VALIDATION", "本地AI结果证据校验完成。")
            self._emit_progress(progress_callback, started, "MERGING", "正在合并规则与本地AI解析结果…")
            combined_drafts = self._deduplicate_combined_candidates([*generic_drafts, *semantic_result.drafts])
            self._emit_progress(progress_callback, started, "DEDUPLICATION", "解析结果去重完成。")
            run.metadata_json = {
                **run.metadata_json,
                "rule_candidate_count": sum(item.extraction_method != "LLM" for item in combined_drafts),
                "llm_candidate_count": sum(item.extraction_method == "LLM" for item in combined_drafts),
            }
            self._emit_progress(progress_callback, started, "SAVING", "正在保存解析结果…")
            self._persist_candidates(session, run, document, combined_drafts)
            run.llm_used = semantic_result.used
            run.llm_enabled = semantic_result.enabled
            run.llm_available = semantic_result.available
            run.llm_provider = semantic_result.provider
            run.llm_model = semantic_result.model
            run.llm_status = semantic_result.status
            run.llm_call_count = semantic_result.call_count
            run.llm_success_count = semantic_result.success_count
            run.llm_failure_count = semantic_result.failure_count
            run.llm_total_duration_ms = semantic_result.total_duration_ms
            run.llm_processed_sections = semantic_result.processed_sections
            run.llm_failure_reason = semantic_result.failure_reason
            if semantic_result.used:
                run.parser_version = f"{PARSER_VERSION}+local-llm-fallback"
            # A local model is optional.  Rule candidates and their human-review
            # workflow remain complete even when semantic enrichment is offline;
            # the persisted LLM status makes that limitation explicit in the UI.
            run.status = "PARTIAL_SUCCESS" if semantic_result.failure_count else "COMPLETED"
            document.status = "PENDING_HUMAN_REVIEW"
        run.completed_at = utc_now()
        session.add(AuditLog(patient_id=document.patient_id, actor=actor, actor_role="health_manager", action="reparsed_health_check_report" if is_reparse else "parsed_health_check_report", entity_type="ReportExtractionRun", entity_id=str(run.id), detail_json={"page_count": context.page_count, "candidate_count": run.candidate_count, "scanned": context.is_probably_scanned, "llm_status": run.llm_status, "llm_call_count": run.llm_call_count, "llm_total_duration_ms": run.llm_total_duration_ms, "llm_processed_sections": run.llm_processed_sections, "is_reparse": is_reparse}))
        completion_message = "体检报告已部分解析，本地AI失败的检查内容已保留供人工确认。" if run.status == "PARTIAL_SUCCESS" else "体检报告解析完成。"
        self._emit_progress(
            progress_callback,
            started,
            "COMPLETED",
            completion_message,
            rule_candidate_count=run.rule_candidate_count,
            candidate_count=run.candidate_count,
            finding_count=sum(item.candidate_type == "FINDING" for item in combined_drafts) if not context.is_probably_scanned else 0,
            followup_count=sum(item.candidate_type == "FOLLOWUP" for item in combined_drafts) if not context.is_probably_scanned else 0,
            llm_call_count=run.llm_call_count,
            llm_success_count=run.llm_success_count,
            llm_failure_count=run.llm_failure_count,
        )
        return run

    @staticmethod
    def _emit_progress(callback: ProgressCallback | None, started: float, stage: str, message: str, *, current: int | None = None, total: int | None = None, section_name: str | None = None, rule_candidate_count: int | None = None, call_duration_ms: int | None = None, candidate_count: int | None = None, finding_count: int | None = None, followup_count: int | None = None, llm_call_count: int | None = None, llm_success_count: int | None = None, llm_failure_count: int | None = None) -> None:
        if callback is not None:
            callback(ReportParseProgress(
                stage=stage,
                message=message,
                current=current,
                total=total,
                section_name=section_name,
                rule_candidate_count=rule_candidate_count,
                call_duration_ms=call_duration_ms,
                elapsed_ms=round((clock.perf_counter() - started) * 1000),
                candidate_count=candidate_count,
                finding_count=finding_count,
                followup_count=followup_count,
                llm_call_count=llm_call_count,
                llm_success_count=llm_success_count,
                llm_failure_count=llm_failure_count,
            ))

    @classmethod
    def _forward_progress(cls, callback: ProgressCallback | None, started: float, event: ReportParseProgress) -> None:
        cls._emit_progress(callback, started, event.stage, event.message, current=event.current, total=event.total, section_name=event.section_name, rule_candidate_count=event.rule_candidate_count, call_duration_ms=event.call_duration_ms)

    def _persist_candidates(self, session: Session, run: ReportExtractionRun, document: Document, drafts: list[CandidateDraft]) -> None:
        for draft in drafts:
            quality = str(draft.structured_data.get("candidate_quality", "COMPLETE"))
            status = "PENDING_REVIEW" if quality == "COMPLETE" and draft.candidate_type != "INCOMPLETE" else "NEEDS_MANUAL_REVIEW"
            session.add(ReportExtractionCandidate(extraction_run_id=run.id, document_id=document.id, patient_id=run.patient_id, candidate_type=draft.candidate_type, canonical_code=draft.canonical_code, raw_name=draft.raw_name, raw_value=draft.raw_value, normalized_value=draft.normalized_value, unit=draft.unit, reference_range=draft.reference_range, abnormal_flag=draft.abnormal_flag, summary=draft.summary, structured_data_json=draft.structured_data, confidence=draft.confidence, extraction_method=draft.extraction_method, source_page=draft.source_page, source_section=draft.source_section, evidence_text=draft.evidence_text, status=status))
            run.candidate_count += 1
            if draft.confidence == "HIGH": run.high_confidence_count += 1
            elif draft.confidence == "MEDIUM": run.medium_confidence_count += 1
            else: run.low_confidence_count += 1

    @staticmethod
    def _deduplicate_combined_candidates(drafts: list[CandidateDraft]) -> list[CandidateDraft]:
        """Merge rule/LLM duplicates while retaining distinct, atomic findings."""
        unique: dict[tuple[str, str], CandidateDraft] = {}
        for draft in drafts:
            evidence = _normalised_evidence(draft.evidence_text)
            if draft.candidate_type == "OBSERVATION":
                key = (draft.candidate_type, f"{draft.canonical_code or draft.raw_name}:{draft.normalized_value or draft.raw_value}:{draft.unit or ''}")
            elif draft.candidate_type == "FOLLOWUP":
                data = draft.structured_data
                key = (draft.candidate_type, _normalised_evidence("|".join((draft.summary or "", str(data.get("department", "")), str(data.get("interval_text", ""))))))
            elif draft.candidate_type == "FINDING":
                key = (draft.candidate_type, evidence or _normalised_evidence(draft.summary or ""))
            else:
                key = (draft.candidate_type, evidence)
            previous = unique.get(key)
            if previous is None:
                unique[key] = draft
                continue
            methods = sorted({previous.extraction_method, draft.extraction_method})
            preferred = draft if draft.extraction_method == "LLM" else previous
            # Preserve the compact legacy draft shape for callers that provide
            # no structured payload; persisted candidates retain both sources.
            structured = {**preferred.structured_data, "extraction_methods": methods} if preferred.structured_data else preferred.structured_data
            unique[key] = CandidateDraft(
                preferred.candidate_type, preferred.raw_name, preferred.raw_value, preferred.canonical_code,
                preferred.normalized_value, preferred.unit, preferred.reference_range, preferred.abnormal_flag,
                preferred.summary, structured, preferred.confidence, preferred.extraction_method,
                preferred.source_page, preferred.source_section, preferred.evidence_text,
            )
        return list(unique.values())

    def candidates(self, session: Session, document_id: UUID, extraction_run_id: UUID | None = None) -> list[ReportExtractionCandidate]:
        statement = select(ReportExtractionCandidate).where(ReportExtractionCandidate.document_id == document_id)
        if extraction_run_id is not None:
            statement = statement.where(ReportExtractionCandidate.extraction_run_id == extraction_run_id)
        return list(session.scalars(statement.order_by(ReportExtractionCandidate.source_page, ReportExtractionCandidate.created_at)))

    def runs(self, session: Session, document_id: UUID) -> list[ReportExtractionRun]:
        return list(session.scalars(select(ReportExtractionRun).where(ReportExtractionRun.document_id == document_id).order_by(ReportExtractionRun.created_at.desc())))

    def possible_duplicate_observation(self, session: Session, candidate: ReportExtractionCandidate) -> Observation | None:
        """Return a matching confirmed report observation without changing either record."""
        if candidate.candidate_type != "OBSERVATION" or not candidate.canonical_code or not candidate.normalized_value or not candidate.unit:
            return None
        run = session.get(ReportExtractionRun, candidate.extraction_run_id)
        if run is None:
            return None
        try:
            value = Decimal(candidate.normalized_value)
        except InvalidOperation:
            return None
        statement = select(Observation).where(
            Observation.patient_id == candidate.patient_id,
            Observation.metric_code == candidate.canonical_code,
            Observation.value_numeric == value,
            Observation.unit == candidate.unit,
            Observation.source == "confirmed_health_check_report",
        )
        # Reports without a reliable examination date are deliberately handled
        # conservatively: the same metric/value/unit is shown for review rather
        # than silently added a second time.  A reviewer can still keep it if it
        # truly belongs to a different report.
        if run.detected_report_date is not None:
            observed_at = datetime.combine(run.detected_report_date, time(12), tzinfo=timezone.utc)
            statement = statement.where(Observation.observed_at == observed_at)
        return session.scalar(statement.order_by(Observation.created_at.desc()))

    def confirm_candidate(self, session: Session, candidate: ReportExtractionCandidate, actor: str) -> Observation | None:
        if candidate.status not in {"PENDING_REVIEW", "CORRECTED"}:
            raise ValueError("只有待确认或已修正的资料可以确认入档。")
        observation: Observation | None = None
        if candidate.candidate_type == "OBSERVATION":
            if not candidate.canonical_code or not candidate.normalized_value or not candidate.unit:
                raise ValueError("尚未匹配标准指标，不能直接写入健康档案。")
            try:
                value = Decimal(candidate.normalized_value)
            except InvalidOperation as error:
                raise ValueError("候选数值无效。") from error
            if self.possible_duplicate_observation(session, candidate) is not None:
                raise ValueError("该健康指标可能已经入档。为避免重复写入，请先查看已有记录。")
            code = canonical_code(candidate.canonical_code)
            quality, note = quality_for(code, value) if code else ("suspect", "unmapped canonical code")
            run = session.get(ReportExtractionRun, candidate.extraction_run_id)
            observed_at = datetime.combine(run.detected_report_date, time(12), tzinfo=timezone.utc) if run and run.detected_report_date else datetime.now(timezone.utc)
            observation = Observation(patient_id=candidate.patient_id, observed_at=observed_at, metric_code=candidate.canonical_code, value_numeric=value, unit=candidate.unit, source="confirmed_health_check_report", quality_flag=quality, source_record_id=str(candidate.id), quality_notes=f"报告页 {candidate.source_page}; {note or '人工确认'}")
            session.add(observation)
            session.flush()
            # Candidate data has become a persisted, human-confirmed
            # Observation.  Only now may deterministic governed risk rules run.
            RiskEvaluationService().evaluate_observation_safely(session, observation.id)
            ManagementRoutingService().evaluate_observation(session, observation.id)
        candidate.status, candidate.reviewed_by, candidate.reviewed_at = "CONFIRMED", actor, utc_now()
        session.add(AuditLog(patient_id=candidate.patient_id, actor=actor, actor_role="health_manager", action="confirmed_report_candidate", entity_type="ReportExtractionCandidate", entity_id=str(candidate.id), detail_json={"candidate_type": candidate.candidate_type, "document_id": str(candidate.document_id)}))
        session.flush()
        return observation

    def correct_candidate(self, session: Session, candidate: ReportExtractionCandidate, actor: str, *, canonical: str | None, value: str | None, unit: str | None, reason: str) -> None:
        before = {"canonical_code": candidate.canonical_code, "normalized_value": candidate.normalized_value, "unit": candidate.unit}
        if canonical: candidate.canonical_code = canonical
        if value: candidate.normalized_value = value
        if unit: candidate.unit = unit
        candidate.status, candidate.reviewed_by, candidate.reviewed_at = "CORRECTED", actor, utc_now()
        after = {"canonical_code": candidate.canonical_code, "normalized_value": candidate.normalized_value, "unit": candidate.unit}
        session.add(AuditLog(patient_id=candidate.patient_id, actor=actor, actor_role="health_manager", action="corrected_report_candidate", entity_type="ReportExtractionCandidate", entity_id=str(candidate.id), detail_json={"before": before, "after": after, "reason": reason}))
        # Persist only a compact, de-identified correction sample. The raw
        # report and full prompt stay in their existing evidence boundary.
        from executive_health_ai.services.ai_feedback import FeedbackService
        FeedbackService().capture_report_correction(
            session, candidate=candidate, actor=actor, before=before, after=after, reason=reason,
        )

    def reject_candidate(self, session: Session, candidate: ReportExtractionCandidate, actor: str, reason: str) -> None:
        candidate.status, candidate.reviewed_by, candidate.reviewed_at = "REJECTED", actor, utc_now()
        session.add(AuditLog(patient_id=candidate.patient_id, actor=actor, actor_role="health_manager", action="rejected_report_candidate", entity_type="ReportExtractionCandidate", entity_id=str(candidate.id), detail_json={"reason": reason}))

    def action_finding(self, session: Session, candidate: ReportExtractionCandidate, actor: str, action: str) -> HealthProblem | None:
        if candidate.candidate_type != "FINDING" or candidate.status not in {"PENDING_REVIEW", "CORRECTED"}: raise ValueError("该检查结论尚未通过完整原文核对，不能执行后续操作。")
        candidate.status, candidate.reviewed_by, candidate.reviewed_at = "CONFIRMED", actor, utc_now()
        problem = None
        if action in {"MANAGE", "DOCTOR_REVIEW"}:
            problem = HealthProblem(patient_id=candidate.patient_id, title=candidate.summary or "报告检查结论待跟进", description=f"来源于已人工确认的体检报告第 {candidate.source_page} 页；不是系统自动诊断。", severity="MEDIUM", responsible_role="doctor" if action == "DOCTOR_REVIEW" else "health_manager", source="confirmed_report_finding")
            session.add(problem)
        session.add(AuditLog(patient_id=candidate.patient_id, actor=actor, actor_role="health_manager", action="actioned_report_finding", entity_type="ReportExtractionCandidate", entity_id=str(candidate.id), detail_json={"action": action, "health_problem_created": action in {"MANAGE", "DOCTOR_REVIEW"}}))
        session.flush(); return problem

    def create_followup_task(self, session: Session, candidate: ReportExtractionCandidate, actor: str) -> Task:
        if candidate.candidate_type != "FOLLOWUP" or candidate.status not in {"PENDING_REVIEW", "CORRECTED"}: raise ValueError("该随访建议尚未通过完整原文核对，不能创建任务。")
        candidate.status, candidate.reviewed_by, candidate.reviewed_at = "CONFIRMED", actor, utc_now()
        task = Task(patient_id=candidate.patient_id, title=candidate.summary or "完成报告建议的后续安排", instruction=f"来源于已确认体检报告第 {candidate.source_page} 页：{candidate.evidence_text}", priority="MEDIUM", assignee=actor, responsible_role="health_manager", source=f"confirmed_report_followup:{candidate.document_id}")
        session.add(task); session.flush()
        session.add(AuditLog(patient_id=candidate.patient_id, actor=actor, actor_role="health_manager", action="created_followup_from_report", entity_type="ReportExtractionCandidate", entity_id=str(candidate.id), detail_json={"task_id": str(task.id)}))
        return task
